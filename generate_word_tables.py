"""
Generate SCI Paper Tables as Word Documents (.docx)

Creates 5 statistical tables organized in folders:
  data/output/word_tables/
  ├── table1_odi_summary/
  │   └── Table1_ODI_Statistical_Summary.docx
  ├── table2_weight_comparison/
  │   └── Table2_Weight_Comparison.docx
  ├── table3_error_analysis/
  │   └── Table3_Error_Analysis.docx
  ├── table4_parameter_range/
  │   └── Table4_Parameter_Range.docx
  └── table5_level_zoning/
      └── Table5_ODI_Level_Zoning.docx
"""

import json
import os
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

DEMO_DIR = Path(r"D:\xiangmu\miningplan\mining-plan\frontend\public\demo")
OUT_BASE = Path(r"D:\xiangmu\miningplan\data\output\word_tables")

DEMO_FILES = sorted(DEMO_DIR.glob("*.miningplan.json"))

DEMO_NAMES = {
    "0-地表下沉.miningplan": "Case 0: Surface Subsidence",
    "1-含水层扰动预评价.miningplan": "Case 1: Aquifer Pre-Evaluation",
    "2-含水层扰动评价.miningplan": "Case 2: Aquifer Disturbance Eval",
    "3-采区规划案例.miningplan": "Case 3: Mining Area Planning",
    "4-协同调控-突水点.miningplan": "Case 4: Coordinated Control",
    "5-采掘接续.miningplan": "Case 5: Mining Succession",
    "6-全覆岩扰动.miningplan": "Case 6: Full Overburden Disturbance",
}

DEMO_SHORT = {
    "0-地表下沉.miningplan": "Case 0",
    "1-含水层扰动预评价.miningplan": "Case 1",
    "2-含水层扰动评价.miningplan": "Case 2",
    "3-采区规划案例.miningplan": "Case 3",
    "4-协同调控-突水点.miningplan": "Case 4",
    "5-采掘接续.miningplan": "Case 5",
    "6-全覆岩扰动.miningplan": "Case 6",
}

LEVEL_LABELS = ["I (Stable)", "II (Slight)", "III (Moderate)", "IV (Strong)", "V (Severe)"]
PARAM_KEYS = ["Ti", "Hi", "Di", "Mi"]
PARAM_NAMES = {
    "Ti": "Stratum thickness $T_i$ (m)",
    "Hi": "Distance to aquifer $H_i$ (m)",
    "Di": "Depth $D_i$ (m)",
    "Mi": "Mining height $M_i$ (m)",
}


# ═══════════════════════════════════════════════════════════
#  STYLE HELPERS
# ═══════════════════════════════════════════════════════════
def set_cell_font(cell, text, bold=False, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(str(text))
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shading_elem)


def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "0",
            qn("w:color"): "000000",
        })
        borders.append(el)
    tblPr.append(borders)


def add_table_title(doc, title, caption=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if caption:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run2 = p2.add_run(caption)
        run2.font.name = "Times New Roman"
        run2.font.size = Pt(8)
        run2.font.italic = True
        run2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def header_row(table, texts, row_idx=0):
    for i, t in enumerate(texts):
        set_cell_font(table.rows[row_idx].cells[i], t, bold=True, size=9)
        set_cell_shading(table.rows[row_idx].cells[i], "D9E2F3")


# ═══════════════════════════════════════════════════════════
#  DATA LOADING
# ═══════════════════════════════════════════════════════════
def load_all_demos():
    demos = []
    for fp in DEMO_FILES:
        with open(fp, "r", encoding="utf-8") as f:
            data = json.load(f)
        demos.append((fp.name, data))
    return demos


def get_odi_points(data, tab_id):
    sp = data.get("scenarioParamsById", {})
    tab = sp.get(tab_id, {})
    odi_result = tab.get("odiResult")
    if odi_result and odi_result.get("points"):
        return odi_result["points"], odi_result
    if tab_id == "aquifer":
        cc = data.get("cocontrol", {})
        union = cc.get("results", {}).get("odiUnionResult")
        if union and union.get("points"):
            return union["points"], union
    return [], {}


def get_param_points(data, tab_id):
    sp = data.get("scenarioParamsById", {})
    tab = sp.get(tab_id, {})
    pr = tab.get("paramExtractionResult")
    return pr.get("points", []) if pr else []


def get_active_tab(data):
    sp = data.get("scenarioParamsById", {})
    for tab_id in ["surface", "aquifer", "upward", "full"]:
        td = sp.get(tab_id, {})
        if not isinstance(td, dict):
            continue
        odi_pts, _ = get_odi_points(data, tab_id)
        if odi_pts:
            return tab_id
    return None


def get_level_ranges(td):
    mzr = td.get("measuredZoningResult")
    if mzr and mzr.get("bins") and len(mzr["bins"]) == 5:
        return [(float(b["odiLo"]), float(b["odiHi"])) for b in mzr["bins"]]
    return [(0.0, 0.2), (0.2, 0.4), (0.4, 0.6), (0.6, 0.8), (0.8, 1.0)]


# ═══════════════════════════════════════════════════════════
#  TABLE 1: ODI STATISTICAL SUMMARY
# ═══════════════════════════════════════════════════════════
def generate_table1(demos):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(9)

    add_table_title(doc, "Table 1", "Statistical summary of ODI values across all engineering cases.")

    headers = ["Case", "Tab", "N", "Min", "Max", "Mean", "Std", "P50", "P90", "Skewness"]
    ncols = len(headers)
    nrows = 1  # header

    rows_data = []
    for fn, data in demos:
        sp = data.get("scenarioParamsById", {})
        for tab_id in ["surface", "aquifer", "upward", "full"]:
            td = sp.get(tab_id, {})
            if not isinstance(td, dict):
                continue
            odi_pts, _ = get_odi_points(data, tab_id)
            if not odi_pts:
                continue
            vals = [p.get("odiNorm", p.get("odi", 0)) for p in odi_pts]
            vals = [v for v in vals if v is not None and np.isfinite(v)]
            if not vals:
                continue

            rows_data.append([
                DEMO_SHORT.get(fn, fn),
                tab_id.capitalize(),
                len(vals),
                f"{min(vals):.4f}",
                f"{max(vals):.4f}",
                f"{np.mean(vals):.4f}",
                f"{np.std(vals):.4f}",
                f"{np.median(vals):.4f}",
                f"{np.percentile(vals, 90):.4f}",
                f"{float(np.mean((vals - np.mean(vals))**3) / np.std(vals)**3):.3f}" if np.std(vals) > 0 else "N/A",
            ])

    table = doc.add_table(rows=1 + len(rows_data), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table)
    header_row(table, headers)

    for i, row in enumerate(rows_data):
        for j, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_font(table.rows[i + 1].cells[j], val, size=9, alignment=align)

    out_dir = OUT_BASE / "table1_odi_summary"
    out_dir.mkdir(parents=True, exist_ok=True)
    doc.save(out_dir / "Table1_ODI_Statistical_Summary.docx")
    print(f"  [OK] Table1_ODI_Statistical_Summary.docx ({len(rows_data)} rows)")


# ═══════════════════════════════════════════════════════════
#  TABLE 2: WEIGHT COMPARISON
# ═══════════════════════════════════════════════════════════
def generate_table2(demos):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(9)

    add_table_title(doc, "Table 2",
                    "Comparison of ODI weight coefficients (geological $w_d$, mining $w_o$, composite $w_f$) across cases.")

    headers = ["Case", "Tab", "N points", "$w_d$ (Geology)", "$w_o$ (Mining)", "$w_f$ (Composite)"]
    ncols = len(headers)

    rows_data = []
    for fn, data in demos:
        sp = data.get("scenarioParamsById", {})
        for tab_id in ["surface", "aquifer", "upward", "full"]:
            td = sp.get(tab_id, {})
            if not isinstance(td, dict):
                continue

            odi_pts, odi_result = get_odi_points(data, tab_id)
            weights = odi_result.get("weights", {})

            if not weights:
                continue

            wd = weights.get("wd")
            wo = weights.get("wo")
            wf = weights.get("wf")

            if isinstance(wd, (int, float)) and isinstance(wo, (int, float)) and isinstance(wf, (int, float)):
                rows_data.append([
                    DEMO_SHORT.get(fn, fn),
                    tab_id.capitalize(),
                    len(odi_pts),
                    f"{wd:.2f}",
                    f"{wo:.2f}",
                    f"{wf:.2f}",
                ])
            elif isinstance(wd, dict):
                # Dict format: show as detailed sub-weights
                rows_data.append([
                    DEMO_SHORT.get(fn, fn),
                    tab_id.capitalize(),
                    len(odi_pts),
                    str({k: f"{v:.2f}" for k, v in wd.items()}) if isinstance(wd, dict) else str(wd),
                    str({k: f"{v:.2f}" for k, v in wo.items()}) if isinstance(wo, dict) else str(wo),
                    str({k: f"{v:.2f}" for k, v in wf.items()}) if isinstance(wf, dict) else str(wf),
                ])

    table = doc.add_table(rows=1 + len(rows_data), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table)
    header_row(table, headers)

    for i, row in enumerate(rows_data):
        for j, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if j <= 1 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_font(table.rows[i + 1].cells[j], val, size=9, alignment=align)

    out_dir = OUT_BASE / "table2_weight_comparison"
    out_dir.mkdir(parents=True, exist_ok=True)
    doc.save(out_dir / "Table2_Weight_Comparison.docx")
    print(f"  [OK] Table2_Weight_Comparison.docx ({len(rows_data)} rows)")


# ═══════════════════════════════════════════════════════════
#  TABLE 3: ERROR ANALYSIS
# ═══════════════════════════════════════════════════════════
def generate_table3(demos):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(9)

    add_table_title(doc, "Table 3",
                    "Error analysis between ODI predictions and measured values (Case 0, Surface Subsidence).")

    headers = ["Survey Line", "N", "MAE (m)", "RMSE (m)", "Mean Error Ratio", "Max Error Ratio",
               "Measured Range (m)", "ODI Renorm Range"]
    ncols = len(headers)

    rows_data = []
    for fn, data in demos:
        sp = data.get("scenarioParamsById", {})
        for tab_id in sp:
            td = sp[tab_id]
            if not isinstance(td, dict):
                continue
            err = td.get("errorAnalysisByLineId", {})
            if not err:
                continue

            line_idx = 1
            for lk, ld in err.items():
                err_pts = ld.get("data", [])
                if not err_pts:
                    continue

                measured = [d.get("measured", 0) for d in err_pts]
                odi_re = [d.get("odiRenorm", 0) for d in err_pts]
                ratios = [abs(d.get("errorRatioChart", d.get("errorRatio", 0))) for d in err_pts]

                ss_res = sum((m - o) ** 2 for m, o in zip(measured, odi_re))
                ss_tot = sum((m - np.mean(measured)) ** 2 for m in measured)
                r2 = 1 - ss_res / ss_tot if ss_tot > 0 else 0

                rmse = np.sqrt(np.mean([(m - o) ** 2 for m, o in zip(measured, odi_re)]))
                mae = np.mean([abs(m - o) for m, o in zip(measured, odi_re)])

                rows_data.append([
                    f"Line {line_idx}",
                    len(err_pts),
                    f"{mae:.4f}",
                    f"{rmse:.4f}",
                    f"{np.mean(ratios):.4f}",
                    f"{max(ratios):.4f}",
                    f"[{min(measured):.3f}, {max(measured):.3f}]",
                    f"[{min(odi_re):.3f}, {max(odi_re):.3f}]",
                ])
                line_idx += 1

    table = doc.add_table(rows=1 + len(rows_data), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table)
    header_row(table, headers)

    for i, row in enumerate(rows_data):
        for j, val in enumerate(row):
            set_cell_font(table.rows[i + 1].cells[j], val, size=9)

    out_dir = OUT_BASE / "table3_error_analysis"
    out_dir.mkdir(parents=True, exist_ok=True)
    doc.save(out_dir / "Table3_Error_Analysis.docx")
    print(f"  [OK] Table3_Error_Analysis.docx ({len(rows_data)} rows)")


# ═══════════════════════════════════════════════════════════
#  TABLE 4: PARAMETER RANGE
# ═══════════════════════════════════════════════════════════
def generate_table4(demos):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(9)

    add_table_title(doc, "Table 4",
                    "Statistical summary of geological and mining parameters across all cases.")

    # Build a table with: Case | Param | Min | Max | Mean | Std
    headers = ["Case", "Tab", "Parameter", "Min (m)", "Max (m)", "Mean (m)", "Std (m)"]
    ncols = len(headers)

    rows_data = []
    for fn, data in demos:
        sp = data.get("scenarioParamsById", {})
        for tab_id in ["surface", "aquifer", "upward", "full"]:
            td = sp.get(tab_id, {})
            if not isinstance(td, dict):
                continue
            param_pts = get_param_points(data, tab_id)
            if not param_pts:
                continue

            for key in PARAM_KEYS:
                vals = [p.get(key, 0) for p in param_pts]
                vals = [v for v in vals if v is not None]
                if not vals or np.std(vals) < 1e-10:
                    continue
                rows_data.append([
                    DEMO_SHORT.get(fn, fn),
                    tab_id.capitalize(),
                    f"{key} ({PARAM_NAMES.get(key, key).split('$')[0].strip()})",
                    f"{min(vals):.2f}",
                    f"{max(vals):.2f}",
                    f"{np.mean(vals):.2f}",
                    f"{np.std(vals):.2f}",
                ])

    table = doc.add_table(rows=1 + len(rows_data), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table)
    header_row(table, headers)

    for i, row in enumerate(rows_data):
        for j, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if j <= 2 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_font(table.rows[i + 1].cells[j], val, size=9, alignment=align)

    out_dir = OUT_BASE / "table4_parameter_range"
    out_dir.mkdir(parents=True, exist_ok=True)
    doc.save(out_dir / "Table4_Parameter_Range.docx")
    print(f"  [OK] Table4_Parameter_Range.docx ({len(rows_data)} rows)")


# ═══════════════════════════════════════════════════════════
#  TABLE 5: ODI LEVEL ZONING
# ═══════════════════════════════════════════════════════════
def generate_table5(demos):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(9)

    add_table_title(doc, "Table 5",
                    "ODI level zoning distribution across all engineering cases (Levels I-V).")

    headers = ["Case", "Tab", "N",
               "I (Stable)", "II (Slight)", "III (Moderate)", "IV (Strong)", "V (Severe)"]
    ncols = len(headers)

    rows_data = []
    for fn, data in demos:
        sp = data.get("scenarioParamsById", {})
        for tab_id in ["surface", "aquifer", "upward", "full"]:
            td = sp.get(tab_id, {})
            if not isinstance(td, dict):
                continue

            odi_pts, _ = get_odi_points(data, tab_id)
            if not odi_pts:
                continue

            ranges = get_level_ranges(td)
            vals = [p.get("odiNorm", 0) for p in odi_pts]
            vals = [v for v in vals if np.isfinite(v)]
            total = len(vals)

            counts = []
            for i, (lo, hi) in enumerate(ranges):
                inc = (i == 4)
                if inc:
                    c = sum(1 for v in vals if lo <= v <= hi)
                else:
                    c = sum(1 for v in vals if lo <= v < hi)
                counts.append(c)

            pcts = [f"{c / total * 100:.1f}%" for c in counts]
            level_strs = [f"{c} ({p})" for c, p in zip(counts, pcts)]

            rows_data.append([
                DEMO_SHORT.get(fn, fn),
                tab_id.capitalize(),
                str(total),
            ] + level_strs)

    table = doc.add_table(rows=1 + len(rows_data), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table)
    header_row(table, headers)

    level_colors = ["C6EFCE", "FFEB9C", "FCE4D6", "F8CBAD", "FFC7CE"]
    for i, row in enumerate(rows_data):
        for j, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if j <= 1 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_font(table.rows[i + 1].cells[j], val, size=9, alignment=align)
            if j >= 3:
                set_cell_shading(table.rows[i + 1].cells[j], level_colors[j - 3])

    out_dir = OUT_BASE / "table5_level_zoning"
    out_dir.mkdir(parents=True, exist_ok=True)
    doc.save(out_dir / "Table5_ODI_Level_Zoning.docx")
    print(f"  [OK] Table5_ODI_Level_Zoning.docx ({len(rows_data)} rows)")


# ═══════════════════════════════════════════════════════════
#  MAIN
# ═══════════════════════════════════════════════════════════
def main():
    OUT_BASE.mkdir(parents=True, exist_ok=True)
    demos = load_all_demos()
    print(f"Loaded {len(demos)} demo files\n")

    generate_table1(demos)
    generate_table2(demos)
    generate_table3(demos)
    generate_table4(demos)
    generate_table5(demos)

    print(f"\nAll tables saved to: {OUT_BASE}")


if __name__ == "__main__":
    main()
