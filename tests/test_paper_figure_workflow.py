from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
import matplotlib.pyplot as plt

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

import paper_figure_workflow as workflow


def write_demo_png(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(2, 1.2))
    ax.plot([0, 1], [0, 1], color="#2563eb", linewidth=2)
    ax.set_axis_off()
    fig.savefig(path, format="png", dpi=120, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def test_apply_boundary_override_replaces_aquifer_boundary() -> None:
    scene = {
        "scenarioParamsById": {
            "aquifer": {
                "boundaryData": [
                    {"id": "old-1", "x": 1.0, "y": 1.0},
                    {"id": "old-2", "x": 2.0, "y": 2.0},
                ]
            },
            "surface": {"boundaryData": []},
        }
    }
    boundary = [
        {"id": "B-1", "x": 10.0, "y": 20.0},
        {"id": "B-2", "x": 30.0, "y": 40.0},
        {"id": "B-3", "x": 50.0, "y": 60.0},
    ]

    updated = workflow.apply_boundary_override(scene, boundary, tab_ids=("aquifer",))

    assert updated["scenarioParamsById"]["aquifer"]["boundaryData"] == boundary
    assert updated["scenarioParamsById"]["surface"]["boundaryData"] == []


def test_insert_figures_into_docx_writes_image_and_caption(tmp_path: Path) -> None:
    source_docx = tmp_path / "source.docx"
    target_docx = tmp_path / "target.docx"
    image_path = tmp_path / "demo.png"
    write_demo_png(image_path)

    doc = Document()
    doc.add_heading("测试文档", level=1)
    doc.add_paragraph("这是引言段。")
    doc.add_paragraph("这里是锚点段落。")
    doc.save(source_docx)

    placements = [
        workflow.FigurePlacement(
            figure_number=1,
            title="测试插图",
            image_path=image_path,
            anchor_text="这里是锚点段落。",
            anchor_mode="after",
        )
    ]

    output_path = workflow.insert_figures_into_docx(source_docx, target_docx, placements)

    assert output_path == target_docx
    out_doc = Document(output_path)
    assert len(out_doc.inline_shapes) == 1
    assert any(p.text.strip() == "图1 测试插图" for p in out_doc.paragraphs)


def test_insert_figures_removes_legacy_placeholder_and_previous_drawing(tmp_path: Path) -> None:
    source_docx = tmp_path / "legacy_source.docx"
    target_docx = tmp_path / "legacy_target.docx"
    image_path = tmp_path / "legacy.png"
    write_demo_png(image_path)

    doc = Document()
    doc.add_paragraph("前文段落。")
    legacy_para = doc.add_paragraph()
    legacy_para.add_run().add_picture(str(image_path))
    doc.add_paragraph("图5 多场景ODI风险分布结果图")
    doc.add_paragraph("锚点段落。")
    doc.save(source_docx)

    placements = [
        workflow.FigurePlacement(
            figure_number=4,
            title="多场景ODI分布结果对比图",
            image_path=image_path,
            anchor_text="锚点段落。",
            anchor_mode="after",
        )
    ]

    output_path = workflow.insert_figures_into_docx(source_docx, target_docx, placements)

    out_doc = Document(output_path)
    texts = [p.text.strip() for p in out_doc.paragraphs if p.text.strip()]
    assert "图5 多场景ODI风险分布结果图" not in texts
    assert "图4 多场景ODI分布结果对比图" in texts
    assert len(out_doc.inline_shapes) == 1


def test_four_panel_compare_keeps_each_panel_extent_independent() -> None:
    short_points = [
        {"x": 0.0, "y": 0.0, "odiNorm": 0.1},
        {"x": 10.0, "y": 0.0, "odiNorm": 0.3},
        {"x": 0.0, "y": 8.0, "odiNorm": 0.6},
        {"x": 10.0, "y": 8.0, "odiNorm": 0.8},
    ]
    long_points = [
        {"x": 0.0, "y": 0.0, "odiNorm": 0.1},
        {"x": 80.0, "y": 0.0, "odiNorm": 0.3},
        {"x": 0.0, "y": 40.0, "odiNorm": 0.6},
        {"x": 80.0, "y": 40.0, "odiNorm": 0.8},
    ]

    fig = workflow.plot_multi_scenario_odi_compare_four_panels(
        [
            ("(a)", short_points, [], []),
            ("(b)", long_points, [], []),
            ("(c)", long_points, [], []),
            ("(d)", short_points, [], []),
        ]
    )

    axes = fig.axes[:4]
    ax0_xlim = axes[0].get_xlim()
    ax3_xlim = axes[3].get_xlim()
    assert ax0_xlim[1] <= 11.0
    assert ax3_xlim[1] <= 11.0
    plt.close(fig)
