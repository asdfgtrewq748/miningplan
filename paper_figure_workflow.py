from __future__ import annotations

import argparse
import copy
import csv
import json
import shutil
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

import export_scene_visuals as exporter


@dataclass(frozen=True)
class FigurePlacement:
    figure_number: int
    title: str
    image_path: Path
    anchor_text: str
    anchor_mode: str = "after"
    width_cm: float = 15.8
    caption: str | None = None


@dataclass(frozen=True)
class PaperFigureAsset:
    figure_number: int
    title: str
    stem: Path
    png_path: Path
    svg_path: Path | None = None
    pdf_path: Path | None = None


def build_default_output_dir(base_dir: Path | None = None) -> Path:
    root = Path(base_dir or Path.cwd())
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return root / "output" / "paper_insert_package" / stamp


def read_boundary_csv(csv_path: Path) -> list[dict]:
    with csv_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))
    boundary = []
    for index, row in enumerate(rows, start=1):
        x = float(row["x"])
        y = float(row["y"])
        boundary.append({"id": row.get("id") or f"B-{index}", "x": x, "y": y})
    return boundary


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def load_scene_by_prefix(scene_dir: Path, prefix: str) -> tuple[Path, dict]:
    path = next(scene_dir.glob(f"{prefix}-*.miningplan.json"))
    return path, load_json(path)


def apply_boundary_override(scene: dict, boundary: list[dict], tab_ids: Iterable[str] = ("aquifer",)) -> dict:
    updated = copy.deepcopy(scene)
    scenario_params = updated.get("scenarioParamsById") or {}
    boundary_copy = [dict(point) for point in boundary]
    for tab_id in tab_ids:
        tab = scenario_params.get(tab_id)
        if isinstance(tab, dict):
            tab["boundaryData"] = [dict(point) for point in boundary_copy]
    return updated


def _set_axis_titles(fig: plt.Figure, title: str = "") -> None:
    if getattr(fig, "_suptitle", None) is not None:
        fig._suptitle.set_text(title)
    for axis in fig.axes:
        axis.set_title(title)


def _save_figure_bundle(fig: plt.Figure | None, stem: Path, formats: tuple[str, ...] = ("svg", "pdf", "png")) -> PaperFigureAsset:
    written = exporter.save_figure(fig, stem, formats)
    png_path = stem.with_suffix(".png")
    svg_path = stem.with_suffix(".svg") if "svg" in formats else None
    pdf_path = stem.with_suffix(".pdf") if "pdf" in formats else None
    if png_path not in written:
        raise RuntimeError(f"Figure export failed for {stem}")
    return PaperFigureAsset(
        figure_number=0,
        title=stem.name,
        stem=stem,
        png_path=png_path,
        svg_path=svg_path if svg_path and svg_path.exists() else None,
        pdf_path=pdf_path if pdf_path and pdf_path.exists() else None,
    )


def _draw_boundary(ax: plt.Axes, boundary: list[dict], *, color: str = "#1d4ed8", linewidth: float = 1.2, alpha: float = 0.9) -> None:
    if not boundary:
        return
    xs = [float(point["x"]) for point in boundary] + [float(boundary[0]["x"])]
    ys = [float(point["y"]) for point in boundary] + [float(boundary[0]["y"])]
    ax.plot(xs, ys, color=color, linewidth=linewidth, alpha=alpha, zorder=3)


def _draw_drillholes(ax: plt.Axes, drillholes: list[dict], *, size: float = 20.0) -> None:
    if not drillholes:
        return
    ax.scatter(
        [float(point["x"]) for point in drillholes],
        [float(point["y"]) for point in drillholes],
        s=size,
        c="#111827",
        marker="o",
        zorder=4,
        label="钻孔",
    )


def _draw_plan_loops(ax: plt.Axes, plan_loops: list[dict], *, facecolor: str = "#f472b6", edgecolor: str = "#be185d", alpha: float = 0.16) -> None:
    for index, face in enumerate(plan_loops, start=1):
        loop = face.get("loop") or []
        if len(loop) < 3:
            continue
        xs = [float(point["x"]) for point in loop] + [float(loop[0]["x"])]
        ys = [float(point["y"]) for point in loop] + [float(loop[0]["y"])]
        ax.fill(xs, ys, color=facecolor, alpha=alpha, zorder=5)
        ax.plot(xs, ys, color=edgecolor, linewidth=1.4, zorder=6)
        cx = sum(float(point["x"]) for point in loop) / len(loop)
        cy = sum(float(point["y"]) for point in loop) / len(loop)
        ax.text(cx, cy, f"No.{face.get('faceIndex', index)}", ha="center", va="center", fontsize=8, color=edgecolor, zorder=7)


def plot_mechanism_schematic() -> plt.Figure:
    fig, ax = plt.subplots(figsize=(9.2, 4.6))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    blocks = [
        (0.04, 0.2, 0.18, 0.6, "#e0f2fe", "输入层", ["采区边界", "钻孔样点", "分层与参数场"]),
        (0.30, 0.2, 0.18, 0.6, "#dbeafe", "约束层", ["多场景风险", "ODI 统一表达", "阈值与分级"]),
        (0.56, 0.2, 0.18, 0.6, "#dcfce7", "方案生成层", ["候选布局生成", "多目标评价", "模式化比选"]),
        (0.82, 0.2, 0.14, 0.6, "#fee2e2", "输出层", ["规划布局", "采掘接续", "经济评价"]),
    ]
    for x, y, w, h, color, title, lines in blocks:
        rect = plt.Rectangle((x, y), w, h, facecolor=color, edgecolor="#334155", linewidth=1.2)
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h - 0.08, title, ha="center", va="center", fontsize=11, fontweight="bold", color="#0f172a")
        for idx, line in enumerate(lines):
            ax.text(x + w / 2, y + h - 0.19 - idx * 0.13, line, ha="center", va="center", fontsize=9, color="#1e293b")

    arrow_specs = [
        (0.22, 0.5, 0.30, 0.5, "地质输入转为可计算约束"),
        (0.48, 0.5, 0.56, 0.5, "风险前置进入规划生成"),
        (0.74, 0.5, 0.82, 0.5, "布局结果继续向下游传递"),
    ]
    for x1, y1, x2, y2, text in arrow_specs:
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1), arrowprops={"arrowstyle": "->", "linewidth": 1.5, "color": "#475569"})
        ax.text((x1 + x2) / 2, y1 + 0.07, text, ha="center", va="bottom", fontsize=8.5, color="#475569")

    ax.text(0.5, 0.07, "规划变量通过扰动响应形成约束反馈，约束结果再反向作用于布局、接续与经济决策。", ha="center", va="center", fontsize=9.2, color="#334155")
    fig.tight_layout()
    return fig


def plot_input_map(drillholes: list[dict], boundary: list[dict]) -> plt.Figure:
    fig, ax = plt.subplots(figsize=(7.2, 5.6))
    _draw_boundary(ax, boundary, color="#1e3a8a", linewidth=1.5)
    _draw_drillholes(ax, drillholes, size=22.0)
    exporter.finalize_map_axes(ax, "")
    ax.legend(loc="lower left", frameon=False)
    fig.tight_layout()
    return fig


def plot_parameter_fields(param_points: list[dict], drillholes: list[dict], boundary: list[dict]) -> plt.Figure:
    fig, axes = plt.subplots(2, 2, figsize=(8.4, 7.0))
    keys = ("Ti", "Hi", "Di", "Mi")
    labels = ("(a) 目标层厚度 Ti", "(b) 层间距 Hi", "(c) 埋深 Di", "(d) 煤层厚度 Mi")
    cmap = exporter.viridis_cmap()
    for axis, key, label in zip(axes.flat, keys, labels):
        grid = exporter.interpolate_grid(param_points, key, resolution=120)
        if grid is None:
            axis.text(0.5, 0.5, "数据不足", ha="center", va="center", transform=axis.transAxes)
            axis.set_axis_off()
            continue
        xi_grid, yi_grid, zi = grid
        contour = axis.contourf(xi_grid, yi_grid, zi, levels=12, cmap=cmap, antialiased=True)
        _draw_boundary(axis, boundary, color="#94a3b8", linewidth=0.9, alpha=0.85)
        _draw_drillholes(axis, drillholes, size=14.0)
        axis.set_title(label, fontsize=9, fontweight="bold")
        axis.set_aspect("equal", adjustable="box")
        axis.grid(color="#e2e8f0", linewidth=0.4, linestyle=":", alpha=0.7)
        fig.colorbar(contour, ax=axis, pad=0.01, shrink=0.82)
    fig.tight_layout()
    return fig


def plot_multi_scenario_odi_compare(panels: list[tuple[str, list[dict], list[dict], list[dict]]], boundary: list[dict]) -> plt.Figure:
    fig, axes = plt.subplots(1, 3, figsize=(12.8, 4.2))
    levels = np.linspace(0.0, 1.0, 11)
    contour_ref = None
    for axis, (label, odi_points, drillholes, boundary_override) in zip(axes, panels):
        grid = exporter.interpolate_grid(odi_points, "odiNorm", resolution=130)
        if grid is None:
            axis.text(0.5, 0.5, "数据不足", ha="center", va="center", transform=axis.transAxes)
            axis.set_axis_off()
            continue
        xi_grid, yi_grid, zi = grid
        contour_ref = axis.contourf(xi_grid, yi_grid, zi, levels=levels, cmap=exporter.blue_red_cmap(), antialiased=True)
        _draw_boundary(axis, boundary_override or boundary, color="#1e3a8a", linewidth=1.0, alpha=0.85)
        _draw_drillholes(axis, drillholes, size=12.0)
        axis.set_title(label, fontsize=10, fontweight="bold")
        axis.set_xlabel("X / m")
        axis.set_ylabel("Y / m")
        axis.set_aspect("equal", adjustable="box")
        axis.grid(color="#e2e8f0", linewidth=0.4, linestyle=":", alpha=0.7)
    cbar = fig.colorbar(contour_ref, ax=axes, fraction=0.025, pad=0.02)
    cbar.set_label("ODI（归一化）")
    fig.tight_layout()
    return fig


def plot_layout_overlay(odi_points: list[dict], drillholes: list[dict], boundary: list[dict], plan_loops: list[dict]) -> plt.Figure:
    fig, ax = plt.subplots(figsize=(7.6, 5.8))
    grid = exporter.interpolate_grid(odi_points, "odiNorm", resolution=140)
    if grid is not None:
        xi_grid, yi_grid, zi = grid
        ax.contourf(xi_grid, yi_grid, zi, levels=np.linspace(0.0, 1.0, 11), cmap=exporter.blue_red_cmap(), antialiased=True)
    _draw_boundary(ax, boundary, color="#64748b", linewidth=0.9, alpha=0.9)
    _draw_drillholes(ax, drillholes, size=18.0)
    _draw_plan_loops(ax, plan_loops, facecolor="#fda4af", edgecolor="#db2777", alpha=0.18)
    exporter.finalize_map_axes(ax, "")
    fig.tight_layout()
    return fig


def plot_layout_only(plan_loops: list[dict], boundary: list[dict]) -> plt.Figure:
    fig, ax = plt.subplots(figsize=(7.2, 5.4))
    _draw_boundary(ax, boundary, color="#1e3a8a", linewidth=1.4, alpha=0.9)
    palette = ["#bfdbfe", "#93c5fd", "#60a5fa", "#3b82f6", "#1d4ed8"]
    for index, face in enumerate(plan_loops, start=1):
        loop = face.get("loop") or []
        if len(loop) < 3:
            continue
        xs = [float(point["x"]) for point in loop] + [float(loop[0]["x"])]
        ys = [float(point["y"]) for point in loop] + [float(loop[0]["y"])]
        color = palette[(index - 1) % len(palette)]
        ax.fill(xs, ys, color=color, alpha=0.32, zorder=3)
        ax.plot(xs, ys, color="#1e3a8a", linewidth=1.0, zorder=4)
        cx = sum(float(point["x"]) for point in loop) / len(loop)
        cy = sum(float(point["y"]) for point in loop) / len(loop)
        ax.text(cx, cy, f"No.{face.get('faceIndex', index)}", ha="center", va="center", fontsize=8, color="#0f172a", zorder=5)
    exporter.finalize_map_axes(ax, "")
    fig.tight_layout()
    return fig


def _retitle_existing_figure(fig: plt.Figure | None) -> plt.Figure | None:
    if fig is None:
        return None
    _set_axis_titles(fig, "")
    fig.tight_layout()
    return fig


def build_main_paper_figures(
    scene_dir: Path,
    boundary_csv_path: Path,
    output_dir: Path,
    formats: tuple[str, ...] = ("svg", "pdf", "png"),
) -> list[PaperFigureAsset]:
    boundary = read_boundary_csv(boundary_csv_path)
    _, surface_scene = load_scene_by_prefix(scene_dir, "0")
    _, aquifer_scene = load_scene_by_prefix(scene_dir, "2")
    _, planning_scene_raw = load_scene_by_prefix(scene_dir, "3")
    _, succession_scene_raw = load_scene_by_prefix(scene_dir, "5")
    _, full_scene = load_scene_by_prefix(scene_dir, "6")

    planning_scene = apply_boundary_override(planning_scene_raw, boundary, tab_ids=("aquifer",))
    succession_scene = apply_boundary_override(succession_scene_raw, boundary, tab_ids=("aquifer",))
    aquifer_scene = apply_boundary_override(aquifer_scene, boundary, tab_ids=("aquifer",))
    surface_scene = apply_boundary_override(surface_scene, boundary, tab_ids=("surface",))
    full_scene = apply_boundary_override(full_scene, boundary, tab_ids=("full",))

    output_dir.mkdir(parents=True, exist_ok=True)
    assets: list[PaperFigureAsset] = []

    def save(number: int, title: str, fig: plt.Figure | None) -> None:
        stem = output_dir / f"{number:02d}-{title}"
        exporter.save_figure(fig, stem, formats)
        assets.append(
            PaperFigureAsset(
                figure_number=number,
                title=title,
                stem=stem,
                png_path=stem.with_suffix(".png"),
                svg_path=stem.with_suffix(".svg") if "svg" in formats else None,
                pdf_path=stem.with_suffix(".pdf") if "pdf" in formats else None,
            )
        )

    planning_aquifer = (planning_scene.get("scenarioParamsById") or {}).get("aquifer") or {}
    aquifer_tab = (aquifer_scene.get("scenarioParamsById") or {}).get("aquifer") or {}
    surface_tab = (surface_scene.get("scenarioParamsById") or {}).get("surface") or {}
    full_tab = (full_scene.get("scenarioParamsById") or {}).get("full") or {}
    plan_loops = (planning_scene.get("workfacePlan") or {}).get("plannedWorkfaceLoopsWorld") or []

    save(1, "覆岩扰动约束机理示意图", plot_mechanism_schematic())
    save(2, "研究区钻孔空间分布图", plot_input_map(list(planning_aquifer.get("drillholeData") or []), boundary))
    save(
        3,
        "研究区主要地质参数场分布图",
        plot_parameter_fields(
            exporter.get_param_points(planning_scene, "aquifer"),
            list(planning_aquifer.get("drillholeData") or []),
            boundary,
        ),
    )
    save(
        4,
        "多场景ODI分布结果对比图",
        plot_multi_scenario_odi_compare(
            [
                ("(a) 地表沉陷场景", exporter.get_odi_points(surface_scene, "surface"), list(surface_tab.get("drillholeData") or []), list(surface_tab.get("boundaryData") or []) or boundary),
                ("(b) 含水层扰动场景", exporter.get_odi_points(aquifer_scene, "aquifer"), list(aquifer_tab.get("drillholeData") or []), boundary),
                ("(c) 全覆岩扰动场景", exporter.get_odi_points(full_scene, "full"), list(full_tab.get("drillholeData") or []), list(full_tab.get("boundaryData") or []) or boundary),
            ],
            boundary,
        ),
    )
    save(
        5,
        "含水层扰动约束下最终规划布局与ODI场叠置图",
        plot_layout_overlay(
            exporter.get_odi_points(planning_scene, "aquifer"),
            list(planning_aquifer.get("drillholeData") or []),
            boundary,
            plan_loops,
        ),
    )
    save(6, "采区规划布局结果图", plot_layout_only(plan_loops, boundary))
    save(7, "四模式规划指标对比图", _retitle_existing_figure(exporter.plot_planning_mode_scores(planning_scene, "")))

    succession_artifacts = exporter.compute_succession_artifacts(succession_scene)
    if succession_artifacts is None:
        raise RuntimeError("采掘接续场景未生成接续结果，无法输出图8-图10。")
    plan = succession_artifacts["plan"]
    target = succession_artifacts["targetTonsPerMonth"]
    save(8, "采掘接续时序安排图", _retitle_existing_figure(exporter.plot_schedule_gantt(plan, "")))
    save(9, "月产量变化曲线图", _retitle_existing_figure(exporter.plot_monthly_production(plan, target, "")))
    economics_result = exporter.compute_economics_from_plan(plan, None, succession_scene.get("economicsParams") or {})
    save(10, "采掘接续方案现金流分析图", _retitle_existing_figure(exporter.plot_cashflow(economics_result, "")))

    manifest = {
        "generatedAt": datetime.now().isoformat(),
        "boundaryCsv": str(boundary_csv_path),
        "sceneDir": str(scene_dir),
        "figures": [
            {
                "number": asset.figure_number,
                "title": asset.title,
                "png": str(asset.png_path),
                "svg": str(asset.svg_path) if asset.svg_path else "",
                "pdf": str(asset.pdf_path) if asset.pdf_path else "",
            }
            for asset in assets
        ],
    }
    (output_dir / "paper_figures_manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return assets


def _find_paragraph(document: Document, anchor_text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == anchor_text.strip():
            return paragraph
    for paragraph in document.paragraphs:
        if anchor_text.strip() in paragraph.text.strip():
            return paragraph
    raise ValueError(f"未找到锚点段落：{anchor_text}")


def _new_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _new_paragraph_before(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def _remove_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._element.getparent()
    parent.remove(paragraph._element)


def _set_caption_style(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if paragraph.runs:
        run = paragraph.runs[0]
    else:
        run = paragraph.add_run()
    run.font.size = Pt(10.5)
    run.font.bold = False


def insert_figures_into_docx(source_docx: Path, target_docx: Path, placements: list[FigurePlacement]) -> Path:
    shutil.copyfile(source_docx, target_docx)
    document = Document(str(target_docx))

    for placeholder in ("图3 数据-模型-决策分层架构图", "图5 多场景ODI风险分布结果图"):
        try:
            _remove_paragraph(_find_paragraph(document, placeholder))
        except ValueError:
            pass

    for placement in placements:
        anchor = _find_paragraph(document, placement.anchor_text)
        if placement.anchor_mode == "before":
            image_paragraph = _new_paragraph_before(anchor)
        else:
            image_paragraph = _new_paragraph_after(anchor)
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.add_run().add_picture(str(placement.image_path), width=Cm(placement.width_cm))

        caption_text = placement.caption or f"图{placement.figure_number} {placement.title}"
        caption_paragraph = _new_paragraph_after(image_paragraph)
        caption_paragraph.text = caption_text
        _set_caption_style(caption_paragraph)

    document.save(str(target_docx))
    return target_docx


def build_default_placements(figure_assets: list[PaperFigureAsset]) -> list[FigurePlacement]:
    figure_by_number = {asset.figure_number: asset for asset in figure_assets}
    return [
        FigurePlacement(1, figure_by_number[1].title, figure_by_number[1].png_path, "采区规划的本质是在几何可实施性、资源回收与覆岩扰动控制之间寻求可接受折中。当采动引起的覆岩变形、裂隙发育和应力重分布超过安全阈值时，工作面布置范围、推进方式及回采参数将受到明显约束；反之，若一味采取保守布局，虽然有利于扰动控制，却可能压缩可采空间、降低资源回收水平并削弱生产组织效率。因此，采区规划并不是单纯的几何布局问题，而是一个受覆岩扰动约束驱动的多目标协同决策问题。图1可将该关系概括为“采区规划决策变量—覆岩扰动响应特征—规划决策约束结果”的传递链，即规划变量通过布置方式、尺度参数与推进组织影响覆岩扰动的空间分布、强度水平和演化特征，扰动响应再反向作用于空间约束、强度约束和过程约束，最终参与规划方案的形成与筛选。", width_cm=16.2),
        FigurePlacement(2, figure_by_number[2].title, figure_by_number[2].png_path, "研究区输入对象主要包括采区边界、钻孔样点和规划控制参数3类。其中，采区边界用于界定研究范围并提供后续有效布置域提取的几何基础；钻孔样点用于构建煤层厚度等连续参数场；规划控制参数则包括边界煤柱、区段煤柱、工作面尺度、推进方向及局部保护距离等约束条件。根据当前样例数据，研究区共布置15个钻孔样点，编号为 ZK-101—ZK-115，煤层厚度样本最小值为2.8 m，最大值为4.8 m，平均值为3.793 3 m。上述数据既构成连续参数场插值的样本基础，也为后续资源评价与候选方案比选提供底层依据。", width_cm=15.6),
        FigurePlacement(3, figure_by_number[3].title, figure_by_number[3].png_path, "采区规划首先需要解决离散地质信息向连续规划约束转换的问题。对于研究区样例而言，钻孔数据本质上属于有限采样点上的离散观测值，若不经过连续化处理，难以直接参与工作面布局、资源评价和风险筛选。为此，本文采用基于钻孔样点的空间插值方法，将煤层厚度等关键属性组织为规划域上的连续参数场，并以此作为候选方案生成与目标评价的共享底图。", width_cm=16.3),
        FigurePlacement(4, figure_by_number[4].title, figure_by_number[4].png_path, "从研究区样例结果看，基于 export_package 批量导出的多场景 ODI 结果已经能够形成相互对应的空间风险分布图件，并在同一研究区边界、钻孔参照和色标口径下进行表达。当前样例中，地表沉陷、含水层扰动、采掘接续相关风险以及全覆岩综合扰动等结果均可组织为 ODI 风险图。这说明 ODI 已不再停留于概念定义阶段，而是进入了可计算、可导出、可比较的规划应用阶段。对于论文论证而言，这组结果支撑的不是“存在多个风险场景”这一常识性结论，而是“多场景风险已被统一组织为规划链路中的同类对象”这一更关键的对象链结论。", width_cm=16.5),
        FigurePlacement(5, figure_by_number[5].title, figure_by_number[5].png_path, "在有效布置域、连续参数场和 ODI 风险场共同作用下，本文进一步开展候选方案生成与多模式规划比选。与传统只给出单一布局结果的做法不同，本文将候选方案组织为工程效率优先、资源回收优先、覆岩扰动优先和综合权衡优化4种规划模式。四种模式并非4套彼此独立的算法，而是面向同一候选方案池的不同评价视角，其目的在于为采区规划提供可比、可选和可解释的方案集合。", width_cm=15.8),
        FigurePlacement(6, figure_by_number[6].title, figure_by_number[6].png_path, "从布局结果本身看，工作面与巷道对象并不是直接贴附在原始边界外轮廓上的简单几何划分，而是在有效布置域内形成必要的保护距离和工程边界。这意味着采区边界在规划语境中不能简单理解为“外轮廓包络”，而应理解为“原始边界经煤柱约束、保护条件和几何合法性处理后的可布置范围”。换言之，规划模块不是直接在静态图形上作图，而是在约束条件、参数场和风险场共同作用下形成具有工程语义的空间对象集合。", width_cm=15.8),
        FigurePlacement(7, figure_by_number[7].title, figure_by_number[7].png_path, "从当前样例的验证口径看，本文更强调四模式规划“可形成、可比较、可解释”，而不是急于给出某一模式在真实矿井条件下绝对最优的结论。现阶段结果已经足以说明：在统一候选池上，通过调整工程效率、资源回收与扰动控制的权重，规划结果能够呈现出具有差异性的目标侧重和空间组织特征。这意味着所提方法不仅能够生成方案，还能够为不同目标偏好下的决策提供明确的比较框架。对于偏工程方法与应用验证的论文而言，这种“多方案可比较”的组织方式比单一静态方案更具说服力。", width_cm=15.4),
        FigurePlacement(8, figure_by_number[8].title, figure_by_number[8].png_path, "在采掘接续层面，规划阶段形成的工作面边界、巷道结构和推进关系可进一步映射为接续任务对象，并围绕产量、风险和工期构建候选接续方案。当前样例结果已经给出了由采区规划空间结果向采掘接续空间结果的传递路径，说明规划阶段形成的对象并不是静态终点，而是后续组织与评价的上游输入。相较于只在规划模块内停留的方案结果，这种“对象可传递、链路可延伸”的特征更符合工程实际中由布局设计向组织实施逐步推进的决策过程。对于论文论证而言，这一部分支撑了本文“规划—接续”一体化方法链的核心命题。", width_cm=16.0),
        FigurePlacement(9, figure_by_number[9].title, figure_by_number[9].png_path, "在调控层面，规划结果还可进一步转化为工作面级控制变量，并通过采高、工作面宽度或区段煤柱等参数的调整，观察 ODI 统计指标及风险暴露程度的变化。当前分析表明，在续采工作面调控过程中，采高降低能够显著压低 ODI 均值及高分位统计量，这说明规划结果不仅能够被动进入后续评价环节，还能够主动支撑参数调控与局部优化。由此，规划阶段形成的空间对象不再是不可更改的静态结果，而是可围绕风险控制目标进一步优化的决策基础。", width_cm=15.6),
        FigurePlacement(10, figure_by_number[10].title, figure_by_number[10].png_path, "在工程经济评价层面，规划与接续结果可继续进入收入、成本、风险联动成本和净现值分析过程，形成从月度净现金流到综合经济指标的闭环。当前样例虽然尚未形成真实矿井条件下的大样本经济对照结果，但已经建立了规划结果—接续组织—经济评价的传递逻辑，说明空间布局、风险约束与经济指标能够在统一链路内进行耦合。对于采区规划研究而言，这一点具有重要意义，因为它意味着规划结果不再只是局部几何解，而能够继续作为效益分析和方案比较的上游对象参与更长链条的工程决策。", width_cm=15.8),
    ]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build paper-ready figures with corrected boundary and insert them into the manuscript.")
    parser.add_argument("--scene-dir", default=str(Path("mining-plan") / "frontend" / "public" / "demo"), help="Scene JSON directory.")
    parser.add_argument("--boundary-csv", default=str(Path("data") / "采区边界_敏东.csv"), help="Correct boundary CSV path.")
    parser.add_argument("--paper-docx", required=True, help="Source paper DOCX path.")
    parser.add_argument("--output-dir", help="Output directory. Defaults to output/paper_insert_package/<timestamp>.")
    parser.add_argument("--formats", default="svg,pdf,png", help="Comma-separated export formats for paper figures.")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    output_dir = Path(args.output_dir) if args.output_dir else build_default_output_dir()
    output_dir.mkdir(parents=True, exist_ok=True)
    scene_dir = Path(args.scene_dir)
    boundary_csv = Path(args.boundary_csv)
    paper_docx = Path(args.paper_docx)
    formats = tuple(part.strip().lower() for part in args.formats.split(",") if part.strip())

    figure_dir = output_dir / "paper_main_figures"
    assets = build_main_paper_figures(scene_dir, boundary_csv, figure_dir, formats=formats)
    placements = build_default_placements(assets)
    inserted_docx = output_dir / f"{paper_docx.stem}_插图版.docx"
    insert_figures_into_docx(paper_docx, inserted_docx, placements)

    print(f"Output directory: {output_dir}")
    print(f"Paper figures: {figure_dir}")
    print(f"Inserted paper: {inserted_docx}")
    for asset in assets:
        print(f"  - 图{asset.figure_number}: {asset.png_path.name}")


if __name__ == "__main__":
    main()
