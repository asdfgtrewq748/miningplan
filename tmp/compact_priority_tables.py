import json
import math
import statistics
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


BASE = Path("D:/xiangmu/miningplan")
DOCX_PATH = BASE / "论文" / "重构工作区" / "06_投稿包" / "最新版论文4.16_插图版_第一优先级修订.docx"
DATA_DIR = BASE / "论文" / "重构工作区" / "05_支撑材料" / "接口结果"


def load_json_with_keys(*keys):
    for path in DATA_DIR.glob("*.json"):
        data = json.loads(path.read_text(encoding="utf-8"))
        if isinstance(data, dict) and all(k in data for k in keys):
            return data
    raise FileNotFoundError(f"No JSON with keys: {keys}")


def percentile(values, pct):
    vals = sorted(values)
    k = (len(vals) - 1) * pct / 100
    f = math.floor(k)
    c = math.ceil(k)
    return vals[f] * (c - k) + vals[c] * (k - f)


def style_cell(cell, text, font_size=8.5):
    cell.text = str(text)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
            run.font.name = "Times New Roman"


def replace_table(doc, old_table, rows, font_size=8.5):
    new_table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        new_table.style = old_table.style
    except Exception:
        pass
    new_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    new_table.autofit = True
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            style_cell(new_table.rows[r].cells[c], value, font_size=font_size)
    old_table._element.addnext(new_table._element)
    old_table._element.getparent().remove(old_table._element)


def main():
    design = load_json_with_keys("panels", "roadways", "designParams", "miningRules")
    odi = load_json_with_keys("field", "weights", "stats")

    boreholes = design["boreholes"]
    thickness = [float(b["coalThickness"]) for b in boreholes]
    design_params = design["designParams"]
    stats = design["stats"]
    roadways = design["roadways"]
    roadway_total = sum(float(r.get("length", 0)) for r in roadways)

    field = [float(x) for row in odi["field"] for x in row]
    odi_mean = statistics.mean(field)
    odi_p90 = percentile(field, 90)
    exceed_070 = sum(v > 0.70 for v in field) / len(field) * 100
    exceed_080 = sum(v > 0.80 for v in field) / len(field) * 100
    odi_short = f"均值{odi_mean:.4f}；P90={odi_p90:.4f}；ODI>0.70占{exceed_070:.2f}%"

    doc = Document(DOCX_PATH)

    table3_rows = [
        ["模式", "数据状态", "工程布置指标", "资源/覆盖指标", "ODI风险指标", "接续与经济解释"],
        ["工程效率优先", "当前接口未单独导出完整数值", "评价口径为连续性、巷道组织与推进长度", "未单独导出", "未单独导出", "作为后续模式对比口径，不写实证排序"],
        ["资源回收优先", "当前接口未单独导出完整数值", "未单独导出", f"煤厚均值{statistics.mean(thickness):.4f} m；以厚度场和可采面积为口径", "未单独导出", "当前仅有参数场支撑"],
        ["覆岩扰动优先", "当前接口未单独导出完整数值", "未单独导出", "未单独导出", odi_short, "当前仅有ODI场支撑"],
        [
            "综合权衡/当前样例输出",
            "已导出结构化样例",
            f"工作面{stats['count']}个；巷道{len(roadways)}条；巷道总长{roadway_total:.2f} m",
            f"布置面积{float(stats['totalArea']):.2f} m²；覆盖率69.17%",
            odi_short,
            f"平均评分{float(stats['avgScore']):.1f}；接续与NPV接口已接通但未导出数值",
        ],
    ]
    replace_table(doc, doc.tables[2], table3_rows, font_size=8.5)

    table4_rows = [
        ["调控对象", "当前样例/接口状态", "ODI统计", "资源指标", "可支持的分析", "本版处理"],
        ["采高", "接口支持，当前DOCX未导出取值序列", odi_short, "未单独导出", "可开展采高序列敏感性分析", "保留为后续参数，不写强结论"],
        ["工作面宽度", f"{design_params['workfaceWidth']:.1f} m", odi_short, "未单独导出", "可开展宽度序列对比", "作为当前样例参数说明"],
        ["区段煤柱宽度", f"{design_params['pillarWidth']:.1f} m", odi_short, "未单独导出", "可分析区段隔离与有效区域变化", "作为当前样例参数说明"],
        ["边界煤柱宽度", f"{design_params['boundaryMargin']:.1f} m", odi_short, "未单独导出", "可分析有效布置域内缩影响", "作为当前样例参数说明"],
        ["综合ODI场", "80×56网格，4480个栅格", f"{odi_short}；ODI>0.80占{exceed_080:.2f}%", "不适用", "支撑风险场描述与阈值统计", "不替代实矿参数标定"],
    ]
    replace_table(doc, doc.tables[3], table4_rows, font_size=8.5)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
