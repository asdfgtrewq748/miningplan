from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement


def insert_paragraph_after(paragraph, text: str = "", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para


def normalize_para(para, font_size: float = 10.5) -> None:
    for run in para.runs:
        run.font.size = Pt(font_size)


def main() -> None:
    docx_path = Path(os.environ["DOCX_PATH"])
    figure_path = Path(os.environ["FIGURE_PATH"])
    doc = Document(docx_path)

    # Replace remaining software-log wording in section 3.1.
    for para in doc.paragraphs:
        if "旧disturbance保存结果中的候选C_old" in para.text:
            para.text = para.text.replace(
                "旧disturbance保存结果中的候选C_old",
                "早期扰动控制模式候选C_old",
            )
            normalize_para(para)

    anchor = None
    for para in doc.paragraphs:
        if para.text.startswith("从方案合理性看"):
            anchor = para
            break
    if anchor is None:
        raise RuntimeError("Insert anchor not found")

    note = insert_paragraph_after(
        anchor,
        "图5进一步给出了A、B、C三类候选方案在同一研究区边界、钻孔点和ODI高值区条件下的空间布局差异，可直观反映不同目标偏好对工作面数量、布置方向和风险暴露位置的影响。",
        anchor.style,
    )
    normalize_para(note)

    fig_para = insert_paragraph_after(note, "", anchor.style)
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig_para.add_run()
    run.add_picture(str(figure_path), width=Inches(6.25))

    cap = insert_paragraph_after(
        fig_para,
        "图5 A/B/C候选方案工作面布局与ODI高值区叠置对比图\nFig.5 Comparison of A/B/C candidate layouts overlaid with high-ODI zones",
        anchor.style,
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    normalize_para(cap, 9.0)

    doc.save(docx_path)


if __name__ == "__main__":
    main()
