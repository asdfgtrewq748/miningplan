from pathlib import Path
import sys

from docx import Document


DOCX = Path(sys.argv[1])

REPLACEMENTS = {
    58: (
        "不同偏好模式通过λ_e、λ_r和λ_m体现：工程效率优先模式提高λ_e，资源回收优先模式提高λ_r，扰动控制优先模式提高λ_m，综合权衡模式则采用相对均衡的权重。本文把这些权重作为统一候选集合内的排序偏好参数，而非现场统计标定常数；因此，A、B、C方案分别代表不同偏好下的推荐候选，用于揭示覆盖率、资源回收和扰动暴露之间的权衡关系。"
    ),
    126: (
        "从可复现角度看，本文样例至少需保留4类中间数据：研究区边界与钻孔样点、连续参数场栅格、ODI分量及综合ODI栅格、候选方案几何对象及方案级统计表。只有这些对象可逐级复算，ODI均值、P90、超阈值暴露比例和A/B/C方案排序才具备可审查性。本文在结果表中保留模式有效候选数、全域栅格数量和统一ODI统计口径，以降低仅展示最终方案图带来的不可复核风险。"
    ),
}


def set_para_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


doc = Document(DOCX)
for idx, text in REPLACEMENTS.items():
    set_para_text(doc.paragraphs[idx], text)
doc.save(DOCX)
print(f"changed={len(REPLACEMENTS)}")
