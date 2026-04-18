from copy import deepcopy
from pathlib import Path
import os

from docx import Document
from docx.text.paragraph import Paragraph


def insert_after(paragraph, text, style=None):
    new_p = deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.clear()
    if style is not None:
        inserted.style = style
    inserted.add_run(text)
    return inserted


def replace_start(doc, prefix, replacement):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            p.clear()
            p.add_run(replacement)
            return p
    return None


def main():
    docx_path = Path(os.environ.get("DOCX_PATH", r"E:\xiangmu\miningplan\煤科投稿\最新版论文4.16_插图版_煤科格式_大修工作稿_20260418.docx"))
    doc = Document(docx_path)

    p80 = replace_start(
        doc,
        "为验证所提方法在采区规划场景中的适用性与链路贯通能力",
        (
            "为验证所提方法在采区规划场景中的适用性与链路贯通能力，选取敏东研究区样例开展分析。"
            "研究区工程背景资料指向内蒙古呼伦贝尔敏东一矿，相关岩性力学资料显示，该类煤层开采条件下顶底板岩性主要包括泥岩、砂质泥岩、炭质泥岩、煤和细粒砂岩，"
            "整理资料中的代表性埋深约为300 m。该背景说明研究对象具有厚煤层、软弱围岩和多源扰动约束并存的工程特征。"
            "本文当前计算采用样例化规划域，以6个边界点构成采区边界、15个钻孔样点和规划控制参数为基础输入，原始边界面积为255250.00 m²。"
            "当前验证目标限定为样例条件下的数据组织、风险表征与规划对象生成，不将样例结果解释为真实矿井条件下的工业级优选结论。"
        ),
    )
    if p80 is not None:
        insert_after(
            p80,
            (
                "从工程约束看，本文样例已显式纳入边界煤柱、区段煤柱、工作面宽度、推进长度校核阈值、走向长壁后退式开采方向和ODI统计阈值等规划参数。"
                "但现有输入尚未完整纳入断层展布、既有巷道系统、通风运输能力、设备能力和生产制度等现场闭环约束，"
                "因此案例定位为方法验证型工程样例，而非直接用于矿井定案的完整设计文件。"
            ),
            style=p80.style,
        )
    else:
        print("P80 anchor not found")

    replace_start(
        doc,
        "研究区输入对象包括采区边界、钻孔样点和规划控制参数3类。",
        (
            "研究区输入对象包括采区边界、钻孔样点、岩性力学背景资料和规划控制参数4类。"
            "钻孔样点用于构建煤层厚度等连续参数场；边界对象用于界定有效布置域；岩性力学背景资料用于限定案例的工程语境；"
            "边界煤柱、区段煤柱、工作面宽度、推进长度校核阈值和ODI统计阈值则共同构成后续规划求解的约束与校核条件。"
        ),
    )

    replace_start(
        doc,
        "因此，本节的作用是把研究区输入压实为可复核的参数体系",
        (
            "因此，本节的作用是把研究区输入压实为可复核的参数体系：边界提供空间范围，钻孔提供连续参数场样本，岩性与埋深资料提供工程背景，"
            "设计参数提供几何与安全约束，ODI阈值提供风险统计口径。后续结果分析均以该统一输入口径为基础。"
        ),
    )

    doc.save(docx_path)
    print(f"Strengthened case background in {docx_path}")


if __name__ == "__main__":
    main()
