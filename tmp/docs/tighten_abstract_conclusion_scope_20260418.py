from pathlib import Path
import os

from docx import Document


def replace_start(doc, prefix, replacement):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            p.clear()
            p.add_run(replacement)
            return True
    return False


def main():
    docx_path = Path(os.environ.get("DOCX_PATH", r"E:\xiangmu\miningplan\煤科投稿\最新版论文4.16_插图版_煤科格式_大修工作稿_20260418.docx"))
    doc = Document(docx_path)

    replace_start(
        doc,
        "摘要：针对采区规划中离散地质信息难以连续参与空间布置",
        (
            "摘要：针对采区规划中离散地质信息难以连续参与空间布置、多场景覆岩扰动风险约束口径不统一以及规划结果难以向采掘接续和经济评价传递等问题，"
            "提出一种基于覆岩扰动指数（overburden disturbance index，ODI）约束的采区候选方案生成与多目标比选方法。"
            "该方法以采区边界、钻孔样点和设计参数为输入，先构建有效布置域与连续参数场，再将地表沉陷、含水层扰动和上行开采等风险分量归一化并加权为统一ODI场，"
            "进一步以ODI均值、90%分位数和超阈值暴露比例刻画候选方案风险水平。"
            "在此基础上，建立包含工作面数量、推进方向、煤柱宽度、工作面宽度和方案选择变量的候选方案池，并采用工程效率、资源回收和扰动控制目标进行非支配排序与综合筛选。"
            "样例中复核得到80×56栅格ODI场，共4480个栅格，均值为0.4669，P90为0.7474，ODI>0.70比例为15.89%。"
            "在同一ODI场下，工程效率优先方案A、资源回收优先方案B和ODI筛选方案C的ODI均值分别为0.4463、0.4552和0.4416，P90分别为0.6407、0.6462和0.6353，"
            "ODI>0.70比例分别为0.44%、0.56%和1.22%。权重和阈值敏感性结果表明，在当前样例和候选池范围内，C方案风险综合得分均低于A、B方案。"
            "研究结果说明，ODI前置约束可将风险图层转化为可复核的方案级统计指标，为采区规划阶段的多目标比选和后续对象传递提供方法支撑；"
            "真实矿井条件下的工程优选仍需结合现场约束和多案例验证。"
        ),
    )

    replace_start(
        doc,
        "Abstract: To address the difficulty of using discrete geological information",
        (
            "Abstract: To address the difficulty of using discrete geological information in continuous mining-district layout, the inconsistent expression of multi-scenario overburden disturbance risks, "
            "and the weak linkage between planning outputs and downstream evaluation, this study proposes an ODI-constrained candidate generation and multi-objective selection method for mining-district planning. "
            "The method constructs an effective layout domain and continuous parameter fields from boundary, borehole and design inputs, normalizes surface subsidence, aquifer disturbance and upward-mining risk components "
            "into a unified overburden disturbance index (ODI) field, and describes candidate-level risk exposure using mean ODI, the 90th percentile and exceedance ratio. "
            "A constrained candidate pool is then ranked by engineering-efficiency, resource-recovery and disturbance-control objectives. "
            "The verified ODI field contains 80×56 grids, with a mean of 0.4669, P90 of 0.7474 and ODI>0.70 ratio of 15.89%. "
            "Under the same ODI field, the mean ODI values of the engineering-efficiency scheme A, resource-recovery scheme B and ODI-screened scheme C are 0.4463, 0.4552 and 0.4416, respectively; "
            "their P90 values are 0.6407, 0.6462 and 0.6353, and their ODI>0.70 ratios are 0.44%, 0.56% and 1.22%. "
            "For the present sample and candidate pool, weight and threshold sensitivity analyses show that scheme C has the lowest composite risk score under ±10% weight perturbations and the aquifer-special weight setting. "
            "The results indicate that ODI pre-constraint can transform risk layers into reproducible scheme-level statistics for multi-objective comparison in mining-district planning, "
            "whereas engineering optimization under real mine conditions still requires site-specific constraints and multi-case validation."
        ),
    )

    replace_start(
        doc,
        "4）当前结果主要支撑样例条件下的方法贯通",
        (
            "4）当前结果主要支撑样例条件下的方法贯通、对象链有效性和初步方案比选。权重敏感性表明，在当前样例和候选池范围内，"
            "基准权重、±10%相对扰动和含水层专项权重条件下C方案均保持最低风险综合得分；阈值敏感性表明，当阈值由0.65提高至0.80时，"
            "各方案超限比例均下降，但方案间风险差异仍可通过均值、P90和超限比例共同识别。真实矿井条件下的工程优选结论仍需结合实矿案例检验。"
        ),
    )

    doc.save(docx_path)
    print(f"Tightened abstract and conclusion scope in {docx_path}")


if __name__ == "__main__":
    main()
