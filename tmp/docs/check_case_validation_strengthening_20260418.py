from __future__ import annotations

import sys
import zipfile
from pathlib import Path

from docx import Document


CHECKS = {
    "controlled_comparison": "\u76f8\u540c\u7814\u7a76\u533a\u8fb9\u754c\u3001\u76f8\u540c\u94bb\u5b54\u53c2\u6570\u573a\u3001\u76f8\u540cODI\u98ce\u9669\u573a\u548c\u76f8\u540c\u51e0\u4f55\u7ea6\u675f",
    "not_manual_baseline": "\u800c\u4e0d\u662f\u72ec\u7acb\u4eba\u5de5\u7ecf\u9a8c\u65b9\u6848\u4e0e\u81ea\u52a8\u65b9\u6848\u7684\u73b0\u573a\u5bf9\u7167",
    "reasonableness_check": "\u8fd9\u79cd\u7ed3\u679c\u7b26\u5408\u91c7\u533a\u89c4\u5212\u4e2d\u8d44\u6e90\u5229\u7528\u4e0e\u98ce\u9669\u63a7\u5236\u76f8\u4e92\u5236\u7ea6\u7684\u5de5\u7a0b\u5e38\u8bc6",
    "validation_boundary": "\u672c\u6848\u4f8b\u7684\u9a8c\u8bc1\u5f3a\u5ea6\u5b9a\u4f4d\u4e3a\u201c\u65b9\u6cd5\u94fe\u53ef\u8fd0\u884c\u6027\u4e0e\u5185\u90e8\u5bf9\u7167\u9a8c\u8bc1\u201d",
    "section_33": "3.3 \u89c4\u5212\u7ed3\u679c\u4f20\u9012\u4e0e\u8bc4\u4ef7\u8fb9\u754c",
}


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: check_case_validation_strengthening_20260418.py <docx>")
        return 2

    path = Path(sys.argv[1])
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
    if bad is not None:
        raise RuntimeError(f"Bad DOCX zip member: {bad}")

    doc = Document(str(path))
    text = "\n".join(p.text for p in doc.paragraphs)
    for name, phrase in CHECKS.items():
        print(f"{name}={phrase in text}")
    print(f"paragraphs={len(doc.paragraphs)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
