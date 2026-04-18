from __future__ import annotations

import sys
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


def insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.text = text
    return inserted


def delete_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def find_exact(paragraphs: list[Paragraph], text: str) -> int:
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == text:
            return idx
    raise RuntimeError(f"Paragraph not found: {text[:60]}")


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: strengthen_multiobjective_model_20260418.py <docx>")
        return 2

    path = Path(sys.argv[1])
    doc = Document(str(path))

    p50 = doc.paragraphs[find_exact(doc.paragraphs, "π={N,W_f,θ,B_b,B_s,A_π,L_π,R_π,y_π}，y_π∈{0,1}。                         （8）")]
    after_50 = (
        "式中，N为工作面数量，W_f为工作面宽度，θ为推进方向，B_b和B_s分别为边界煤柱与区段煤柱宽度，"
        "A_π为方案π中所有工作面占用区域的集合，L_π为各工作面推进长度集合，R_π为巷道连接与服务关系，"
        "y_π为是否选取该候选方案的0-1变量。该变量定义使规划对象由单一几何图形扩展为包含数量、尺度、方向、"
        "煤柱、工作面区域和巷道连接关系的组合对象，从而能够同时承接几何约束、风险统计和后续接续评价。"
    )
    if after_50 not in [p.text for p in doc.paragraphs]:
        insert_after(p50, after_50)

    p54 = doc.paragraphs[find_exact(doc.paragraphs, "候选池生成采用规则枚举与工程过滤相结合的方式。首先根据采区主控方向、煤层走向和边界长轴方向形成有限推进方向集合Θ，再在θ∈Θ、边界煤柱宽度B_b、区段煤柱宽度B_s、工作面宽度W_f和边界内缩距离的参数组合下生成平行工作面条带。随后剔除越界、重叠、推进长度不足、煤柱宽度不满足要求和连通关系异常的方案，形成可比候选池。因此，推进方向不是黑箱“自动确定”，而是在工程允许方向集合内随几何参数共同枚举，并由约束过滤和多目标排序共同确定推荐结果。")]
    after_54 = (
        "在实现流程上，候选池生成可分为参数取值、几何构造、硬约束过滤、指标计算和去重合并5个步骤。"
        "首先读取边界、钻孔参数场和规划控制参数，确定Θ、W_f、B_b、B_s及内缩距离的有限取值；"
        "其次按参数组合切分有效布置域并生成工作面区域、推进长度和巷道连接对象；随后按照式（9）进行必要约束过滤，"
        "并对合格候选计算覆盖率、煤厚覆盖、巷道组织、ODI均值、P90和超阈值暴露比例；最后以方向、煤柱、工作面数量、"
        "工作面边界和连接关系形成候选签名，对重复几何方案进行合并。由此得到的Π是有限候选集合，后续排序只在该候选池范围内成立。"
    )
    if after_54 not in [p.text for p in doc.paragraphs]:
        insert_after(p54, after_54)

    p65 = doc.paragraphs[find_exact(doc.paragraphs, "为降低单一加权排序对权重的依赖，本文在加权排序前引入非支配排序和拥挤距离选择形成Top-K候选方案。非支配排序以S_e、S_r和S_m为目标向量，优先保留不存在其他方案同时优于它的候选；当同一非支配层内候选数量超过输出数量时，采用拥挤距离保持解集多样性，再依据综合权重F(π)进行推荐排序。工程效率优先、资源回收优先、扰动控制优先和综合权衡模式均来自同一候选池，仅代表不同评价偏好，不应被表述为真实矿井条件下的全局最优解。")]
    after_65a = (
        "非支配关系具体定义为：若候选方案π_i在S_e、S_r和S_m三个目标上均不劣于π_j，且至少一个目标严格优于π_j，"
        "则π_i支配π_j。第一非支配层由不被任何其他候选支配的方案组成，第二层由剔除第一层后新的非支配方案组成，"
        "依此类推。拥挤距离用于保留同一非支配层内目标空间分布较分散的候选，避免输出方案集中在单一偏好附近。"
    )
    after_65b = (
        "不同偏好模式通过λ_e、λ_r和λ_m体现：工程效率优先模式提高λ_e，资源回收优先模式提高λ_r，扰动控制优先模式提高λ_m，"
        "综合权衡模式则采用相对均衡的权重。本文把这些权重作为同一候选池内部的排序偏好参数，而非现场统计标定常数；"
        "因此，A、B、C方案分别代表不同偏好下的推荐候选，用于揭示覆盖率、资源回收和扰动暴露之间的权衡关系。"
    )
    existing = [p.text for p in doc.paragraphs]
    if after_65a not in existing:
        inserted = insert_after(p65, after_65a)
    else:
        inserted = p65
    if after_65b not in [p.text for p in doc.paragraphs]:
        insert_after(inserted, after_65b)

    duplicate_text = "本文方法的一个重要特点，是不把采区规划的几何布局视为终点，而是强调规划结果可以继续作为采掘接续和工程经济评价的上游输入。在本样例中，工作面边界、巷道路径、推进长度、布置面积和ODI统计量均已形成可传递的规划结果，后续环节可据此组织生产顺序、产量核算和经济评价。"
    duplicate_indices = [idx for idx, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == duplicate_text]
    for idx in reversed(duplicate_indices):
        delete_paragraph(doc.paragraphs[idx])

    doc.save(str(path))
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
    if bad is not None:
        raise RuntimeError(f"Bad DOCX zip member: {bad}")

    print(f"DOCX OK: {path}")
    print(f"Inserted model paragraphs: 4")
    print(f"Deleted duplicate transfer paragraphs: {len(duplicate_indices)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
