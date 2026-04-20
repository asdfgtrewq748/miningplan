from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.shared import Pt


def replace_para(doc: Document, old_prefix: str, new_text: str) -> int:
    hits = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith(old_prefix):
            para.text = new_text
            try:
                para.style = "10正文"
            except KeyError:
                pass
            for run in para.runs:
                run.font.size = Pt(10.5)
            hits += 1
    if hits != 1:
        raise RuntimeError(f"Expected 1 paragraph for prefix {old_prefix!r}, found {hits}")
    return hits


def set_cell_text(table, row: int, col: int, text: str) -> None:
    table.rows[row].cells[col].text = text


def main() -> None:
    docx_path = Path(os.environ["DOCX_PATH"])
    doc = Document(docx_path)

    replace_para(
        doc,
        "设候选方案集合为Π",
        "设候选方案集合为Π，候选方案π由工作面数量N、工作面倾向宽度W_f、推进方向θ、边界煤柱宽度B_b、区段煤柱宽度B_s、工作面起止边界、各工作面走向推进长度集合L_π、巷道连接关系和方案选择变量y_π共同描述，可写为",
    )
    replace_para(
        doc,
        "式中，N为工作面数量，W_f为工作面宽度",
        "式中，N为工作面数量，W_f为工作面倾向宽度，θ为推进方向，B_b和B_s分别为边界煤柱与区段煤柱宽度，A_π为方案π中工作面占用区域集合，L_π为各工作面走向推进长度集合，R_π为巷道连接与服务关系，y_π为候选方案选择变量。该变量定义将规划对象由单一几何图形扩展为包含数量、尺度、方向、煤柱、工作面区域、推进长度和巷道关系的组合对象，可同时承接几何约束、风险统计和后续接续评价。",
    )
    replace_para(
        doc,
        "式中，A_π⊂Ω_e为边界约束",
        "式中，A_π⊂Ω_e为边界约束，B_b和B_s为煤柱约束，W_f为工作面倾向宽度约束，A_i∩A_j=∅为工作面不重叠约束，C_L表示单工作面几何长度推荐区间与方案推进长度校核条件，I_ODI表示ODI暴露统计参与方案筛选与风险排序的指标口径。当前样例将边界、煤柱、工作面宽度和不重叠作为硬约束，将推进长度校核阈值与ODI暴露阈值作为候选方案筛选和排序中的推荐性约束，以避免样例验证阶段因现场闭环约束不完整而过度收缩候选空间。",
    )
    replace_para(
        doc,
        "本文首先在统一基础输入",
        "本文首先在统一基础输入、统一几何底线约束和统一ODI场条件下形成可合并的总候选集合Π_all=Π_e∪Π_r∪Π_m，再在该集合上分别按工程效率、资源回收和扰动控制偏好进行排序与筛选，输出工程效率优先方案A、资源回收优先方案B和扰动控制优先方案C。其中Π_e、Π_r和Π_m并非在不同基础输入下独立生成，而是在同一输入、同一几何底线约束和同一ODI场条件下按不同目标偏好保留的模式候选子集，其并集构成统一比较入口。推进方向不作黑箱“自动确定”，而是在工程允许方向集合Θ内与煤柱宽度、工作面宽度和边界内缩距离共同枚举，并由约束过滤和多目标排序确定推荐结果。",
    )
    replace_para(
        doc,
        "对候选方案π，本文把工程效率、资源回收和扰动控制分别写成可复核的分项评分",
        "对候选方案π，本文把工程效率、资源回收和扰动控制分别写成可复核的分项评分。为保证不同量纲指标可比较，覆盖率、资源贡献、工程组织和扰动风险等子项均先转换到[0,1]或0—100的统一评分区间；其中正向指标按式（3）归一化，成本型或风险型指标按式（4）转化为正向评分。",
    )
    replace_para(
        doc,
        "式中，C_cov为有效覆盖率",
        "式中，C_cov为有效覆盖率，取值为[0,1]；P_N为工作面数量复杂度惩罚，CV_L为各工作面走向推进长度变异系数，P_short为短推进惩罚项。P_N、CV_L项和P_short均作为0—100工程效率评分中的扣分项。当前样例取P_N=0.20N，CV_L项系数取10，短推进惩罚系数取5，短推进参考值取100 m；上述系数仅用于样例内部排序敏感性校核，不作为跨矿区通用经验常数。",
    )
    replace_para(
        doc,
        "式中，R_ton为吨煤量或煤厚场覆盖贡献",
        "式中，R_ton、R_area和S_eng均已归一化到[0,1]，分别表示吨煤量或煤厚场覆盖贡献、有效布置面积贡献和工程组织辅助得分。式（11）中的0.55/0.30/0.15为资源回收分项内部权重；外层0.45+0.55(·)为线性区间映射，用于将样例候选集合内的资源回收评分平移至45—100区间，以避免极端低值压缩方案间区分度。故89.76表示资源回收分项评分，而非综合排序总分。",
    )
    replace_para(
        doc,
        "为降低单一加权排序对权重的依赖",
        "为降低单一加权排序对权重的依赖，本文在加权排序前引入非支配排序和拥挤距离选择形成有限推荐候选集；当前样例保留前10个候选用于复核与比较。非支配排序以S_e、S_r和S_m为目标向量，优先保留不存在其他方案同时优于它的候选；同一非支配层候选数量超过输出数量时，采用拥挤距离保持解集多样性，再依据F(π)推荐排序。A、B、C三种偏好模式中，λ_e/λ_r/λ_m分别取1/0/0、0/1/0和0/0/1；综合权衡模式备用权重取0.34/0.33/0.33，本文不将其未达标回退结果作为最终推荐方案。",
    )
    replace_para(
        doc,
        "在规划语境下，ODI的价值体现在两个层面",
        "在规划语境下，ODI的价值体现在两个层面：一是描述研究区高扰动敏感区的空间背景，二是提供候选方案区域内的风险暴露统计。在同一ODI场下，方案A、B、C的ODI均值分别为0.4463、0.4552和0.4416，P90分别为0.6407、0.6462和0.6353，ODI>0.70比例分别为0.44%、0.56%和1.22%。其中，C方案为在统一候选集合内按统一ODI场重新筛选得到的联合判据下低扰动候选。",
    )
    replace_para(
        doc,
        "ODI前置约束的优势不在于替代地表沉陷",
        "ODI前置约束的优势不在于替代地表沉陷、含水层扰动或上行开采等专项分析，而在于为多源风险结果提供同一比较口径。本文的增量不在于提出新的覆岩力学机理模型，而在于把专项风险分析结果组织为采区规划阶段可统一表达、可排序比较的方案级约束信息。在同一ODI场下，方案可以同时报告均值、P90和超阈值暴露比例，从而避免仅凭单一风险图或单一阈值判断方案优劣。本样例中C方案的超阈值比例并非在所有阈值下都最低，但其均值、P90和风险综合得分较低，说明方案风险需要用多统计量联合解释。",
    )

    # Table 1: clarify geometry terms without changing the formula symbols.
    table1 = doc.tables[15]
    set_cell_text(table1, 8, 1, "工作面倾向宽度")
    set_cell_text(table1, 8, 5, "当前样例采用的单面面宽")
    set_cell_text(table1, 9, 1, "单工作面几何长度规则")
    set_cell_text(table1, 9, 5, "用于候选条带生成的单工作面几何长度规则")
    set_cell_text(table1, 10, 1, "方案推进长度校核阈值")
    set_cell_text(table1, 10, 5, "推荐/校核阈值；用于方案层面推进长度提示，不作为硬性达标结果")

    # Table 3: remove software-log wording from the legacy candidate row.
    table3 = doc.tables[17]
    set_cell_text(table3, 4, 1, "早期扰动控制模式候选")
    set_cell_text(table3, 4, 4, "早期口径候选，不作为本文C方案")

    # Table 4: move exact file names out of the main text table.
    table4 = doc.tables[18]
    set_cell_text(table4, 0, 3, "数据对象")
    set_cell_text(table4, 1, 3, "综合ODI场中间数据")
    set_cell_text(table4, 2, 3, "A/B/C方案统一ODI统计表")
    set_cell_text(table4, 3, 3, "阈值敏感性统计表")
    set_cell_text(table4, 4, 3, "权重敏感性统计表")
    set_cell_text(table4, 5, 3, "规划对象与接续输入表")

    doc.save(docx_path)


if __name__ == "__main__":
    main()
