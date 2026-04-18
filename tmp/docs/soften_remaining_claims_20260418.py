from __future__ import annotations

import sys
import zipfile
from pathlib import Path

from docx import Document


REPLACEMENTS = {
    "权重敏感性分析进一步表明，在基准权重、地表沉陷分量±10%、含水层扰动分量±10%、上行开采分量±10%以及含水层专项权重条件下，C方案的风险综合得分均为三方案最低，A方案次之，B方案最高。基准权重下A、B、C风险综合得分分别为0.4481、0.4546和0.4449；在含水层专项权重0.15/0.25/0.60下，三者分别为0.4505、0.4564和0.4498。":
    "权重敏感性分析进一步表明，在基准权重、地表沉陷分量±10%、含水层扰动分量±10%、上行开采分量±10%以及含水层专项权重条件下，C方案的风险综合得分均低于A、B方案。基准权重下A、B、C风险综合得分分别为0.4481、0.4546和0.4449；在含水层专项权重0.15/0.25/0.60下，三者分别为0.4505、0.4564和0.4498。",

    "需要注意，权重敏感性所证明的是当前候选池和扰动范围内的排序稳定，而不是ODI权重具有普适工程常数属性。因此，本文把敏感性分析作为方案排序可靠性的内部校核，用于回应权重和阈值人为设定可能导致结论不稳定的问题。":
    "需要注意，权重敏感性所表明的是当前候选池和扰动范围内的排序一致性，而不是ODI权重具有普适工程常数属性。因此，本文把敏感性分析作为方案排序可靠性的内部校核，用于回应权重和阈值人为设定可能导致结论波动的问题。",

    "综上，研究区样例表明，本文方法已完成“边界与钻孔输入—参数场构建—ODI风险组织—候选方案池—方案级统计—后续评价输入”的连续运行过程。当前结果能够证明方法链和对象链已经贯通，但仍需通过更多实矿案例和现场约束闭环检验工程适用性。":
    "综上，研究区样例表明，本文方法已完成“边界与钻孔输入—参数场构建—ODI风险组织—候选方案池—方案级统计—后续评价输入”的连续运行过程。当前结果表明方法链和对象链在样例条件下具备贯通性，但仍需通过更多实矿案例和现场约束闭环检验工程适用性。",
}


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: soften_remaining_claims_20260418.py <docx>")
        return 2

    path = Path(sys.argv[1])
    doc = Document(str(path))
    texts = [p.text.strip() for p in doc.paragraphs]
    changed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in REPLACEMENTS:
            paragraph.text = REPLACEMENTS[text]
            changed += 1
    if changed != len(REPLACEMENTS):
        refreshed = [p.text.strip() for p in doc.paragraphs]
        missing = [old[:60] for old in REPLACEMENTS if old in texts and REPLACEMENTS[old] not in refreshed]
        if missing:
            raise RuntimeError(f"Failed replacements: {missing}")

    doc.save(str(path))
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
    if bad is not None:
        raise RuntimeError(f"Bad DOCX zip member: {bad}")

    print(f"DOCX OK: {path}")
    print(f"Softened remaining paragraphs: {changed}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
