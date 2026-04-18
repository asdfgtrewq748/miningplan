from pathlib import Path
import sys

from docx import Document


DOCX = Path(sys.argv[1])

ABSTRACTS = {
    3: (
        "摘要：针对采区规划中离散地质信息难以连续参与布置、多场景覆岩扰动风险约束口径不统一以及规划结果难以向接续和经济评价传递等问题，提出一种基于覆岩扰动指数（overburden disturbance index，ODI）约束的采区候选方案生成与多目标比选方法。该方法以采区边界、钻孔样点和设计参数为输入，构建有效布置域和连续参数场，将地表沉陷、含水层扰动、上行开采等风险分量归一化并加权为统一ODI场，并以ODI均值、90%分位数和超阈值暴露比例表征方案风险。在此基础上，建立包含工作面数量、推进方向、煤柱宽度、工作面宽度和方案选择变量的候选方案池，采用工程效率、资源回收和扰动控制目标进行非支配排序与综合筛选。样例复核得到80×56栅格ODI场，均值为0.4669，P90为0.7474，ODI>0.70比例为15.89%。在同一ODI场下，工程效率优先方案A、资源回收优先方案B和ODI筛选方案C的ODI均值分别为0.4463、0.4552和0.4416，P90分别为0.6407、0.6462和0.6353，ODI>0.70比例分别为0.44%、0.56%和1.22%。敏感性分析表明，在当前样例和候选池范围内，C方案风险综合得分低于A、B方案。结果说明，ODI前置约束可将风险图层转化为可复核的方案级统计指标，为采区规划多目标比选和后续对象传递提供支撑；真实矿井优选仍需结合现场约束和多案例验证。"
    ),
    9: (
        "Abstract: To address the limited use of discrete geological data in continuous mining-district layout, the inconsistent expression of multi-scenario overburden disturbance risks, and the weak linkage between planning outputs and downstream evaluation, an ODI-constrained candidate generation and multi-objective selection method is proposed. Boundary, borehole and design inputs are used to construct the effective layout domain and continuous parameter fields. Surface subsidence, aquifer disturbance and upward-mining risk components are normalized and weighted into a unified overburden disturbance index (ODI) field, and candidate-level risk is described by mean ODI, the 90th percentile and exceedance ratio. A constrained candidate pool considering face number, advance direction, pillar width, face width and scheme-selection variables is then ranked using engineering-efficiency, resource-recovery and disturbance-control objectives. The verified sample ODI field contains 80×56 grids, with a mean of 0.4669, P90 of 0.7474 and ODI>0.70 ratio of 15.89%. Under the same ODI field, schemes A, B and C have mean ODI values of 0.4463, 0.4552 and 0.4416, P90 values of 0.6407, 0.6462 and 0.6353, and ODI>0.70 ratios of 0.44%, 0.56% and 1.22%, respectively. Sensitivity analysis indicates that scheme C has a lower composite risk score than schemes A and B within the present sample and candidate pool. The results show that ODI pre-constraint can convert risk layers into reproducible scheme-level statistics for multi-objective comparison, while engineering selection under real mine conditions still requires site-specific constraints and multi-case validation."
    ),
}


def set_paragraph_text(paragraph, text):
    if paragraph.runs:
        first = paragraph.runs[0]
        first.text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def main():
    doc = Document(DOCX)
    for idx, text in ABSTRACTS.items():
        set_paragraph_text(doc.paragraphs[idx], text)
    doc.save(DOCX)
    print("changed=2")


if __name__ == "__main__":
    main()
