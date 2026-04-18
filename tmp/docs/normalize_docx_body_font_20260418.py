import os
from pathlib import Path

from docx import Document
from docx.shared import Pt


path = Path(os.environ["DOCX_PATH"])
doc = Document(path)

formula_prefixes = (
    "ODI(",
    "w_s+",
    "X_i'",
    "E_π=",
    "Ω_e=",
    "z(x)=",
    "π={",
    "A_π⊂",
    "F(π)=",
    "G(s)=",
    "NCF_t=",
)

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    if i <= 12:
        continue
    target_size = Pt(9) if text.startswith(formula_prefixes) else Pt(10.5)
    if text.startswith("图") or text.startswith("Fig.") or text.startswith("表") or text.startswith("Table"):
        target_size = Pt(9)
    for run in para.runs:
        run.font.size = target_size

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(7)

doc.save(path)
print(f"normalized body fonts {path}")
