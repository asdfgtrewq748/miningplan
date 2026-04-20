from __future__ import annotations

import copy
import shutil
import sys
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

DOCX = Path(sys.argv[1])
BACKUP = Path(sys.argv[2])
shutil.copy2(DOCX, BACKUP)

def set_cell_text(cell, text, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = text
    for para in cell.paragraphs:
        para.alignment = align
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        for run in para.runs:
            run.font.size = Pt(size)

def insert_row_after(table, row_idx: int, values):
    tr = table.rows[row_idx]._tr
    new_tr = copy.deepcopy(tr)
    tr.addnext(new_tr)
    new_row = table.rows[row_idx + 1]
    for cell, val in zip(new_row.cells, values):
        set_cell_text(cell, val)
    return new_row

doc = Document(DOCX)
body_style = next((p.style for p in doc.paragraphs if p.style.name == '10正文'), None)
if body_style is None:
    raise RuntimeError('10正文 style not found')

# 1, 2, 4, 5: paragraph-level terminology and final identity unification.
para_repls = {
    '候选池筛选与排序环节': '候选方案集合筛选与排序环节',
    '当前候选池和扰动范围内': '当前候选集合和扰动范围内',
    '候选池生成': '候选集合生成',
    '输出工程效率优先方案A、资源回收优先方案B和扰动控制优先方案C': '输出工程效率优先方案A、资源回收优先方案B和联合判据下低扰动候选C',
    'C方案代表ODI风险约束偏好，强调低扰动暴露。': 'C方案为联合判据下低扰动候选，强调ODI均值、P90和风险综合得分的综合控制。',
    'ODI筛选方案C来自统一候选集合复核，覆盖率为80.67%，工程效率分项评分为79.27，ODI均值为0.4416，P90为0.6353，ODI>0.70比例为1.22%。': '联合判据下低扰动候选C来自统一候选集合复核，覆盖率为80.67%，H_m=0.4449，ODI均值为0.4416，P90为0.6353，ODI>0.70比例为1.22%。',
    '表现较优的低扰动候选': '联合判据下低扰动候选',
    'ODI筛选方案分别体现出覆盖率、资源利用和风险暴露之间的权衡': '联合判据下低扰动候选分别体现出覆盖率、资源利用和风险暴露之间的权衡',
}
for para in doc.paragraphs:
    txt = para.text
    new = txt
    for old, val in para_repls.items():
        new = new.replace(old, val)
    if new != txt:
        para.text = new

# Add scoring-boundary sentence after scoring explanation if not present.
needle = '上述评分函数用于当前样例条件下候选方案的内部排序与敏感性比较'
if not any(needle in p.text for p in doc.paragraphs):
    target = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.startswith('上述评分给出了87.82、89.76等方案分数的计算口径'):
            target = para
            break
    if target is not None:
        new_p = target.insert_paragraph_before('')
        # Move inserted paragraph after target by XML manipulation.
        target._p.addnext(new_p._p)
        new_p.text = '上述评分函数用于当前样例条件下候选方案的内部排序与敏感性比较，其系数和映射区间服务于样例级区分度控制，不解释为跨矿区普适经验参数。'
        new_p.style = body_style
        for run in new_p.runs:
            run.font.size = Pt(10.5)

# 3: split Table 1 candidate counts.
for table in doc.tables:
    text = '\n'.join(cell.text for row in table.rows for cell in row.cells)
    if '模式有效候选数' in text and '统一候选集合' in text:
        for ri, row in enumerate(table.rows):
            if row.cells[1].text.strip() == '模式有效候选数':
                vals = ['候选集合', '工程效率/资源回收偏好有效候选数', 'N_e/N_r', '2417/1149', '个', '统一基础输入下工程效率和资源回收偏好筛选后的有效候选记录', '规划结果']
                for c, val in enumerate(vals):
                    set_cell_text(row.cells[c], val)
                insert_row_after(table, ri, ['候选集合', '扰动控制偏好有效候选数', 'N_m', '374', '个', '统一基础输入下扰动控制偏好筛选后的有效候选记录', '规划结果'])
                insert_row_after(table, ri + 1, ['候选集合', '联合判据下低扰动候选数', 'N_m,low', '88', '个', '按ODI均值、P90和风险综合得分联合判据筛选', '规划结果'])
                break
        break

# 1: Table 3 add dominant sorting indicator and fix C row to H_m.
for table in doc.tables:
    if len(table.rows) >= 5 and table.cell(0, 0).text.strip() == '方案' and 'ODI统计' in table.cell(0, 3).text:
        # Add a new trailing column unless it already exists.
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if '主导排序指标' not in headers:
            table.add_column(Inches(0.75))
        # Refresh rows after add_column.
        set_cell_text(table.rows[0].cells[-1], '主导排序指标', align=WD_ALIGN_PARAGRAPH.CENTER)
        data = {
            'A': ['A', '工程效率偏好', '5个工作面；覆盖率89.34%；S_e=87.82', '均值0.4463；P90=0.6407；E_0.70=0.44%', '按工程效率分项排序', 'S_e'],
            'B': ['B', '资源回收偏好', '9个工作面；覆盖率98.52%；S_r=89.76', '均值0.4552；P90=0.6462；E_0.70=0.56%', '按资源回收分项排序', 'S_r'],
            'C': ['C', '联合判据下低扰动候选', '4个工作面；覆盖率80.67%；H_m=0.4449', '均值0.4416；P90=0.6353；E_0.70=1.22%', '联合判据下低扰动候选', 'H_m'],
            'C0': ['C0', '早期扰动控制偏好', '13个工作面；覆盖率75.06%；S_e=71.17', '均值0.4560；P90=0.6472；E_0.70=0.80%', '早期口径候选，不纳入本文C方案', '不纳入'],
        }
        for row in table.rows[1:]:
            key = row.cells[0].text.strip()
            if key in data:
                for c, val in enumerate(data[key]):
                    set_cell_text(row.cells[c], val)
        break

# Table wording consistency.
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                txt = para.text
                new = txt.replace('低扰动较优', '低扰动候选')
                if new != txt:
                    para.text = new
                    for run in para.runs:
                        run.font.size = Pt(9)

# Normalize data table alignment.
for table in doc.tables:
    table_text = '\n'.join(cell.text for row in table.rows for cell in row.cells)
    is_data = any(k in table_text for k in ['参数名称', '候选方案对比', 'ODI统计口径', '阈值敏感性', '权重敏感性', '经济评价'])
    if not is_data:
        continue
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if ri == 0 else WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.font.size = Pt(9)

doc.save(DOCX)
print('done')
