from docx import Document
from docx.shared import Pt
import sys
p=sys.argv[1]
doc=Document(p)
repls={
    '本文对低扰动候选的判定并非依据': '本文对联合判据下低扰动候选的判定并非依据',
    '本文对低扰动候选的判定依据为': '本文对联合判据下低扰动候选的判定依据为',
    '不作为统一口径下的ODI筛选方案': '不作为本文统一口径下的C方案',
}
changed=0
for para in doc.paragraphs:
    txt=para.text
    new=txt
    for old,val in repls.items(): new=new.replace(old,val)
    if new != txt:
        para.text=new
        for run in para.runs:
            if para.style.name == '10正文': run.font.size=Pt(10.5)
        changed+=1
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                txt=para.text
                new=txt
                for old,val in repls.items(): new=new.replace(old,val)
                if new != txt:
                    para.text=new
                    for run in para.runs: run.font.size=Pt(9)
                    changed+=1

doc.save(p)
print('changed', changed)
