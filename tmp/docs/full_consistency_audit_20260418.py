from __future__ import annotations

import re
import sys
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document


EXPECTED_HEADINGS = [
    "\u0030 \u5f15\u8a00",
    "\u0031 \u57fa\u4e8e\u8986\u5ca9\u6270\u52a8\u6307\u6570\u7ea6\u675f\u7684\u91c7\u533a\u591a\u76ee\u6807\u534f\u540c\u89c4\u5212\u65b9\u6cd5",
    "\u0031\u002e\u0031 \u9762\u5411\u91c7\u533a\u89c4\u5212\u7684 ODI \u98ce\u9669\u7ea6\u675f\u5b9a\u4e49",
    "\u0031\u002e\u0032 \u6709\u6548\u5e03\u7f6e\u57df\u4e0e\u8fde\u7eed\u53c2\u6570\u573a\u6784\u5efa",
    "\u0031\u002e\u0033 \u5019\u9009\u65b9\u6848\u751f\u6210\u4e0e\u591a\u76ee\u6807\u6bd4\u9009\u6a21\u578b",
    "\u0031\u002e\u0034 \u65b9\u6848\u4f20\u9012\u4e0e\u8bc4\u4ef7\u6d41\u7a0b",
    "\u0032 \u5de5\u7a0b\u6848\u4f8b\u4e0e\u7ed3\u679c\u5206\u6790",
    "\u0032\u002e\u0031 \u7814\u7a76\u533a\u6982\u51b5\u4e0e\u8f93\u5165\u6570\u636e",
    "\u0032\u002e\u0032 \u8fde\u7eed\u53c2\u6570\u573a\u6784\u5efa\u7ed3\u679c\u53ca\u7164\u539a\u573a\u793a\u4f8b",
    "\u0032\u002e\u0033 \u591a\u573a\u666f ODI \u98ce\u9669\u8868\u5f81\u7ed3\u679c",
    "\u0033 \u5019\u9009\u65b9\u6848\u5bf9\u6bd4\u4e0e\u98ce\u9669\u7edf\u8ba1",
    "\u0033\u002e\u0031 \u5019\u9009\u65b9\u6848\u751f\u6210\u7ed3\u679c\u4e0e\u5bf9\u6bd4",
    "\u0033\u002e\u0032 \u9608\u503c\u4e0e\u6743\u91cd\u654f\u611f\u6027\u5206\u6790",
    "\u0033\u002e\u0033 \u89c4\u5212\u7ed3\u679c\u4f20\u9012\u4e0e\u8bc4\u4ef7\u8fb9\u754c",
    "\u0034 \u8ba8\u8bba",
    "\u0034\u002e\u0031 ODI \u524d\u7f6e\u7ea6\u675f\u5bf9\u91c7\u533a\u89c4\u5212\u7684\u4f5c\u7528\u4e0e\u8fb9\u754c",
    "\u0034\u002e\u0032 \u53c2\u6570\u573a\u9a71\u52a8\u7684\u89c4\u5212\u65b9\u6cd5\u76f8\u5bf9\u4e8e\u4f20\u7edf\u7ecf\u9a8c\u5e03\u7f6e\u7684\u610f\u4e49",
    "\u0034\u002e\u0033 \u540e\u7eed\u6df1\u5316\u65b9\u5411",
    "\u0035 \u7ed3\u8bba",
    "\u53c2\u8003\u6587\u732e",
]

EXPECTED_NUMBERS = {
    "field_grid": ["80\u00d756", "4480"],
    "field_stats": ["0.4669", "0.7474", "15.89%", "3.55%"],
    "abc_mean": ["0.4463", "0.4552", "0.4416"],
    "abc_p90": ["0.6407", "0.6462", "0.6353"],
    "abc_exceed": ["0.44%", "0.56%", "1.22%"],
    "thresholds": ["0.65", "0.70", "0.75", "0.80"],
}

OLD_OR_RISKY = [
    "0.3872",
    "0.7414",
    "14.25",
    "911343",
    "NPV",
    "\u9000\u4fee",
    "\u5f85\u590d\u6838",
    "\u4ecd\u9700\u8865\u7b97",
    "qualified=false",
    "\u7a33\u5b9a\u4f20\u9012",
    "\u53ef\u7528\u4e8e\u5de5\u7a0b\u5b9a\u6848",
    "\u7a33\u5b9a\u964d\u4f4e",
    "\u4e3b\u8981\u8bc1\u660e",
]


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: full_consistency_audit_20260418.py <docx> <report-md>")
        return 2

    docx_path = Path(sys.argv[1])
    report_path = Path(sys.argv[2])
    with zipfile.ZipFile(docx_path) as zf:
        bad = zf.testzip()
    if bad is not None:
        raise RuntimeError(f"Bad DOCX zip member: {bad}")

    doc = Document(str(docx_path))
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    text = "\n".join(paragraphs)

    heading_positions = {}
    for expected in EXPECTED_HEADINGS:
        heading_positions[expected] = next((i for i, p in enumerate(paragraphs) if p == expected), None)

    captions = []
    for i, p in enumerate(paragraphs):
        if (
            re.match(r"^\u56fe\d+\s", p)
            or p.startswith("Fig.")
            or re.match(r"^\u8868\d+\s", p)
            or p.startswith("Table")
        ):
            captions.append((i, p))

    cn_abstract = next((p for p in paragraphs if p.startswith("\u6458\u8981\uff1a")), "")
    en_abstract = next((p for p in paragraphs if p.startswith("Abstract:")), "")
    contribution = next((p for p in paragraphs if p.startswith("\u672c\u6587\u7684\u4e3b\u8981\u8d21\u732e")), "")
    conclusions = [p for p in paragraphs if re.match(r"^[1-4]\uff09", p)]

    number_presence = {
        group: {item: text.count(item) for item in values}
        for group, values in EXPECTED_NUMBERS.items()
    }
    old_presence = {item: text.count(item) for item in OLD_OR_RISKY}

    formula_paras = []
    for i, p in enumerate(paragraphs):
        if re.search(r"\uff08\d+\uff09\s*$", p):
            formula_paras.append((i, p[-100:]))

    table_summaries = []
    for ti, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns)
        header = " | ".join(cell.text.strip().replace("\n", " / ") for cell in table.rows[0].cells) if rows else ""
        table_summaries.append((ti, rows, cols, header))

    issues = []
    missing_headings = [h for h, pos in heading_positions.items() if pos is None]
    if missing_headings:
        issues.append(f"Missing headings: {missing_headings}")
    for group, counts in number_presence.items():
        missing = [k for k, v in counts.items() if v == 0]
        if missing:
            issues.append(f"Missing expected numbers in {group}: {missing}")
    risky_nonzero = {k: v for k, v in old_presence.items() if v > 0}
    allowed_boundary = {"C_old"}
    if risky_nonzero:
        issues.append(f"Old/risky term counts need review: {risky_nonzero}")

    lines = [
        "# 全文一致性总审记录",
        "",
        f"- 检查时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"- 文档: `{docx_path}`",
        f"- 段落数: {len(paragraphs)}",
        f"- 表格数: {len(doc.tables)}",
        f"- DOCX压缩结构: 通过",
        "",
        "## 章节结构",
        "",
    ]
    for h, pos in heading_positions.items():
        lines.append(f"- {h}: {pos if pos is not None else '缺失'}")

    lines += ["", "## 图表题名", ""]
    for pos, cap in captions:
        lines.append(f"- 段落{pos}: {cap}")

    lines += ["", "## 表格结构", ""]
    for ti, rows, cols, header in table_summaries:
        lines.append(f"- 表{ti + 1}: {rows}行 x {cols}列；表头: {header}")

    lines += ["", "## 关键数值出现次数", ""]
    for group, counts in number_presence.items():
        lines.append(f"- {group}: {counts}")

    lines += ["", "## 旧口径/风险词检查", ""]
    for item, count in old_presence.items():
        lines.append(f"- {item}: {count}")

    lines += ["", "## 摘要-引言-结论口径", ""]
    lines.append(f"- 中文摘要含ODI/候选方案/敏感性: {all(k in cn_abstract for k in ['ODI', '\u5019\u9009\u65b9\u6848', '\u654f\u611f\u6027'])}")
    lines.append(f"- 英文摘要含ODI/candidate/sensitivity: {all(k in en_abstract for k in ['ODI', 'candidate', 'sensitivity'])}")
    lines.append(f"- 引言贡献条数: {len(re.findall(r'[1-3]\uff09', contribution))}")
    lines.append(f"- 结论条数: {len(conclusions)}")

    lines += ["", "## 公式段落", ""]
    for pos, formula in formula_paras:
        lines.append(f"- 段落{pos}: ...{formula}")

    lines += ["", "## 审核结论", ""]
    if issues:
        lines.extend([f"- 需复核: {issue}" for issue in issues])
    else:
        lines.append("- 未发现章节、关键数值或旧口径残留的硬性不一致。")
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text("\n".join(lines), encoding="utf-8")

    print(report_path)
    print(f"issues={len(issues)}")
    for issue in issues:
        print(issue)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
