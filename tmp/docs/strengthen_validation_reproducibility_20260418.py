from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.style = paragraph.style
    run = inserted.add_run(text)
    if paragraph.runs:
        src = paragraph.runs[0]
        run.font.name = src.font.name
        run.font.size = src.font.size
        run.bold = src.font.bold
        run.italic = src.font.italic
        try:
            run._element.rPr.rFonts.set(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
                src._element.rPr.rFonts.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"
                )
                or "宋体",
            )
        except Exception:
            pass
    inserted.paragraph_format.first_line_indent = paragraph.paragraph_format.first_line_indent
    inserted.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing
    inserted.paragraph_format.space_after = paragraph.paragraph_format.space_after
    inserted.paragraph_format.space_before = paragraph.paragraph_format.space_before
    return inserted


def find_paragraph(doc: Document, needle: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            return paragraph
    raise ValueError(f"Paragraph not found: {needle}")


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: strengthen_validation_reproducibility_20260418.py <docx>")
        return 2

    path = Path(sys.argv[1])
    doc = Document(str(path))

    insertions = [
        (
            "因此案例定位为方法验证型工程样例",
            "为避免验证目标泛化，本文将样例验证限定为3个可复核问题：第一，离散钻孔和边界约束能否转化为统一坐标下的连续参数场和有效布置域；第二，多场景扰动风险能否转化为同一ODI统计口径，并在全域栅格和候选方案区域内复算；第三，不同目标偏好下的候选方案能否在同一约束、同一风险场和同一评分口径下进行比较。后文A、B、C方案对比、阈值敏感性和权重敏感性均围绕上述3个问题展开。",
        ),
        (
            "该对照属于同一候选池内的目标偏好对照",
            "为保证对比的可复核性，A、B、C方案不改变基础地质输入和风险场，只改变候选池排序偏好或筛选目标；因此，方案差异主要反映工程效率、资源回收和扰动控制目标之间的权衡，而不是输入数据、网格尺度或边界条件变化造成的差异。该设置能够支撑方法内部有效性判断，但不能替代真实人工经验方案与本文方法之间的外部工程对照。",
        ),
        (
            "当前结果表明方法链和对象链在样例条件下具备贯通性",
            "从可复现角度看，本文样例结果至少需要保留4类中间数据：研究区边界与钻孔样点、连续参数场栅格、ODI分量及综合ODI栅格、候选方案几何对象及其方案级统计表。只有这些中间对象能够被逐级复算，ODI均值、P90、超阈值暴露比例和A/B/C方案排序才具备可审查性。本文在结果表中保留候选池规模、全域栅格数量和统一ODI统计口径，目的即在于降低单纯展示最终方案图带来的不可复核风险。",
        ),
        (
            "只有完成上述验证后，本文方法才能从“可复核的规划链路”进一步发展为“可辅助工程方案论证的决策工具”。",
            "因此，对本文当前阶段更合适的评价标准不是“是否已经给出某矿最终开采设计”，而是“是否把采区规划中原本分散的地质参数、风险约束、候选方案和评价输入转化为可计算、可比较、可复查的对象链”。在此边界下，本文补充的变量定义、约束集合、候选池生成、方案级ODI统计和敏感性分析共同构成论文的主要证据链。",
        ),
    ]

    inserted = 0
    for needle, text in insertions:
        paragraph = find_paragraph(doc, needle)
        if text not in [p.text for p in doc.paragraphs]:
            insert_after(paragraph, text)
            inserted += 1

    doc.save(str(path))
    print(f"inserted={inserted}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
