from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def set_run_font(run, font_name: str, size_pt: float) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = r_pr._add_rFonts()
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)


def clear_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_cell_width(cell, width) -> None:
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.tcW
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width.twips)))


def delete_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: format_equations_borderless_tables_20260418.py <docx>")
        return 2

    path = Path(sys.argv[1])
    doc = Document(str(path))

    targets = []
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        match = re.fullmatch(r"\t?(.+?)\t?\s*\uff08(\d+)\uff09", text)
        if not match:
            continue
        number = int(match.group(2))
        if 1 <= number <= 15 and len(text) < 240:
            expr = match.group(1).strip()
            if expr.endswith("\u3002"):
                expr = expr[:-1]
            targets.append((idx, paragraph, expr, number))

    if len(targets) != 15:
        raise RuntimeError(f"Expected 15 formula paragraphs, found {len(targets)}")

    created = []
    for idx, paragraph, expr, number in reversed(targets):
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        clear_table_borders(table)
        table._tbl.getparent().remove(table._tbl)
        paragraph._p.addnext(table._tbl)

        widths = [Inches(0.35), Inches(5.55), Inches(0.65)]
        for col_idx, cell in enumerate(table.rows[0].cells):
            set_cell_width(cell, widths[col_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            cell_p = cell.paragraphs[0]
            cell_p.paragraph_format.space_before = Pt(0)
            cell_p.paragraph_format.space_after = Pt(0)

        eq_p = table.cell(0, 1).paragraphs[0]
        eq_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eq_run = eq_p.add_run(expr)
        set_run_font(eq_run, "Cambria Math", 10.5)

        num_p = table.cell(0, 2).paragraphs[0]
        num_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        num_run = num_p.add_run(f"\uff08{number}\uff09")
        set_run_font(num_run, "Times New Roman", 10.5)

        delete_paragraph(paragraph)
        created.append((idx, number, expr))

    doc.save(str(path))
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
    if bad is not None:
        raise RuntimeError(f"Bad DOCX zip member: {bad}")

    print(f"DOCX OK: {path}")
    print(f"Formula tables created: {len(created)}")
    print("numbers=" + ",".join(str(n) for _, n, _ in sorted(created)))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
