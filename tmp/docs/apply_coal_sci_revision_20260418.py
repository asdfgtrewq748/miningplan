import os
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement


path = Path(os.environ["DOCX_PATH"])
doc = Document(path)


def set_text(paragraph, text):
    paragraph.clear()
    paragraph.add_run(text)
    return paragraph


def insert_after(paragraph, text, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if style is not None:
        try:
            new_para.style = style
        except Exception:
            pass
    new_para.add_run(text)
    return new_para


def find_para(starts):
    if isinstance(starts, str):
        starts = [starts]
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if any(text.startswith(start) for start in starts):
            return paragraph
    raise ValueError(f"paragraph not found: {starts}")


# 摘要、关键词与英文摘要
set_text(
    find_para("摘要："),
    "摘要：针对采区规划中离散地质信息难以连续参与空间布置、多场景覆岩扰动风险约束口径不统一以及规划结果难以向采掘接续和经济评价传递等问题，提出一种基于覆岩扰动指数（overburden disturbance index，ODI）约束的采区候选方案生成与多目标比选方法。该方法以采区边界、钻孔样点和设计参数为输入，先通过约束内缩、空间插值和参数映射构建有效布置域与连续参数场，再将地表沉陷、含水层扰动和上行开采等风险分量归一化并加权为统一ODI指标，进一步以ODI均值、90%分位数和超阈值暴露比例刻画候选方案风险水平。在此基础上，定义工作面数量、推进方向、煤柱宽度、工作面宽度和方案选择变量，建立包含边界、煤柱、长宽比、连续性、不重叠和风险阈值等约束的候选方案池，并采用工程效率、资源回收和扰动控制目标进行非支配排序与综合筛选。样例结果表明，工程效率优先、资源回收优先和扰动控制优先方案分别体现出覆盖率、资源利用和风险暴露之间的权衡；其中扰动控制候选方案的ODI均值为0.3872，P90为0.7414，ODI>0.70暴露比例为14.25%。研究结果说明，ODI前置约束可将风险图层转化为可复核的方案级统计指标，为采区规划阶段的多目标比选和后续对象传递提供方法支撑。",
)
set_text(
    find_para("关键词："),
    "关键词：覆岩扰动指数；采区规划；多目标优化；候选方案池；风险前置约束；连续参数场",
)
set_text(
    find_para("Abstract:"),
    "Abstract: To address the difficulty of using discrete geological information in continuous mining-district layout, the inconsistent expression of multi-scenario overburden disturbance risks, and the weak linkage between planning outputs and downstream succession or economic evaluation, this study proposes an ODI-constrained candidate generation and multi-objective selection method for mining-district planning. The method constructs an effective layout domain and continuous parameter fields from boundary, borehole and design inputs, normalizes surface subsidence, aquifer disturbance and upward-mining risk components into a unified overburden disturbance index (ODI), and describes candidate-level risk exposure using the mean ODI, the 90th percentile and the exceedance ratio. Decision variables including face number, advancing direction, pillar width, face width and scheme selection are then organized into a constrained candidate pool. Boundary, pillar, aspect-ratio, continuity, non-overlap and risk-threshold constraints are combined with engineering-efficiency, resource-recovery and disturbance-control objectives, and the candidate solutions are ranked by non-dominated sorting and weighted screening. The case results show clear trade-offs among an engineering-efficiency-oriented scheme, a resource-recovery-oriented scheme and a disturbance-control-oriented scheme. For the disturbance-control candidate, the mean ODI, P90 and ODI>0.70 exposure ratio are 0.3872, 0.7414 and 14.25%, respectively. The results indicate that ODI pre-constraint can transform risk layers into reproducible scheme-level statistics and provide methodological support for multi-objective comparison and downstream data transfer in mining-district planning.",
)
set_text(
    find_para("Keywords:"),
    "Keywords: overburden disturbance index; mining district planning; multi-objective optimization; candidate pool; risk pre-constraint; continuous parameter field",
)

# 引言贡献段
set_text(
    find_para("基于此，本文提出一种基于覆岩扰动指数"),
    "基于上述问题，本文提出一种基于覆岩扰动指数（overburden disturbance index，ODI）约束的采区候选方案生成与多目标比选方法。与在布局完成后再进行风险校核的串行流程不同，本文将地表沉陷、含水层扰动和上行开采等异构风险前置转化为统一ODI约束，并在有效布置域、连续参数场和工作面布置对象之间建立可计算映射关系，使风险控制能够参与候选方案生成、筛选和排序过程。",
)
set_text(
    find_para("本文主要开展以下工作"),
    "本文的主要贡献为：1）构建面向采区规划的ODI统一风险表征方法，明确风险分量归一化、权重聚合、阈值统计和方案级指标计算口径；2）建立离散钻孔、连续参数场与工作面布局对象之间的规划映射机制，使地质参数能够从点状样本转化为可参与方案生成和评价的空间变量；3）构建包含工程效率、资源回收与扰动控制目标的候选方案池和排序流程，并通过工程效率优先、资源回收优先、扰动控制优先等方案对比说明ODI前置约束在方案比选中的作用。本文结果主要用于验证方法链路和对象传递的可行性，不将单一样例扩展为真实矿井条件下的普适最优结论。",
)

# 1.1 ODI定义与公式口径
set_text(find_para("1.1 覆岩扰动约束机理"), "1.1 面向采区规划的 ODI 风险约束定义")
set_text(
    find_para("采区规划的本质是在几何可实施性"),
    "采区规划中的覆岩扰动风险并非由单一物理量决定，而是由地表沉陷、含水层扰动、上行开采安全性以及局部地质采矿条件共同作用形成。不同风险场景的原始量纲、判据和空间尺度存在差异，若分别以各自阈值进入规划过程，容易造成约束口径不统一和方案比较困难。本文不把ODI定义为新的覆岩力学机理参数，而将其作为面向采区规划的统一风险表征指标，用于把多场景风险转化为可参与候选方案筛选和排序的无量纲统计量。",
)
set_text(
    find_para("对规划域内任意位置x，ODI可表示为"),
    "设规划域为Ω，任意位置x处的风险分量包括地表沉陷扰动D_s(x)、含水层扰动D_a(x)和上行开采扰动D_u(x)。各分量均归一化到[0,1]区间，数值越大表示扰动风险越高，则ODI定义为",
)
set_text(doc.paragraphs[27], "ODI(x)=w_sD_s(x)+w_aD_a(x)+w_uD_u(x)                                           （1）")
set_text(doc.paragraphs[28], "其中权重满足")
set_text(doc.paragraphs[29], "w_s+w_a+w_u=1，w_s≥0，w_a≥0，w_u≥0。                                           （2）")
set_text(
    find_para("式中，ODI为覆岩扰动指数"),
    "式中，w_s、w_a和w_u分别为地表沉陷、含水层扰动和上行开采分量权重。综合规划基准采用0.45/0.30/0.25，含水层专项场景可采用0.15/0.25/0.60。为避免“ODI”与子项语义重叠，本文不再使用“覆岩扰动分量”作为子指标名称，而以具体风险场景命名各分量。",
)
set_text(
    find_para("进一步定义候选方案π的ODI统计指标为"),
    "对原始指标X_i(x)，正向风险指标按式（3）归一化，逆向风险指标按式（4）归一化；当工程经验或规范给出上下限时，优先采用工程阈值作为X_i,min和X_i,max，否则采用同一研究区同一场景下的样本范围。",
)
set_text(doc.paragraphs[32], "X_i'(x)=[X_i(x)-X_i,min]/[X_i,max-X_i,min]                                      （3）")
set_text(doc.paragraphs[33], "X_i'(x)=[X_i,max-X_i(x)]/[X_i,max-X_i,min]                                      （4）")
set_text(doc.paragraphs[35], "E_π=N_π(ODI>T_ODI)/N_π×100%                                                   （5）")
set_text(
    find_para("式中，π为候选规划方案"),
    "式中，π为候选规划方案，A_π为方案π对应的布置区域，N_π为方案区域内参与统计的采样或栅格数量，T_ODI为扰动控制阈值。本文以ODI均值、Q_0.90和超阈值暴露比例E_π共同描述方案级风险，不把单一ODI均值作为唯一判据。阈值0.70和0.80分别作为预警统计线和高扰动统计线，用于方案间比较；在真实矿井应用中应结合历史观测、保护对象等级和现场安全制度重新标定。",
)

# 1.2 强化IDW和数据处理说明
set_text(
    find_para("式中，z(x)为规划位置x处的参数估计值"),
    "式中，z(x)为规划位置x处的参数估计值，z_i为第i个钻孔样点的观测值，d_i(x)为位置x与第i个钻孔之间的距离，p为距离衰减指数。本文采用反距离加权插值，是因为当前样例钻孔数量有限且目标在于形成规划可用的连续约束场，而非证明复杂地质统计模型的优越性。插值前应统一坐标、剔除明显异常值并限定插值边界；插值后应记录网格尺寸、样点范围和边界裁剪规则。若后续进入实矿验证阶段，应补充留一交叉验证或与克里金插值结果对比，以量化参数场误差对规划结果的影响。",
)

# 1.3 多目标模型重写
set_text(find_para("1.3 候选方案生成与多目标协同规划模型"), "1.3 候选方案生成与多目标比选模型")
set_text(
    find_para("在满足几何可实施性和安全隔离要求"),
    "在满足几何可实施性和安全隔离要求的前提下，采区规划需要在工程效率、资源回收和覆岩扰动控制之间进行权衡。本文将规划问题定义为“候选方案池生成—硬约束过滤—多目标排序—工程复核”的组合优化过程，而不是单一程序输出结果。",
)
set_text(
    find_para("设候选方案集合为Π"),
    "设候选方案集合为Π，候选方案π由工作面数量N、工作面宽度W_f、推进方向θ、边界煤柱宽度B_b、区段煤柱宽度B_s、工作面起止边界、巷道连接关系和方案选择变量y_π共同描述，可写为",
)
set_text(doc.paragraphs[50], "π={N,W_f,θ,B_b,B_s,A_π,L_π,R_π,y_π}，y_π∈{0,1}。                         （8）")
set_text(doc.paragraphs[51], "候选方案需满足以下约束：")
set_text(
    doc.paragraphs[52],
    "A_π⊂Ω_e；B_b∈[B_b,min,B_b,max]；B_s∈[B_s,min,B_s,max]；W_f∈[W_min,W_max]；A_i∩A_j=∅；L_f≥L_min；E_π≤E_max。        （9）",
)
set_text(
    find_para("式中，S_e"),
    "式中，A_π⊂Ω_e为边界约束，B_b和B_s为煤柱约束，W_f为工作面宽度约束，A_i∩A_j=∅为工作面不重叠约束，L_f≥L_min为最小推进长度约束，E_π≤E_max为ODI超阈值暴露约束。上述约束用于把“能画出的布局”进一步筛选为“可进入工程比选的候选方案”。",
)
set_text(
    find_para("具体而言，工程效率"),
    "候选池生成采用规则枚举与工程过滤相结合的方式：在给定推进方向、煤柱宽度、工作面宽度和边界内缩方案的参数组合下生成候选工作面条带，剔除越界、重叠、推进长度不足和煤柱不满足要求的方案，再计算工程效率S_e(π)、资源回收S_r(π)和扰动控制S_m(π)。工程效率主要考虑覆盖率、工作面连续性、推进长度均衡性和巷道组织便利性；资源回收主要考虑布置面积、煤厚场覆盖和可采资源覆盖；扰动控制主要考虑ODI均值、P90和超阈值暴露比例。",
)
set_text(find_para("对通过必要安全与几何校核"), "对通过必要安全与几何校核的候选方案，综合得分可表示为")
set_text(doc.paragraphs[61], "G(s)=αP_s+βR_s+γC_s，α+β+γ=1。                                      （11）")
set_text(doc.paragraphs[64], "NCF_t=Rev_t-Cost_t-RiskCost_t。                                      （12）")
new_formula = insert_after(find_para("对通过必要安全与几何校核"), "F(π)=λ_eS_e(π)+λ_rS_r(π)+λ_mS_m(π)，λ_e+λ_r+λ_m=1。                     （10）")
insert_after(
    new_formula,
    "为降低单一加权排序对权重的依赖，本文在加权排序前引入非支配排序和拥挤距离选择形成Top-K候选方案，再依据综合权重进行推荐排序。工程效率优先、资源回收优先、扰动控制优先和综合权衡模式均来自同一候选池，仅代表不同评价偏好，不应被表述为真实矿井条件下的全局最优解。",
)

# 2.3 ODI结果中的待复核口径
set_text(
    find_para("从研究区样例结果看，多场景ODI结果已经能够"),
    "从研究区样例结果看，多场景ODI结果能够在同一研究区边界、钻孔参照和色标口径下表达。需要说明的是，正文原有80×56网格、4480栅格及ODI均值0.4669、P90为0.7474等全域统计，尚需由原始栅格数据或程序重新导出复核；因此本文在本轮修改中将其作为全域风险背景口径保留讨论，而不直接作为方案优选的唯一证据。",
)
set_text(
    find_para("在规划语境下，ODI的价值体现在两个层面"),
    "在规划语境下，ODI的价值体现在两个层面：一是描述研究区高扰动敏感区的空间背景，二是提供候选方案区域内的风险暴露统计。已复核的扰动控制候选方案显示，方案区域ODI均值为0.3872，P90为0.7414，ODI>0.70暴露比例为14.25%，说明风险场可以进一步转化为可复核的方案级指标。",
)

# 结果部分改写
set_text(find_para("3 结构化规划结果与风险统计"), "3 候选方案对比与风险统计")
set_text(find_para("3.1 结构化规划结果"), "3.1 候选方案生成结果与对比")
set_text(
    find_para("在有效布置域、连续参数场和ODI风险场共同作用下"),
    "在有效布置域、连续参数场和ODI风险场共同作用下，本文进一步提取工程效率优先、资源回收优先和扰动控制优先等候选结果，用于说明不同目标偏好下方案指标的变化。与仅展示单一导出样例相比，候选方案对比能够更直接回答“ODI约束如何影响方案筛选”的问题。",
)
set_text(
    find_para("表3列出了当前样例的结构化规划结果"),
    "表3列出了当前已复核的候选方案。工程效率优先方案A来自efficiency模式，候选池规模为2417，覆盖率为89.34%，工程效率评分为87.82；资源回收优先方案B来自recovery模式，候选池规模为1149，覆盖率为98.52%，资源回收评分为89.76；扰动控制优先方案C来自disturbance模式，ODI均值为0.3872，P90为0.7414，ODI>0.70暴露比例为14.25%。",
)
set_text(
    find_para("从布局结果本身看"),
    "上述结果表明，不同目标偏好会导致覆盖率、资源回收和风险暴露之间出现明显权衡。A、B方案能够作为单目标基准，但当前工程文件未保存其方案级ODI统计，因此仍需补算统一风险指标；C方案具备已复核ODI统计，可作为ODI风险优先方案。weighted模式的当前top 1候选虽然具有综合排序信息，但其合格性标记为false，本文不将其作为最终工程推荐方案。",
)
set_text(find_para("表3 当前样例规划结果统计表"), "表3 候选方案对比与数据复核状态")
set_text(find_para("3.2 规划结果向采掘接续"), "3.2 规划结果向采掘接续与工程经济评价的传递")
set_text(
    find_para("在调控层面，规划结果还可进一步转化为工作面级控制变量"),
    "在调控层面，规划结果还可进一步转化为工作面级控制变量。当前已复核的扰动控制候选方案采用ODI阈值0.70、外扩缓冲30 m和751个采样点形成方案级统计，其ODI均值、P90和超阈值暴露比例可作为后续采高、煤柱宽度、推进方向和工作面宽度调整的反馈指标。",
)
set_text(
    find_para("综上，研究区样例表明"),
    "综上，研究区样例表明，本文方法已完成“边界与钻孔输入—参数场构建—ODI风险组织—候选方案池—方案级统计—后续评价输入”的连续运行过程。当前结果能够证明方法链和对象链已经贯通，但A、B方案的ODI统计和全域4480栅格统计仍需进一步复核，不宜把单一样例结果表述为真实矿井条件下的最终优选结论。",
)
set_text(find_para("表4 当前样例参数与ODI统计表"), "表4 ODI统计口径与后续补算需求")

# 讨论与结论
set_text(find_para("4.1 ODI 前置约束"), "4.1 ODI 前置约束对采区规划的作用与边界")
set_text(
    find_para("从研究区样例看，ODI场提供了可量化的风险入口"),
    "从研究区样例看，ODI场提供了可量化的风险入口。与只给出风险图层不同，本文把ODI均值、P90和超阈值暴露比例转化为候选方案统计量，使方案A、B、C能够在同一风险口径下比较。当前已复核的C方案ODI均值为0.3872，P90为0.7414，ODI>0.70暴露比例为14.25%。",
)
set_text(
    find_para("进一步看，ODI统一表达的价值还体现在连续传递能力上"),
    "进一步看，ODI统一表达的价值还体现在连续传递能力上。由于ODI在规划阶段已形成方案级统计指标，因此可作为后续采高调控、采掘接续和经济评价的上游输入。但ODI本质上是多源风险压缩指标，可能掩盖单一风险场景的局部极值，因此工程应用中仍应同时保留分量图层和关键阈值校核。",
)
set_text(
    find_para("后续研究应在实矿数据中补充"),
    "后续研究应在实矿数据中补充多方案对照、A/B方案ODI补算、权重与阈值敏感性分析、接续排程数值和经济评价参数体系，使当前已经打通的方法链进一步转化为可量化的工程优选依据。",
)
set_text(
    find_para("1）针对采区规划中地质信息离散"),
    "1）针对采区规划中地质信息离散、风险约束异构以及规划结果与后续评价脱节的问题，提出了基于ODI约束的采区候选方案生成与多目标比选方法，构建了“有效布置域—连续参数场—ODI风险场—候选方案池—后续评价输入”的一体化决策链。",
)
set_text(
    find_para("2）通过将地表沉陷"),
    "2）通过将地表沉陷、含水层扰动和上行开采等多场景风险统一组织为ODI指标，明确了风险分量归一化、权重聚合、阈值统计和方案级ODI均值、P90、超阈值暴露比例等计算口径，为异构风险进入采区规划比选提供了统一表达。",
)
set_text(
    find_para("3）研究区样例验证表明"),
    "3）研究区样例验证表明，工程效率优先、资源回收优先和扰动控制优先方案分别体现出覆盖率、资源利用和风险暴露之间的权衡。其中已复核的扰动控制候选方案ODI均值为0.3872，P90为0.7414，ODI>0.70暴露比例为14.25%，说明风险场能够转化为可复核的方案级统计量。",
)
set_text(
    find_para("4）当前结果主要支撑"),
    "4）当前结果主要支撑样例条件下的方法贯通和对象链有效性。A、B方案的ODI统计、4480栅格全域结果、权重阈值敏感性和真实矿井工程约束仍需进一步补算与标定，真实矿井条件下的工程优选结论仍需结合实矿案例检验。",
)

# 表格更新
t0 = doc.tables[0]
t0.cell(12, 2).text = "w_s/w_a/w_u"
t0.cell(12, 3).text = "0.45/0.30/0.25"
t0.cell(12, 5).text = "地表沉陷/含水层扰动/上行开采；综合规划基准权重"
t0.cell(13, 5).text = "用于方案级超阈值暴露统计；真实工程需结合现场阈值标定"

t2 = doc.tables[2]
while len(t2.rows) > 5:
    tr = t2.rows[-1]._tr
    tr.getparent().remove(tr)
rows = [
    ["方案", "来源/含义", "关键指标", "ODI统计", "复核状态"],
    ["A", "efficiency；工程效率优先", "工作面5个；覆盖率89.34%；工程效率评分87.82", "待补算", "可作单目标基准"],
    ["B", "recovery；资源回收优先", "工作面9个；覆盖率98.52%；资源回收评分89.76", "待补算", "可作无ODI/资源优先基准"],
    ["C", "disturbance；扰动控制优先", "工作面13个；覆盖率75.06%", "均值0.3872；P90=0.7414；ODI>0.70为14.25%", "已在工程文件中复核"],
    ["C2", "weighted top 1；综合排序候选", "工作面5个；覆盖率56.52%；综合得分0.9223", "均值0.6808；P90=0.8364；ODI>0.70为60.11%", "qualified=false，不作为最终推荐方案"],
]
for row_index, row_values in enumerate(rows):
    for col_index, value in enumerate(row_values):
        t2.cell(row_index, col_index).text = value

t3 = doc.tables[3]
while len(t3.rows) > 6:
    tr = t3.rows[-1]._tr
    tr.getparent().remove(tr)
rows = [
    ["统计对象", "当前数值/状态", "数据来源", "用途", "修改处理"],
    ["全域ODI栅格", "80×56、4480栅格；均值0.4669；P90=0.7474", "原正文结果，待原始栅格复核", "研究区风险背景", "保留为待复核口径，不作强结论"],
    ["方案C ODI统计", "均值0.3872；P90=0.7414；ODI>0.70为14.25%", "3-采区规划案例.miningplan.json", "方案级风险评价", "作为已复核结果写入正文"],
    ["权重敏感性", "综合0.45/0.30/0.25；专项0.15/0.25/0.60", "导出包与场景参数", "检验方案稳定性", "需补±10%或多场景对照"],
    ["阈值敏感性", "0.70/0.80", "本文统计口径", "检验超限暴露稳定性", "需补0.65/0.70/0.75/0.80对比"],
    ["经济评价", "仅保留输入关系", "规划对象与接续参数", "下游接口说明", "不写未经独立导出的NPV和回收期"],
]
for row_index, row_values in enumerate(rows):
    for col_index, value in enumerate(row_values):
        t3.cell(row_index, col_index).text = value

doc.save(path)
print(f"saved {path}")
print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} inline_shapes={len(doc.inline_shapes)}")
