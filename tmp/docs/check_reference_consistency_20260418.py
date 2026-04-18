from __future__ import annotations

import csv
import re
import sys
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document


MAX_REF_NO = 50


def expand_citation_group(raw: str) -> set[int]:
    text = (
        raw.replace("，", ",")
        .replace("－", "-")
        .replace("–", "-")
        .replace("—", "-")
        .replace(" ", "")
    )
    if not text:
        return set()

    parts = [p for p in text.split(",") if p]
    if not parts:
        return set()

    expanded: set[int] = set()
    for part in parts:
        if re.fullmatch(r"\d+", part):
            num = int(part)
            if not 1 <= num <= MAX_REF_NO:
                return set()
            expanded.add(num)
            continue
        match = re.fullmatch(r"(\d+)-(\d+)", part)
        if match:
            start, end = int(match.group(1)), int(match.group(2))
            if not (1 <= start <= end <= MAX_REF_NO):
                return set()
            expanded.update(range(start, end + 1))
            continue
        return set()

    return expanded


def find_reference_heading(paragraphs) -> int:
    for idx, para in enumerate(paragraphs):
        text = para.text.strip()
        if text == "参考文献" or ("参考" in text and "文献" in text and len(text) <= 10):
            return idx
    raise RuntimeError("Reference heading not found")


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: check_reference_consistency_20260418.py <docx> <output-md>")
        return 2

    docx_path = Path(sys.argv[1])
    report_path = Path(sys.argv[2])
    doc = Document(str(docx_path))

    with zipfile.ZipFile(docx_path) as zf:
        bad = zf.testzip()
    if bad is not None:
        raise RuntimeError(f"Bad DOCX zip member: {bad}")

    ref_idx = find_reference_heading(doc.paragraphs)
    body_text = "\n".join(p.text for p in doc.paragraphs[:ref_idx])
    ref_text = "\n".join(p.text for p in doc.paragraphs[ref_idx + 1 :])

    ref_numbers = sorted({int(m.group(1)) for m in re.finditer(r"^\[(\d+)\]", ref_text, re.M)})
    doi_values = re.findall(r"doi:\s*([^\s;。；]+)", ref_text, flags=re.I)

    used: set[int] = set()
    raw_groups: list[str] = []
    ignored_groups: list[str] = []
    for match in re.finditer(r"\[([0-9０-９ ,，\-－–—]+)\]", body_text):
        raw = match.group(1)
        normalized = (
            raw.replace("０", "0")
            .replace("１", "1")
            .replace("２", "2")
            .replace("３", "3")
            .replace("４", "4")
            .replace("５", "5")
            .replace("６", "6")
            .replace("７", "7")
            .replace("８", "8")
            .replace("９", "9")
        )
        expanded = expand_citation_group(normalized)
        if expanded:
            used.update(expanded)
            raw_groups.append(f"[{raw}]")
        else:
            ignored_groups.append(f"[{raw}]")

    missing_in_body = [n for n in ref_numbers if n not in used]
    missing_in_refs = [n for n in sorted(used) if n not in ref_numbers]
    duplicate_refs = [n for n in ref_numbers if ref_numbers.count(n) > 1]
    expected_missing = [n for n in range(1, MAX_REF_NO + 1) if n not in ref_numbers]

    report_path.parent.mkdir(parents=True, exist_ok=True)
    lines = [
        "# 参考文献正文引用一致性检查",
        "",
        f"- 检查时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"- 文档: `{docx_path}`",
        f"- 参考文献标题段落索引: {ref_idx}",
        f"- 参考文献条目数: {len(ref_numbers)}",
        f"- 参考文献编号范围: {min(ref_numbers) if ref_numbers else 'NA'}-{max(ref_numbers) if ref_numbers else 'NA'}",
        f"- DOI 数量: {len(doi_values)}",
        f"- DOI 去重数量: {len(set(doi_values))}",
        f"- 正文识别到的引用编号数: {len(used)}",
        f"- 正文引用覆盖范围: {min(used) if used else 'NA'}-{max(used) if used else 'NA'}",
        "",
        "## 检查结果",
        "",
        f"- 参考文献表中存在但正文未引用: {missing_in_body}",
        f"- 正文引用但参考文献表缺失: {missing_in_refs}",
        f"- 参考文献编号缺号: {expected_missing}",
        f"- 参考文献编号重复: {sorted(set(duplicate_refs))}",
        "",
        "## 识别到的正文引用组",
        "",
        ", ".join(raw_groups) if raw_groups else "无",
        "",
        "## 被忽略的方括号组",
        "",
        "以下方括号组未按参考文献处理，主要用于避免把归一化区间等数学表达误判为引用。",
        "",
        ", ".join(sorted(set(ignored_groups))) if ignored_groups else "无",
        "",
    ]
    report_path.write_text("\n".join(lines), encoding="utf-8")

    summary_csv = report_path.with_suffix(".csv")
    with summary_csv.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["metric", "value"])
        writer.writerow(["reference_count", len(ref_numbers)])
        writer.writerow(["doi_count", len(doi_values)])
        writer.writerow(["unique_doi_count", len(set(doi_values))])
        writer.writerow(["used_body_count", len(used)])
        writer.writerow(["missing_in_body", ";".join(map(str, missing_in_body))])
        writer.writerow(["missing_in_refs", ";".join(map(str, missing_in_refs))])
        writer.writerow(["missing_reference_numbers", ";".join(map(str, expected_missing))])
        writer.writerow(["duplicate_reference_numbers", ";".join(map(str, sorted(set(duplicate_refs))))])

    print(report_path)
    print(f"reference_count={len(ref_numbers)} doi_count={len(doi_values)} used_body_count={len(used)}")
    print(f"missing_in_body={missing_in_body}")
    print(f"missing_in_refs={missing_in_refs}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
