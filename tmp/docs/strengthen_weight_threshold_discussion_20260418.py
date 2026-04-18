from copy import deepcopy
from pathlib import Path
import os

from docx import Document


def insert_paragraph_after(paragraph, text, style=None):
    new_p = deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.paragraphs[
        list(paragraph._parent._element).index(new_p)
    ] if False else None
    # python-docx has no public insert-after API; build a blank paragraph by
    # cloning paragraph XML then replacing runs through the wrapper below.
    from docx.text.paragraph import Paragraph

    inserted = Paragraph(new_p, paragraph._parent)
    inserted.clear()
    if style is not None:
        inserted.style = style
    inserted.add_run(text)
    return inserted


def replace_paragraph_starting(doc, prefix, replacement):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            p.clear()
            p.add_run(replacement)
            return True
    return False


def find_paragraph_starting(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None


def main():
    docx_path = Path(os.environ.get("DOCX_PATH", r"E:\xiangmu\miningplan\煤科投稿\最新版论文4.16_插图版_煤科格式_大修工作稿_20260418.docx"))
    doc = Document(docx_path)

    replacements = {
        "式中，w_s、w_a和w_u分别为": (
            "式中，w_s、w_a和w_u分别为地表沉陷、含水层扰动和上行开采分量权重。"
            "本文采用“场景目标驱动的专家规则赋权+敏感性校核”的确定方式，而不把权重解释为已完成现场统计标定的经验常数。"
            "综合规划基准权重取0.45/0.30/0.25，强调一般采区规划中地表沉陷控制、含水层扰动控制和上行开采扰动控制的综合平衡；"
            "含水层专项场景权重取0.15/0.25/0.60，用于检验当覆岩导水裂隙带发育及含水层保护约束被强化时，方案排序是否发生方向性改变。"
            "为避免“ODI”与子项语义重叠，本文不再使用“覆岩扰动分量”作为子指标名称，而以具体风险场景命名各分量。"
            "因此，本文给出的权重属于规划阶段的可复核假设，后续工程应用需结合矿区历史沉陷、水文地质观测和保护对象等级重新标定。"
        ),
        "式中，π为候选规划方案": (
            "式中，π为候选规划方案，A_π为方案π对应的布置区域，N_π为方案区域内参与统计的采样或栅格数量，T_ODI为扰动控制阈值。"
            "本文以ODI均值、Q_0.90和超阈值暴露比例E_π共同描述方案级风险，其中ODI均值反映方案区域的总体扰动水平，"
            "Q_0.90反映高值尾部风险，E_π反映超过统计阈值的空间暴露比例，不把单一ODI均值或单一超限比例作为唯一判据。"
            "阈值0.70和0.80分别作为预警统计线和高扰动统计线，用于样例内部的相对比较，而非行业安全红线或法规阈值。"
            "为降低自定义阈值带来的结论依赖，本文进一步设置0.65、0.70、0.75和0.80四组阈值进行敏感性分析；"
            "在真实矿井应用中，上述阈值仍应结合历史观测、含水层保护等级、地表建构筑物敏感性和现场安全制度重新标定。"
        ),
        "从研究区样例看，ODI场提供了可量化的风险入口。": (
            "从研究区样例看，ODI场提供了可量化的风险入口。与只给出风险图层不同，本文把ODI均值、P90和超阈值暴露比例转化为候选方案统计量，"
            "使方案A、B、C能够在同一风险口径下比较。权重敏感性结果表明，在基准权重、各分量±10%相对扰动以及含水层专项权重条件下，"
            "候选方案的风险综合得分均表现为C最低、A次之、B最高，说明当前样例中的ODI筛选结果对权重小幅扰动具有一定稳定性。"
            "但这一稳定性只针对本文数据、候选池和权重扰动范围成立，不能直接外推为所有采区规划条件下的普适排序。"
        ),
    }

    missing = []
    for prefix, replacement in replacements.items():
        if not replace_paragraph_starting(doc, prefix, replacement):
            missing.append(prefix)

    discussion_anchor = find_paragraph_starting(doc, "从研究区样例看，ODI场提供了可量化的风险入口。")
    if discussion_anchor is not None:
        follow_text = (
            "需要指出的是，本文确定的C方案并不是覆盖率最大的方案，而是在合格候选集中风险综合得分最低的方案。"
            "A方案和B方案分别偏向工程效率与资源回收，覆盖率达到89.34%和98.52%；C方案覆盖率为80.67%，牺牲了部分资源覆盖，"
            "但其ODI均值为0.4416、P90为0.6353，均低于A方案和B方案。"
            "在T_ODI=0.70时，C方案超阈值暴露比例为1.22%，略高于A方案的0.44%和B方案的0.56%，这说明单一超限比例不能完全代表方案风险。"
            "因此，本文采用均值、P90和超限比例的联合判据识别风险-收益折中关系，并把C方案解释为当前样例下的低扰动候选，而非工程定案方案。"
        )
        already = any(p.text.strip().startswith("需要指出的是，本文确定的C方案并不是覆盖率最大的方案") for p in doc.paragraphs)
        if not already:
            insert_paragraph_after(discussion_anchor, follow_text, style=discussion_anchor.style)
    else:
        missing.append("discussion anchor")

    # Table 4 wording: avoid ambiguous symbolic ordering in the manuscript text.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "风险得分排序均为C>A>B" in cell.text:
                    cell.text = cell.text.replace("风险得分排序均为C>A>B", "风险综合得分表现为C最低、A次之、B最高")

    if missing:
        print("Missing anchors:", missing)
    doc.save(docx_path)
    print(f"Updated {docx_path}")


if __name__ == "__main__":
    main()
