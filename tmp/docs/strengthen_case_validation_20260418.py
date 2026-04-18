from __future__ import annotations

import sys
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


def insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.text = text
    return inserted


def insert_before(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    paragraph._p.addprevious(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.text = text
    return inserted


def find_exact(paragraphs: list[Paragraph], text: str) -> int:
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == text:
            return idx
    raise RuntimeError(f"Paragraph not found: {text[:80]}")


P117 = (
    "在有效布置域、连续参数场和ODI风险场共同作用下，本文进一步提取工程效率优先、资源回收优先和扰动控制优先等候选结果，"
    "用于说明不同目标偏好下方案指标的变化。与仅展示单一导出样例相比，候选方案对比能够更直接回答“ODI约束如何影响方案筛选”的问题。"
)

INSERT_AFTER_117 = (
    "为增强案例验证的对照性，本文将A、B、C三类方案置于相同研究区边界、相同钻孔参数场、相同ODI风险场和相同几何约束下进行比较。"
    "其中，A方案代表工程效率偏好，重点体现工作面数量、推进长度均衡性和组织效率；B方案代表资源回收偏好，重点体现覆盖率和煤厚场利用；"
    "C方案代表ODI风险约束偏好，重点体现低扰动暴露。该对照属于同一候选池内的目标偏好对照，而不是独立人工经验方案与自动方案的现场对照。"
)

P119 = (
    "上述结果表明，不同目标偏好会导致覆盖率、资源回收和风险暴露之间出现明显权衡。与A方案相比，C方案覆盖率降低8.67个百分点，"
    "但ODI均值降低1.07%，P90降低0.86%，风险综合得分降低0.70%；与B方案相比，C方案覆盖率降低17.85个百分点，"
    "但ODI均值降低3.00%，P90降低1.69%，风险综合得分降低2.12%。B方案相对于A方案覆盖率提高9.18个百分点，"
    "但ODI均值和P90分别提高1.99%和0.85%。这说明资源覆盖增益可能伴随扰动风险统计值上升，"
    "ODI前置约束能够把这种风险-收益取舍显式化。旧disturbance保存结果中的候选C_old在统一ODI场下均值为0.4560，"
    "P90为0.6472，因此本文不再将其作为统一口径下的ODI筛选方案。"
)

INSERT_AFTER_119 = (
    "从方案合理性看，A方案覆盖率和工程效率较高，说明几何组织较为充分；B方案覆盖率最高且资源回收评分较高，说明连续参数场能够推动方案向资源富集区域扩展；"
    "C方案覆盖率较低，但ODI均值、P90和风险综合得分均低于A、B方案，说明其牺牲部分覆盖面积以降低总体扰动暴露。"
    "这种结果符合采区规划中资源利用与风险控制相互制约的工程常识，也说明本文方法能够把隐含在图层中的风险差异转化为方案级统计证据。"
)

P127 = (
    "综上，研究区样例表明，统一ODI场不仅可以生成候选方案的均值、P90和超阈值暴露比例，还能够支撑阈值敏感性和权重敏感性复核。"
    "当前结果已补齐A、B、C方案的统一ODI统计和全域4480栅格统计，但仍不宜把单一样例结果表述为真实矿井条件下的最终优选结论。"
)

INSERT_AFTER_127 = (
    "因此，本案例的验证强度定位为“方法有效性与内部对照验证”：它能够证明离散钻孔、连续参数场、ODI风险场和候选方案对象可以在同一口径下完成计算与比较，"
    "并能揭示不同目标偏好下的风险-收益权衡；但它尚不能替代真实矿井条件下的人工经验方案对照、现场生产约束闭环和开采后实测反馈校核。"
)

P129 = (
    "本文方法的一个重要特点，是不把采区规划的几何布局视为终点，而是强调规划结果可以继续作为采掘接续和工程经济评价的上游输入。"
    "煤矿智能化技术体系与三维地质建模研究已经将工作面接续设计、储量计算和设备配套纳入数据化流程[28,31-32]，"
    "煤矿地下水库储水空间、库容和调度研究也说明，采动空间对象可继续服务于后续水资源调控和风险管理[42-47]。"
    "在本样例中，工作面边界、巷道路径、推进长度、布置面积和ODI统计量均已形成可传递的规划结果，后续环节可据此组织生产顺序、产量核算和经济评价。"
)

HEADING_33 = "3.3 规划结果传递与评价边界"


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: strengthen_case_validation_20260418.py <docx>")
        return 2

    path = Path(sys.argv[1])
    doc = Document(str(path))
    existing = "\n".join(p.text for p in doc.paragraphs)

    if INSERT_AFTER_117 not in existing:
        insert_after(doc.paragraphs[find_exact(doc.paragraphs, P117)], INSERT_AFTER_117)
    if INSERT_AFTER_119 not in existing:
        insert_after(doc.paragraphs[find_exact(doc.paragraphs, P119)], INSERT_AFTER_119)
    if INSERT_AFTER_127 not in existing:
        insert_after(doc.paragraphs[find_exact(doc.paragraphs, P127)], INSERT_AFTER_127)
    if HEADING_33 not in existing:
        insert_before(doc.paragraphs[find_exact(doc.paragraphs, P129)], HEADING_33)

    doc.save(str(path))
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
    if bad is not None:
        raise RuntimeError(f"Bad DOCX zip member: {bad}")

    print(f"DOCX OK: {path}")
    print("Inserted case-validation paragraphs: 3")
    print("Inserted heading: 3.3")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
