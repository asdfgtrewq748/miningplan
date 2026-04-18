from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Pt


def set_run_font(run, font_name: str, size_pt: float) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = r_pr._add_rFonts()
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: format_equations_right_numbered_20260418.py <docx>")
        return 2

    path = Path(sys.argv[1])
    doc = Document(str(path))
    section = doc.sections[0]
    usable_width = section.page_width - section.left_margin - section.right_margin
    center_tab = int(usable_width / 2)
    right_tab = int(usable_width)

    formatted = []
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        match = re.fullmatch(r"(.+?)\s*\uff08(\d+)\uff09", text)
        if not match:
            continue
        number = int(match.group(2))
        if not 1 <= number <= 20 or len(text) >= 220:
            continue

        expr = match.group(1).strip()
        if expr.endswith("\u3002"):
            expr = expr[:-1]
        num_text = f"\uff08{number}\uff09"

        paragraph.clear()
        fmt = paragraph.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.first_line_indent = Pt(0)
        fmt.left_indent = Pt(0)
        fmt.right_indent = Pt(0)
        fmt.space_before = Pt(3)
        fmt.space_after = Pt(3)
        fmt.tab_stops.clear_all()
        fmt.tab_stops.add_tab_stop(center_tab, WD_TAB_ALIGNMENT.CENTER)
        fmt.tab_stops.add_tab_stop(right_tab, WD_TAB_ALIGNMENT.RIGHT)

        paragraph.add_run("\t")
        eq_run = paragraph.add_run(expr)
        set_run_font(eq_run, "Cambria Math", 10.5)
        paragraph.add_run("\t")
        num_run = paragraph.add_run(num_text)
        set_run_font(num_run, "Times New Roman", 10.5)
        formatted.append((idx, number, expr))

    if len(formatted) != 15:
        raise RuntimeError(f"Expected 15 equations, formatted {len(formatted)}: {formatted}")

    doc.save(str(path))
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
    if bad is not None:
        raise RuntimeError(f"Bad DOCX zip member: {bad}")

    print(f"DOCX OK: {path}")
    print(f"Formatted equations: {len(formatted)}")
    for idx, number, expr in formatted:
        print(f"{idx}: ({number}) {expr}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
