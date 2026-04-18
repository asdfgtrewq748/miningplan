from pathlib import Path
import sys

from docx import Document


DOCX = Path(sys.argv[1])

PARA_REPLACEMENTS = {
    3: (
        "摘要：针对采区规划中离散地质信息难以连续参与布置、多场景覆岩扰动风险约束口径不统一以及规划结果难以向接续和经济评价传递等问题，提出一种基于覆岩扰动指数（overburden disturbance index，ODI）约束的采区候选方案生成与多目标比选方法。该方法以采区边界、钻孔样点和设计参数为输入，构建有效布置域和连续参数场；当前样例以煤层厚度场为主要展示对象，将地表沉陷、含水层扰动、上行开采等风险分量归一化并加权为统一ODI场，以ODI均值、90%分位数和超阈值暴露比例表征方案风险。在此基础上，建立包含工作面数量、推进方向、煤柱宽度、工作面宽度和方案选择变量的候选方案集合，采用工程效率、资源回收和扰动控制目标进行非支配排序与综合筛选。样例复核得到80×56栅格ODI场，均值为0.4669，P90为0.7474，ODI>0.70比例为15.89%。在同一ODI场下，方案A、B、C的ODI均值分别为0.4463、0.4552和0.4416，P90分别为0.6407、0.6462和0.6353，ODI>0.70比例分别为0.44%、0.56%和1.22%。本文对低扰动候选的判定依据为ODI均值、P90和风险综合得分的联合判据，而非单一阈值下的超限比例最小。敏感性分析表明，在当前样例和候选集合范围内，C方案风险综合得分低于A、B方案。结果说明，ODI前置约束可将风险图层转化为可复核的方案级统计指标，为采区规划多目标比选和后续对象传递提供支撑；真实矿井优选仍需结合现场约束和多案例验证。"
    ),
    4: "关键词：覆岩扰动指数；采区规划；多目标比选；候选方案集合；风险前置约束；连续参数场",
    9: (
        "Abstract: To address the limited use of discrete geological data in continuous mining-district layout, the inconsistent expression of multi-scenario overburden disturbance risks, and the weak linkage between planning outputs and downstream evaluation, an ODI-constrained candidate generation and multi-objective selection method is proposed. Boundary, borehole and design inputs are used to construct the effective layout domain and continuous parameter fields; in the current sample, coal-seam thickness is used as the main example showing how a continuous parameter field enters the planning chain. Surface subsidence, aquifer disturbance and upward-mining risk components are normalized and weighted into a unified overburden disturbance index (ODI) field, and candidate-level risk is described by mean ODI, the 90th percentile and exceedance ratio. A constrained candidate set considering face number, advance direction, pillar width, face width and scheme-selection variables is ranked using engineering-efficiency, resource-recovery and disturbance-control objectives. The verified sample ODI field contains 80×56 grids, with a mean of 0.4669, P90 of 0.7474 and ODI>0.70 ratio of 15.89%. Under the same ODI field, schemes A, B and C have mean ODI values of 0.4463, 0.4552 and 0.4416, P90 values of 0.6407, 0.6462 and 0.6353, and ODI>0.70 ratios of 0.44%, 0.56% and 1.22%, respectively. The low-disturbance candidate is identified by the joint criterion of mean ODI, P90 and composite risk score rather than by the minimum exceedance ratio under a single threshold. Sensitivity analysis indicates that scheme C has a lower composite risk score than schemes A and B within the present sample and candidate set. The results show that ODI pre-constraint can convert risk layers into reproducible scheme-level statistics for multi-objective comparison, while engineering selection under real mine conditions still requires site-specific constraints and multi-case validation."
    ),
    50: (
        "对候选方案π，本文把工程效率、资源回收和扰动控制分别写成可复核的分项评分。工程效率评分采用覆盖率基础分扣减组织复杂度惩罚的形式；资源回收和扰动控制中的正向指标按式（3）归一化，成本型或风险型指标按式（4）转化为正向评分。"
    ),
    94: (
        "从研究区样例结果看，多场景ODI结果能够在同一研究区边界、钻孔参照和色标口径下表达。本轮复核确认，综合ODI场采用80×56网格，共4480个栅格；归一化后ODI均值为0.4669，中位数为0.4542，P90为0.7474；其中ODI>0.70的栅格占15.89%，ODI>0.80的栅格占3.55%。上述结果表明，风险信息已由概念性指标转化为可计算、可统计并可进入规划比选的空间场。"
    ),
    108: (
        "不同目标偏好会导致覆盖率、资源回收和风险暴露出现权衡。与A方案相比，C方案覆盖率降低8.67个百分点，但ODI均值、P90和风险综合得分分别降低1.07%、0.86%和0.70%；与B方案相比，C方案覆盖率降低17.85个百分点，但上述3项风险指标分别降低3.00%、1.69%和2.12%。B方案较A方案覆盖率提高9.18个百分点，但ODI均值和P90分别提高1.99%和0.85%。结果表明，资源覆盖增益可能伴随扰动风险统计值上升，ODI前置约束可显式呈现风险-收益取舍。旧disturbance保存结果中的候选C_old在统一ODI场下均值为0.4560、P90为0.6472，不作为统一口径下的ODI筛选方案。"
    ),
}

TABLE_UPDATES = [
    (15, 15, 0, "候选集合"),
    (15, 16, 0, "候选集合"),
    (17, 3, 4, "联合判据下低扰动较优"),
]


def set_para_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell_text(cell, text):
    if cell.paragraphs:
        set_para_text(cell.paragraphs[0], text)
        for paragraph in cell.paragraphs[1:]:
            set_para_text(paragraph, "")
    else:
        cell.text = text


doc = Document(DOCX)
for idx, text in PARA_REPLACEMENTS.items():
    set_para_text(doc.paragraphs[idx], text)
for table_idx, row_idx, col_idx, text in TABLE_UPDATES:
    set_cell_text(doc.tables[table_idx].rows[row_idx].cells[col_idx], text)
doc.save(DOCX)
print(f"changed={len(PARA_REPLACEMENTS) + len(TABLE_UPDATES)}")
