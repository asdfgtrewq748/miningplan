from __future__ import annotations

import copy
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from latex2mathml.converter import convert
from lxml import etree


DOCX = Path(sys.argv[1])
XSL = Path(r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL")

FORMULAS = [
    r"\mathrm{ODI}(x)=w_sD_s(x)+w_aD_a(x)+w_uD_u(x)",
    r"w_s+w_a+w_u=1,\quad w_s\ge0,\quad w_a\ge0,\quad w_u\ge0",
    r"X'_i(x)=\frac{X_i(x)-X_{i,\min}}{X_{i,\max}-X_{i,\min}}",
    r"X'_i(x)=\frac{X_{i,\max}-X_i(x)}{X_{i,\max}-X_{i,\min}}",
    r"E_{\pi}(T_{ODI})=\frac{n_{\pi}\left(\mathrm{ODI}(x)>T_{ODI}\right)}{N_{\pi}}\times100\%",
    r"\Omega_e=\Omega_0\setminus\left(B_b\cup B_s\cup D_p\right)",
    r"z(x)=\frac{\sum_{i=1}^{n}z_id_i(x)^{-p}}{\sum_{i=1}^{n}d_i(x)^{-p}}",
    r"\pi=\{N,W_f,\theta,B_b,B_s,A_{\pi},L_{\pi},R_{\pi},y_{\pi}\},\quad y_{\pi}\in\{0,1\}",
    r"A_{\pi}\subset\Omega_e;\ B_b\in[B_{b,\min},B_{b,\max}];\ B_s\in[B_{s,\min},B_{s,\max}];\ W_f\in[W_{\min},W_{\max}];\ A_i\cap A_j=\varnothing;\ L_f\in C_L;\ E_{\pi}\in I_{ODI}",
    r"S_e(\pi)=100C_{\mathrm{cov}}(\pi)-[P_N(\pi)+10\mathrm{CV}_L(\pi)+5P_{\mathrm{short}}(\pi)]",
    r"S_r(\pi)=100[0.45+0.55(0.55R_{\mathrm{ton}}(\pi)+0.30R_{\mathrm{area}}(\pi)+0.15S_{\mathrm{eng}}(\pi))]",
    r"H_m(\pi)=c_1\mathrm{ODI}_{\mathrm{mean}}(\pi)+c_2Q_{0.90}(\pi)+c_3E_{\pi}(T_{ODI}),\quad S_m(\pi)=100[1-H_m(\pi)]",
    r"F(\pi)=\lambda_eS_e(\pi)+\lambda_rS_r(\pi)+\lambda_mS_m(\pi),\quad \lambda_e+\lambda_r+\lambda_m=1",
    r"G(s)=\alpha P_s+\beta R_s+\gamma C_s,\quad \alpha+\beta+\gamma=1",
    r"\mathrm{NCF}_t=\mathrm{Rev}_t-\mathrm{Cost}_t-\mathrm{RiskCost}_t",
]

P32_TEXT = (
    "式中，π为候选规划方案，A_π为方案π对应的布置区域，N_π为方案区域内参与统计的采样或栅格总数，"
    "n_π(ODI(x)>T_ODI)为满足ODI超过统计阈值条件的采样或栅格数量，T_ODI为扰动控制阈值。"
    "本文以ODI均值、Q_0.90和超阈值暴露比例E_π共同描述方案级风险：均值反映总体扰动水平，"
    "Q_0.90反映高值尾部风险，E_π反映超过统计阈值的空间暴露比例。阈值0.70和0.80分别作为"
    "样例内部预警统计线和高扰动统计线，并非行业安全红线。为降低阈值设定对结论的影响，本文设置"
    "0.65、0.70、0.75和0.80四组阈值开展敏感性分析；真实矿井应用中仍需结合历史观测、"
    "含水层保护等级、地表建构筑物敏感性和安全制度重新标定。"
)


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("}pPr"):
            continue
        p.remove(child)


def set_para_text(paragraph, text):
    clear_paragraph(paragraph)
    paragraph.add_run(text)


def latex_to_omml(latex: str, transform):
    mathml = etree.fromstring(convert(latex).encode("utf-8"))
    omml_tree = transform(mathml)
    return copy.deepcopy(omml_tree.getroot())


def main() -> None:
    if not XSL.exists():
        raise FileNotFoundError(XSL)

    transform = etree.XSLT(etree.parse(str(XSL)))
    doc = Document(str(DOCX))

    if len(doc.tables) < 15:
        raise RuntimeError("Expected at least 15 formula tables")

    for idx, latex in enumerate(FORMULAS):
        table = doc.tables[idx]
        cell = table.rows[0].cells[1]
        paragraph = cell.paragraphs[0]
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph._p.append(latex_to_omml(latex, transform))
        # Keep any extra paragraphs empty so Word does not retain stale formula text.
        for extra in cell.paragraphs[1:]:
            clear_paragraph(extra)

    set_para_text(doc.paragraphs[32], P32_TEXT)
    doc.save(str(DOCX))
    print(f"omml_formulas={len(FORMULAS)}")


if __name__ == "__main__":
    main()
