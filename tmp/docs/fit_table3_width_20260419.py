from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import sys

def set_cell_width(cell, width_twips):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.first_child_found_in('w:tcW')
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')

def set_table_width(table, width_twips):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(width_twips))
    tblW.set(qn('w:type'), 'dxa')

p=sys.argv[1]
doc=Document(p)
for table in doc.tables:
    if len(table.rows) >= 5 and table.cell(0,0).text.strip() == '方案' and '主导排序指标' in table.rows[0].cells[-1].text:
        table.autofit = False
        widths = [420, 1500, 1850, 1850, 1700, 780]  # total 8100 twips = 5.625 in
        set_table_width(table, sum(widths))
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                if ci < len(widths):
                    set_cell_width(cell, widths[ci])
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if row == table.rows[0] else WD_ALIGN_PARAGRAPH.LEFT
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    for run in para.runs:
                        run.font.size = Pt(8)
        break

doc.save(p)
print('done')
