from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, name: str = "宋体", size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    set_run_font(run, "黑体", 14 if level == 1 else 12, True)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, "宋体", 10.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "D9EAF7")
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, "黑体", 10.5, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    set_run_font(run, "宋体", 9.5)


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: create_revision_response_doc_20260418.py <output-docx>")
        return 2

    out_path = Path(sys.argv[1])
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("《覆岩扰动约束下采区协同规划方法》大修说明与逐条回应")
    set_run_font(run, "黑体", 16, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("供《煤炭科学技术》投稿/返修说明使用的工作底稿")
    set_run_font(run, "宋体", 10.5)

    add_heading(doc, "一、总体修改说明", 1)
    add_para(
        doc,
        "针对模拟审稿意见中提出的“论文更像系统说明、创新点不够硬、ODI定义不完整、多目标模型不严格、案例验证不足、经济评价支撑偏弱”等问题，本轮修改将稿件定位由程序导出说明调整为“ODI前置约束下的采区候选方案生成与多目标比选方法”。修改重点不是语言润色，而是补充方法定义、变量与约束、候选池生成逻辑、风险统计口径、内部对照和敏感性分析。",
    )
    add_para(
        doc,
        "修改后稿件保留样例验证的边界，不将单一样例推断为真实矿井条件下的普适最优结论；对经济评价部分也从完整经济结论收缩为下游接口与月度净现金流口径说明，避免给出缺乏独立参数支撑的现金流或NPV结论。",
    )

    summary_rows = [
        ["论文定位", "将“系统流程说明”改为“风险前置约束+候选方案池+多目标排序”的方法论文定位。", "引言末尾、1章开头、4章讨论、5章结论"],
        ["ODI定义", "补充风险分量、归一化、权重来源、阈值含义、方案级统计和敏感性分析。", "1.1、2.3、3.2、4.1"],
        ["多目标模型", "补充决策变量、几何约束、煤柱约束、不重叠约束、风险约束、候选池生成与非支配排序。", "1.3"],
        ["案例验证", "补充A/B/C三类候选方案对比、统一ODI统计、风险-收益权衡解释。", "3.1、3.2"],
        ["经济评价", "删去或弱化未经独立参数支撑的经济结果，保留Rev、Cost、RiskCost与NCF传递口径。", "1.4、3.3"],
        ["参考文献", "联网核验并扩充至50条，正文引用覆盖1-50号。", "参考文献与全文引用"],
        ["公式格式", "15个主公式改为右编号版式，编号（1）至（15）连续。", "全文公式"],
    ]
    add_table(doc, ["修改模块", "主要处理", "对应位置"], summary_rows)

    add_heading(doc, "二、逐条回应", 1)
    rows = [
        [
            "1. 论文定位不清，更像系统说明而非学术论文。",
            "接受。原稿确有较强系统流程说明色彩，方法贡献凝练不足。",
            "已在引言末尾明确3点贡献：ODI统一风险表征、离散钻孔—连续参数场—工作面对象映射、候选方案池与后续评价传递。正文中同步弱化“程序导出样例”表述，强调方法链路和候选方案比选。",
            "0引言；1章；4讨论；5结论",
        ],
        [
            "2. ODI科学性与可复核性不足，权重和阈值依据不清。",
            "接受。ODI是全文核心，必须从概念性指标改为可复核指标。",
            "已补充ODI风险分量D_s、D_a、D_u，说明正向/逆向指标归一化方法；明确权重为“场景目标驱动的专家规则赋权+敏感性校核”，不作为现场统计常数；阈值0.70/0.80定位为样例内部统计线而非安全红线；补充0.65、0.70、0.75、0.80阈值敏感性以及权重扰动分析。",
            "1.1；2.3；3.2；4.1",
        ],
        [
            "3. 多目标协同规划模型没有真正建立，变量、约束和求解机制不清。",
            "接受。原稿模型层表达不足，容易被认为只是评价排序。",
            "已将规划问题定义为“候选方案池生成—硬约束过滤—多目标排序—工程复核”的组合优化过程；补充候选方案π的变量组成，包括N、W_f、θ、B_b、B_s、A_π、L_π、R_π、y_π；补充边界、煤柱、宽度、不重叠、最小推进长度和ODI超阈值暴露约束；说明推进方向在有限工程允许方向集合Θ内枚举，不再表述为黑箱自动确定；补充非支配排序和拥挤距离选择。",
            "1.3",
        ],
        [
            "4. 工程效率评分、资源回收评分等不可复核。",
            "接受。评分必须给出口径，否则数值不可重复。",
            "已补充工程效率评分S_e、资源回收评分S_r、扰动控制评分S_m和综合得分F(π)的表达式，说明覆盖率、推进长度均衡性、巷道可达性、巷道工程量惩罚、煤厚覆盖和可采资源回收等分项含义；明确87.82、89.76等分数仅用于同一候选池内部比较。",
            "1.3；3.1",
        ],
        [
            "5. 案例验证强度不足，缺少对比对象。",
            "部分接受。当前数据条件尚不足以补成完整实矿验证，但可增强内部对照。",
            "已补充A、B、C三类候选方案对比：A代表工程效率偏好，B代表资源回收偏好，C代表ODI风险约束偏好。三者在同一研究区边界、同一钻孔参数场、同一ODI场和同一几何约束下比较，报告覆盖率、ODI均值、P90、ODI>0.70比例和风险综合得分，并明确该对照属于同一候选池内部目标偏好对照，不替代真实人工方案现场对照。",
            "3.1；3.2；4.3",
        ],
        [
            "6. 研究区工程背景交代不充分。",
            "接受。原稿样例背景偏弱。",
            "已补充敏东研究区样例的工程语境，包括内蒙古呼伦贝尔敏东一矿相关背景、代表性埋深约300 m、顶底板岩性类型、边界面积、钻孔样点数、边界煤柱、区段煤柱、工作面宽度和走向长壁后退式开采等输入参数。同时明确断层、既有巷道、通风运输能力、设备能力和生产制度等现场闭环约束尚未完整纳入。",
            "2.1；表1；4.3",
        ],
        [
            "7. 经济评价部分容易被质疑为拉长战线、证据不足。",
            "接受。经济评价不宜半展开。",
            "已将经济评价部分收缩为下游接口说明，仅保留月度净现金流NCF_t=Rev_t-Cost_t-RiskCost_t等传递口径；不再给出未经独立参数体系支撑的NPV、投资回收期、总收入和总成本等结论性数值。",
            "1.4；3.3；4.3",
        ],
        [
            "8. 讨论部分重复结果，缺少边界、不确定性和适用性讨论。",
            "接受。讨论应解释方法意义和风险，而不是复述结果。",
            "已重写讨论逻辑，分别讨论ODI前置约束的作用与边界、参数场驱动规划相对于传统经验布置的意义、连续参数场质量对规划结果的影响、ODI信息压缩风险、权重/阈值敏感性以及后续实矿验证方向。",
            "4.1；4.2；4.3",
        ],
        [
            "9. 参考文献需要聚焦煤炭采区规划、覆岩扰动、保水采煤和智能化矿山。",
            "接受。",
            "已联网核验原有参考文献并补充中文核心方向文献，参考文献扩充至50条，均保留DOI；正文引用覆盖1-50号，未发现正文未引或参考文献缺号问题。",
            "参考文献；全文引用",
        ],
        [
            "10. 公式格式和编号需符合期刊论文表达。",
            "接受。",
            "已将15个主公式整理为右编号版式，编号（1）至（15）连续。另测试了Word/MathType OMML自动转换方案，但复杂公式自动BuildUp后版面质量下降，因此正式稿保留更稳定的右编号表格结构，后续可逐式替换为MathType对象。",
            "全文公式；公式处理记录",
        ],
    ]
    add_table(doc, ["审稿意见", "回应态度", "已完成修改", "稿件位置"], rows)

    add_heading(doc, "三、仍需说明的边界", 1)
    add_para(
        doc,
        "本轮修改已显著增强方法定义、可复核性和内部对照，但稿件仍不应宣称已经完成真实矿井条件下的工程定案验证。当前A/B/C对比属于同一候选池内部目标偏好对照，尚不等同于传统人工经验方案、无ODI方案和真实生产方案之间的完整工程对照。",
    )
    add_para(
        doc,
        "若后续需要进一步提高录用把握，建议继续补充三类材料：一是现场专家给出的经验布置方案作为外部基准；二是IDW与克里金或不同网格尺度下的参数场敏感性；三是真实煤价、成本、税费、产能约束和风险停产参数支撑下的经济评价。",
    )

    add_heading(doc, "四、当前质量检查结果", 1)
    qa_rows = [
        ["参考文献一致性", "reference_count=50；doi_count=50；used_body_count=50；missing_in_body=[]；missing_in_refs=[]"],
        ["全文一致性审计", "issues=0"],
        ["公式编号检查", "formula_tables=15；formula_numbers=1-15；formula_sequence_ok=True"],
        ["PDF预览", "已导出公式右编号预览PDF，用于检查编号位置和公式版式"],
    ]
    add_table(doc, ["检查项", "结果"], qa_rows)

    doc.save(str(out_path))
    print(out_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
