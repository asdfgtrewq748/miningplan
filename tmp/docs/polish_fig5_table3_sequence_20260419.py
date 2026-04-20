from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def set_text(para, text: str, style=None, size: float = 10.5, align=None) -> None:
    para.text = text
    if style is not None:
        para.style = style
    if align is not None:
        para.alignment = align
    for run in para.runs:
        run.font.size = Pt(size)


def main() -> None:
    docx_path = Path(os.environ["DOCX_PATH"])
    out_path = Path(os.environ.get("OUTPUT_DOCX_PATH", str(docx_path)))
    doc = Document(docx_path)
    body_style = doc.paragraphs[104].style
    fig_style = None
    for para in doc.paragraphs:
        if para.text.startswith("图1 ODI风险约束逻辑示意图"):
            fig_style = para.style
            break
    if fig_style is None:
        raise RuntimeError("Figure caption style not found")

    for para in doc.paragraphs:
        if para.text.startswith("表3列出了当前已复核的候选方案"):
            set_text(
                para,
                "图5和表3分别给出了当前已复核候选方案的空间布局和统计结果。工程效率优先方案A来自efficiency模式，覆盖率为89.34%，工程效率分项评分为87.82；资源回收优先方案B来自recovery模式，覆盖率为98.52%，资源回收分项评分为89.76；ODI筛选方案C来自统一候选集合复核，覆盖率为80.67%，工程效率分项评分为79.27，ODI均值为0.4416，P90为0.6353，ODI>0.70比例为1.22%。需要指出，本文对低扰动候选的判定并非依据单一阈值下的超限比例最小，而是依据ODI均值、P90与风险综合得分的联合判据。",
                body_style,
                10.5,
            )
        elif para.text.startswith("图5 A/B/C候选方案工作面布局"):
            set_text(
                para,
                "图5 A/B/C候选方案工作面布局与ODI高值区叠置对比图\nFig.5 Comparison of A/B/C candidate layouts overlaid with high-ODI zones",
                fig_style,
                9.0,
                WD_ALIGN_PARAGRAPH.CENTER,
            )

    doc.save(out_path)


if __name__ == "__main__":
    main()
