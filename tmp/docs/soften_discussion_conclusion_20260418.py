from __future__ import annotations

import sys
import zipfile
from pathlib import Path

from docx import Document


REPLACEMENTS = {
    "因此，本案例的验证强度定位为“方法有效性与内部对照验证”：它能够证明离散钻孔、连续参数场、ODI风险场和候选方案对象可以在同一口径下完成计算与比较，并能揭示不同目标偏好下的风险-收益权衡；但它尚不能替代真实矿井条件下的人工经验方案对照、现场生产约束闭环和开采后实测反馈校核。":
    "因此，本案例的验证强度定位为“方法链可运行性与内部对照验证”：它表明离散钻孔、连续参数场、ODI风险场和候选方案对象可以在同一口径下完成计算与比较，并能揭示不同目标偏好下的风险-收益权衡；但它尚不能替代真实矿井条件下的人工经验方案对照、现场生产约束闭环和开采后实测反馈校核。",

    "本文当前结果主要证明所提方法在样例条件下的链路贯通能力，即能够实现采区边界、钻孔样点、连续参数场、ODI风险场、候选方案池和规划对象之间的连续组织与传递。A、B、C候选方案已在同一ODI场下完成均值、P90和超阈值暴露比例补算，能够支撑初步方案对比，但仍不足以形成真实矿井条件下的工程优选结论。":
    "本文当前结果主要表明，所提方法在样例条件下具备链路贯通能力，即能够实现采区边界、钻孔样点、连续参数场、ODI风险场、候选方案池和规划对象之间的连续组织与传递。A、B、C候选方案已在同一ODI场下完成均值、P90和超阈值暴露比例补算，可支撑初步方案对比，但仍不足以形成真实矿井条件下的工程优选结论。",

    "后续研究应从样例级验证转向多矿井对照验证：一方面补充传统经验方案、无ODI方案和不同插值方案的基准对比，另一方面在真实接续排程和经济参数下检验ODI前置约束是否能够稳定降低风险暴露并保持可接受的资源回收水平。只有完成上述验证后，本文方法才能从“可复核的规划链路”进一步发展为“可用于工程定案的优选工具”。":
    "后续研究应从样例级验证转向多矿井对照验证：一方面补充传统经验方案、无ODI方案和不同插值方案的基准对比，另一方面在真实接续排程和经济参数下检验ODI前置约束对风险暴露和资源回收水平的影响。只有完成上述验证后，本文方法才能从“可复核的规划链路”进一步发展为“可辅助工程方案论证的决策工具”。",

    "3）研究区样例验证表明，工程效率优先、资源回收优先和ODI筛选方案分别体现出覆盖率、资源利用和风险暴露之间的权衡。方案A、B、C的ODI均值分别为0.4463、0.4552和0.4416，P90分别为0.6407、0.6462和0.6353，ODI>0.70比例分别为0.44%、0.56%和1.22%，说明风险场能够转化为可复核的方案级统计量。":
    "3）研究区样例结果表明，工程效率优先、资源回收优先和ODI筛选方案分别体现出覆盖率、资源利用和风险暴露之间的权衡。方案A、B、C的ODI均值分别为0.4463、0.4552和0.4416，P90分别为0.6407、0.6462和0.6353，ODI>0.70比例分别为0.44%、0.56%和1.22%，说明风险场能够转化为可复核的方案级统计量。",

    "4）当前结果主要支撑样例条件下的方法贯通、对象链有效性和初步方案比选。权重敏感性表明，在当前样例和候选池范围内，基准权重、±10%相对扰动和含水层专项权重条件下C方案均保持最低风险综合得分；阈值敏感性表明，当阈值由0.65提高至0.80时，各方案超限比例均下降，但方案间风险差异仍可通过均值、P90和超限比例共同识别。真实矿井条件下的工程优选结论仍需结合实矿案例检验。":
    "4）当前结果主要支撑样例条件下的方法贯通、对象链可传递性和初步方案比选。权重敏感性表明，在当前样例和候选池范围内，基准权重、±10%相对扰动和含水层专项权重条件下C方案风险综合得分均低于A、B方案；阈值敏感性表明，当阈值由0.65提高至0.80时，各方案超限比例均下降，但方案间风险差异仍可通过均值、P90和超限比例共同识别。真实矿井条件下的工程优选结论仍需结合实矿案例检验。",
}


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: soften_discussion_conclusion_20260418.py <docx>")
        return 2

    path = Path(sys.argv[1])
    doc = Document(str(path))

    changed = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in REPLACEMENTS:
            paragraph.text = REPLACEMENTS[text]
            changed.append(text[:30])

    missing = [old[:50] for old in REPLACEMENTS if old not in [p.text.strip() for p in doc.paragraphs] and REPLACEMENTS[old] not in [p.text.strip() for p in doc.paragraphs]]
    if missing:
        raise RuntimeError(f"Some target paragraphs were not found: {missing}")

    doc.save(str(path))
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
    if bad is not None:
        raise RuntimeError(f"Bad DOCX zip member: {bad}")

    print(f"DOCX OK: {path}")
    print(f"Softened paragraphs: {len(changed)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
