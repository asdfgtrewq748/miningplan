from __future__ import annotations

from pathlib import Path
import re
import shutil
import sys
from docx import Document

DOCX = Path(sys.argv[1])
BACKUP = Path(sys.argv[2])
shutil.copy2(DOCX, BACKUP)

doc = Document(DOCX)
body_style = None
for p in doc.paragraphs:
    if p.style.name == '10正文':
        body_style = p.style
        break
if body_style is None:
    raise RuntimeError('Body style 10正文 not found')

# Fix body paragraphs accidentally styled as headings. Keep numbered headings and reference heading unchanged.
for p in doc.paragraphs:
    t = p.text.strip()
    if not t:
        continue
    if p.style.name in {'11一级标题', '12二级标题'} and not re.match(r'^\d+(\.\d+)?\s', t) and t != '参考文献':
        p.style = body_style

# Paragraph-level language polish.
replacements = {
    '单一程序输出结果': '单一计算结果',
    '当前导出记录中，工程效率、资源回收和扰动控制模式的有效候选数分别为2417、1149和374，其中扰动控制模式合格候选数为88。': '当前候选集合统计中，工程效率、资源回收和扰动控制偏好下的有效候选数分别为2417、1149和374，其中扰动控制偏好下满足低扰动筛选条件的候选数为88。',
    '工程效率优先方案A来自efficiency模式': '工程效率优先方案A来自工程效率偏好模式',
    '资源回收优先方案B来自recovery模式': '资源回收优先方案B来自资源回收偏好模式',
    '早期扰动控制模式候选C_old': '早期扰动控制偏好候选C0',
    '旧disturbance保存结果中的候选C_old': '早期扰动控制偏好候选C0',
    '本轮复核确认': '样例计算结果显示',
    '当前结果已补齐A、B、C方案的统一ODI统计和全域4480栅格统计': '当前结果已给出A、B、C方案的统一ODI统计和全域4480栅格统计',
    '从可复现角度看，本文样例至少需保留4类中间数据': '从可复现角度看，本文样例至少需保留4类过程对象',
    '工程效率优先方案A来自efficiency模式，覆盖率为89.34%，工程效率分项评分为87.82；资源回收优先方案B来自recovery模式，覆盖率为98.52%，资源回收分项评分为89.76；': '工程效率优先方案A来自工程效率偏好模式，覆盖率为89.34%，工程效率分项评分为87.82；资源回收优先方案B来自资源回收偏好模式，覆盖率为98.52%，资源回收分项评分为89.76；',
}
for p in doc.paragraphs:
    if not p.text:
        continue
    new = p.text
    for old, val in replacements.items():
        new = new.replace(old, val)
    if new != p.text:
        p.text = new

# Table-level wording cleanup.
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if not p.text:
                    continue
                new = p.text
                cell_repls = {
                    'efficiency；工程效率优先': '工程效率偏好；工程效率优先',
                    'recovery；资源回收优先': '资源回收偏好；资源回收优先',
                    '工程效率评分': '工程效率分项评分',
                    '资源回收评分': '资源回收分项评分',
                    'C_old': 'C0',
                    '早期扰动控制模式候选': '早期扰动控制偏好候选',
                    '早期口径候选，不作为本文C方案': '早期口径候选，不纳入本文C方案',
                    '综合ODI场中间数据': '综合ODI场统计结果',
                    '权重敏感性统计表': '权重敏感性统计结果',
                    '规划对象与接续输入表': '规划对象与接续输入关系',
                }
                for old, val in cell_repls.items():
                    new = new.replace(old, val)
                if new != p.text:
                    p.text = new

# Normalize font size for paragraphs that were restyled from headings by assigning body style already; leave run formatting to template.
doc.save(DOCX)
print(DOCX)
