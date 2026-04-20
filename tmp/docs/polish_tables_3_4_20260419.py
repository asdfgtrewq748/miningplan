from docx import Document
import sys
from pathlib import Path

p=Path(sys.argv[1])
doc=Document(p)

# Find Table 3 by first cell header.
for table in doc.tables:
    txt='\n'.join(cell.text for row in table.rows for cell in row.cells)
    if '候选方案对比' in txt or ('方案' in table.cell(0,0).text and 'ODI统计' in txt and 'C0' in txt):
        if len(table.rows) >= 5 and table.cell(0,0).text.strip().startswith('方案'):
            rows = {
                'A': ['A', '工程效率偏好', '5个工作面；覆盖率89.34%；S_e=87.82', '均值0.4463；P90=0.6407；E_0.70=0.44%', '按工程效率分项排序'],
                'B': ['B', '资源回收偏好', '9个工作面；覆盖率98.52%；S_r=89.76', '均值0.4552；P90=0.6462；E_0.70=0.56%', '按资源回收分项排序'],
                'C': ['C', '联合判据筛选', '4个工作面；覆盖率80.67%；S_e=79.27', '均值0.4416；P90=0.6353；E_0.70=1.22%', '联合判据下低扰动较优'],
                'C0': ['C0', '早期扰动控制偏好', '13个工作面；覆盖率75.06%；S_e=71.17', '均值0.4560；P90=0.6472；E_0.70=0.80%', '早期口径候选，不纳入本文C方案'],
            }
            for r in range(1, len(table.rows)):
                key = table.cell(r,0).text.strip()
                if key in rows:
                    for c, val in enumerate(rows[key]):
                        table.cell(r,c).text = val
            break

# Keep Table 4 title with its table by starting it on a new page.
for para in doc.paragraphs:
    if para.text.strip().startswith('表4 ODI统计口径与敏感性复核状态'):
        para.paragraph_format.page_break_before = True
        break

doc.save(p)
print('done')
