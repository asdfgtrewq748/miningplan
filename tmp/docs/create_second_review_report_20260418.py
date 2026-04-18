from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def font(run, name="宋体", size=10.5, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    r = p.add_run(text)
    font(r, "黑体", 14 if level == 1 else 12, True)


def para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    font(r)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    font(r)


def shade(cell, fill="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        shade(cell)
        for p in cell.paragraphs:
            for r in p.runs:
                font(r, "黑体", 10.5, True)
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                p.paragraph_format.line_spacing = 1.15
                for r in p.runs:
                    font(r, "宋体", 9.5)


def main():
    if len(sys.argv) != 2:
        print("Usage: create_second_review_report_20260418.py <output-docx>")
        return 2
    out = Path(sys.argv[1])
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Pt(72)
    sec.bottom_margin = Pt(72)
    sec.left_margin = Pt(72)
    sec.right_margin = Pt(72)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("模拟二审意见与剩余风险清单")
    font(r, "黑体", 16, True)

    para(
        doc,
        "审查对象为当前大修工作稿。审查口径按《煤炭科学技术》工程技术论文的二审标准，重点检查大修意见是否被实质性回应、方法定义是否可复核、案例验证是否过度外推、结论是否与证据强度一致。",
    )

    heading(doc, "一、总体判断", 1)
    para(
        doc,
        "与初稿相比，当前稿件已经从“系统说明+样例输出”明显转向“ODI前置约束下的候选方案生成与多目标比选方法”。ODI定义、变量约束、候选池生成、评分口径、A/B/C内部对照和敏感性分析均有实质补充。若按模拟二审口径，稿件已不再处于退修边缘，但仍属于“修改后可继续外审/小修偏大”的状态。",
    )

    heading(doc, "二、已明显改善的部分", 1)
    for item in [
        "创新点已从流程串联改写为ODI统一风险表征、参数场到规划对象映射、候选池与评价传递3个可审查贡献。",
        "ODI已补充组成项、归一化、权重来源、阈值定位和敏感性分析，避免把自定义指标写成未经标定的安全红线。",
        "多目标模型已补充决策变量、约束集合、候选池生成、非支配排序和综合排序逻辑。",
        "案例验证已从单一导出样例扩展为A/B/C内部对照，并明确控制变量和验证边界。",
        "经济评价已收缩为下游接口和净现金流口径，避免缺少煤价、税费、成本参数时给出过强经济结论。",
        "参考文献已扩充至50条且正文引用全覆盖，中文煤炭领域文献权重明显提高。",
    ]:
        bullet(doc, item)

    heading(doc, "三、模拟二审仍可能提出的问题", 1)
    rows = [
        [
            "P1",
            "公式仍不是严格MathType/OMML对象",
            "当前右编号版式稳定，但公式主体仍主要是线性文本。若编辑部严格要求可编辑公式对象，仍需逐式替换为MathType或Word公式。",
            "格式风险，不必改变研究结论；投稿前可由作者逐式处理。",
        ],
        [
            "P1",
            "外部基准方案仍不足",
            "A/B/C属于同一候选池内部对照，不是传统人工经验方案、无ODI方案和本文方案的完整外部对照。",
            "这是剩余最大科学验证风险；若有条件，应补1个专家经验方案或无ODI方案。",
        ],
        [
            "P2",
            "数据可复现材料未独立成附录",
            "正文已说明需保留4类中间数据，但尚未提供独立数据清单、字段说明或样例表。",
            "建议后续补“数据与参数复核表”或附录。",
        ],
        [
            "P2",
            "插值方法仍偏样例化",
            "IDW选择已有解释，但未给留一交叉验证、克里金对比或网格尺度敏感性。",
            "目前已在讨论中列为后续工作；若篇幅允许，可补一段误差复核计划。",
        ],
        [
            "P2",
            "图件分辨率和图中文字可能影响期刊排版",
            "图1至图4承载较多信息，若投稿系统压缩图片，图中文字可能偏小。",
            "属于排版与制图风险；后续格式阶段处理。",
        ],
        [
            "P3",
            "结论第3条中C方案的ODI>0.70比例高于A/B，需避免读者误解",
            "正文已解释C不是所有单项指标最优，而是均值、P90和风险综合得分较优；摘要和结论仍需保持这种联合判据表述。",
            "当前基本可接受，后续可继续微调结论措辞。",
        ],
    ]
    table(doc, ["优先级", "问题", "二审可能意见", "处理建议"], rows)

    heading(doc, "四、建议的下一轮处理顺序", 1)
    for item in [
        "优先处理公式对象：保留当前右编号表格结构，逐式替换中列公式为MathType/Word公式对象。",
        "补一个数据与参数复核附表：列出边界、钻孔、网格、ODI分量、候选方案对象和统计字段。",
        "若有现场或程序数据，补无ODI方案或专家经验方案作为外部基准。",
        "最后再做全文语言压缩与图件清晰度处理。",
    ]:
        bullet(doc, item)

    heading(doc, "五、模拟二审结论", 1)
    para(
        doc,
        "模拟二审结论为：大修意见已获得实质回应，稿件的研究边界比初稿清楚，核心方法链条基本可审。若不补外部基准和严格公式对象，仍可能被要求继续修改；若补充数据附表和公式对象，进入小修或继续外审的可能性会明显提高。",
    )

    doc.save(str(out))
    print(out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
