from docx import Document
import sys
from pathlib import Path
from docx.shared import Pt

p=Path(sys.argv[1])
doc=Document(p)
repls = {
    '本文对低扰动候选的判定依据为': '本文对联合判据下低扰动候选的判定依据为',
    '本文对联合判据下联合判据下低扰动候选的判定依据为': '本文对联合判据下低扰动候选的判定依据为',
    '而是在ODI均值、P90和风险综合得分联合判据下联合判据下低扰动候选。': '而是依据ODI均值、P90和风险综合得分确定的联合判据下低扰动候选。',
    '工程效率优先、资源回收优先和联合判据下低扰动候选分别体现出': '工程效率优先方案、资源回收优先方案和联合判据下低扰动候选分别体现出',
}
changed=0
for para in doc.paragraphs:
    txt=para.text
    new=txt
    for old,val in repls.items():
        new=new.replace(old,val)
    if new != txt:
        para.text=new
        for run in para.runs:
            if para.style.name == '10正文':
                run.font.size = Pt(10.5)
        changed += 1
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                txt=para.text
                new=txt
                for old,val in repls.items():
                    new=new.replace(old,val)
                if new != txt:
                    para.text=new
                    for run in para.runs:
                        run.font.size = Pt(9)
                    changed += 1

doc.save(p)
print('changed', changed)
