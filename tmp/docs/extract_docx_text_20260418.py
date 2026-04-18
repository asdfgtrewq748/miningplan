from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: extract_docx_text_20260418.py <docx> <out-md>")
        return 2

    docx_path = Path(sys.argv[1])
    out_path = Path(sys.argv[2])
    doc = Document(str(docx_path))

    lines: list[str] = []
    for idx, paragraph in enumerate(doc.paragraphs):
        text = " ".join(paragraph.text.split())
        if text:
            lines.append(f"{idx}: {text}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(lines), encoding="utf-8")
    print(out_path)
    print(f"paragraphs={len(doc.paragraphs)}")
    print(f"nonempty={len(lines)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
