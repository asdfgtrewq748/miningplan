from __future__ import annotations

import json
import math
import statistics
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill, Side, Border
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "煤科投稿"
DOCS = ROOT / "docs" / "plans"


def read_csv(path: Path) -> pd.DataFrame:
    try:
        return pd.read_csv(path, encoding="utf-8-sig")
    except UnicodeDecodeError:
        return pd.read_csv(path, encoding="gbk")


def find_one(pattern: str) -> Path | None:
    matches = list(ROOT.glob(pattern))
    return matches[0] if matches else None


def find_rglob(name: str) -> Path | None:
    matches = list(ROOT.rglob(name))
    return matches[0] if matches else None


def flatten_field_values(obj) -> list[float]:
    candidates = []
    if isinstance(obj, dict):
        for key in ["field", "values", "data", "grid", "odiField", "odi_field"]:
            if key in obj:
                candidates.append(obj[key])
    else:
        candidates.append(obj)

    values: list[float] = []

    def walk(x):
        if isinstance(x, dict):
            if "value" in x and isinstance(x["value"], (int, float)):
                values.append(float(x["value"]))
            else:
                for v in x.values():
                    walk(v)
        elif isinstance(x, list):
            for item in x:
                walk(item)
        elif isinstance(x, (int, float)) and not isinstance(x, bool):
            if math.isfinite(float(x)):
                values.append(float(x))

    for candidate in candidates:
        walk(candidate)
    return values


def pctl(values: list[float], q: float) -> float:
    if not values:
        return float("nan")
    xs = sorted(values)
    pos = (len(xs) - 1) * q
    lo = math.floor(pos)
    hi = math.ceil(pos)
    if lo == hi:
        return xs[lo]
    return xs[lo] * (hi - pos) + xs[hi] * (pos - lo)


def add_df_sheet(wb: Workbook, title: str, df: pd.DataFrame, source: str | None = None) -> None:
    ws = wb.create_sheet(title[:31])
    row_offset = 1
    if source:
        ws.cell(row=1, column=1, value="数据来源")
        ws.cell(row=1, column=2, value=source)
        row_offset = 3

    for c_idx, col in enumerate(df.columns, 1):
        cell = ws.cell(row=row_offset, column=c_idx, value=str(col))
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E78")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for r_idx, row in enumerate(df.itertuples(index=False), row_offset + 1):
        for c_idx, value in enumerate(row, 1):
            if pd.isna(value):
                value = ""
            cell = ws.cell(row=r_idx, column=c_idx, value=value)
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    style_sheet(ws)


def style_sheet(ws) -> None:
    thin = Side(style="thin", color="D9E2F3")
    for row in ws.iter_rows():
        for cell in row:
            cell.border = Border(bottom=thin)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    for col in range(1, ws.max_column + 1):
        values = [str(ws.cell(r, col).value or "") for r in range(1, min(ws.max_row, 80) + 1)]
        width = min(max(max((len(v) for v in values), default=8) + 2, 10), 42)
        ws.column_dimensions[get_column_letter(col)].width = width
    ws.freeze_panes = "A2"


def add_kv_sheet(wb: Workbook, title: str, rows: list[tuple[str, object, str]]) -> None:
    df = pd.DataFrame(rows, columns=["复核项", "复核值", "说明/来源"])
    add_df_sheet(wb, title, df)


def font(run, name="宋体", size=10.5, bold=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def shade(cell, fill="D9EAF7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_doc_table(doc: Document, headers: list[str], rows: list[list[object]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        shade(cell)
        for p in cell.paragraphs:
            for r in p.runs:
                font(r, "黑体", 9.5, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = "" if value is None else str(value)
            for p in cells[i].paragraphs:
                p.paragraph_format.line_spacing = 1.1
                for r in p.runs:
                    font(r, "宋体", 9)


def main() -> int:
    xlsx_out = OUT_DIR / "数据与参数复核附表_20260418.xlsx"
    docx_out = OUT_DIR / "数据与参数复核附表说明_20260418.docx"

    boundary_path = ROOT / "data" / "采区边界_敏东.csv"
    borehole_path = ROOT / "data" / "钻孔坐标_敏东.csv"
    field_path = find_rglob("000_mindong_layout_odi_field.json")
    candidate_path = DOCS / "coal_sci_abc_odi_unified_stats_20260418.csv"
    threshold_path = DOCS / "coal_sci_threshold_sensitivity_candidates_20260418.csv"
    weight_path = DOCS / "coal_sci_weight_sensitivity_candidates_20260418.csv"
    weight_field_path = DOCS / "coal_sci_weight_sensitivity_field_20260418.csv"
    old_candidate_path = DOCS / "coal_sci_abc_candidate_summary_20260418.csv"

    wb = Workbook()
    wb.remove(wb.active)

    inventory_rows = [
        ["采区边界", str(boundary_path), boundary_path.exists(), "用于复核规划域边界输入"],
        ["钻孔坐标", str(borehole_path), borehole_path.exists(), "用于复核钻孔样点输入"],
        ["统一ODI场", str(field_path) if field_path else "", field_path is not None, "用于复核80x56网格及全域ODI统计"],
        ["A/B/C方案统计", str(candidate_path), candidate_path.exists(), "用于复核方案级ODI统计"],
        ["阈值敏感性", str(threshold_path), threshold_path.exists(), "用于复核0.65/0.70/0.75/0.80阈值影响"],
        ["权重敏感性-方案", str(weight_path), weight_path.exists(), "用于复核不同权重下A/B/C风险综合得分"],
        ["权重敏感性-全域", str(weight_field_path), weight_field_path.exists(), "用于复核全域ODI场权重扰动"],
        ["旧候选方案口径", str(old_candidate_path), old_candidate_path.exists(), "用于说明C_old不作为正文C方案"],
    ]
    add_df_sheet(wb, "00_数据源清单", pd.DataFrame(inventory_rows, columns=["类别", "路径", "是否存在", "用途"]))

    if boundary_path.exists():
        add_df_sheet(wb, "01_采区边界输入", read_csv(boundary_path), str(boundary_path))
    if borehole_path.exists():
        add_df_sheet(wb, "02_钻孔坐标输入", read_csv(borehole_path), str(borehole_path))

    field_rows: list[tuple[str, object, str]] = []
    if field_path and field_path.exists():
        obj = json.loads(field_path.read_text(encoding="utf-8"))
        values = flatten_field_values(obj)
        grid_w = obj.get("gridW") or obj.get("grid_w") or obj.get("width")
        grid_h = obj.get("gridH") or obj.get("grid_h") or obj.get("height")
        field_rows.extend(
            [
                ("网格宽度", grid_w, str(field_path)),
                ("网格高度", grid_h, str(field_path)),
                ("栅格数", len(values), "由field数值展平后计算"),
                ("ODI均值", round(statistics.fmean(values), 6), "由统一ODI场计算"),
                ("ODI中位数", round(statistics.median(values), 6), "由统一ODI场计算"),
                ("ODI P90", round(pctl(values, 0.9), 6), "由统一ODI场计算"),
                ("ODI>0.70比例/%", round(sum(v > 0.70 for v in values) / len(values) * 100, 4), "由统一ODI场计算"),
                ("ODI>0.80比例/%", round(sum(v > 0.80 for v in values) / len(values) * 100, 4), "由统一ODI场计算"),
            ]
        )
    else:
        field_rows.append(("统一ODI场", "未找到", "需检查支撑材料目录"))
    add_kv_sheet(wb, "03_ODI全域统计复核", field_rows)

    for title, path in [
        ("04_ABC方案统计", candidate_path),
        ("05_阈值敏感性", threshold_path),
        ("06_权重敏感性_方案", weight_path),
        ("07_权重敏感性_全域", weight_field_path),
        ("08_旧候选口径说明", old_candidate_path),
    ]:
        if path.exists():
            add_df_sheet(wb, title, read_csv(path), str(path))

    checklist_rows = [
        ["边界输入", "采区边界点文件存在，可追溯", "已纳入附表"],
        ["钻孔输入", "钻孔坐标文件存在，可追溯", "已纳入附表"],
        ["全域ODI场", "可复核80x56、4480栅格、均值、P90、超阈值比例", "已纳入附表"],
        ["方案级统计", "A/B/C在统一ODI场下复算均值、P90、ODI>0.70比例", "已纳入附表"],
        ["阈值敏感性", "0.65、0.70、0.75、0.80四组阈值", "已纳入附表"],
        ["权重敏感性", "基准、分量扰动及含水层专项权重", "已纳入附表"],
        ["外部基准方案", "传统人工经验方案或无ODI方案", "仍需后续补充"],
        ["插值误差验证", "IDW留一交叉验证或Kriging对比", "仍需后续补充"],
        ["严格公式对象", "MathType/OMML逐式替换", "仍需格式阶段处理"],
    ]
    add_df_sheet(wb, "09_复核状态清单", pd.DataFrame(checklist_rows, columns=["复核项", "状态/要求", "处理状态"]))

    wb.save(xlsx_out)

    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.left_margin = Pt(54)
    sec.right_margin = Pt(54)
    sec.top_margin = Pt(54)
    sec.bottom_margin = Pt(54)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("数据与参数复核附表说明")
    font(r, "黑体", 16, True)
    p = doc.add_paragraph()
    r = p.add_run("用于支撑《覆岩扰动约束下采区协同规划方法》大修稿的数据可复核性说明。")
    font(r, "宋体", 10.5)

    rows = [[a, b, c, d] for a, b, c, d in inventory_rows]
    add_doc_table(doc, ["类别", "路径", "是否存在", "用途"], rows)
    doc.add_paragraph()
    add_doc_table(doc, ["复核项", "复核值", "说明/来源"], [[a, b, c] for a, b, c in field_rows])
    doc.add_paragraph()
    add_doc_table(doc, ["复核项", "状态/要求", "处理状态"], checklist_rows)

    doc.save(docx_out)

    print(f"XLSX={xlsx_out}")
    print(f"DOCX={docx_out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
