from copy import deepcopy
from pathlib import Path
import os

from docx import Document
from docx.text.paragraph import Paragraph


REMOVE_STARTS = [
    "3.3 规划结果向采掘接续与工程经济评价的传递",
    "本文方法的一个重要特点，是不把采区规划的几何布局视为终点",
    "在采掘接续层面，规划阶段形成的工作面边界和推进关系可映射为接续任务对象。",
    "在调控层面，规划结果还可进一步转化为工作面级控制变量。",
    "在工程经济评价层面，规划与接续结果可以继续进入收入、成本、风险联动成本和现金流分析过程。",
    "综上，研究区样例表明，本文方法已完成“边界与钻孔输入",
]


TRANSMISSION = [
    "3.3 规划结果向采掘接续与工程经济评价的传递",
    (
        "本文方法的一个重要特点，是不把采区规划的几何布局视为终点，而是强调规划结果可以继续作为采掘接续和工程经济评价的上游输入。"
        "在本样例中，工作面边界、巷道路径、推进长度、布置面积和ODI统计量均已形成可传递的规划结果，后续环节可据此组织生产顺序、产量核算和经济评价。"
    ),
    (
        "在采掘接续层面，规划阶段形成的工作面边界和推进关系可映射为接续任务对象。当前样例仅保留对象传递口径说明，"
        "不插入缺乏独立数据支撑的接续图件；其证据重点是“对象可传递、链路可延伸”，而不是未经独立导出的量化优选结论。"
    ),
    (
        "在调控层面，规划结果还可进一步转化为工作面级控制变量。当前已复核的A、B、C方案均采用同一ODI场和4500个采样点形成方案级统计，"
        "其ODI均值、P90和超阈值暴露比例可作为后续采高、煤柱宽度、推进方向和工作面宽度调整的反馈指标。"
    ),
    (
        "在工程经济评价层面，规划与接续结果可以继续进入收入、成本、风险联动成本和现金流分析过程。当前样例已完成规划对象向后续评价口径的传递，"
        "但尚不具备独立导出的真实矿井经济对照数据，因此本文仅说明工程经济评价的输入关系与计算口径，不给出未经数据支撑的现金流曲线或经济优选结论。"
    ),
    (
        "综上，研究区样例表明，本文方法已完成“边界与钻孔输入—参数场构建—ODI风险组织—候选方案池—方案级统计—后续评价输入”的连续运行过程。"
        "当前结果能够证明方法链和对象链已经贯通，但仍需通过更多实矿案例和现场约束闭环检验工程适用性。"
    ),
]


def insert_before(paragraph, text, style=None):
    new_p = deepcopy(paragraph._p)
    paragraph._p.addprevious(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.clear()
    if style is not None:
        inserted.style = style
    inserted.add_run(text)
    return inserted


def main():
    docx_path = Path(os.environ.get("DOCX_PATH", r"E:\xiangmu\miningplan\煤科投稿\最新版论文4.16_插图版_煤科格式_大修工作稿_20260418.docx"))
    doc = Document(docx_path)

    # Remove the misordered 3.3 block.
    for p in list(doc.paragraphs):
        t = p.text.strip()
        if any(t.startswith(s) for s in REMOVE_STARTS):
            p._element.getparent().remove(p._element)

    discussion = None
    for p in doc.paragraphs:
        if p.text.strip() == "4 讨论":
            discussion = p
            break
    if discussion is None:
        raise RuntimeError("Discussion anchor not found")

    for text in TRANSMISSION:
        insert_before(discussion, text, style=discussion.style)

    doc.save(docx_path)
    print(f"Fixed 3.3 order in {docx_path}")


if __name__ == "__main__":
    main()
