import os
from pathlib import Path

from docx import Document


path = Path(os.environ["DOCX_PATH"])
doc = Document(path)


def set_text(paragraph, text):
    paragraph.clear()
    paragraph.add_run(text)


def find_para(starts):
    if isinstance(starts, str):
        starts = [starts]
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if any(text.startswith(start) for start in starts):
            return paragraph
    raise ValueError(f"paragraph not found: {starts}")


set_text(
    find_para("摘要："),
    "摘要：针对采区规划中离散地质信息难以连续参与空间布置、多场景覆岩扰动风险约束口径不统一以及规划结果难以向采掘接续和经济评价传递等问题，提出一种基于覆岩扰动指数（overburden disturbance index，ODI）约束的采区候选方案生成与多目标比选方法。该方法以采区边界、钻孔样点和设计参数为输入，先构建有效布置域与连续参数场，再将地表沉陷、含水层扰动和上行开采等风险分量归一化并加权为统一ODI场，进一步以ODI均值、90%分位数和超阈值暴露比例刻画候选方案风险水平。在此基础上，建立包含工作面数量、推进方向、煤柱宽度、工作面宽度和方案选择变量的候选方案池，并采用工程效率、资源回收和扰动控制目标进行非支配排序与综合筛选。样例中复核得到80×56栅格ODI场，共4480个栅格，均值为0.4669，P90为0.7474，ODI>0.70比例为15.89%。在同一ODI场下，工程效率优先方案A、资源回收优先方案B和ODI筛选方案C的ODI均值分别为0.4463、0.4552和0.4416，P90分别为0.6407、0.6462和0.6353，ODI>0.70比例分别为0.44%、0.56%和1.22%。权重和阈值敏感性结果表明，在基准权重±10%扰动及含水层专项权重条件下，C方案风险综合得分均低于A、B方案。研究结果说明，ODI前置约束可将风险图层转化为可复核的方案级统计指标，为采区规划阶段的多目标比选和后续对象传递提供方法支撑。",
)
set_text(
    find_para("Abstract:"),
    "Abstract: To address the difficulty of using discrete geological information in continuous mining-district layout, the inconsistent expression of multi-scenario overburden disturbance risks, and the weak linkage between planning outputs and downstream evaluation, this study proposes an ODI-constrained candidate generation and multi-objective selection method for mining-district planning. The method constructs an effective layout domain and continuous parameter fields from boundary, borehole and design inputs, normalizes surface subsidence, aquifer disturbance and upward-mining risk components into a unified overburden disturbance index (ODI) field, and describes candidate-level risk exposure using mean ODI, the 90th percentile and exceedance ratio. A constrained candidate pool is then ranked by engineering-efficiency, resource-recovery and disturbance-control objectives. The verified ODI field contains 80×56 grids, with a mean of 0.4669, P90 of 0.7474 and ODI>0.70 ratio of 15.89%. Under the same ODI field, the mean ODI values of the engineering-efficiency scheme A, resource-recovery scheme B and ODI-screened scheme C are 0.4463, 0.4552 and 0.4416, respectively; their P90 values are 0.6407, 0.6462 and 0.6353, and their ODI>0.70 ratios are 0.44%, 0.56% and 1.22%. Weight and threshold sensitivity analyses show that scheme C maintains the lowest composite risk score under ±10% weight perturbations and the aquifer-special weight setting. The results indicate that ODI pre-constraint can transform risk layers into reproducible scheme-level statistics for multi-objective comparison in mining-district planning.",
)

set_text(
    find_para("从研究区样例结果看，多场景ODI结果能够"),
    "从研究区样例结果看，多场景ODI结果能够在同一研究区边界、钻孔参照和色标口径下表达。本轮复核确认，综合ODI场采用80×56网格，共4480个栅格；归一化后ODI均值为0.4669，中位数为0.4542，P90为0.7474；其中ODI>0.70的栅格占15.89%，ODI>0.80的栅格占3.55%。这说明风险信息已经从概念性指标转化为可计算、可统计和可进入规划比选的空间场。",
)
set_text(
    find_para("在规划语境下，ODI的价值体现在两个层面"),
    "在规划语境下，ODI的价值体现在两个层面：一是描述研究区高扰动敏感区的空间背景，二是提供候选方案区域内的风险暴露统计。在同一ODI场下，方案A、B、C的ODI均值分别为0.4463、0.4552和0.4416，P90分别为0.6407、0.6462和0.6353，ODI>0.70比例分别为0.44%、0.56%和1.22%。其中C方案为在已保存候选池内按统一ODI场重新筛选得到的合格低风险候选。",
)
set_text(
    find_para("表3列出了当前已复核的候选方案"),
    "表3列出了当前已复核的候选方案。工程效率优先方案A来自efficiency模式，候选池规模为2417，覆盖率为89.34%，工程效率评分为87.82；资源回收优先方案B来自recovery模式，候选池规模为1149，覆盖率为98.52%，资源回收评分为89.76；ODI筛选方案C来自统一候选池复核，覆盖率为80.67%，工程效率评分为79.27，ODI均值为0.4416，P90为0.6353，ODI>0.70比例为1.22%。",
)
set_text(
    find_para("上述结果表明，不同目标偏好会导致覆盖率"),
    "上述结果表明，不同目标偏好会导致覆盖率、资源回收和风险暴露之间出现明显权衡。A方案覆盖率较高且超阈值比例较低，B方案资源回收评分最高但ODI均值和P90相对较高；C方案以牺牲部分覆盖率为代价，取得最低的ODI均值、P90和风险综合得分。旧disturbance保存结果中的候选C_old在统一ODI场下均值为0.4560，P90为0.6472，因此本文不再将其作为统一口径下的ODI筛选方案。",
)
set_text(
    find_para("从研究区样例看，ODI场提供了可量化的风险入口"),
    "从研究区样例看，ODI场提供了可量化的风险入口。与只给出风险图层不同，本文把ODI均值、P90和超阈值暴露比例转化为候选方案统计量，使方案A、B、C能够在同一风险口径下比较。权重敏感性结果表明，在基准权重、各分量±10%相对扰动以及含水层专项权重条件下，方案风险综合得分排序均保持为C>A>B，说明当前样例中的ODI筛选结果对权重小幅扰动具有一定稳定性。",
)
set_text(
    find_para("后续研究应在实矿数据中补充"),
    "后续研究应在实矿数据中补充更多真实工程约束、现场风险标定、接续排程数值和经济评价参数体系，并在更多采区案例中检验ODI权重、阈值和网格尺度对方案排序的影响，使当前已经打通的方法链进一步转化为可量化的工程优选依据。",
)
set_text(
    find_para("3）研究区样例验证表明"),
    "3）研究区样例验证表明，工程效率优先、资源回收优先和ODI筛选方案分别体现出覆盖率、资源利用和风险暴露之间的权衡。方案A、B、C的ODI均值分别为0.4463、0.4552和0.4416，P90分别为0.6407、0.6462和0.6353，ODI>0.70比例分别为0.44%、0.56%和1.22%，说明风险场能够转化为可复核的方案级统计量。",
)
set_text(
    find_para("4）当前结果主要支撑"),
    "4）当前结果主要支撑样例条件下的方法贯通、对象链有效性和初步方案比选。权重敏感性表明，在基准权重、±10%相对扰动和含水层专项权重条件下，C方案均保持最低风险综合得分；阈值敏感性表明，当阈值由0.65提高至0.80时，各方案超限比例均下降，但方案间风险差异仍可通过均值、P90和超限比例共同识别。真实矿井条件下的工程优选结论仍需结合实矿案例检验。",
)

table3 = doc.tables[2]
rows = [
    ["方案", "来源/含义", "关键指标", "ODI统计", "复核状态"],
    ["A", "efficiency；工程效率优先", "工作面5个；覆盖率89.34%；工程效率评分87.82", "均值0.4463；P90=0.6407；ODI>0.70为0.44%", "统一ODI场已复核"],
    ["B", "recovery；资源回收优先", "工作面9个；覆盖率98.52%；工程效率98.52；资源回收89.76", "均值0.4552；P90=0.6462；ODI>0.70为0.56%", "统一ODI场已复核"],
    ["C", "统一ODI筛选；合格低风险候选", "工作面4个；覆盖率80.67%；工程效率评分79.27", "均值0.4416；P90=0.6353；ODI>0.70为1.22%", "风险综合得分最低"],
    ["C_old", "旧disturbance保存候选", "工作面13个；覆盖率75.06%；工程效率评分71.17", "均值0.4560；P90=0.6472；ODI>0.70为0.80%", "旧口径候选，不作为本文C方案"],
]
for r, vals in enumerate(rows):
    for c, val in enumerate(vals):
        table3.cell(r, c).text = val

table4 = doc.tables[3]
while len(table4.rows) > 6:
    tr = table4.rows[-1]._tr
    tr.getparent().remove(tr)
rows = [
    ["项目", "结果", "用途", "数据文件", "说明"],
    ["全域ODI场", "80×56；4480栅格；均值0.4669；P90=0.7474；ODI>0.70为15.89%", "研究区风险背景", "000_mindong_layout_odi_field.json", "已复核"],
    ["统一方案对比", "A/B/C均值为0.4463/0.4552/0.4416；P90为0.6407/0.6462/0.6353", "方案级风险比较", "coal_sci_abc_odi_unified_stats_20260418.csv", "同一ODI场采样"],
    ["阈值敏感性", "阈值0.65/0.70/0.75/0.80下，C超限率为6.31%/1.22%/0.13%/0.00%", "检验阈值影响", "coal_sci_threshold_sensitivity_candidates_20260418.csv", "A、B同步计算"],
    ["权重敏感性", "基准、±10%扰动和含水层专项权重下，风险得分排序均为C>A>B", "检验权重稳定性", "coal_sci_weight_sensitivity_candidates_20260418.csv", "风险得分越小越优"],
    ["经济评价", "仅保留输入关系", "下游接口说明", "规划对象与接续参数", "不写未经独立导出的NPV和回收期"],
]
for r, vals in enumerate(rows):
    for c, val in enumerate(vals):
        table4.cell(r, c).text = val

doc.save(path)
print(f"updated {path}")
