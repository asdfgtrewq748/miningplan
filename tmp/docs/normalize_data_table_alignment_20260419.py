from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import sys
from pathlib import Path

p=Path(sys.argv[1])
doc=Document(p)
# Data tables are non-formula tables. Formula tables have equation-number structure or OMML content; use known table captions nearby less critical.
for table in doc.tables:
    table_text='\n'.join(cell.text for row in table.rows for cell in row.cells)
    is_data = any(k in table_text for k in ['候选方案对比', 'ODI统计口径', '参数类别', '参数名称', '阈值敏感性', '权重敏感性', '经济评价'])
    if not is_data:
        continue
    for ri,row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if ri == 0 else WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.font.size = Pt(9)

doc.save(p)
print('done')
