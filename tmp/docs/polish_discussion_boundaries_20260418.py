from pathlib import Path
import os

from docx import Document


def replace_between(doc, start_text, end_text, new_paragraphs):
    start_idx = end_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if start_idx is None and t == start_text:
            start_idx = i
        elif start_idx is not None and t == end_text:
            end_idx = i
            break
    if start_idx is None or end_idx is None:
        raise RuntimeError(f"Cannot find block {start_text!r} -> {end_text!r}")

    # Keep the section heading paragraph, replace content until before next heading.
    paras = doc.paragraphs
    heading = paras[start_idx]
    for p in list(paras[start_idx + 1:end_idx]):
        p._element.getparent().remove(p._element)

    # Insert new paragraphs after heading in reverse order.
    from copy import deepcopy
    from docx.text.paragraph import Paragraph

    anchor = heading
    for text in reversed(new_paragraphs):
        new_p = deepcopy(anchor._p)
        anchor._p.addnext(new_p)
        inserted = Paragraph(new_p, anchor._parent)
        inserted.clear()
        inserted.style = heading.style
        inserted.add_run(text)


def replace_start(doc, prefix, replacement):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            p.clear()
            p.add_run(replacement)
            return True
    return False


def main():
    docx_path = Path(os.environ.get("DOCX_PATH", r"E:\xiangmu\miningplan\煤科投稿\最新版论文4.16_插图版_煤科格式_大修工作稿_20260418.docx"))
    doc = Document(docx_path)

    replace_between(
        doc,
        "4.1 ODI 前置约束对采区规划的作用与边界",
        "4.2 参数场驱动的规划方法相对于传统经验布置的意义",
        [
            (
                "传统采区规划多采用“先形成布局方案、再开展风险校核”的串行方式。该流程的主要问题在于，"
                "风险信息只有在布局完成后才参与修正，一旦高扰动区与工作面边界发生冲突，就需要重新调整边界、煤柱和巷道组织。"
                "本文将ODI转化为候选方案生成阶段即可调用的空间统计量，使风险控制从后验校核前移到候选池筛选与排序环节。"
            ),
            (
                "ODI前置约束的优势不在于替代地表沉陷、含水层扰动或上行开采等专项分析，而在于为多源风险结果提供同一比较口径。"
                "在同一ODI场下，方案可以同时报告均值、P90和超阈值暴露比例，从而避免仅凭单一风险图或单一阈值判断方案优劣。"
                "本样例中C方案的超阈值比例并非在所有阈值下都最低，但其均值、P90和风险综合得分较低，说明方案风险需要用多统计量联合解释。"
            ),
            (
                "ODI也存在信息压缩带来的误判风险。线性加权会把不同风险场景压缩为单一指标，可能掩盖某一分量的局部高值；"
                "阈值和权重的选择也会影响超限暴露比例和排序结果。因此，工程应用中不宜只输出ODI综合图层，而应同步保留地表沉陷、含水层扰动、上行开采等分量图层，"
                "并将权重敏感性、阈值敏感性和关键保护对象校核作为必要的复核步骤。"
            ),
        ],
    )

    replace_between(
        doc,
        "4.2 参数场驱动的规划方法相对于传统经验布置的意义",
        "4.3 后续深化方向",
        [
            (
                "采区规划结果的可靠性不仅取决于边界几何条件，还与地质属性的空间连续变化密切相关。"
                "传统经验式布置能够快速形成初始形态，但对钻孔控制下的煤厚变化、资源富集区和低厚度区响应不足，"
                "容易造成资源评价与几何布置脱节。连续参数场的作用，是把离散钻孔样点转化为可与有效布置域、ODI场和工作面对象叠加的空间变量。"
            ),
            (
                "本样例中，煤厚样点由2.8 m变化到4.8 m，平均值为3.7933 m。厚煤区在资源覆盖评价中具有更高权重，"
                "低厚度区则提示边界调整、煤柱保留或低优先级布置的必要性。由此可见，参数场不是单纯的制图结果，而是候选方案评分和资源回收评价的输入条件。"
            ),
            (
                "同时，连续参数场质量会直接影响规划结果。钻孔数量、空间分布、插值方法、网格尺度和异常值处理都会改变参数场形态，"
                "进而影响工作面覆盖率、资源回收评分和ODI统计。本文当前采用样例数据验证方法链路，尚未比较IDW、克里金等不同插值方法对方案排序的影响；"
                "后续若用于工程定案，应补充交叉验证、插值误差分析和网格尺度敏感性分析。"
            ),
        ],
    )

    replace_between(
        doc,
        "4.3 后续深化方向",
        "5 结论",
        [
            (
                "本文当前结果主要证明所提方法在样例条件下的链路贯通能力，即能够实现采区边界、钻孔样点、连续参数场、ODI风险场、候选方案池和规划对象之间的连续组织与传递。"
                "A、B、C候选方案已在同一ODI场下完成均值、P90和超阈值暴露比例补算，能够支撑初步方案对比，但仍不足以形成真实矿井条件下的工程优选结论。"
            ),
            (
                "方法的适用边界主要包括三方面：第一，ODI权重和阈值需要结合矿区历史观测、保护对象等级和现场安全制度重新标定；"
                "第二，候选方案池需要纳入断层、既有巷道、通风运输能力、设备能力和生产制度等现场闭环约束；"
                "第三，经济评价需要真实煤价、成本、税费、投资和风险停产参数支撑，不能仅凭接口传递关系给出现金流优选。"
            ),
            (
                "后续研究应从样例级验证转向多矿井对照验证：一方面补充传统经验方案、无ODI方案和不同插值方案的基准对比，"
                "另一方面在真实接续排程和经济参数下检验ODI前置约束是否能够稳定降低风险暴露并保持可接受的资源回收水平。"
                "只有完成上述验证后，本文方法才能从“可复核的规划链路”进一步发展为“可用于工程定案的优选工具”。"
            ),
        ],
    )

    replace_start(
        doc,
        "1）针对采区规划中地质信息离散、风险约束异构以及规划结果与后续评价脱节的问题",
        (
            "1）针对采区规划中地质信息离散、风险约束异构以及规划结果与后续评价脱节的问题，"
            "提出了基于ODI约束的采区候选方案生成与多目标比选方法，构建了“有效布置域—连续参数场—ODI风险场—候选方案池—后续评价输入”的方法链路。"
        ),
    )

    doc.save(docx_path)
    print(f"Polished discussion boundaries in {docx_path}")


if __name__ == "__main__":
    main()
