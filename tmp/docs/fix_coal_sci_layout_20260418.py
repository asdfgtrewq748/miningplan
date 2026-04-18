import os
from pathlib import Path

from docx import Document
from docx.shared import Pt


path = Path(os.environ["DOCX_PATH"])
doc = Document(path)

formula_prefixes = (
    "ODI(",
    "w_s+",
    "X_i'",
    "E_π=",
    "Ω_e=",
    "z(x)=",
    "π={",
    "A_π⊂",
    "F(π)=",
    "G(s)=",
    "NCF_t=",
)

for para in doc.paragraphs:
    text = para.text.strip()
    if text.startswith(formula_prefixes):
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(3)
        para.paragraph_format.line_spacing = 1.0
        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)

for table in doc.tables:
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.0
                for run in para.runs:
                    run.font.size = Pt(7)

doc.save(path)
print(f"fixed paragraph/table fonts {path}")
