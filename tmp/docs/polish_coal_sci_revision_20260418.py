import os
from pathlib import Path

from docx import Document


path = Path(os.environ["DOCX_PATH"])
doc = Document(path)


def set_text(paragraph, text):
    paragraph.clear()
    paragraph.add_run(text)


def find_para(starts):
    if isinstance(starts, str):
        starts = [starts]
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if any(text.startswith(start) for start in starts):
            return paragraph
    raise ValueError(f"paragraph not found: {starts}")


set_text(
    find_para("图1 覆岩扰动约束机理示意图"),
    "图1 ODI风险约束逻辑示意图\nFig.1 Schematic diagram of the ODI risk-constraint logic",
)
set_text(
    find_para("式中，π为候选规划方案"),
    "式中，π为候选规划方案，A_π为方案π对应的布置区域，N_π为方案区域内参与统计的采样或栅格数量，T_ODI为扰动控制阈值。本文以ODI均值、Q_0.90和超阈值暴露比例E_π共同描述方案级风险，其中ODI均值为方案区域内采样值的算术平均，Q_0.90为方案区域ODI序列的90%分位数，E_π为超过阈值T_ODI的采样比例，不把单一ODI均值作为唯一判据。阈值0.70和0.80分别作为预警统计线和高扰动统计线，用于方案间比较；在真实矿井应用中应结合历史观测、保护对象等级和现场安全制度重新标定。",
)
set_text(doc.paragraphs[41], "Ω_e=Ω_0\\(B_b∪B_s∪D_p)。                                                        （6）")
set_text(doc.paragraphs[44], "z(x)=Σ[z_i d_i(x)^(-p)] / Σ[d_i(x)^(-p)]。                                      （7）")
set_text(
    find_para("本文当前结果主要证明所提方法"),
    "本文当前结果主要证明所提方法在样例条件下的链路贯通能力，即能够实现采区边界、钻孔样点、连续参数场、ODI风险场、候选方案池和规划对象之间的连续组织与传递。当前已提取的A、B、C候选方案说明不同目标偏好会改变覆盖率、资源回收和扰动暴露结果，但仍需补算A、B方案的ODI统计并开展更多实矿对照，才能形成更强的工程优选证据。",
)

doc.save(path)
print(f"polished {path}")
