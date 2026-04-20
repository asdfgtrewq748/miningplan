from docx import Document
from pathlib import Path
import sys
p=Path(sys.argv[1])
out=Path(sys.argv[2])
doc=Document(p)
with out.open('w',encoding='utf-8') as f:
    for i,para in enumerate(doc.paragraphs):
        t=para.text.strip()
        if t:
            f.write(f'P{i:04d}\t{para.style.name}\t{t}\n')
    for ti,table in enumerate(doc.tables):
        f.write(f'\n[TABLE {ti}]\n')
        for row in table.rows:
            cells=[' '.join(c.text.split()) for c in row.cells]
            f.write('\t'.join(cells)+'\n')
print(out)
