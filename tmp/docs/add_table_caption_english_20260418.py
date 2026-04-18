from pathlib import Path
import os

from docx import Document


def replace_exact(doc, old, new):
    for p in doc.paragraphs:
        if p.text.strip() == old:
            p.clear()
            p.add_run(new)
            return True
    return False


def main():
    docx_path = Path(os.environ.get("DOCX_PATH", r"E:\xiangmu\miningplan\煤科投稿\最新版论文4.16_插图版_煤科格式_大修工作稿_20260418.docx"))
    doc = Document(docx_path)
    replacements = {
        "表3 候选方案对比与数据复核状态": "表3 候选方案对比与数据复核状态\nTable 3 Candidate-scheme comparison and data verification status",
        "表4 ODI统计口径与敏感性复核状态": "表4 ODI统计口径与敏感性复核状态\nTable 4 ODI statistical scope and sensitivity verification status",
    }
    missing = []
    for old, new in replacements.items():
        if not replace_exact(doc, old, new):
            missing.append(old)
    if missing:
        print("Missing captions:", missing)
    doc.save(docx_path)
    print(f"Updated table captions in {docx_path}")


if __name__ == "__main__":
    main()
