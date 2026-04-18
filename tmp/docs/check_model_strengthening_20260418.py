from __future__ import annotations

import sys
import zipfile
from pathlib import Path

from docx import Document


CHECKS = {
    "variable_definition": "\u5f0f\u4e2d\uff0cN\u4e3a\u5de5\u4f5c\u9762\u6570\u91cf",
    "generation_flow": "\u5019\u9009\u6c60\u751f\u6210\u53ef\u5206\u4e3a\u53c2\u6570\u53d6\u503c\u3001\u51e0\u4f55\u6784\u9020\u3001\u786c\u7ea6\u675f\u8fc7\u6ee4\u3001\u6307\u6807\u8ba1\u7b97\u548c\u53bb\u91cd\u5408\u5e765\u4e2a\u6b65\u9aa4",
    "dominance_definition": "\u975e\u652f\u914d\u5173\u7cfb\u5177\u4f53\u5b9a\u4e49\u4e3a",
    "preference_weights": "\u4e0d\u540c\u504f\u597d\u6a21\u5f0f\u901a\u8fc7\u03bb_e\u3001\u03bb_r\u548c\u03bb_m\u4f53\u73b0",
}


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: check_model_strengthening_20260418.py <docx>")
        return 2

    path = Path(sys.argv[1])
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
    if bad is not None:
        raise RuntimeError(f"Bad DOCX zip member: {bad}")

    doc = Document(str(path))
    text = "\n".join(p.text for p in doc.paragraphs)
    for name, phrase in CHECKS.items():
        print(f"{name}={phrase in text}")
    print(
        "duplicate_transfer_count="
        + str(text.count("\u672c\u6587\u65b9\u6cd5\u7684\u4e00\u4e2a\u91cd\u8981\u7279\u70b9\uff0c\u662f\u4e0d\u628a\u91c7\u533a\u89c4\u5212\u7684\u51e0\u4f55\u5e03\u5c40\u89c6\u4e3a\u7ec8\u70b9"))
    )
    print(f"paragraphs={len(doc.paragraphs)}")
    print(f"tables={len(doc.tables)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
