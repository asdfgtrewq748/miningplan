from copy import deepcopy
from pathlib import Path
import os

from docx import Document
from docx.text.paragraph import Paragraph


def insert_before(paragraph, text, style=None):
    new_p = deepcopy(paragraph._p)
    paragraph._p.addprevious(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.clear()
    if style is not None:
        inserted.style = style
    inserted.add_run(text)
    return inserted


def replace_exact_or_start(doc, old_start, new_text):
    for p in doc.paragraphs:
        if p.text.strip().startswith(old_start):
            p.clear()
            p.add_run(new_text)
            return p
    return None


def main():
    docx_path = Path(os.environ.get("DOCX_PATH", r"E:\xiangmu\miningplan\煤科投稿\最新版论文4.16_插图版_煤科格式_大修工作稿_20260418.docx"))
    doc = Document(docx_path)

    replace_exact_or_start(
        doc,
        "上述结果表明，不同目标偏好会导致覆盖率、资源回收和风险暴露之间出现明显权衡。",
        (
            "上述结果表明，不同目标偏好会导致覆盖率、资源回收和风险暴露之间出现明显权衡。"
            "与A方案相比，C方案覆盖率降低8.67个百分点，但ODI均值降低1.07%，P90降低0.86%，风险综合得分降低0.70%；"
            "与B方案相比，C方案覆盖率降低17.85个百分点，但ODI均值降低3.00%，P90降低1.69%，风险综合得分降低2.12%。"
            "B方案相对于A方案覆盖率提高9.18个百分点，但ODI均值和P90分别提高1.99%和0.85%。"
            "这说明资源覆盖增益可能伴随扰动风险统计值上升，ODI前置约束能够把这种风险-收益取舍显式化。"
            "旧disturbance保存结果中的候选C_old在统一ODI场下均值为0.4560，P90为0.6472，因此本文不再将其作为统一口径下的ODI筛选方案。"
        ),
    )

    # Convert old 3.2 transmission section into sensitivity section.
    replace_exact_or_start(doc, "3.2 规划结果向采掘接续与工程经济评价的传递", "3.2 阈值与权重敏感性分析")
    replace_exact_or_start(
        doc,
        "本文方法的一个重要特点，是不把采区规划的几何布局视为终点",
        (
            "为检验ODI阈值设定对方案判读的影响，本文在0.65、0.70、0.75和0.80四组阈值下重新统计A、B、C方案的超阈值暴露比例。"
            "当阈值为0.65时，A、B、C方案超限比例分别为7.42%、9.09%和6.31%；当阈值为0.70时，三者分别为0.44%、0.56%和1.22%；"
            "当阈值提高至0.75和0.80时，各方案超限比例均明显下降。该结果说明，超限比例对阈值设定较敏感，不能脱离ODI均值和P90单独作为方案优选依据。"
        ),
    )
    replace_exact_or_start(
        doc,
        "在采掘接续层面，规划阶段形成的工作面边界和推进关系可映射为接续任务对象。",
        (
            "权重敏感性分析进一步表明，在基准权重、地表沉陷分量±10%、含水层扰动分量±10%、上行开采分量±10%以及含水层专项权重条件下，"
            "C方案的风险综合得分均为三方案最低，A方案次之，B方案最高。基准权重下A、B、C风险综合得分分别为0.4481、0.4546和0.4449；"
            "在含水层专项权重0.15/0.25/0.60下，三者分别为0.4505、0.4564和0.4498。"
        ),
    )
    replace_exact_or_start(
        doc,
        "在调控层面，规划结果还可进一步转化为工作面级控制变量。",
        (
            "需要注意，权重敏感性所证明的是当前候选池和扰动范围内的排序稳定，而不是ODI权重具有普适工程常数属性。"
            "因此，本文把敏感性分析作为方案排序可靠性的内部校核，用于回应权重和阈值人为设定可能导致结论不稳定的问题。"
        ),
    )
    replace_exact_or_start(
        doc,
        "在工程经济评价层面，规划与接续结果可以继续进入收入、成本、风险联动成本和现金流分析过程。",
        (
            "综合来看，C方案并不是覆盖率最大或超限比例在所有阈值下都最低的方案，而是在ODI均值、P90和风险综合得分联合判据下表现较优的低扰动候选。"
            "这一结果强化了本文方法的核心作用：把不同目标方案放入同一风险统计框架中比较，并把单一指标无法解释的风险-收益取舍显式呈现。"
        ),
    )
    replace_exact_or_start(
        doc,
        "综上，研究区样例表明，本文方法已完成“边界与钻孔输入—参数场构建—ODI风险组织—候选方案池—方案级统计—后续评价输入”的连续运行过程。",
        (
            "综上，研究区样例表明，统一ODI场不仅可以生成候选方案的均值、P90和超阈值暴露比例，还能够支撑阈值敏感性和权重敏感性复核。"
            "当前结果已补齐A、B、C方案的统一ODI统计和全域4480栅格统计，但仍不宜把单一样例结果表述为真实矿井条件下的最终优选结论。"
        ),
    )
    replace_exact_or_start(doc, "表4 ODI统计口径与后续补算需求", "表4 ODI统计口径与敏感性复核状态")

    # Insert 3.3 transmission section before Discussion.
    discussion = None
    for p in doc.paragraphs:
        if p.text.strip() == "4 讨论":
            discussion = p
            break
    if discussion is not None:
        transmission = [
            "3.3 规划结果向采掘接续与工程经济评价的传递",
            (
                "本文方法的一个重要特点，是不把采区规划的几何布局视为终点，而是强调规划结果可以继续作为采掘接续和工程经济评价的上游输入。"
                "在本样例中，工作面边界、巷道路径、推进长度、布置面积和ODI统计量均已形成可传递的规划结果，后续环节可据此组织生产顺序、产量核算和经济评价。"
            ),
            (
                "在采掘接续层面，规划阶段形成的工作面边界和推进关系可映射为接续任务对象。当前样例仅保留对象传递口径说明，"
                "不插入缺乏独立数据支撑的接续图件；其证据重点是“对象可传递、链路可延伸”，而不是未经独立导出的量化优选结论。"
            ),
            (
                "在调控层面，规划结果还可进一步转化为工作面级控制变量。当前已复核的A、B、C方案均采用同一ODI场和4500个采样点形成方案级统计，"
                "其ODI均值、P90和超阈值暴露比例可作为后续采高、煤柱宽度、推进方向和工作面宽度调整的反馈指标。"
            ),
            (
                "在工程经济评价层面，规划与接续结果可以继续进入收入、成本、风险联动成本和现金流分析过程。当前样例已完成规划对象向后续评价口径的传递，"
                "但尚不具备独立导出的真实矿井经济对照数据，因此本文仅说明工程经济评价的输入关系与计算口径，不给出未经数据支撑的现金流曲线或经济优选结论。"
            ),
            (
                "综上，研究区样例表明，本文方法已完成“边界与钻孔输入—参数场构建—ODI风险组织—候选方案池—方案级统计—后续评价输入”的连续运行过程。"
                "当前结果能够证明方法链和对象链已经贯通，但仍需通过更多实矿案例和现场约束闭环检验工程适用性。"
            ),
        ]
        # Insert in reverse order so final order is preserved.
        for text in reversed(transmission):
            insert_before(discussion, text, style=discussion.style)
    else:
        print("Discussion anchor not found")

    replace_exact_or_start(
        doc,
        "本文当前结果主要证明所提方法在样例条件下的链路贯通能力",
        (
            "本文当前结果主要证明所提方法在样例条件下的链路贯通能力，即能够实现采区边界、钻孔样点、连续参数场、ODI风险场、候选方案池和规划对象之间的连续组织与传递。"
            "当前已提取的A、B、C候选方案说明不同目标偏好会改变覆盖率、资源回收和扰动暴露结果，且A、B、C方案已在同一ODI场下完成均值、P90和超阈值暴露比例补算。"
            "但这些结果仍主要属于样例级对比，仍需更多实矿对照和现场阈值标定，才能形成更强的工程优选证据。"
        ),
    )

    doc.save(docx_path)
    print(f"Strengthened comparison and sensitivity results in {docx_path}")


if __name__ == "__main__":
    main()
