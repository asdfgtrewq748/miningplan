from pathlib import Path
import sys

from docx import Document


DOCX = Path(sys.argv[1])


PARA_REPLACEMENTS = {
    3: (
        "摘要：针对采区规划中离散地质信息难以连续参与布置、多场景覆岩扰动风险约束口径不统一以及规划结果难以向接续和经济评价传递等问题，提出一种基于覆岩扰动指数（overburden disturbance index，ODI）约束的采区候选方案生成与多目标比选方法。该方法以采区边界、钻孔样点和设计参数为输入，构建有效布置域和连续参数场；当前样例以煤层厚度场为主要展示对象，将地表沉陷、含水层扰动、上行开采等风险分量归一化并加权为统一ODI场，并以ODI均值、90%分位数和超阈值暴露比例表征方案风险。在此基础上，建立包含工作面数量、推进方向、煤柱宽度、工作面宽度和方案选择变量的候选方案集合，采用工程效率、资源回收和扰动控制目标进行非支配排序与综合筛选。样例复核得到80×56栅格ODI场，均值为0.4669，P90为0.7474，ODI>0.70比例为15.89%。在同一ODI场下，工程效率优先方案A、资源回收优先方案B和ODI筛选方案C的ODI均值分别为0.4463、0.4552和0.4416，P90分别为0.6407、0.6462和0.6353，ODI>0.70比例分别为0.44%、0.56%和1.22%。本文对低扰动候选的判定依据为ODI均值、P90和风险综合得分的联合判据，而非单一阈值下的超限比例最小。敏感性分析表明，在当前样例和候选集合范围内，C方案风险综合得分低于A、B方案。结果说明，ODI前置约束可将风险图层转化为可复核的方案级统计指标，为采区规划多目标比选和后续对象传递提供支撑；真实矿井优选仍需结合现场约束和多案例验证。"
    ),
    9: (
        "Abstract: To address the limited use of discrete geological data in continuous mining-district layout, the inconsistent expression of multi-scenario overburden disturbance risks, and the weak linkage between planning outputs and downstream evaluation, an ODI-constrained candidate generation and multi-objective selection method is proposed. Boundary, borehole and design inputs are used to construct the effective layout domain and continuous parameter fields; in the current sample, coal-seam thickness is used as the main example showing how a continuous parameter field enters the planning chain. Surface subsidence, aquifer disturbance and upward-mining risk components are normalized and weighted into a unified overburden disturbance index (ODI) field, and candidate-level risk is described by mean ODI, the 90th percentile and exceedance ratio. A constrained candidate set considering face number, advance direction, pillar width, face width and scheme-selection variables is then ranked using engineering-efficiency, resource-recovery and disturbance-control objectives. The verified sample ODI field contains 80×56 grids, with a mean of 0.4669, P90 of 0.7474 and ODI>0.70 ratio of 15.89%. Under the same ODI field, schemes A, B and C have mean ODI values of 0.4463, 0.4552 and 0.4416, P90 values of 0.6407, 0.6462 and 0.6353, and ODI>0.70 ratios of 0.44%, 0.56% and 1.22%, respectively. The low-disturbance candidate is identified by the joint criterion of mean ODI, P90 and composite risk score rather than by the minimum exceedance ratio under a single threshold. Sensitivity analysis indicates that scheme C has a lower composite risk score than schemes A and B within the present sample and candidate set. The results show that ODI pre-constraint can convert risk layers into reproducible scheme-level statistics for multi-objective comparison, while engineering selection under real mine conditions still requires site-specific constraints and multi-case validation."
    ),
    18: (
        "基于上述问题，本文提出一种基于覆岩扰动指数（overburden disturbance index，ODI）约束的采区候选方案生成与多目标比选方法。与在布局完成后再进行风险校核的串行流程不同，本文将地表沉陷、含水层扰动和上行开采等异构风险前置转化为统一ODI约束，并在有效布置域、连续参数场和工作面布置对象之间建立可计算映射关系，使风险控制能够参与候选方案生成、筛选和排序过程。本文侧重方法论验证与内部对照，不以输出矿井定案设计参数为直接目标。"
    ),
    28: (
        "式中，w_s、w_a和w_u分别为地表沉陷、含水层扰动和上行开采分量权重。本文采用“场景目标驱动的专家规则赋权+敏感性校核”确定权重，不将其解释为现场统计标定常数。综合规划基准权重取0.45/0.30/0.25；为检验上行开采约束强化时方案排序是否改变，上行开采专项权重取0.15/0.25/0.60。为避免总指标与子指标语义重叠，本文以具体风险场景命名各分量；后续工程应用仍需结合矿区历史沉陷、水文地质观测和保护对象等级重新标定权重。"
    ),
    43: (
        "在满足几何可实施性和安全隔离要求的前提下，采区规划需要在工程效率、资源回收和覆岩扰动控制之间进行权衡。本文将规划问题定义为“候选方案集合生成—底线约束过滤—多目标排序—工程复核”的组合优化过程，而不是单一程序输出结果。"
    ),
    47: (
        "式中，A_π⊂Ω_e为边界约束，B_b和B_s为煤柱约束，W_f为工作面宽度约束，A_i∩A_j=∅为工作面不重叠约束，C_L表示推进长度推荐区间或校核条件，I_ODI表示ODI暴露统计参与方案筛选与风险排序的指标口径。当前样例将边界、煤柱、工作面宽度和不重叠作为硬约束，将推进长度校核阈值与ODI暴露阈值作为候选方案筛选和排序中的推荐性约束，以避免样例验证阶段因现场闭环约束不完整而过度收缩候选空间。"
    ),
    48: (
        "本文首先在统一基础输入、统一几何底线约束和统一ODI场条件下形成可合并的总候选集合Π_all=Π_e∪Π_r∪Π_m，再在该集合上分别按工程效率、资源回收和扰动控制偏好进行排序与筛选，输出工程效率优先方案A、资源回收优先方案B和扰动控制优先方案C。推进方向不作黑箱“自动确定”，而是在工程允许方向集合Θ内与煤柱宽度、工作面宽度和边界内缩距离共同枚举，并由约束过滤和多目标排序确定推荐结果。"
    ),
    49: (
        "候选集合生成包括参数取值、几何构造、底线约束过滤、指标计算和去重合并5步：读取边界、钻孔参数场和规划控制参数，确定Θ、W_f、B_b、B_s及内缩距离取值；切分有效布置域并生成工作面区域、推进长度和巷道连接对象；按式（9）剔除越界、煤柱不足、宽度不符和工作面重叠等几何不合格方案，并记录推进长度与ODI暴露校核状态；计算覆盖率、煤厚覆盖、巷道组织、ODI均值、P90和超阈值暴露比例；最后依据方向、煤柱、工作面数量、边界和连接关系形成候选签名并合并重复方案。后续排序仅在该有限候选集合Π_all内成立。"
    ),
    51: (
        "式中，C_cov为有效覆盖率，P_N为工作面数量复杂度惩罚，CV_L为工作面推进长度变异系数，P_short为短推进惩罚项。当前样例工程效率评分采用程序导出的覆盖率基础分扣减组织复杂度惩罚口径，其中P_N=0.20N，CV_L项系数取10，短推进惩罚系数取5，短推进参考值取100 m。"
    ),
    52: (
        "式中，R_ton为吨煤量或煤厚场覆盖贡献，R_area为布置面积或有效覆盖面积得分，S_eng为工程组织辅助得分。当前样例资源回收评分采用0.55/0.30/0.15的吨煤量、覆盖率和工程组织权重，并进行显示标定；因此，89.76表示资源回收分项评分，而非综合排序总分。"
    ),
    54: (
        "上述评分给出了87.82、89.76等方案分数的计算口径：87.82为A方案工程效率分项评分，89.76为B方案资源回收分项评分；B方案导出的工程效率辅助值为98.52，与覆盖率98.52%数值相同，但B方案排序依据为资源回收评分。A、B、C推荐方案的最终比较还需结合非支配排序后的偏好加权结果，相关分数仅用于当前候选集合内部比较，不作为跨矿区通用评分标准。"
    ),
    56: (
        "为降低单一加权排序对权重的依赖，本文在加权排序前引入非支配排序和拥挤距离选择形成Top-K候选方案，当前样例Top-K取10。非支配排序以S_e、S_r和S_m为目标向量，优先保留不存在其他方案同时优于它的候选；同一非支配层候选数量超过输出数量时，采用拥挤距离保持解集多样性，再依据F(π)推荐排序。A、B、C三种偏好模式中，λ_e/λ_r/λ_m分别取1/0/0、0/1/0和0/0/1；综合权衡模式备用权重取0.34/0.33/0.33，本文不将其未达标回退结果作为最终推荐方案。"
    ),
    73: (
        "为避免验证目标泛化，本文将样例验证限定为3个可复核问题：离散钻孔和边界约束能否转化为统一坐标下的连续参数场和有效布置域；多场景扰动风险能否转化为同一ODI统计口径，并在全域栅格和候选方案区域内复算；不同目标偏好下的候选方案能否在同一基础输入、同一几何底线约束、同一ODI风险场和同一评分口径下比较。后文A、B、C方案对比、阈值敏感性和权重敏感性均围绕上述问题展开。"
    ),
    82: "2.2 连续参数场构建结果及煤厚场示例",
    105: (
        "为增强案例对照性，本文将A、B、C三类方案置于相同研究区边界、钻孔参数场、ODI风险场和几何底线约束下比较。A方案代表工程效率偏好，强调工作面数量、推进长度均衡性和组织效率；B方案代表资源回收偏好，强调覆盖率和煤厚场利用；C方案代表ODI风险约束偏好，强调低扰动暴露。该对照属于统一候选集合内的目标偏好对照，而非独立人工经验方案与自动方案的现场对照。"
    ),
    106: (
        "为保证可复核性，A、B、C方案不改变基础地质输入和风险场，仅改变候选排序偏好或筛选目标；方案差异主要反映工程效率、资源回收和扰动控制目标之间的权衡，而非输入数据、网格尺度或边界条件变化。当前导出记录中，工程效率、资源回收和扰动控制模式的有效候选数分别为2417、1149和374，其中扰动控制模式合格候选数为88。该设置可支撑方法内部有效性判断，但不能替代真实人工经验方案与本文方法之间的外部工程对照。"
    ),
    107: (
        "表3列出了当前已复核的候选方案。工程效率优先方案A来自efficiency模式，覆盖率为89.34%，工程效率分项评分为87.82；资源回收优先方案B来自recovery模式，覆盖率为98.52%，资源回收分项评分为89.76；ODI筛选方案C来自统一候选集合复核，覆盖率为80.67%，工程效率分项评分为79.27，ODI均值为0.4416，P90为0.6353，ODI>0.70比例为1.22%。需要指出，本文对低扰动候选的判定并非依据单一阈值下的超限比例最小，而是依据ODI均值、P90与风险综合得分的联合判据。"
    ),
    109: (
        "从方案合理性看，A方案覆盖率和工程效率较高，几何组织较充分；B方案覆盖率最高且资源回收评分较高，说明连续参数场推动方案向资源富集区域扩展；C方案覆盖率较低，且在T_ODI=0.70下的超限比例并非最小，但其ODI均值、P90和风险综合得分均低于A、B方案，说明联合判据下总体扰动暴露较低。该结果符合采区规划中资源利用与风险控制相互制约的工程常识，也说明本文方法能够将图层风险差异转化为方案级统计证据。"
    ),
    110: (
        "表3 候选方案对比与数据复核状态（低扰动判定基于ODI均值、P90和风险综合得分联合判据）\nTable 3 Candidate-scheme comparison and data verification status"
    ),
    114: (
        "权重敏感性分析进一步表明，在基准权重、地表沉陷分量±10%、含水层扰动分量±10%、上行开采分量±10%以及上行开采专项权重条件下，C方案的风险综合得分均低于A、B方案。基准权重下A、B、C风险综合得分分别为0.4481、0.4546和0.4449；在上行开采专项权重0.15/0.25/0.60下，三者分别为0.4505、0.4564和0.4498。"
    ),
    135: (
        "连续参数场质量会直接影响规划结果。钻孔数量、空间分布、插值方法、网格尺度和异常值处理均会改变参数场形态，进而影响覆盖率、资源回收评分和ODI统计。智能精准开采工作面地质模型研究强调多源信息逐级融合和动态更新，以提高工作面尺度地质表达精度[31-32]。本文当前样例以煤厚场作为连续参数场的主要展示对象，岩石硬度、瓦斯含量和涌水量等属性尚未作为独立图层展开验证；工程定案前应补充交叉验证、插值误差分析、网格尺度敏感性分析以及不同插值方法对方案排序的对照。"
    ),
    142: (
        "1）针对采区规划中地质信息离散、风险约束异构以及规划结果与后续评价脱节的问题，提出了基于ODI约束的采区候选方案生成与多目标比选方法，构建了“有效布置域—连续参数场—ODI风险场—候选方案集合—后续评价输入”的方法链路。"
    ),
    144: (
        "3）研究区样例结果表明，工程效率优先、资源回收优先和ODI筛选方案分别体现出覆盖率、资源利用和风险暴露之间的权衡。方案A、B、C的ODI均值分别为0.4463、0.4552和0.4416，P90分别为0.6407、0.6462和0.6353，ODI>0.70比例分别为0.44%、0.56%和1.22%。本文对低扰动候选的判定依据为ODI均值、P90和风险综合得分的联合判据，而非单一阈值下的超限比例最小；因此，上述结果说明风险场能够转化为可复核的方案级统计量。"
    ),
    145: (
        "4）当前结果主要支撑样例条件下的方法贯通、对象链可传递性和初步方案比选。权重敏感性表明，在当前样例和候选集合范围内，基准权重、±10%相对扰动和上行开采专项权重条件下C方案风险综合得分均低于A、B方案；阈值敏感性表明，当阈值由0.65提高至0.80时，各方案超限比例均下降，但方案间风险差异仍可通过均值、P90和超限比例共同识别。因此，当前结果更适合支撑样例条件下的方法贯通性、对象链可传递性与内部方案比选的可复核性，而不宜直接外推为真实矿井定案条件下的最优工程结论。"
    ),
}


TABLE_UPDATES = [
    (8, 0, 1, "A_π⊂Ω_e；B_b∈[B_b,min,B_b,max]；B_s∈[B_s,min,B_s,max]；W_f∈[W_min,W_max]；A_i∩A_j=∅；L_f∈C_L；E_π∈I_ODI"),
    (9, 0, 1, "S_e(π)=100C_cov(π)-[P_N(π)+10CV_L(π)+5P_short(π)]"),
    (10, 0, 1, "S_r(π)=100[0.45+0.55(0.55R_ton(π)+0.30R_area(π)+0.15S_eng(π))]"),
    (15, 14, 0, "评分口径"),
    (15, 14, 1, "资源回收评分权重"),
    (15, 14, 2, "w_t/w_c/w_e"),
    (15, 14, 3, "0.55/0.30/0.15"),
    (15, 14, 4, "—"),
    (15, 14, 5, "吨煤量/覆盖率/工程组织；用于资源回收分项评分"),
    (15, 14, 6, "planningResults"),
    (15, 15, 0, "候选池"),
    (15, 15, 1, "模式有效候选数"),
    (15, 15, 2, "N_e/N_r/N_m"),
    (15, 15, 3, "2417/1149/374（扰动合格88）"),
    (15, 15, 4, "个"),
    (15, 15, 5, "统一基础输入下工程效率/资源回收/扰动控制偏好筛选后的有效候选记录"),
    (15, 15, 6, "planningResults"),
    (17, 1, 4, "统一ODI场已复核；A按工程效率分项排序"),
    (17, 2, 2, "工作面9个；覆盖率98.52%；资源回收评分89.76"),
    (17, 2, 4, "统一ODI场已复核；B按资源回收分项排序"),
    (17, 3, 1, "统一ODI筛选；联合判据下低扰动候选"),
    (17, 3, 4, "联合判据下低扰动最优"),
    (18, 4, 1, "基准、±10%扰动和上行开采专项权重下，风险综合得分表现为C最低、A次之、B最高"),
]


def set_para_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell_text(cell, text):
    # Preserve the first paragraph style where possible.
    if cell.paragraphs:
        set_para_text(cell.paragraphs[0], text)
        for p in cell.paragraphs[1:]:
            set_para_text(p, "")
    else:
        cell.text = text


def main():
    doc = Document(DOCX)
    changed = 0
    for idx, text in PARA_REPLACEMENTS.items():
        if idx >= len(doc.paragraphs):
            raise IndexError(f"paragraph {idx} missing")
        set_para_text(doc.paragraphs[idx], text)
        changed += 1

    for table_idx, row_idx, col_idx, text in TABLE_UPDATES:
        cell = doc.tables[table_idx].rows[row_idx].cells[col_idx]
        set_cell_text(cell, text)
        changed += 1

    # Add one explicit row for the unified candidate-set notation if it does not already exist.
    table = doc.tables[15]
    existing = "\n".join(cell.text for row in table.rows for cell in row.cells)
    if "Π_all" not in existing:
        row = table.add_row()
        values = [
            "候选池",
            "统一候选集合",
            "Π_all",
            "Π_e∪Π_r∪Π_m",
            "—",
            "统一输入、几何底线约束和ODI场下的模式候选集合并入口径",
            "planningResults",
        ]
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value)
        changed += 1

    doc.save(DOCX)
    print(f"changed={changed}")


if __name__ == "__main__":
    main()
