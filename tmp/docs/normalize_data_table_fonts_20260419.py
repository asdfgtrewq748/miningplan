from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.shared import Pt


def main() -> None:
    docx_path = Path(os.environ["DOCX_PATH"])
    doc = Document(docx_path)
    # Formula tables are 0-14; data tables are 15-18 in the current manuscript.
    for table_idx in range(15, min(19, len(doc.tables))):
        table = doc.tables[table_idx]
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9.0)
    doc.save(docx_path)


if __name__ == "__main__":
    main()
