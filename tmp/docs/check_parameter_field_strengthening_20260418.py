from __future__ import annotations

import sys
import zipfile
from pathlib import Path

from docx import Document


CHECKS = {
    "multi_attribute_extension": "\u5176\u4ed6\u5c5e\u6027\u4f5c\u4e3a\u540c\u4e00\u6d41\u7a0b\u4e0b\u7684\u53ef\u6269\u5c55\u8f93\u5165\u4fdd\u7559",
    "independent_layers_boundary": "\u5c1a\u672a\u5c55\u5f00\u4e3a\u72ec\u7acb\u8bc4\u4ef7\u56fe\u5c42",
    "grid_scope": "80\u00d756\u7f51\u683c\u3001\u51714480\u4e2a\u6805\u683c",
}


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: check_parameter_field_strengthening_20260418.py <docx>")
        return 2

    path = Path(sys.argv[1])
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
    if bad is not None:
        raise RuntimeError(f"Bad DOCX zip member: {bad}")

    doc = Document(str(path))
    text = "\n".join(p.text for p in doc.paragraphs)
    for name, phrase in CHECKS.items():
        print(f"{name}={phrase in text}")
    print(f"paragraphs={len(doc.paragraphs)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
