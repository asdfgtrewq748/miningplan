from docx import Document
import sys
from pathlib import Path

p = Path(sys.argv[1])
doc = Document(p)
repls = {
    '本轮复核的ODI风险场': '本文计算的ODI风险场',
    '单一程序输出结果': '单一计算结果',
    '单一导出样例': '单一方案结果',
    '不能仅凭接口传递关系给出现金流优选': '不能仅凭对象传递关系给出现金流优选',
    '下游接口说明': '下游输入关系说明',
    'designParams/miningRules': '设计参数/规划规则',
    'designParams': '设计参数',
    'miningRules': '规划规则',
    'planningResults': '规划结果',
}
changed = 0
for para in doc.paragraphs:
    txt = para.text
    new = txt
    for old, val in repls.items():
        new = new.replace(old, val)
    if new != txt:
        para.text = new
        changed += 1
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                txt = para.text
                new = txt
                for old, val in repls.items():
                    new = new.replace(old, val)
                if new != txt:
                    para.text = new
                    changed += 1

doc.save(p)
print('changed', changed)
