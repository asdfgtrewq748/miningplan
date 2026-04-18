from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document


FORMULAS = [
    "ODI(x)=w_s D_s(x)+w_a D_a(x)+w_u D_u(x)",
    "w_s+w_a+w_u=1, w_s≥0, w_a≥0, w_u≥0",
    "X_i'(x)=(X_i(x)-X_(i,min))/(X_(i,max)-X_(i,min))",
    "X_i'(x)=(X_(i,max)-X_i(x))/(X_(i,max)-X_(i,min))",
    "E_π=N_π(ODI>T_(ODI))/N_π×100%",
    "Ω_e=Ω_0∖(B_b∪B_s∪D_p)",
    "z(x)=(∑ z_i d_i(x)^(-p))/(∑ d_i(x)^(-p))",
    "π={N,W_f,θ,B_b,B_s,A_π,L_π,R_π,y_π}, y_π∈{0,1}",
    "A_π⊂Ω_e; B_b∈[B_(b,min),B_(b,max)]; B_s∈[B_(s,min),B_(s,max)]; W_f∈[W_(min),W_(max)]; A_i∩A_j=∅; L_f≥L_(min); E_π≤E_(max)",
    "S_e(π)=100[a_1 C_(cov)(π)+a_2 B_(bal)(π)+a_3 B_(conn)(π)+a_4(1-R_(road)(π))]",
    "S_r(π)=100[b_1 R_(area)(π)+b_2 R_(thick)(π)+b_3 R_(rec)(π)]",
    "H_m(π)=c_1 ODI_(mean)(π)+c_2 Q_(0.90)(π)+c_3 E_π(T_(ODI)), S_m(π)=100[1-H_m(π)]",
    "F(π)=λ_e S_e(π)+λ_r S_r(π)+λ_m S_m(π), λ_e+λ_r+λ_m=1",
    "G(s)=αP_s+βR_s+γC_s, α+β+γ=1",
    "NCF_t=Rev_t-Cost_t-RiskCost_t",
]

NUM_RE = re.compile(r"^\uff08(\d+)\uff09$")


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: prepare_formula_linear_omath_text_20260418.py <docx>")
        return 2

    path = Path(sys.argv[1])
    doc = Document(str(path))
    replaced = 0

    for table in doc.tables:
        if len(table.rows) != 1 or len(table.columns) != 3:
            continue
        num_text = table.cell(0, 2).text.strip()
        match = NUM_RE.fullmatch(num_text)
        if not match:
            continue
        num = int(match.group(1))
        if not 1 <= num <= len(FORMULAS):
            continue
        cell = table.cell(0, 1)
        cell.text = FORMULAS[num - 1]
        for paragraph in cell.paragraphs:
            paragraph.alignment = 1
        replaced += 1

    if replaced:
        doc.save(str(path))

    print(f"replaced_formula_text={replaced}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
