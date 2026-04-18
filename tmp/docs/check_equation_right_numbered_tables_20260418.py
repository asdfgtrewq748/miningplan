from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

from docx import Document


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: check_equation_right_numbered_tables_20260418.py <docx>")
        return 2

    path = Path(sys.argv[1])
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
    if bad is not None:
        raise RuntimeError(f"Bad DOCX zip member: {bad}")

    doc = Document(str(path))
    formula_numbers = []
    data_tables = 0
    formula_tables = 0
    for table in doc.tables:
        rows = len(table.rows)
        cols = len(table.columns)
        texts = [cell.text.strip() for row in table.rows for cell in row.cells]
        nums = [re.fullmatch(r"\uff08(\d+)\uff09", t) for t in texts]
        matched_nums = [int(m.group(1)) for m in nums if m]
        if rows == 1 and cols == 3 and matched_nums:
            formula_tables += 1
            formula_numbers.extend(matched_nums)
        else:
            data_tables += 1

    formula_numbers = sorted(formula_numbers)
    print(f"total_tables={len(doc.tables)}")
    print(f"data_tables={data_tables}")
    print(f"formula_tables={formula_tables}")
    print("formula_numbers=" + ",".join(map(str, formula_numbers)))
    print(f"formula_sequence_ok={formula_numbers == list(range(1, 16))}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
