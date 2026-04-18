from __future__ import annotations

import sys
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


def insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.text = text
    return inserted


def find_exact(paragraphs: list[Paragraph], text: str) -> int:
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == text:
            return idx
    raise RuntimeError(f"Paragraph not found: {text[:80]}")


METHOD_ANCHOR = (
    "式中，z(x)为规划位置x处的参数估计值，z_i为第i个钻孔样点的观测值，d_i(x)为位置x与第i个钻孔之间的距离，"
    "p为距离衰减指数。本文采用反距离加权插值，是因为当前样例钻孔数量有限且目标在于形成规划可用的连续约束场，"
    "而非证明复杂地质统计模型的优越性。插值前应统一坐标、剔除明显异常值并限定插值边界；插值后应记录网格尺寸、"
    "样点范围和边界裁剪规则。若后续进入实矿验证阶段，应补充留一交叉验证或与克里金插值结果对比，以量化参数场误差对规划结果的影响。"
)

METHOD_INSERT = (
    "为保证不同图层能够进入同一规划计算，煤层厚度、岩石硬度、瓦斯含量和涌水量等参数原则上应分别插值并统一到相同坐标、"
    "边界裁剪和网格参照下。本文样例中以煤层厚度场作为连续参数场的主要展示对象，其他属性作为同一流程下的可扩展输入保留；"
    "当某一属性样点缺失或异常时，不直接参与该属性场插值，而应记录缺失状态并在方案评分中降级为不确定输入。"
    "本轮复核的ODI风险场采用80×56网格、共4480个栅格，该网格用于样例内部风险统计和方案叠加，不被解释为不同矿区的固定网格尺度。"
)

RESULT_ANCHOR = (
    "采区规划首先需要解决离散地质信息向连续规划约束转换的问题。对于研究区样例而言，15个钻孔样点提供了煤层厚度、岩石硬度、"
    "瓦斯含量和涌水量等离散属性；通过空间插值和统一边界裁剪处理，这些点状信息被转换为可与规划域、ODI风险场和工作面布局叠加分析的连续参数场。"
    "该处理思路与透明工作面、智能精准开采地质模型和采煤工作面三维模型动态修正研究所强调的多源地质信息统一表达具有一致性[30-32]。"
)

RESULT_INSERT = (
    "需要说明的是，本文结果图重点展示煤层厚度场，是因为煤厚直接参与资源覆盖和工作面布置评价，且样点统计口径已经完成复核；"
    "岩石硬度、瓦斯含量和涌水量等属性在当前样例中主要用于说明连续参数场的可扩展输入结构，尚未展开为独立评价图层。"
    "因此，本文不把单一煤厚场结果外推为完整地质建模成果，而将其作为离散钻孔信息进入采区规划链路的示范性参数场。"
)


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: strengthen_parameter_field_20260418.py <docx>")
        return 2

    path = Path(sys.argv[1])
    doc = Document(str(path))
    existing = "\n".join(p.text for p in doc.paragraphs)

    if METHOD_INSERT not in existing:
        insert_after(doc.paragraphs[find_exact(doc.paragraphs, METHOD_ANCHOR)], METHOD_INSERT)
    if RESULT_INSERT not in existing:
        insert_after(doc.paragraphs[find_exact(doc.paragraphs, RESULT_ANCHOR)], RESULT_INSERT)

    doc.save(str(path))
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
    if bad is not None:
        raise RuntimeError(f"Bad DOCX zip member: {bad}")

    print(f"DOCX OK: {path}")
    print("Inserted parameter-field paragraphs: 2")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
