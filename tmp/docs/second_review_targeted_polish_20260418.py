from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


REPLACEMENTS = {
    "多目标优化；候选方案池": "多目标比选；候选方案池",
    "multi-objective optimization; candidate pool": "multi-objective selection; candidate pool",
    "当前已复核的A、B、C方案均采用同一ODI场和4500个采样点形成方案级统计": "当前已复核的A、B、C方案均采用同一ODI场和统一方案区栅格形成方案级统计",
    "研究区工程背景资料指向内蒙古呼伦贝尔敏东一矿": "研究区以内蒙古呼伦贝尔敏东一矿相关工程背景资料为参照",
}


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: second_review_targeted_polish_20260418.py <docx>")
        return 2

    path = Path(sys.argv[1])
    doc = Document(str(path))
    changed = 0

    for paragraph in doc.paragraphs:
        for old, new in REPLACEMENTS.items():
            if old in paragraph.text:
                for run in paragraph.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
                        changed += 1
                if old in paragraph.text:
                    # Fallback for paragraphs split across runs. Preserve paragraph
                    # style, but rebuild run text when the target spans multiple runs.
                    text = paragraph.text.replace(old, new)
                    for run in paragraph.runs:
                        run.text = ""
                    if paragraph.runs:
                        paragraph.runs[0].text = text
                    else:
                        paragraph.add_run(text)
                    changed += 1

    doc.save(str(path))
    print(f"changed={changed}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
