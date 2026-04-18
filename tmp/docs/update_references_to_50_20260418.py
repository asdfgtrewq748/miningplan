from __future__ import annotations

import csv
import sys
import zipfile
from pathlib import Path

from docx import Document


ADDED_REFS = [
    {
        "ref": "[26]\t许家林. 煤矿绿色开采20年研究及进展[J]. 煤炭科学技术, 2020, 48(9): 1-15. DOI: 10.13199/j.cnki.cst.2020.09.001.",
        "authors": "许家林",
        "title": "煤矿绿色开采20年研究及进展",
        "journal": "煤炭科学技术",
        "year": "2020",
        "doi": "10.13199/j.cnki.cst.2020.09.001",
        "source_url": "https://www.mtkxjs.com.cn/fileMTKXJS/journal/article/mtkxjs/html/2020/9/%E7%85%A4%E7%82%AD%E7%A7%91%E5%AD%A6%E6%8A%80%E6%9C%AF202009001.html",
        "note": "期刊网页给出题名、作者、期刊、年卷期页码和 DOI。",
    },
    {
        "ref": "[27]\t王国法, 杜毅博. 煤矿智能化标准体系框架与建设思路[J]. 煤炭科学技术, 2020, 48(1): 1-9. DOI: 10.13199/j.cnki.cst.2020.01.001.",
        "authors": "王国法; 杜毅博",
        "title": "煤矿智能化标准体系框架与建设思路",
        "journal": "煤炭科学技术",
        "year": "2020",
        "doi": "10.13199/j.cnki.cst.2020.01.001",
        "source_url": "https://www.mtkxjs.com.cn/fileMTKXJS/journal/article/mtkxjs/html/2020/1/%E7%85%A4%E7%82%AD%E7%A7%91%E5%AD%A6%E6%8A%80%E6%9C%AF202001001.html",
        "note": "期刊网页给出题名、作者、期刊、年卷期页码和 DOI。",
    },
    {
        "ref": "[28]\t王国法, 任怀伟, 庞义辉, 等. 煤矿智能化(初级阶段)技术体系研究与工程进展[J]. 煤炭科学技术, 2020, 48(7): 1-27. DOI: 10.13199/j.cnki.cst.2020.07.001.",
        "authors": "王国法; 任怀伟; 庞义辉; 等",
        "title": "煤矿智能化(初级阶段)技术体系研究与工程进展",
        "journal": "煤炭科学技术",
        "year": "2020",
        "doi": "10.13199/j.cnki.cst.2020.07.001",
        "source_url": "https://www.mtkxjs.com.cn/fileMTKXJS/journal/article/mtkxjs/html/2020/7/%E7%85%A4%E7%82%AD%E7%A7%91%E5%AD%A6%E6%8A%80%E6%9C%AF202007001.html",
        "note": "期刊网页给出题名、作者、期刊、年卷期页码和 DOI。",
    },
    {
        "ref": "[29]\t李首滨. 智能化开采研究进展与发展趋势[J]. 煤炭科学技术, 2019, 47(10): 102-110. DOI: 10.13199/j.cnki.cst.2019.10.012.",
        "authors": "李首滨",
        "title": "智能化开采研究进展与发展趋势",
        "journal": "煤炭科学技术",
        "year": "2019",
        "doi": "10.13199/j.cnki.cst.2019.10.012",
        "source_url": "https://www.mtkxjs.com.cn/fileMTKXJS/journal/article/mtkxjs/html/2019/10/%E7%85%A4%E7%82%AD%E7%A7%91%E5%AD%A6%E6%8A%80%E6%9C%AF201910012.html",
        "note": "期刊网页给出题名、作者、期刊、年卷期页码和 DOI。",
    },
    {
        "ref": "[30]\t王存飞, 荣耀. 透明工作面的概念、架构与关键技术[J]. 煤炭科学技术, 2019, 47(7): 156-163. DOI: 10.13199/j.cnki.cst.2019.07.019.",
        "authors": "王存飞; 荣耀",
        "title": "透明工作面的概念、架构与关键技术",
        "journal": "煤炭科学技术",
        "year": "2019",
        "doi": "10.13199/j.cnki.cst.2019.07.019",
        "source_url": "https://www.mtkxjs.com.cn/fileMTKXJS/journal/article/mtkxjs/html/2019/7/%E7%85%A4%E7%82%AD%E7%A7%91%E5%AD%A6%E6%8A%80%E6%9C%AF201907019.html",
        "note": "期刊网页给出题名、作者、期刊、年卷期页码和 DOI。",
    },
    {
        "ref": "[31]\t程建远, 朱梦博, 王云宏, 等. 煤炭智能精准开采工作面地质模型梯级构建及其关键技术[J]. 煤炭学报, 2019, 44(8): 2285-2295. DOI: 10.13225/j.cnki.jccs.KJ19.0510.",
        "authors": "程建远; 朱梦博; 王云宏; 等",
        "title": "煤炭智能精准开采工作面地质模型梯级构建及其关键技术",
        "journal": "煤炭学报",
        "year": "2019",
        "doi": "10.13225/j.cnki.jccs.KJ19.0510",
        "source_url": "https://www.mtxb.com.cn/fileMTXB/journal/article/mtxb/html/2019/8/20190916%E7%85%A4%E7%82%AD%E5%AD%A6%E6%8A%A5201908002.html",
        "note": "期刊中文引用行使用 KJ19.0510；同页英文行疑有 KJ18.0510 录入差异，正文采用中文引用行 DOI。",
    },
    {
        "ref": "[32]\t刘万里, 张学亮, 王世博. 采煤工作面煤层三维模型构建及动态修正技术[J]. 煤炭学报, 2020, 45(6): 1973-1983. DOI: 10.13225/j.cnki.jccs.ZN20.0364.",
        "authors": "刘万里; 张学亮; 王世博",
        "title": "采煤工作面煤层三维模型构建及动态修正技术",
        "journal": "煤炭学报",
        "year": "2020",
        "doi": "10.13225/j.cnki.jccs.ZN20.0364",
        "source_url": "https://mtxb.com.cn/fileMTXB/journal/article/mtxb/html/2020/6/%E7%85%A4%E7%82%AD%E5%AD%A6%E6%8A%A5202006006.html",
        "note": "期刊网页给出题名、作者、期刊、年卷期页码和 DOI。",
    },
    {
        "ref": "[33]\t徐智敏, 孙亚军, 高尚, 等. 干旱矿区采动顶板导水裂隙的演化规律及保水采煤意义[J]. 煤炭学报, 2019, 44(3): 767-776. DOI: 10.13225/j.cnki.jccs.2018.6041.",
        "authors": "徐智敏; 孙亚军; 高尚; 等",
        "title": "干旱矿区采动顶板导水裂隙的演化规律及保水采煤意义",
        "journal": "煤炭学报",
        "year": "2019",
        "doi": "10.13225/j.cnki.jccs.2018.6041",
        "source_url": "https://www.mtxb.com.cn/fileMTXB/journal/article/mtxb/html/2019/3/20190412%E7%85%A4%E7%82%AD%E5%AD%A6%E6%8A%A5201903012.html",
        "note": "期刊网页给出题名、作者、期刊、年卷期页码和 DOI。",
    },
    {
        "ref": "[34]\t曹志国, 鞠金峰, 许家林. 采动覆岩导水裂隙主通道分布模型及其水流动特性[J]. 煤炭学报, 2019, 44(12): 3719-3728. DOI: 10.13225/j.cnki.jccs.SH19.0446.",
        "authors": "曹志国; 鞠金峰; 许家林",
        "title": "采动覆岩导水裂隙主通道分布模型及其水流动特性",
        "journal": "煤炭学报",
        "year": "2019",
        "doi": "10.13225/j.cnki.jccs.SH19.0446",
        "source_url": "https://www.mtxb.com.cn/fileMTXB/journal/article/mtxb/html/2019/12/%E7%85%A4%E7%82%AD%E5%AD%A6%E6%8A%A5201912013.html",
        "note": "期刊网页给出题名、作者、期刊、年卷期页码和 DOI。",
    },
    {
        "ref": "[35]\t郭文兵, 娄高中. 覆岩破坏充分采动程度定义及判别方法[J]. 煤炭学报, 2019, 44(3): 755-766. DOI: 10.13225/j.cnki.jccs.2018.6038.",
        "authors": "郭文兵; 娄高中",
        "title": "覆岩破坏充分采动程度定义及判别方法",
        "journal": "煤炭学报",
        "year": "2019",
        "doi": "10.13225/j.cnki.jccs.2018.6038",
        "source_url": "https://www.mtxb.com.cn/fileMTXB/journal/article/mtxb/html/2019/3/20190412%E7%85%A4%E7%82%AD%E5%AD%A6%E6%8A%A5201903011.html",
        "note": "期刊网页给出题名、作者、期刊、年卷期页码和 DOI。",
    },
    {
        "ref": "[36]\t鞠金峰, 马祥, 赵富强, 等. 东胜煤田导水裂缝发育及其分区特征研究[J]. 煤炭科学技术, 2022, 50(2): 202-212. DOI: 10.13199/j.cnki.cst.2021-0169.",
        "authors": "鞠金峰; 马祥; 赵富强; 等",
        "title": "东胜煤田导水裂缝发育及其分区特征研究",
        "journal": "煤炭科学技术",
        "year": "2022",
        "doi": "10.13199/j.cnki.cst.2021-0169",
        "source_url": "https://www.mtkxjs.com.cn/fileMTKXJS/journal/article/mtkxjs/html/2022/2/20220326MTKJ202202022.html",
        "note": "期刊网页给出题名、作者、期刊、年卷期页码和 DOI。",
    },
    {
        "ref": "[37]\t郭小铭, 王皓, 周麟晟. 煤层顶板巨厚基岩含水层空间富水性评价[J]. 煤炭科学技术, 2021, 49(9): 167-175. DOI: 10.13199/j.cnki.cst.2021.09.024.",
        "authors": "郭小铭; 王皓; 周麟晟",
        "title": "煤层顶板巨厚基岩含水层空间富水性评价",
        "journal": "煤炭科学技术",
        "year": "2021",
        "doi": "10.13199/j.cnki.cst.2021.09.024",
        "source_url": "https://www.mtkxjs.com.cn/fileMTKXJS/journal/article/mtkxjs/html/2021/9/%E7%85%A4%E7%82%AD%E7%A7%91%E5%AD%A6%E6%8A%80%E6%9C%AF202109024.html",
        "note": "期刊网页给出题名、作者、期刊、年卷期页码和 DOI。",
    },
    {
        "ref": "[38]\t王苏健, 金声尧. 类孤岛工作面复采方案优选分析研究[J]. 煤炭科学技术, 2021, 49(9): 9-16. DOI: 10.13199/j.cnki.cst.2021.09.002.",
        "authors": "王苏健; 金声尧",
        "title": "类孤岛工作面复采方案优选分析研究",
        "journal": "煤炭科学技术",
        "year": "2021",
        "doi": "10.13199/j.cnki.cst.2021.09.002",
        "source_url": "https://www.mtkxjs.com.cn/fileMTKXJS/journal/article/mtkxjs/html/2021/9/%E7%85%A4%E7%82%AD%E7%A7%91%E5%AD%A6%E6%8A%80%E6%9C%AF202109002.html",
        "note": "期刊网页给出题名、作者、期刊、年卷期页码和 DOI。",
    },
    {
        "ref": "[39]\t王玉涛, 刘震. 深部煤层非充分采动下覆岩裂隙场可视化探测研究[J]. 煤炭科学技术, 2020, 48(3): 197-204. DOI: 10.13199/j.cnki.cst.2020.03.024.",
        "authors": "王玉涛; 刘震",
        "title": "深部煤层非充分采动下覆岩裂隙场可视化探测研究",
        "journal": "煤炭科学技术",
        "year": "2020",
        "doi": "10.13199/j.cnki.cst.2020.03.024",
        "source_url": "https://www.mtkxjs.com.cn/fileMTKXJS/journal/article/mtkxjs/html/2020/3/%E7%85%A4%E7%82%AD%E7%A7%91%E5%AD%A6%E6%8A%80%E6%9C%AF202003024.html",
        "note": "期刊网页给出题名、作者、期刊、年卷期页码和 DOI。",
    },
    {
        "ref": "[40]\t董江鑫, 王飞. 地质雷达和高密度电法联合探测底板含水性的应用[J]. 煤炭科学技术, 2022, 50(5): 222-231. DOI: 10.13199/j.cnki.cst.2020-0418.",
        "authors": "董江鑫; 王飞",
        "title": "地质雷达和高密度电法联合探测底板含水性的应用",
        "journal": "煤炭科学技术",
        "year": "2022",
        "doi": "10.13199/j.cnki.cst.2020-0418",
        "source_url": "https://www.mtkxjs.com.cn/fileMTKXJS/journal/article/mtkxjs/html/2022/5/20220722MTKJ202205027.html",
        "note": "期刊网页给出题名、作者、期刊、年卷期页码和 DOI。",
    },
    {
        "ref": "[41]\t张玉军, 宋业杰, 樊振丽, 等. 鄂尔多斯盆地侏罗系煤田保水开采技术与应用[J]. 煤炭科学技术, 2021, 49(4): 159-168. DOI: 10.13199/j.cnki.cst.2021.04.019.",
        "authors": "张玉军; 宋业杰; 樊振丽; 等",
        "title": "鄂尔多斯盆地侏罗系煤田保水开采技术与应用",
        "journal": "煤炭科学技术",
        "year": "2021",
        "doi": "10.13199/j.cnki.cst.2021.04.019",
        "source_url": "https://mtkxjs.com.cn/fileMTKXJS/journal/article/mtkxjs/html/2021/4/%E7%85%A4%E7%82%AD%E7%A7%91%E5%AD%A6%E6%8A%80%E6%9C%AF202104019.html",
        "note": "期刊网页和 PDF 检索结果给出题名、期刊、年卷期页码和 DOI。",
    },
    {
        "ref": "[42]\t李全生, 鞠金峰, 曹志国, 等. 基于导水裂隙带高度的地下水库适应性评价[J]. 煤炭学报, 2017, 42(8): 2116-2124. DOI: 10.13225/j.cnki.jccs.2016.1871.",
        "authors": "李全生; 鞠金峰; 曹志国; 等",
        "title": "基于导水裂隙带高度的地下水库适应性评价",
        "journal": "煤炭学报",
        "year": "2017",
        "doi": "10.13225/j.cnki.jccs.2016.1871",
        "source_url": "https://www.mtxb.com.cn/cn/article/pdf/preview/10.13225/j.cnki.jccs.2016.1871.pdf",
        "note": "期刊 PDF 给出题名、作者、期刊、年卷期页码和 DOI。",
    },
    {
        "ref": "[43]\t鞠金峰, 许家林, 朱卫兵. 西部缺水矿区地下水库保水的库容研究[J]. 煤炭学报, 2017, 42(2): 381-387. DOI: 10.13225/j.cnki.jccs.2016.6016.",
        "authors": "鞠金峰; 许家林; 朱卫兵",
        "title": "西部缺水矿区地下水库保水的库容研究",
        "journal": "煤炭学报",
        "year": "2017",
        "doi": "10.13225/j.cnki.jccs.2016.6016",
        "source_url": "https://mtxb.com.cn/cn/article/pdf/preview/10.13225/j.cnki.jccs.2016.6016.pdf",
        "note": "期刊 PDF 给出题名、作者、期刊、年卷期页码和 DOI。",
    },
    {
        "ref": "[44]\t徐智敏, 高尚, 孙亚军, 等. 西部典型侏罗系富煤区含水介质条件与水动力学特征[J]. 煤炭学报, 2017, 42(2): 444-451. DOI: 10.13225/j.cnki.jccs.2016.6024.",
        "authors": "徐智敏; 高尚; 孙亚军; 等",
        "title": "西部典型侏罗系富煤区含水介质条件与水动力学特征",
        "journal": "煤炭学报",
        "year": "2017",
        "doi": "10.13225/j.cnki.jccs.2016.6024",
        "source_url": "https://www.mtxb.com.cn/cn/article/pdf/preview/10.13225/j.cnki.jccs.2016.6024.pdf",
        "note": "期刊 PDF 给出题名、作者、期刊、年卷期页码和 DOI。",
    },
    {
        "ref": "[45]\t董书宁, 姬亚东, 王皓, 等. 鄂尔多斯盆地侏罗纪煤田典型顶板水害防控技术与应用[J]. 煤炭学报, 2020, 45(7): 2367-2375. DOI: 10.13225/j.cnki.jccs.DZ20.0697.",
        "authors": "董书宁; 姬亚东; 王皓; 等",
        "title": "鄂尔多斯盆地侏罗纪煤田典型顶板水害防控技术与应用",
        "journal": "煤炭学报",
        "year": "2020",
        "doi": "10.13225/j.cnki.jccs.DZ20.0697",
        "source_url": "https://www.mtxb.com.cn/fileMTXB/journal/article/mtxb/html/2020/7/435108.html",
        "note": "期刊网页给出题名、作者、期刊、年卷期页码和 DOI。",
    },
    {
        "ref": "[46]\t刘晓丽, 曹志国, 陈苏社, 等. 煤矿分布式地下水库渗流场分析及优化调度[J]. 煤炭学报, 2019, 44(12): 3693-3699. DOI: 10.13225/j.cnki.jccs.SH19.1167.",
        "authors": "刘晓丽; 曹志国; 陈苏社; 等",
        "title": "煤矿分布式地下水库渗流场分析及优化调度",
        "journal": "煤炭学报",
        "year": "2019",
        "doi": "10.13225/j.cnki.jccs.SH19.1167",
        "source_url": "https://www.mtxb.com.cn/fileMTXB/journal/article/mtxb/html/2019/12/%E7%85%A4%E7%82%AD%E5%AD%A6%E6%8A%A5201912010.html",
        "note": "期刊网页给出题名、作者、期刊、年卷期页码和 DOI。",
    },
    {
        "ref": "[47]\t庞义辉, 李全生, 曹光明, 等. 煤矿地下水库储水空间构成分析及计算方法[J]. 煤炭学报, 2019, 44(2): 557-566. DOI: 10.13225/j.cnki.jccs.2018.0417.",
        "authors": "庞义辉; 李全生; 曹光明; 等",
        "title": "煤矿地下水库储水空间构成分析及计算方法",
        "journal": "煤炭学报",
        "year": "2019",
        "doi": "10.13225/j.cnki.jccs.2018.0417",
        "source_url": "https://www.mtxb.com.cn/fileMTXB/journal/article/mtxb/html/2019/2/20190330%E7%85%A4%E7%82%AD%E5%AD%A6%E6%8A%A5201902023.html",
        "note": "期刊网页给出题名、作者、期刊、年卷期页码和 DOI。",
    },
    {
        "ref": "[48]\t范立民, 马雄德, 蒋泽泉, 等. 保水采煤研究30年回顾与展望[J]. 煤炭科学技术, 2019, 47(7): 1-30. DOI: 10.13199/j.cnki.cst.2019.07.001.",
        "authors": "范立民; 马雄德; 蒋泽泉; 等",
        "title": "保水采煤研究30年回顾与展望",
        "journal": "煤炭科学技术",
        "year": "2019",
        "doi": "10.13199/j.cnki.cst.2019.07.001",
        "source_url": "https://www.mtkxjs.com.cn/fileMTKXJS/journal/article/mtkxjs/html/2019/7/%E7%85%A4%E7%82%AD%E7%A7%91%E5%AD%A6%E6%8A%80%E6%9C%AF201907001.html",
        "note": "期刊网页给出题名、作者、期刊、年卷期页码和 DOI。",
    },
    {
        "ref": "[49]\t许家林, 朱卫兵, 王晓振. 基于关键层位置的导水裂隙带高度预计方法[J]. 煤炭学报, 2012, 37(5): 762-769. DOI: 10.13225/j.cnki.jccs.2012.05.002.",
        "authors": "许家林; 朱卫兵; 王晓振",
        "title": "基于关键层位置的导水裂隙带高度预计方法",
        "journal": "煤炭学报",
        "year": "2012",
        "doi": "10.13225/j.cnki.jccs.2012.05.002",
        "source_url": "https://www.mtxb.com.cn/cn/article/pdf/preview/9f920284-c6d1-4852-b8a9-93bbaf288ba8.pdf",
        "note": "期刊 PDF/检索结果给出题名、作者、期刊、年卷期页码和 DOI。",
    },
    {
        "ref": "[50]\t郭小铭, 董书宁. 深埋煤层开采顶板基岩含水层渗流规律及保水技术[J]. 煤炭学报, 2019, 44(3): 804-811. DOI: 10.13225/j.cnki.jccs.2018.6025.",
        "authors": "郭小铭; 董书宁",
        "title": "深埋煤层开采顶板基岩含水层渗流规律及保水技术",
        "journal": "煤炭学报",
        "year": "2019",
        "doi": "10.13225/j.cnki.jccs.2018.6025",
        "source_url": "https://www.mtxb.com.cn/fileMTXB/journal/article/mtxb/html/2019/3/20190412%E7%85%A4%E7%82%AD%E5%AD%A6%E6%8A%A5201903016.html",
        "note": "期刊网页给出题名、作者、期刊、年卷期页码和 DOI。",
    },
]


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: update_references_to_50_20260418.py <docx>")
        return 2

    docx_path = Path(sys.argv[1])
    if not docx_path.exists():
        raise FileNotFoundError(docx_path)

    doc = Document(str(docx_path))
    ref_heading_idx = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "参考文献":
            ref_heading_idx = idx
            break
    if ref_heading_idx is None:
        raise RuntimeError("未找到参考文献标题")

    existing_refs = []
    for paragraph in doc.paragraphs[ref_heading_idx + 1 :]:
        text = paragraph.text.strip()
        if text.startswith("["):
            existing_refs.append(text)
    existing_refs = existing_refs[:25]
    if len(existing_refs) != 25:
        raise RuntimeError(f"现有参考文献不是 25 条: {len(existing_refs)}")

    for paragraph in list(doc.paragraphs[ref_heading_idx + 1 :]):
        delete_paragraph(paragraph)

    for ref_text in existing_refs:
        doc.add_paragraph(ref_text)
    for item in ADDED_REFS:
        doc.add_paragraph(item["ref"])

    doc.save(str(docx_path))

    with zipfile.ZipFile(docx_path) as zf:
        bad = zf.testzip()
    if bad is not None:
        raise RuntimeError(f"DOCX zip check failed at {bad}")

    out_dir = Path("docs/plans")
    out_dir.mkdir(parents=True, exist_ok=True)
    csv_path = out_dir / "coal_sci_reference_verification_added_chinese_20260418.csv"
    with csv_path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=[
                "ref_no",
                "doi",
                "status",
                "source_type",
                "source_url",
                "title",
                "authors",
                "journal",
                "year",
                "note",
            ],
        )
        writer.writeheader()
        for i, item in enumerate(ADDED_REFS, start=26):
            writer.writerow(
                {
                    "ref_no": i,
                    "doi": item["doi"],
                    "status": "VERIFIED_BY_JOURNAL_PAGE",
                    "source_type": "journal_html_or_pdf",
                    "source_url": item["source_url"],
                    "title": item["title"],
                    "authors": item["authors"],
                    "journal": item["journal"],
                    "year": item["year"],
                    "note": item["note"],
                }
            )

    md_path = out_dir / "coal_sci_reference_verification_summary_20260418.md"
    md_lines = [
        "# 参考文献联网核验与扩充记录（2026-04-18）",
        "",
        "## 核验结论",
        "",
        "- 原有 25 条英文/国际文献已通过 Crossref DOI API 核验，结果见 `coal_sci_reference_verification_existing_20260418.csv`。",
        "- 新增 25 条中文文献已通过期刊网页或期刊 PDF 的题名、作者、期刊、年份、页码和 DOI 字段核验，结果见 `coal_sci_reference_verification_added_chinese_20260418.csv`。",
        "- 工作稿参考文献总数已调整为 50 条。",
        "",
        "## 新增中文文献",
        "",
    ]
    for item in ADDED_REFS:
        md_lines.append(f"- {item['ref']}")
    md_path.write_text("\n".join(md_lines) + "\n", encoding="utf-8")

    print(docx_path)
    print(csv_path)
    print(md_path)
    print("existing_refs", len(existing_refs))
    print("added_refs", len(ADDED_REFS))
    print("total_refs", len(existing_refs) + len(ADDED_REFS))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
