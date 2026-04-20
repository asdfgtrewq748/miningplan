from __future__ import annotations

import shutil
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX=Path(sys.argv[1])
BACKUP=Path(sys.argv[2])
shutil.copy2(DOCX, BACKUP)

doc=Document(DOCX)

# Replace equation (6) text in its formula table. Word COM will convert this linear formula to OMML later.
formula_text = 'Ω_e=Conn_max{[Ω_0⊖𝔅(B_b)]⊖𝔅(B_s)⊖𝔅(D_p)}'
for table in doc.tables:
    if len(table.rows)==1 and len(table.columns)==3 and table.cell(0,2).text.strip() == '（6）':
        cell=table.cell(0,1)
        cell.text = formula_text
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = 'Cambria Math'
                run.font.size = Pt(10.5)
        break
else:
    raise RuntimeError('Equation (6) table not found')

old = '式中，Ω_0为原始采区边界，B_b为边界煤柱宽度，B_s为区段煤柱宽度，D_p为局部保护距离，Ω_e为经约束内缩和几何合法性处理后的有效布置域。若内缩后出现多连通域或局部狭长畸变，则保留主连通区域并采用降级内缩策略，以保证后续工作面布置具有几何可解性。'
new = '式中，Ω_0为原始采区边界，B_b为边界煤柱宽度，B_s为区段煤柱宽度，D_p为局部保护距离，𝔅(·)为按给定距离生成的缓冲内缩算子，Conn_max(·)表示保留面积最大的主连通区域，Ω_e为经约束内缩和几何合法性处理后的有效布置域。若内缩后出现多连通域或局部狭长畸变，则优先保留主连通区域，并在不突破安全煤柱底线的前提下采用降级内缩策略，以保证后续工作面布置具有几何可解性。'
changed=False
for para in doc.paragraphs:
    if para.text.strip() == old:
        para.text = new
        para.style = next((p.style for p in doc.paragraphs if p.style.name == '10正文'), para.style)
        for run in para.runs:
            run.font.size = Pt(10.5)
        changed=True
        break
if not changed:
    for para in doc.paragraphs:
        if para.text.strip().startswith('式中，Ω_0为原始采区边界'):
            para.text = new
            for run in para.runs:
                run.font.size = Pt(10.5)
            changed=True
            break
if not changed:
    raise RuntimeError('Explanation paragraph not found')

doc.save(DOCX)
print('patched')
