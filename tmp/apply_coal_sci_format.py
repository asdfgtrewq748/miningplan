from __future__ import annotations

import os
import re
import shutil
import zipfile
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"D:/xiangmu/miningplan")
SRC = ROOT / "论文/重构工作区/06_投稿包/最新版论文4.16_插图版_第四轮审稿优化版.docx"
OUT = ROOT / "论文/重构工作区/06_投稿包/最新版论文4.16_插图版_煤科模板格式版.docx"
REPORT = ROOT / "论文/重构工作区/00_过程文档/煤炭科学技术模板格式化记录.md"


CN_TITLE = "覆岩扰动约束下采区协同规划方法"
EN_TITLE = (
    "Collaborative planning method for mining districts constrained by "
    "overburden disturbance"
)
CN_ABSTRACT = (
    "摘要：针对采区规划中离散地质信息难以连续参与空间布置、多场景覆岩扰动风险约束口径不统一、"
    "规划结果与采掘接续及经济评价衔接不足等问题，提出一种覆岩扰动约束下的采区多目标协同规划方法。"
    "该方法以采区边界、钻孔样点和设计参数为基础，首先构建有效布置域和连续参数场，将离散煤层厚度、"
    "瓦斯含量、涌水量等输入转换为可参与规划求解的空间约束；随后引入overburden disturbance index"
    "（覆岩扰动指数，ODI），将地表沉陷、含水层扰动和上行开采等异构风险统一组织为同尺度指标，"
    "并以ODI均值、高分位值和超限暴露比例描述方案级扰动特征；在此基础上，建立兼顾工程效率、资源覆盖"
    "和扰动控制的候选方案评价口径，形成工作面、巷道、推进长度、布置面积和ODI统计量等可复核对象，"
    "并将规划结果继续传递至采掘接续和工程经济评价入口。以样例研究区进行验证，15个钻孔样点的煤层厚度"
    "为2.8～4.8 m，平均值为3.7933 m；在255250.00 m²原始边界范围内，方法生成3个工作面和11条巷道，"
    "布置面积为176565.72 m²，有效覆盖率为69.17%，巷道总长度为4817.50 m，平均规划评分为68.8。"
    "综合ODI场采用80×56网格，均值为0.4669，P90为0.7474，ODI>0.70的栅格比例为15.89%。"
    "结果表明，该方法能够在统一空间参照和统一风险口径下完成参数场构建、风险前置约束、规划对象生成"
    "和后续评价输入组织，可为采区规划中的多目标权衡和结果传递提供可复核的方法框架；当前结果主要用于"
    "验证方法链贯通能力，工程定案仍需结合实矿尺度、参数标定、接续排程和经济评价数据进一步校核。"
)
EN_ABSTRACT = (
    "Abstract: To address the difficulty of incorporating discrete geological information "
    "into continuous spatial layout, the inconsistent expression of multi-scenario overburden "
    "disturbance constraints, and the weak connection between planning results, mining succession "
    "and economic evaluation, a multi-objective collaborative planning method constrained by "
    "overburden disturbance is proposed for mining districts. Boundary data, borehole samples and "
    "design parameters are used to construct an effective layout domain and continuous parameter "
    "fields, so that discrete inputs such as coal thickness, gas content and water inflow can be "
    "converted into spatial constraints for planning. An overburden disturbance index (ODI) is then "
    "introduced to organize surface subsidence, aquifer disturbance and upward-mining-related risks "
    "into a unified indicator. The mean ODI value, high-percentile value and exceedance ratio are "
    "used to describe scheme-level disturbance characteristics. A candidate-scheme evaluation "
    "framework considering engineering efficiency, resource coverage and disturbance control is "
    "established, and planning objects including working faces, roadways, advancing lengths, layout "
    "areas and ODI statistics are generated and transferred to mining-succession and engineering "
    "economic evaluation inputs. In the sample study area, the coal thickness of 15 boreholes ranges "
    "from 2.8 to 4.8 m with an average of 3.7933 m. Within an original boundary area of 255250.00 m², "
    "3 working faces and 11 roadways are generated, with a layout area of 176565.72 m², an effective "
    "coverage ratio of 69.17%, a total roadway length of 4817.50 m and an average planning score of "
    "68.8. The integrated ODI field contains an 80×56 grid, with a mean value of 0.4669, a P90 value "
    "of 0.7474 and 15.89% of grids exceeding 0.70. The results show that parameter-field construction, "
    "forward risk constraint, planning-object generation and downstream evaluation input organization "
    "can be completed under a unified spatial and risk-evaluation framework. The current results "
    "mainly verify the connectivity of the method chain, and engineering implementation still requires "
    "further checks using real-mine scale data, parameter calibration, succession scheduling and "
    "economic evaluation data."
)

CN_KEYWORDS = "关键词：覆岩扰动指数；采区规划；多目标协同规划；风险前置约束；连续参数场；采掘接续；工程经济评价"
EN_KEYWORDS = (
    "Keywords: overburden disturbance index; mining district planning; multi-objective collaborative "
    "planning; risk pre-constraint; continuous parameter field; mining succession; engineering economic evaluation"
)
CN_UNIT = "（1. 中国矿业大学（北京） 能源与矿业学院，北京 100083）"
EN_UNIT = (
    "(1. School of Energy and Mining Engineering, China University of Mining and Technology-Beijing, "
    "Beijing 100083, China)"
)
FOOTNOTE = (
    "收稿日期：    修回日期：    责任编辑：\n"
    "基金项目：无。\n"
    "作者简介：李杨（1982—），男，河北唐山人，教授，博士，研究方向为矿业工程。\n"
    "通信作者：刘浩天（2001—），男，河北唐山人，博士研究生，研究方向为矿业工程。"
    "Tel：17692508100，E-mail：gqt2500103011@student.cumtb.edu.cn。"
)

FIG_CAPTIONS = {
    26: ("图1 覆岩扰动约束机理示意图", "Fig.1 Schematic diagram of the overburden disturbance constraint mechanism"),
    74: ("图2 研究区钻孔空间分布图", "Fig.2 Spatial distribution of boreholes in the study area"),
    83: ("图3 研究区主要地质参数场分布图", "Fig.3 Distribution of main geological parameter fields in the study area"),
    93: ("图4 多场景ODI分布结果对比图", "Fig.4 Comparison of multi-scenario ODI distribution results"),
    103: ("图5 含水层扰动约束下最终规划布局与ODI场叠置图", "Fig.5 Final planning layout overlaid with the ODI field under aquifer disturbance constraint"),
    107: ("图6 采区规划布局结果图", "Fig.6 Planning layout result of the mining district"),
    114: ("图7 规划结果向采掘接续传递示意图", "Fig.7 Schematic transfer from planning results to mining succession"),
    117: ("图8 规划对象与评价参数联动示意图", "Fig.8 Linkage between planning objects and evaluation parameters"),
    120: ("图9 工程经济评价输入口径示意图", "Fig.9 Input framework for engineering economic evaluation"),
}

TABLE_CAPTIONS = {
    78: ("表1 研究区输入参数与约束条件", "Table 1 Input parameters and constraints of the study area"),
    97: ("表2 多场景ODI风险组织关系", "Table 2 Risk organization relationships under multi-scenario ODI"),
    9701: ("表2 多场景 ODI 风险组织关系", "Table 2 Risk organization relationships under multi-scenario ODI"),
    108: ("表3 当前样例规划结果统计表", "Table 3 Statistics of current sample planning results"),
    122: ("表4 当前样例参数与ODI统计表", "Table 4 Parameters and ODI statistics of the current sample"),
}

FORMULA_RE = re.compile(r"（(?:[1-9]|1[01])）")
CITE_RE = re.compile(r"\[(?:\d+(?:-\d+)?)(?:,\s*\d+(?:-\d+)?)*\]")
REF_RE = re.compile(r"^\[(\d+)\]\s+")


def ensure_style(doc: Document, name: str, size: float, *, bold=False, align=None, first=0.0, left=0.0, line=1.0):
    try:
        st = doc.styles[name]
    except KeyError:
        st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    st.font.name = "宋体"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    st.font.size = Pt(size)
    st.font.bold = bold
    pf = st.paragraph_format
    if align is not None:
        pf.alignment = align
    pf.first_line_indent = Cm(first) if first else None
    pf.left_indent = Cm(left) if left else None
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = line
    return st


def setup_styles(doc: Document) -> None:
    ensure_style(doc, "01 中文题名", 16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    ensure_style(doc, "02 英文题名", 12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    ensure_style(doc, "03 中文作者", 10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    ensure_style(doc, "04 中文单位", 9, align=WD_ALIGN_PARAGRAPH.CENTER)
    ensure_style(doc, "05 中文摘要关键词", 9, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    ensure_style(doc, "06 中图分类号", 9, align=WD_ALIGN_PARAGRAPH.LEFT)
    ensure_style(doc, "07 英文作者", 10, align=WD_ALIGN_PARAGRAPH.CENTER)
    ensure_style(doc, "08 英文单位", 9, align=WD_ALIGN_PARAGRAPH.CENTER)
    ensure_style(doc, "09 英文摘要关键词", 9, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    ensure_style(doc, "10 正文", 9, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first=0.7)
    ensure_style(doc, "11 一级标题", 10.5, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    ensure_style(doc, "12 二级标题", 10, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    ensure_style(doc, "13 三级标题", 9, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    ensure_style(doc, "14 图题", 8, align=WD_ALIGN_PARAGRAPH.CENTER)
    ensure_style(doc, "15 表题", 8, align=WD_ALIGN_PARAGRAPH.CENTER)
    ensure_style(doc, "16 公式编号", 9, align=WD_ALIGN_PARAGRAPH.CENTER)
    ensure_style(doc, "17 参考文献", 8, align=WD_ALIGN_PARAGRAPH.JUSTIFY, left=0.0)
    ensure_style(doc, "18 首页脚注", 7.5, align=WD_ALIGN_PARAGRAPH.LEFT)


def clear_para(p):
    # Preserve paragraph properties (<w:pPr>) so clearing text does not drop the
    # assigned Word style, alignment, tabs, or section break.
    for child in list(p._p):
        if child.tag != qn("w:pPr"):
            p._p.remove(child)


def set_text(p, text: str, style: str | None = None, align=None):
    clear_para(p)
    parts = text.split("\n")
    for i, part in enumerate(parts):
        if i:
            p.add_run().add_break()
        p.add_run(part)
    if style:
        p.style = style
    if align is not None:
        p.alignment = align


def set_author_line(p, names: Iterable[str], style: str):
    clear_para(p)
    for i, name in enumerate(names):
        if i:
            p.add_run("，")
        p.add_run(name)
        r = p.add_run("1")
        r.font.superscript = True
    p.style = style
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def delete_paragraph(p):
    parent = p._element.getparent()
    parent.remove(p._element)


def set_columns(section, num: int) -> None:
    sectPr = section._sectPr
    cols = sectPr.xpath("./w:cols")
    if cols:
        cols = cols[0]
    else:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    if num == 1:
        cols.attrib.pop(qn("w:num"), None)
    else:
        cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), "425")


def set_cell_borders(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        tag = "w:{}".format(edge)
        element = tcBorders.find(qn(tag))
        if edge_data is None:
            if element is not None:
                tcBorders.remove(element)
            continue
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def format_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_i, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                for run in p.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(7.5)
            border = {}
            if r_i == 0:
                border["top"] = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
                border["bottom"] = {"val": "single", "sz": "6", "space": "0", "color": "000000"}
            if r_i == len(table.rows) - 1:
                border["bottom"] = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
            set_cell_borders(cell, **border)


def superscript_citations(p):
    text = p.text
    if not CITE_RE.search(text):
        return
    if text.startswith("["):
        return
    if not p.runs:
        return
    style_name = p.style.name
    align = p.alignment
    clear_para(p)
    pos = 0
    for m in CITE_RE.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        r = p.add_run(m.group(0))
        r.font.superscript = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])
    p.style = style_name
    p.alignment = align


def normalize_table_headers(doc: Document) -> None:
    replacements = {"数值/范围": "数值或范围", "单位/口径": "单位或口径"}
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for src, dst in replacements.items():
                    if src in cell.text:
                        for p in cell.paragraphs:
                            if src in p.text:
                                set_text(p, p.text.replace(src, dst))


def apply_paragraph_styles(doc: Document) -> None:
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if re.match(r"^\d+\s+", text) or text == "参考文献":
            p.style = "11 一级标题"
            p.paragraph_format.first_line_indent = None
        elif re.match(r"^\d+\.\d+\s+", text):
            p.style = "12 二级标题"
            p.paragraph_format.first_line_indent = None
        elif re.match(r"^\d+\.\d+\.\d+\s+", text):
            p.style = "13 三级标题"
            p.paragraph_format.first_line_indent = None
        elif text.startswith("图") and "Fig." in text:
            p.style = "14 图题"
        elif text.startswith("表") and "Table " in text:
            p.style = "15 表题"
        elif FORMULA_RE.search(text) and ("=" in text or "≤" in text or text.startswith(("F(", "ODI", "P90", "E_", "Ω", "z(", "Score", "NCF"))):
            p.style = "16 公式编号"
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.tab_stops.clear_all()
            p.paragraph_format.tab_stops.add_tab_stop(Cm(15.0), WD_ALIGN_PARAGRAPH.RIGHT)
        elif text.startswith("["):
            p.style = "17 参考文献"
        elif p.style.name in {"Normal", "Body Text"}:
            p.style = "10 正文"


def main() -> None:
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)
    setup_styles(doc)

    # Front matter: use existing first 12 paragraphs so the original section break is preserved.
    ps = doc.paragraphs
    set_text(ps[0], CN_TITLE, "01 中文题名", WD_ALIGN_PARAGRAPH.CENTER)
    set_author_line(ps[1], ["李杨", "王楠", "刘浩天"], "03 中文作者")
    set_text(ps[2], CN_UNIT, "04 中文单位", WD_ALIGN_PARAGRAPH.CENTER)
    set_text(ps[3], CN_ABSTRACT, "05 中文摘要关键词")
    set_text(ps[4], CN_KEYWORDS, "05 中文摘要关键词")
    set_text(ps[5], "中图分类号：TD822    文献标志码：A", "06 中图分类号")
    set_text(ps[6], EN_TITLE, "02 英文题名", WD_ALIGN_PARAGRAPH.CENTER)
    set_author_line(ps[7], ["LI Yang", "WANG Nan", "LIU Haotian"], "07 英文作者")
    set_text(ps[8], EN_UNIT, "08 英文单位", WD_ALIGN_PARAGRAPH.CENTER)
    set_text(ps[9], EN_ABSTRACT, "09 英文摘要关键词")
    set_text(ps[10], EN_KEYWORDS, "09 英文摘要关键词")
    set_text(ps[11], FOOTNOTE, "18 首页脚注")

    # Remove leftover blank paragraph after front matter when it has no section break.
    if len(doc.paragraphs) > 12 and not doc.paragraphs[12].text.strip():
        # Keep if it owns a section break; otherwise remove.
        if not doc.paragraphs[12]._p.xpath("./w:pPr/w:sectPr"):
            delete_paragraph(doc.paragraphs[12])

    # Merge the accidental split paragraph in the introduction.
    for i, p in enumerate(list(doc.paragraphs)):
        if p.text.strip().startswith("本文主要开展以下工作"):
            nxt = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else None
            if nxt is not None and nxt.text.strip().startswith("标协同规划模型"):
                set_text(
                    p,
                    "本文主要开展以下工作：构建采区有效布置域与连续参数场，建立多场景 ODI 风险统一表达方法；"
                    "构建面向工程效率、资源回收与覆岩扰动控制的多目标协同规划模型，形成多模式候选方案生成、"
                    "筛选与比选机制；结合研究区案例，对参数场构建结果、ODI 风险分布、结构化规划结果与ODI风险"
                    "统计结果进行分析。研究结果可为采区规划中的多目标权衡、风险前置控制与方案比选提供统一的"
                    "计算框架和技术路径。",
                    "10 正文",
                )
                delete_paragraph(nxt)
            break

    # Merge reference 25 split across two paragraphs.
    for i, p in enumerate(list(doc.paragraphs[:-1])):
        if p.text.strip().startswith("[25]") and doc.paragraphs[i + 1].text.strip().startswith("Heidelberg"):
            set_text(p, p.text.rstrip(",") + ", " + doc.paragraphs[i + 1].text.strip(), "17 参考文献")
            delete_paragraph(doc.paragraphs[i + 1])
            break

    # Bilingual figure and table captions by current paragraph indices after one possible deletion.
    for p in doc.paragraphs:
        text = p.text.strip()
        for cn, en in list(FIG_CAPTIONS.values()):
            if text == cn:
                set_text(p, cn + "\n" + en, "14 图题", WD_ALIGN_PARAGRAPH.CENTER)
        for cn, en in list(TABLE_CAPTIONS.values()):
            if text == cn:
                set_text(p, cn + "\n" + en, "15 表题", WD_ALIGN_PARAGRAPH.CENTER)

    # Reference style and tab after serial number.
    for p in doc.paragraphs:
        text = p.text.strip()
        m = REF_RE.match(text)
        if m:
            set_text(p, REF_RE.sub(lambda x: f"[{x.group(1)}]\t", text), "17 参考文献")

    normalize_table_headers(doc)
    apply_paragraph_styles(doc)

    # Citation superscripts after styles are assigned.
    for p in doc.paragraphs:
        if p.style.name not in {"17 参考文献", "14 图题", "15 表题", "18 首页脚注"}:
            superscript_citations(p)
            for run in p.runs:
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                if p.style.name in {"09 英文摘要关键词", "02 英文题名", "07 英文作者", "08 英文单位", "17 参考文献"}:
                    run.font.name = "Times New Roman"

    # Page and section setup.
    for i, sec in enumerate(doc.sections):
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)
        set_columns(sec, 1 if i == 0 else 2)

    # Keep large images/tables readable by using single-column sections where the existing doc already
    # inserted dedicated section breaks around figure-heavy parts.
    for i in (4, 7, 9):
        if i < len(doc.sections):
            set_columns(doc.sections[i], 1)

    # Tables: three-line table and 6-point Chinese size.
    for table in doc.tables:
        format_table(table)

    doc.save(OUT)

    # Post-process columns in the final section body sectPr if present.
    with zipfile.ZipFile(OUT, "r") as zin:
        items = {name: zin.read(name) for name in zin.namelist()}
    with zipfile.ZipFile(OUT, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in items.items():
            zout.writestr(name, data)

    REPORT.write_text(
        "\n".join(
            [
                "# 煤炭科学技术模板格式化记录",
                "",
                f"- 源文件：`{SRC}`",
                f"- 输出文件：`{OUT}`",
                "- 已按期刊说明创建并应用 01～18 样式：题名、作者、单位、摘要、正文、三级标题、图题、表题、公式、参考文献、首页脚注。",
                "- 已重排首页：中文题名小于22字，补齐中文/英文作者单位、摘要、关键词、中图分类号、文献标志码和首页脚注。",
                "- 已将图1～图9、表1～表4改为中英文对照题名，并应用图题/表题样式。",
                "- 已将参考文献编号后空格改为Tab，并应用参考文献样式。",
                "- 已将表格处理为三线表口径，表中文字设为6号左右、单倍行距，并去除表头中的斜线表达。",
                "- 已尝试按首页单栏、正文双栏、大图大表单栏例外设置分节栏数。",
                "",
                "## 仍需人工确认",
                "",
                "- 若编辑部提供正式 `.dotx/.docx` 模板，应以正式模板的内置样式为准再导入本文内容。",
                "- 公式目前为可见文字编号格式，不是 MathType 对象；若期刊严格要求 MathType，需要在 Word 中进一步公式对象化。",
                "- 收稿日期、修回日期、责任编辑仍按投稿阶段留空。",
            ]
        ),
        encoding="utf-8",
    )
    print(OUT)
    print(REPORT)


if __name__ == "__main__":
    main()
