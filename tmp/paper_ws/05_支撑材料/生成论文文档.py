from __future__ import annotations

import os
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image

try:
    from lxml import etree
except Exception:
    etree = None

try:
    import latex2mathml.converter
except Exception:
    latex2mathml = None
else:
    latex2mathml = latex2mathml.converter


SKIP_FILES = {
    "投稿格式对照表.md",
    "作者与基金信息.md",
    "参考文献候选池.md",
    "引文映射表.md",
    "格式模版.md",
    "投稿待补信息清单.md",
    "敏东矿案例补强提纲.md",
}
XSL_CANDIDATES = [
    Path(r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"),
    Path(r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL"),
    Path(r"C:\Program Files\Microsoft Office\Office16\MML2OMML.XSL"),
]


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_columns(section, count: int) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    col = cols[0] if cols else OxmlElement("w:cols")
    if not cols:
        sect_pr.append(col)
    col.set(qn("w:num"), str(count))
    col.set(qn("w:space"), "425")


def add_bookmark(paragraph, bookmark_name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, text: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    r_pr.append(color)
    vert = OxmlElement("w:vertAlign")
    vert.set(qn("w:val"), "superscript")
    r_pr.append(vert)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def strip_markdown(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def find_main_markdown(manuscript_dir: Path) -> Path:
    for path in sorted(manuscript_dir.glob("*.md")):
        if path.name not in SKIP_FILES:
            return path
    raise FileNotFoundError("未找到主论文 Markdown 文件")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    set_columns(section, 1)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def find_xsl_path() -> Path | None:
    for path in XSL_CANDIDATES:
        if path.exists():
            return path
    return None


def load_xslt():
    if etree is None or latex2mathml is None:
        return None
    xsl_path = find_xsl_path()
    if xsl_path is None:
        return None
    try:
        return etree.XSLT(etree.parse(str(xsl_path)))
    except Exception:
        return None


def latex_to_omml_element(latex: str, transform):
    if transform is None or latex2mathml is None or etree is None:
        return None
    try:
        mathml = latex2mathml.convert(latex)
        mathml_root = etree.fromstring(mathml.encode("utf-8"))
        omml_tree = transform(mathml_root)
        omml_bytes = etree.tostring(omml_tree, encoding="utf-8")
        return parse_xml(omml_bytes)
    except Exception:
        return None


def append_math(paragraph, latex: str, transform) -> bool:
    element = latex_to_omml_element(latex, transform)
    if element is None:
        return False
    tag = element.tag.split("}")[-1]
    if tag == "oMathPara":
        for child in element:
            paragraph._p.append(child)
    else:
        paragraph._p.append(element)
    return True


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    set_east_asia_font(run, "黑体")


def add_meta_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    set_east_asia_font(run, "宋体")


def add_label_paragraph(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    label_run = p.add_run(label)
    label_run.bold = True
    label_run.font.size = Pt(10.5)
    set_east_asia_font(label_run, "黑体")
    body_run = p.add_run(text)
    body_run.font.size = Pt(10.5)
    set_east_asia_font(body_run, "宋体")


def add_heading(doc: Document, level: int, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 2 else 11)
    set_east_asia_font(run, "黑体" if level == 2 else "宋体")


def copy_section_layout(src, dst, columns: int) -> None:
    dst.top_margin = src.top_margin
    dst.bottom_margin = src.bottom_margin
    dst.left_margin = src.left_margin
    dst.right_margin = src.right_margin
    dst.page_width = src.page_width
    dst.page_height = src.page_height
    set_columns(dst, columns)


def add_center_caption(doc: Document, text: str, size: float = 10.0, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    set_east_asia_font(run, "宋体")


def add_list_item(doc: Document, text: str, numbered: bool) -> None:
    style = "List Number" if numbered else "List Bullet"
    p = doc.add_paragraph(style=style)
    run = p.add_run(strip_markdown(text))
    run.font.size = Pt(10.5)
    set_east_asia_font(run, "宋体")


def add_image(doc: Document, md_path: Path, alt: str, rel_path: str) -> None:
    image_path = (md_path.parent / rel_path).resolve()
    if not image_path.exists() or image_path.suffix.lower() not in {".png", ".jpg", ".jpeg", ".bmp", ".webp"}:
        return
    base = doc.sections[-1]
    single = doc.add_section(WD_SECTION.CONTINUOUS)
    copy_section_layout(base, single, 1)
    available_width = single.page_width - single.left_margin - single.right_margin
    width_inches = available_width / 914400
    try:
        with Image.open(image_path) as img:
            img_w, img_h = img.size
        if img_w > 0 and img_h > 0:
            max_width_inches = min(width_inches, 6.2)
            max_height_inches = 3.9
            ratio = min(max_width_inches / img_w, max_height_inches / img_h)
            width_inches = min(max_width_inches, img_w * ratio)
    except Exception:
        width_inches = min(width_inches, 6.0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(image_path), width=Inches(width_inches))
    if alt:
        add_center_caption(doc, alt, size=10.0)
    if base._sectPr.xpath("./w:cols/@w:num") != ["1"]:
        dual = doc.add_section(WD_SECTION.CONTINUOUS)
        copy_section_layout(base, dual, 2)


def add_superscript_text(paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    run.font.superscript = True
    run.font.size = Pt(8)
    set_east_asia_font(run, "宋体")


def is_table_separator(line: str) -> bool:
    candidate = line.strip().replace(" ", "")
    return bool(candidate) and set(candidate) <= {"|", "-", ":"}


def parse_table_row(line: str) -> list[str]:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return [strip_markdown(cell) for cell in cells]


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        edge_data = kwargs.get(edge)
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        if edge_data:
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))
        else:
            element.set(qn("w:val"), "nil")


def format_three_line_table(table) -> None:
    rows = table.rows
    total_rows = len(rows)
    for r_idx, row in enumerate(rows):
        for cell in row.cells:
            top = {"val": "single", "sz": 12, "space": 0, "color": "000000"} if r_idx == 0 else None
            bottom = {"val": "single", "sz": 8 if r_idx == 0 else 12, "space": 0, "color": "000000"} if r_idx in {0, total_rows - 1} else None
            set_cell_border(cell, top=top, bottom=bottom, left=None, right=None)
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.space_before = Pt(0)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(9)
                    if r_idx == 0:
                        run.bold = True
                    set_east_asia_font(run, "宋体")


def add_markdown_table(doc: Document, table_lines: list[str], captions: list[str]) -> None:
    rows = [parse_table_row(line) for line in table_lines if line.strip()]
    if len(rows) < 2:
        return
    header = rows[0]
    body = [row for row in rows[2:] if row]

    base = doc.sections[-1]
    single = doc.add_section(WD_SECTION.CONTINUOUS)
    copy_section_layout(base, single, 1)

    for caption in captions:
        add_center_caption(doc, caption, size=10.0)

    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.autofit = False
    available_width = single.page_width - single.left_margin - single.right_margin
    col_width = int(available_width / len(header))
    for c_idx, text in enumerate(header):
        cell = table.cell(0, c_idx)
        cell.text = text
        cell.width = col_width
    for r_idx, row in enumerate(body, start=1):
        normalized = row + [""] * (len(header) - len(row))
        for c_idx, text in enumerate(normalized[: len(header)]):
            cell = table.cell(r_idx, c_idx)
            cell.text = text
            cell.width = col_width

    format_three_line_table(table)

    if base._sectPr.xpath("./w:cols/@w:num") != ["1"]:
        dual = doc.add_section(WD_SECTION.CONTINUOUS)
        copy_section_layout(base, dual, 2)


def iter_line_segments(text: str) -> list[tuple[str, str]]:
    pattern = re.compile(r"(\$(?!\$).*?(?<!\\)\$|\[(?:\d+(?:\s*[-,，]\s*\d+)*)\])")
    segments: list[tuple[str, str]] = []
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            segments.append(("text", strip_markdown(text[cursor:match.start()])))
        token = match.group(1)
        if token.startswith("$"):
            segments.append(("math", token[1:-1].strip()))
        else:
            segments.append(("cite", token))
        cursor = match.end()
    if cursor < len(text):
        segments.append(("text", strip_markdown(text[cursor:])))
    return [(kind, value) for kind, value in segments if value]


def add_body_paragraph(doc: Document, text: str, transform, math_counter: dict[str, int]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(2)
    for kind, value in iter_line_segments(text):
        if kind == "text":
            run = p.add_run(value)
            run.font.size = Pt(10.5)
            set_east_asia_font(run, "宋体")
        elif kind == "cite":
            inner = value[1:-1]
            add_superscript_text(p, "[")
            for piece in re.split(r"(,|，|-|–)", inner):
                if not piece:
                    continue
                if piece.isdigit():
                    add_internal_hyperlink(p, piece, f"ref-{piece}")
                else:
                    add_superscript_text(p, piece)
            add_superscript_text(p, "]")
        else:
            if append_math(p, value, transform):
                math_counter["count"] += 1
            else:
                run = p.add_run(f"${value}$")
                run.italic = True
                run.font.size = Pt(10.5)
                set_east_asia_font(run, "宋体")


def add_equation_paragraph(doc: Document, latex: str, transform, math_counter: dict[str, int]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if append_math(p, latex, transform):
        math_counter["count"] += 1
        return
    run = p.add_run(f"$${latex}$$")
    run.italic = True
    run.font.size = Pt(10.5)
    set_east_asia_font(run, "宋体")


def add_reference_paragraph(doc: Document, text: str, bookmark_id: int) -> None:
    match = re.match(r"^\[(\d+)\]\s*(.*)$", text)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(-0.74)
    p.paragraph_format.space_after = Pt(2)
    if not match:
        run = p.add_run(strip_markdown(text))
        run.font.size = Pt(9)
        set_east_asia_font(run, "宋体")
        return
    ref_id, body = match.groups()
    add_bookmark(p, f"ref-{ref_id}", bookmark_id)
    run = p.add_run(f"[{ref_id}] {strip_markdown(body)}")
    run.font.size = Pt(9)
    set_east_asia_font(run, "宋体")


def convert(md_path: Path, docx_path: Path) -> tuple[int, bool]:
    doc = Document()
    configure_document(doc)
    transform = load_xslt()
    math_counter = {"count": 0}

    lines = md_path.read_text(encoding="utf-8").splitlines()
    body_started = False
    in_references = False
    bookmark_id = 1
    in_math_block = False
    math_lines: list[str] = []

    pending_table_captions: list[str] = []
    i = 0
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        if in_math_block:
            if stripped == "$$":
                add_equation_paragraph(doc, "\n".join(math_lines).strip(), transform, math_counter)
                in_math_block = False
                math_lines = []
                i += 1
                continue
            else:
                math_lines.append(raw)
                i += 1
                continue

        if not stripped:
            i += 1
            continue

        if stripped == "$$":
            in_math_block = True
            math_lines = []
            i += 1
            continue

        if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) > 4:
            add_equation_paragraph(doc, stripped[2:-2].strip(), transform, math_counter)
            i += 1
            continue

        if stripped.startswith("# "):
            add_title(doc, strip_markdown(stripped[2:]))
            i += 1
            continue

        if not body_started and stripped.startswith("作者："):
            add_meta_line(doc, strip_markdown(stripped))
            i += 1
            continue
        if not body_started and stripped.startswith("单位："):
            add_meta_line(doc, strip_markdown(stripped))
            i += 1
            continue
        if not body_started and stripped.startswith("基金项目："):
            add_meta_line(doc, strip_markdown(stripped))
            i += 1
            continue
        if not body_started and stripped.startswith("中图分类号："):
            add_meta_line(doc, strip_markdown(stripped))
            i += 1
            continue
        if not body_started and stripped.startswith("文献标志码："):
            add_meta_line(doc, strip_markdown(stripped))
            i += 1
            continue

        if stripped.startswith("摘要："):
            add_label_paragraph(doc, "摘要：", strip_markdown(stripped[3:]))
            i += 1
            continue
        if stripped.startswith("关键词："):
            add_label_paragraph(doc, "关键词：", strip_markdown(stripped[4:]))
            section = doc.add_section(WD_SECTION.CONTINUOUS)
            section.top_margin = doc.sections[0].top_margin
            section.bottom_margin = doc.sections[0].bottom_margin
            section.left_margin = doc.sections[0].left_margin
            section.right_margin = doc.sections[0].right_margin
            set_columns(section, 2)
            body_started = True
            i += 1
            continue

        if stripped.startswith("## "):
            heading_text = strip_markdown(stripped[3:])
            in_references = heading_text == "参考文献"
            add_heading(doc, 2, heading_text)
            body_started = True
            i += 1
            continue

        if stripped.startswith("### "):
            add_heading(doc, 3, strip_markdown(stripped[4:]))
            i += 1
            continue

        if stripped.startswith("**表") or stripped.startswith("**Table "):
            pending_table_captions.append(strip_markdown(stripped))
            i += 1
            continue

        if stripped.startswith("|"):
            block: list[str] = []
            j = i
            while j < len(lines) and lines[j].strip().startswith("|"):
                block.append(lines[j].strip())
                j += 1
            if len(block) >= 2 and is_table_separator(block[1]):
                add_markdown_table(doc, block, pending_table_captions)
                pending_table_captions = []
                i = j
                continue

        if stripped.startswith("![") and "](" in stripped and stripped.endswith(")"):
            alt = stripped[2: stripped.index("](")]
            rel = stripped[stripped.index("](") + 2 : -1]
            add_image(doc, md_path, strip_markdown(alt), rel)
            i += 1
            continue

        if stripped.startswith("- "):
            add_list_item(doc, stripped[2:], numbered=False)
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped) and not in_references:
            add_list_item(doc, re.sub(r"^\d+\.\s+", "", stripped), numbered=True)
            i += 1
            continue

        if in_references:
            add_reference_paragraph(doc, stripped, bookmark_id)
            bookmark_id += 1
        else:
            add_body_paragraph(doc, stripped, transform, math_counter)
        i += 1

    doc.save(docx_path)
    return math_counter["count"], transform is not None


if __name__ == "__main__":
    base = Path(__file__).resolve().parents[1]
    manuscript_dir = base / "04_论文稿件"
    md_path = find_main_markdown(manuscript_dir)
    custom_out = os.environ.get("PAPER_DOCX_OUT", "").strip()
    docx_path = Path(custom_out) if custom_out else manuscript_dir / f"{md_path.stem}.docx"
    equation_count, math_enabled = convert(md_path, docx_path)
    print(f"已生成：{docx_path}")
    print(f"公式对象链路：{'已启用' if math_enabled else '未启用'}")
    print(f"写入公式数量：{equation_count}")
