import json
import math
import shutil
import statistics
from pathlib import Path

from docx import Document
from docx.shared import Pt


BASE = Path("D:/xiangmu/miningplan")
DATA_DIR = BASE / "论文" / "重构工作区" / "05_支撑材料" / "接口结果"
IN_DOCX = BASE / "论文" / "重构工作区" / "06_投稿包" / "最新版论文4.16_插图版_第一优先级修订.docx"
OUT_DOCX = BASE / "论文" / "重构工作区" / "06_投稿包" / "最新版论文4.16_插图版_第二轮精修终版.docx"


def load_json(name):
    return json.loads((DATA_DIR / name).read_text(encoding="utf-8"))


def polygon_area(points):
    area = 0.0
    for i, p in enumerate(points):
        q = points[(i + 1) % len(points)]
        area += p["x"] * q["y"] - q["x"] * p["y"]
    return abs(area) / 2


def percentile(values, pct):
    vals = sorted(values)
    k = (len(vals) - 1) * pct / 100
    f = math.floor(k)
    c = math.ceil(k)
    return vals[f] * (c - k) + vals[c] * (k - f)


def set_text(paragraph, text):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run.font.size = Pt(10.5)


def replace_by_index(doc, mapping):
    for idx, text in mapping.items():
        if idx < len(doc.paragraphs):
            set_text(doc.paragraphs[idx], text)


def replace_caption(doc, prefix, text):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            set_text(paragraph, text)
            return


def insert_caption_before_table(doc, table_index, text):
    if any(p.text.strip().startswith(text.split()[0]) for p in doc.paragraphs):
        return
    paragraph = doc.add_paragraph()
    set_text(paragraph, text)
    table = doc.tables[table_index]
    table._element.addprevious(paragraph._p)


def table_cell_text(table, row, col, text):
    cell = table.rows[row].cells[col]
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "宋体"
            run.font.size = Pt(9)


def main():
    shutil.copy2(IN_DOCX, OUT_DOCX)

    design = load_json("采区设计结果.json")
    boundary = load_json("边界数据.json")
    odi = load_json("000_mindong_layout_odi_field.json")

    boreholes = design["boreholes"]
    thickness = [float(b["coalThickness"]) for b in boreholes]
    stats = design["stats"]
    panels = design["panels"]
    roadways = design["roadways"]
    roadway_total = sum(float(r["length"]) for r in roadways)
    boundary_area = polygon_area(boundary["boundary"])
    coverage = float(stats["totalArea"]) / boundary_area * 100
    field = [float(x) for row in odi["field"] for x in row]
    odi_mean = statistics.mean(field)
    odi_p90 = percentile(field, 90)
    odi_p50 = percentile(field, 50)
    odi_gt070 = sum(v > 0.70 for v in field) / len(field) * 100
    odi_gt080 = sum(v > 0.80 for v in field) / len(field) * 100
    max_bh = max(boreholes, key=lambda b: b["coalThickness"])
    min_bh = min(boreholes, key=lambda b: b["coalThickness"])
    panel_ranges = "、".join(f"{p['id']}推进{float(p['advanceLength']):.1f} m" for p in panels)
    panel_areas = "、".join(f"{p['id']}面积{float(p['area']):.2f} m²" for p in panels)

    doc = Document(OUT_DOCX)

    cn_abstract = (
        "摘要：针对采区规划中离散地质信息难以转化为连续规划约束、多场景覆岩扰动风险难以统一表达以及规划结果难以向后续评价传递等问题，"
        "提出一种基于覆岩扰动指数（overburden disturbance index，ODI）约束的采区多目标协同规划方法。该方法以采区边界、钻孔样点和设计参数为输入，"
        "构建有效布置域与连续参数场，将地表沉陷、含水层扰动和上行开采等风险统一为ODI指标，并以前置约束形式参与工作面与巷道布局生成。"
        f"研究区样例结果表明：15个钻孔样点的煤层厚度为{min(thickness):.1f}～{max(thickness):.1f} m，平均{statistics.mean(thickness):.4f} m；"
        f"在{boundary_area:.2f} m²原始边界范围内，样例生成3个工作面和11条巷道，布置面积{float(stats['totalArea']):.2f} m²，"
        f"有效覆盖率{coverage:.2f}%，巷道总长度{roadway_total:.2f} m，平均规划评分{float(stats['avgScore']):.1f}。"
        f"ODI场统计显示，研究区ODI均值为{odi_mean:.4f}，P90为{odi_p90:.4f}，ODI>0.70的栅格占{odi_gt070:.2f}%。"
        "结果说明，该方法能够在统一空间参照和统一风险口径下完成参数场构建、风险前置约束和规划对象生成，可为采区规划中的多目标权衡与后续评价提供可复核的工程方法框架。"
    )
    en_abstract = (
        "Abstract: To address the difficulty of converting discrete geological information into continuous planning constraints, "
        "the heterogeneous expression of multi-scenario overburden disturbance risks, and the weak transferability of planning results to downstream evaluation, "
        "a multi-objective collaborative mining-district planning method constrained by the overburden disturbance index (ODI) is proposed. "
        "The method constructs an effective layout domain and continuous parameter fields from boundary, borehole and design-parameter inputs, "
        "unifies surface subsidence, aquifer disturbance and upward-mining-related risks into the ODI framework, and embeds the risk indicator into working-face and roadway layout generation. "
        f"In the sample study area, the coal thickness of 15 boreholes ranges from {min(thickness):.1f} to {max(thickness):.1f} m with an average of {statistics.mean(thickness):.4f} m. "
        f"Within the original boundary area of {boundary_area:.2f} m2, the generated layout contains 3 working faces and 11 roadways, with a layout area of {float(stats['totalArea']):.2f} m2, "
        f"an effective coverage ratio of {coverage:.2f}%, a total roadway length of {roadway_total:.2f} m, and an average planning score of {float(stats['avgScore']):.1f}. "
        f"The ODI field has a mean value of {odi_mean:.4f}, a P90 value of {odi_p90:.4f}, and {odi_gt070:.2f}% of grids exceeding 0.70. "
        "The results indicate that the proposed method can integrate parameter-field construction, forward risk constraints, and planning-object generation under a unified spatial and risk-evaluation framework."
    )

    replacements = {
        5: cn_abstract,
        10: en_abstract,
        18: (
            "基于此，本文提出一种基于覆岩扰动指数（overburden disturbance index，ODI）约束的采区多目标协同规划方法。"
            "方法主线包括：构建有效布置域与连续参数场，将离散钻孔和边界信息转化为可计算约束；建立ODI统一风险口径，使地表沉陷、含水层扰动和上行开采等风险能够在同一尺度下参与规划；"
            "在工程效率、资源覆盖和扰动控制之间进行多目标权衡，形成可复核的工作面与巷道布局结果；最后，将规划结果作为后续采掘接续和经济评价的上游输入。"
            "本文强调的是样例条件下方法链和对象链的贯通性，不将当前结果扩展为真实矿井条件下的普适最优结论。"
        ),
        31: (
            "式中，ODI为覆岩扰动指数；x_i为第i类风险因子的归一化指标值；w_i为相应权重；n为纳入计算的风险因子数量。"
            "本文主要考虑地表沉陷、含水层扰动和上行开采3类风险场景，并通过权重归一化保证不同风险场景可在同一评价口径下参与规划。"
        ),
        37: (
            "式中，π为候选规划方案；A_π为方案π对应的布置区域；|A_π|为区域面积；Q_90表示ODI的90%分位数；T_ODI为扰动控制阈值；"
            "E_π表示方案区域内ODI超过阈值的暴露比例。上述统计量用于把风险场从图层表达转化为方案级评价指标。"
        ),
        43: (
            "式中，Ω为原始采区边界，B_b为边界煤柱宽度，B_s为区段煤柱宽度，D_p为局部保护距离，Ω_e为经约束内缩和几何合法性处理后的有效布置域。"
            "若内缩后出现多连通域或局部狭长畸变，则保留主连通区域并采用降级内缩策略，以保证后续工作面布置具有几何可解性。"
        ),
        46: (
            "式中，z(x)为规划位置x处的参数估计值；z_i为第i个钻孔样点的观测值；d_i(x)为位置x与第i个钻孔之间的距离；p为距离衰减指数。"
            "该表达体现反距离加权插值思想，目的在于把有限钻孔样点转换为可参与空间规划和风险叠加的连续参数场。"
        ),
        54: (
            "式中，S_e(π)、S_r(π)和S_m(π)分别表示方案π在工程效率、资源覆盖和扰动控制方面的评价值；w_e、w_r和w_m为对应权重，且满足w_e+w_r+w_m=1。"
            "该目标函数用于说明多目标权衡口径，不替代真实矿井条件下的最终工程优选。"
        ),
        63: (
            "式中，P_s为产量组织指标，R_s为风险可控性指标，C_s为工期与组织可控性指标；α、β和γ为对应权重，且满足α+β+γ=1。"
            "该式用于表达规划结果向接续评价传递的指标关系，本文不据此给出未经实证导出的接续评价数值。"
        ),
        66: "式中，Rev_t为第t月收入，Cost_t为常规成本，RiskCost_t为风险联动成本，NCF_t为第t月净现金流。该式用于说明经济评价可接收规划与接续结果形成的月度输入。",
        68: (
            "式中，I_0为初始投入，r_m为月折现率，T为评价期。本文保留经济评价指标的理论传递关系，但不在样例结果中给出未经独立导出的净现值数值。"
        ),
        69: (
            "因此，本文方法可概括为：以有效布置域和连续参数场为基础，以ODI为统一风险约束指标，以多目标综合评价为方案筛选口径，"
            "以采掘接续和经济评价为后续传递方向，形成从“风险前置约束—规划对象生成—方案级统计—后续评价输入”的连续决策流程。"
        ),
        72: (
            "为验证所提方法在采区规划场景中的适用性与链路贯通能力，选取研究区样例开展分析。该研究区以6个边界点构成的采区边界、"
            "15个钻孔样点和规划控制参数为基础输入，原始边界面积为255250.00 m²。当前验证目标限定为样例条件下的数据组织、风险表征与规划对象生成，"
            "不将样例结果解释为真实矿井条件下的工业级优选结论。"
        ),
        73: (
            "研究区输入对象包括采区边界、钻孔样点和规划控制参数3类。钻孔样点用于构建煤层厚度等连续参数场；边界对象用于界定有效布置域；"
            "边界煤柱、区段煤柱、工作面宽度、推进长度和ODI统计阈值则共同构成后续规划求解的约束条件。"
        ),
        76: (
            f"图2显示，15个钻孔样点在研究区内形成了基本可用的空间采样网络。样点煤厚最小值出现在{min_bh['id']}，为{min_bh['coalThickness']:.1f} m；"
            f"最大值出现在{max_bh['id']}，为{max_bh['coalThickness']:.1f} m，样点平均煤厚为{statistics.mean(thickness):.4f} m。"
            "这些离散样点为后续厚度场插值、资源覆盖判断和ODI风险场叠加提供了统一输入口径。"
        ),
        77: (
            "在研究区进入规划求解前，原始边界并不直接作为工作面布置边界使用，而是结合边界煤柱、区段煤柱及局部保护距离进行约束处理。"
            "本样例采用30.0 m边界煤柱、20.0 m区段煤柱和120.0 m工作面宽度，并以走向长壁后退式开采作为布置方式。"
        ),
        78: (
            "因此，本节的作用是把研究区输入压实为可复核的参数体系：边界提供空间范围，钻孔提供连续参数场样本，设计参数提供几何与安全约束，"
            "ODI阈值提供风险统计口径。后续结果分析均以该统一输入口径为基础。"
        ),
        82: (
            "采区规划首先需要解决离散地质信息向连续规划约束转换的问题。对于研究区样例而言，15个钻孔样点提供了煤层厚度、岩石硬度、瓦斯含量和涌水量等离散属性；"
            "通过空间插值和统一边界裁剪处理，这些点状信息被转换为可与规划域、ODI风险场和工作面布局叠加分析的连续参数场。"
        ),
        85: (
            f"从图3可见，研究区煤层厚度样本范围为{min(thickness):.1f}～{max(thickness):.1f} m，平均值为{statistics.mean(thickness):.4f} m。"
            f"高厚度样点主要包括{max_bh['id']}（{max_bh['coalThickness']:.1f} m）及周边样点，低厚度样点以{min_bh['id']}（{min_bh['coalThickness']:.1f} m）为代表。"
            "这种厚度差异说明，若仅依据几何边界进行布置，容易忽略资源条件在空间上的非均匀性。"
        ),
        86: (
            "对于采区规划而言，厚度较高区域在资源覆盖和回收价值评价中具有更高吸引力，而厚度较低区域更适合作为边界调整、煤柱保留或低优先级布置区域。"
            "因此，连续参数场不是单纯的地质展示图，而是后续工作面排序、覆盖率统计和风险叠加分析的基础底图。"
        ),
        87: (
            "需要强调的是，本文并不试图在该部分证明某种复杂地质建模算法的优越性，而是强调连续参数场在采区规划链路中的工程作用："
            "它使离散钻孔信息能够以统一空间参照参与工作面生成、ODI风险计算和方案级统计。"
        ),
        88: (
            "此外，参数场还为ODI风险表征提供了空间载体。由于多场景风险需要在同一边界、同一坐标和同一网格上进行叠加，"
            "参数场、风险场和规划对象之间的一致性直接决定了后续结果是否可比较、可传递和可复核。"
        ),
        92: (
            f"从研究区样例结果看，多场景ODI结果已经能够在同一研究区边界、钻孔参照和色标口径下表达。当前综合ODI场采用80×56网格，共4480个栅格，"
            f"归一化后ODI均值为{odi_mean:.4f}，中位数为{odi_p50:.4f}，P90为{odi_p90:.4f}；其中ODI>0.70的栅格占{odi_gt070:.2f}%，ODI>0.80的栅格占{odi_gt080:.2f}%。"
            "这说明风险信息已经从概念性指标转化为可计算、可统计和可进入规划约束的空间场。"
        ),
        95: (
            "从风险分布特征看，不同场景下的ODI高值区并不完全重合。地表沉陷场景更偏向表征采动引起的地表影响范围，含水层扰动场景更强调导水裂隙带和含水层受扰风险，"
            "上行开采场景则更关注上覆扰动条件下的安全边界。统一为ODI后，不同场景可在同一统计口径下比较，从而避免多指标、多量纲并列时难以进入方案筛选的问题。"
        ),
        96: (
            f"在规划语境下，ODI的价值体现在两个层面：一是识别高扰动敏感区，当前样例中ODI>0.70区域占{odi_gt070:.2f}%，可作为工作面布置时需重点避让或降权处理的区域；"
            f"二是提供方案级风险统计基础，例如均值{odi_mean:.4f}、P90={odi_p90:.4f}和高阈值暴露比例{odi_gt080:.2f}%等指标，可直接进入后续布局评价。"
        ),
        97: (
            "由此，ODI不再是布局完成后的附加说明图层，而是能够在候选方案生成前即参与约束设定和风险排序。"
            "对于研究区样例而言，ODI场的建立完成了从“边界+钻孔输入”到“风险约束可参与规划”的关键转换。"
        ),
        102: (
            "在有效布置域、连续参数场和ODI风险场共同作用下，本文进一步开展采区规划结果生成与风险约束检验。"
            "与传统只给出静态边界划分的做法不同，本文将钻孔样点、连续参数场、保护煤柱约束和ODI风险场纳入同一评价口径，"
            "使规划结果能够以工作面、巷道、覆盖面积和ODI统计量等可复核指标表达。"
        ),
        105: (
            f"从样例结果看，当前规划共形成3个工作面和11条巷道，巷道总长度为{roadway_total:.2f} m，工作面布置面积为{float(stats['totalArea']):.2f} m²，"
            f"相对于原始边界的有效覆盖率为{coverage:.2f}%。3个工作面的推进长度分别为{panel_ranges}，平均推进长度为{float(stats['avgAdvanceLength']):.1f} m。"
        ),
        106: (
            f"从布局结果本身看，工作面并非直接贴附原始边界外轮廓，而是在有效布置域内形成必要的保护距离和工程边界。"
            f"{panel_areas}，说明各工作面在空间上形成了连续且相互隔离的条带式布局；平均规划评分为{float(stats['avgScore']):.1f}，其中WF-03评分最高，为{max(float(p['avgScore']) for p in panels):.1f}。"
        ),
        111: (
            "本文方法的一个重要特点，是不把采区规划的几何布局视为终点，而是强调规划结果可以继续作为采掘接续和工程经济评价的上游输入。"
            "在本样例中，工作面边界、巷道路径、推进长度、布置面积和ODI统计量均已形成可传递的规划结果，后续环节可据此组织生产顺序、产量核算和经济评价。"
        ),
        112: (
            "在采掘接续层面，规划阶段形成的工作面边界和推进关系可映射为接续任务对象。当前样例已经给出由采区规划空间结果向采掘接续空间结果的传递路径，"
            "其证据重点是“对象可传递、链路可延伸”，而不是未经独立导出的接续优选评分。"
        ),
        115: (
            "在调控层面，规划结果还可进一步转化为工作面级控制变量。当前样例明确给出了120.0 m工作面宽度、20.0 m区段煤柱宽度和30.0 m边界煤柱宽度，"
            f"并可与ODI均值{odi_mean:.4f}、P90={odi_p90:.4f}及ODI>0.70比例{odi_gt070:.2f}%等统计量对应，用于后续敏感性分析。"
        ),
        118: (
            "在工程经济评价层面，规划与接续结果可以继续进入收入、成本、风险联动成本和现金流分析过程。当前样例已完成规划对象向后续评价口径的传递，"
            "但尚不具备独立导出的真实矿井经济对照数据，因此本文仅讨论其作为评价输入的可传递性，不给出未经数据支撑的经济优选结论。"
        ),
        121: (
            f"综上，研究区样例表明，本文方法已完成“边界与钻孔输入—参数场构建—ODI风险组织—规划对象生成—后续评价输入”的连续运行过程。"
            f"在15个钻孔、6个边界点和4480个ODI栅格的样例条件下，方法生成3个工作面、11条巷道、{coverage:.2f}%有效覆盖率和{odi_gt070:.2f}%高ODI暴露比例等可复核结果。"
            "其核心价值在于证明方法链和对象链已经贯通，而非证明某一单独模块在真实矿井条件下达到普适最优。"
        ),
        125: (
            "传统采区规划多采用“先形成布局方案、再开展风险校核”的串行方式。该流程容易使地表沉陷、含水层扰动和上行开采等风险因素停留在后验修正环节，"
            "一旦校核不满足要求，往往需要重新调整边界、工作面和巷道布置。本文将ODI作为前置约束，使风险信息在候选布局形成阶段即参与空间筛选和统计评价。"
        ),
        126: (
            f"从研究区样例看，ODI场提供了可量化的风险入口：均值{odi_mean:.4f}反映总体扰动水平，P90={odi_p90:.4f}反映高分位风险，"
            f"ODI>0.70和ODI>0.80比例分别为{odi_gt070:.2f}%和{odi_gt080:.2f}%，可用于识别高扰动敏感区并指导工作面避让或降权。"
        ),
        127: (
            "进一步看，ODI统一表达的价值还体现在连续传递能力上。由于ODI在规划阶段已形成方案级统计指标，因此可直接作为后续采高调控、采掘接续和经济评价的上游输入，"
            "不需要在下游环节重新定义风险口径。"
        ),
        130: (
            "采区规划结果的可靠性不仅取决于边界几何条件，还与地质属性的空间连续变化密切相关。传统经验式布置能够快速形成初始形态，"
            "但对钻孔控制下的厚度差异、资源富集区和低厚度区响应不足，容易造成资源评价与几何布置脱节。"
        ),
        131: (
            f"研究区样例表明，参数场引入后，高低厚度区在规划评价中的作用可以被区分：样点煤厚从{min(thickness):.1f} m变化到{max(thickness):.1f} m，"
            f"平均值为{statistics.mean(thickness):.4f} m。高厚度区在资源覆盖评价中更具吸引力，而低厚度区则提示布局需结合边界约束和风险约束进行权衡。"
        ),
        132: (
            "因此，参数场不是独立的地质建模结果，而是服务规划决策的中间环节。只有当参数场、风险场和几何对象被组织在统一空间参照之下，"
            "规划结果才可能进一步传递至采掘接续与工程经济分析。"
        ),
        134: (
            "本文当前结果主要证明所提方法在样例条件下的链路贯通能力，即能够实现采区边界、钻孔样点、连续参数场、ODI风险场和规划对象之间的连续组织与传递。"
            f"现有结果形成3个工作面、11条巷道和{roadway_total:.2f} m巷道总长度，说明该方法在方法链和对象链层面具有较好的完整性。"
        ),
        135: (
            "同时，本文结论具有明确适用边界。当前验证主要基于样例级输入和既有程序输出，尚未形成多个真实矿井案例下的系统对照，"
            "因此现阶段更适合支撑“方法可行”和“链路贯通”的结论，而不宜扩展为真实矿井条件下的优选幅度判断。"
        ),
        136: (
            "后续研究应在实矿数据中补充多方案对照、调控参数敏感性分析、接续排程数值和经济评价结果，使当前已经打通的方法链进一步转化为可量化的工程优选依据。"
        ),
        138: (
            "1）针对采区规划中地质信息离散、风险约束异构以及规划结果与后续评价脱节的问题，提出了基于ODI约束的采区多目标协同规划方法，"
            "构建了“有效布置域—连续参数场—ODI风险场—规划对象—后续评价输入”的一体化决策链。"
        ),
        139: (
            f"2）通过将地表沉陷、含水层扰动和上行开采等多场景风险统一组织为ODI指标，实现了异构风险在同一尺度下的前置表达。"
            f"样例ODI场均值为{odi_mean:.4f}，P90为{odi_p90:.4f}，ODI>0.70栅格占{odi_gt070:.2f}%，能够为高扰动敏感区识别和方案级风险统计提供依据。"
        ),
        140: (
            f"3）研究区样例验证表明，在15个钻孔样点和6个边界点条件下，方法形成3个工作面和11条巷道，布置面积{float(stats['totalArea']):.2f} m²，"
            f"巷道总长度{roadway_total:.2f} m，有效覆盖率{coverage:.2f}%，平均规划评分{float(stats['avgScore']):.1f}，说明规划结果能够由输入数据稳定传递为工程布置对象。"
        ),
        141: (
            "4）当前结果主要支撑样例条件下的方法贯通和对象链有效性。真实矿井条件下的工程优选结论仍需结合实矿案例、参数标定、接续排程和经济评价结果进一步检验。"
        ),
    }
    replace_by_index(doc, replacements)

    # Remove residual platform-style wording inherited from earlier drafts.
    for paragraph in doc.paragraphs:
        if paragraph.text:
            paragraph.text = (
                paragraph.text
                .replace("统一风险组织接口", "统一风险组织口径")
                .replace("接口", "口径")
            )
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    cell.text = (
                        cell.text
                        .replace("钻孔/设计接口结果", "钻孔/设计结果")
                        .replace("设计接口结果", "设计结果")
                        .replace("接口", "口径")
                    )

    replace_caption(doc, "表3", "表3 当前样例规划结果统计表")
    replace_caption(doc, "表4", "表4 当前样例参数与ODI统计表")
    insert_caption_before_table(doc, 2, "表3 当前样例规划结果统计表")

    # Compact Table 2 wording.
    if len(doc.tables) >= 2:
        t = doc.tables[1]
        rows = [
            ["场景", "关注对象", "输入依据", "规划作用"],
            ["地表沉陷", "地表移动与沉陷风险", "沉陷预测、保护约束", "约束工作面布局范围"],
            ["含水层扰动", "导水裂隙带与含水层受扰", "导水带高度、含水层敏感性", "参与风险筛选与避让"],
            ["上行开采", "上覆岩层扰动与安全边界", "采动应力、覆岩破坏特征", "支撑可行性评价"],
            ["综合ODI", "多场景叠加扰动", "多源风险归一化结果", "形成统一方案级指标"],
        ]
        for r, row in enumerate(rows):
            for c, value in enumerate(row):
                table_cell_text(t, r, c, value)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
