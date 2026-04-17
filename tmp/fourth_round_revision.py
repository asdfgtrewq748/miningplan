from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

import matplotlib
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, Rectangle


ROOT = Path.cwd()
PKG = ROOT / "论文" / "重构工作区" / "06_投稿包"
SRC = PKG / "最新版论文4.16_插图版_第三轮修订版.docx"
OUT = PKG / "最新版论文4.16_插图版_第四轮审稿优化版.docx"
PDF = PKG / "最新版论文4.16_插图版_第四轮审稿优化版.pdf"
FIG_DIR = ROOT / "tmp" / "fourth_round_figures"
REPORT = ROOT / "论文" / "重构工作区" / "00_过程文档" / "第四轮审稿优化与输出图件取舍清单.md"


def set_run_font(run, size=10.5, name="Times New Roman"):
    run.font.name = name
    run.font.size = Pt(size)


def replace_text(doc: Document, needle: str, new: str) -> int:
    n = 0
    for p in doc.paragraphs:
        if needle in p.text:
            p.text = new
            n += 1
    return n


def style_equation(p, text: str):
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, 10.5, "Times New Roman")


def normalize_equations(doc: Document):
    formulas = {
        28: r"ODI(x)=Σ(i=1..n) w_i x_i                                                   （1）",
        30: r"0≤x_i≤1，Σ(i=1..n) w_i=1，w_i≥0                                      （2）",
        33: r"ODI_bar(π)=1/|A_π| ∫_{A_π} ODI(x) dx                                  （3）",
        34: r"P90_π=Q_0.90{ODI(x) | x∈A_π}                                         （4）",
        36: r"E_π=|{x∈A_π | ODI(x)>T_ODI}|/|A_π|                                    （5）",
        42: r"Ω_e=Ω_0 \\ (B_b∪B_s∪D_p)                                               （6）",
        45: r"z(x)=Σ(i=1..n)[d_i(x)^(-p) z_i] / Σ(i=1..n)d_i(x)^(-p)                  （7）",
        51: r"F(π)=w_e S_e(π)+w_r S_r(π)+w_m S_m(π)                                  （8）",
        53: r"w_e+w_r+w_m=1，w_e≥0，w_r≥0，w_m≥0                                    （9）",
        62: r"Score_s=αP_s+βR_s+γC_s                                                 （10）",
        65: r"NCF_t=Rev_t-Cost_t-RiskCost_t                                          （11）",
    }
    paras = doc.paragraphs
    for idx, formula in formulas.items():
        if idx < len(paras):
            style_equation(paras[idx], formula)


def add_score_rows(doc: Document):
    tbl = doc.tables[2]
    existing = "\n".join(cell.text for row in tbl.rows for cell in row.cells)
    if "WF-01" in existing and "67.5" in existing:
        return
    rows = [
        ["工作面明细", "WF-01", "面积52181.16 m²；推进515.1 m", "评分67.5", "支撑平均规划评分计算"],
        ["工作面明细", "WF-02", "面积61822.04 m²；推进536.6 m", "评分67.5", "支撑平均规划评分计算"],
        ["工作面明细", "WF-03", "面积62562.51 m²；推进551.1 m", "评分71.5", "支撑平均规划评分计算"],
    ]
    for values in rows:
        row = tbl.add_row()
        for cell, val in zip(row.cells, values):
            cell.text = val


def draw_economic_input(out: Path):
    plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "SimSun", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False
    fig, ax = plt.subplots(figsize=(2256 / 220, 1172 / 220), dpi=220)
    ax.set_axis_off()
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 6)

    def box(x, y, w, h, text, fc="#eef6f8", fs=9):
        ax.add_patch(Rectangle((x, y), w, h, facecolor=fc, edgecolor="#2f586b", linewidth=1.4))
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fs, color="#17313f")

    def arrow(s, e):
        ax.add_patch(FancyArrowPatch(s, e, arrowstyle="-|>", mutation_scale=13, lw=1.4, color="#53656f"))

    nodes = [
        (0.6, "规划对象\n工作面/巷道/面积"),
        (3.2, "接续对象\n顺序/工期/产量口径"),
        (5.9, "月度输入\n收入/成本/风险联动成本"),
        (8.8, "评价口径\n月度净现金流"),
    ]
    for x, text in nodes:
        box(x, 3.75, 2.2, 0.92, text, "#eef6f8" if x < 5 else "#fff5e4")
    arrow((2.82, 4.2), (3.15, 4.2))
    arrow((5.43, 4.2), (5.82, 4.2))
    arrow((8.15, 4.2), (8.72, 4.2))
    box(1.2, 1.75, 4.0, 0.78, "NCF_t = Rev_t - Cost_t - RiskCost_t", "#f7fbfd", 10)
    box(6.6, 1.75, 4.0, 0.78, "用于后续经济评价\n本样例不报告独立经济评价数值", "#f7fbfd", 9)
    arrow((5.25, 2.14), (6.55, 2.14))
    ax.text(0.6, 0.48, "注：本图仅说明工程经济评价输入口径，不作为现金流结果或经济优选结论。", fontsize=8, color="#5b6870")
    fig.tight_layout(pad=0.25)
    fig.savefig(out, dpi=220)
    plt.close(fig)


def replace_media(docx_path: Path, replacements: dict[str, Path]):
    tmp = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename in replacements:
                data = replacements[item.filename].read_bytes()
            zout.writestr(item, data)
    tmp.replace(docx_path)


def set_a4(doc: Document):
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)


def write_report():
    REPORT.write_text(
        "\n".join(
            [
                "# 第四轮审稿优化与输出图件取舍清单",
                "",
                "## 已审图件目录",
                "- `D:/xiangmu/miningplan/output/scene_visual_exports/20260416_201037`",
                "",
                "## 图件取舍结论",
                "- 不建议直接插入：`03_mining_planning/overview/02-四模式规划指标对比.*`、`03_mining_planning/overview/03-加权优选候选方案.*`。这些图有完整评分和候选方案信息，但与当前论文正文的3个工作面样例、15个钻孔样点和链路验证口径不一致，直接插入会引发新的数据冲突。",
                "- 不建议直接插入：`05_mining_succession/succession/*`和`05_mining_succession/economics/*`。这些图对应No.1至No.5、136个月、NPV等独立接续经济数据，与当前稿件已明确弱化的结果边界不一致。",
                "- 可作为补充材料候选：各场景`01-ODI分布图.*`、`04-ODI频率分布.*`、`05-ODI分级占比.*`。适合在后续扩展版中展示多场景ODI分布，但当前主文已有图4和ODI统计表，暂不新增以避免版面膨胀。",
                "- 可作为后续重写主案例候选：`03_mining_planning/overview/01-采区规划布局.*`、`05_mining_succession/overview/01-采区规划布局.*`。若后续改写为5工作面完整工程案例，可替换当前3工作面样例体系。",
                "",
                "## 本轮实际修改",
                "- 统一全文公式编号为（1）至（11），并将公式转为PDF中可见的编号文本。",
                "- 统一ODI、多目标函数、接续评价和现金流公式的符号解释。",
                "- 将页面设置由Letter调整为A4。",
                "- 重绘图9，删除“本轮未导出”等内部过程语气。",
                "- 在表3追加WF-01至WF-03面积、推进长度和评分明细，支撑平均规划评分68.8。",
                "- 在摘要中增加样例边界说明，避免把当前结果解释为工程定案。",
            ]
        ),
        encoding="utf-8",
    )


def main():
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if OUT.exists():
        OUT.unlink()
    shutil.copy2(SRC, OUT)
    doc = Document(str(OUT))
    set_a4(doc)
    replace_text(
        doc,
        "结果说明，该方法能够在统一空间参照和统一风险口径下完成参数场构建、风险前置约束和规划对象生成，可为采区规划中的多目标权衡与后续评价提供可复核的工程方法框架。",
        "结果说明，该方法能够在统一空间参照和统一风险口径下完成参数场构建、风险前置约束和规划对象生成，可为采区规划中的多目标权衡与后续评价提供可复核的工程方法框架；当前结果主要用于验证方法链贯通能力，工程定案仍需结合矿井尺度与约束参数进一步校核。",
    )
    replace_text(
        doc,
        "对规划域内任意位置 ，ODI 可表示为",
        "对规划域内任意位置x，ODI可表示为",
    )
    replace_text(
        doc,
        "式中，ODI为覆岩扰动指数；x_i为第i类风险因子的归一化指标值；w_i为相应权重；n为纳入计算的风险因子数量。",
        "式中，ODI为覆岩扰动指数；x_i为第i类风险因子的归一化指标值；w_i为相应权重；n为纳入计算的风险因子数量。",
    )
    replace_text(
        doc,
        "进一步定义候选方案 的 ODI 统计指标为",
        "进一步定义候选方案π的ODI统计指标为",
    )
    replace_text(
        doc,
        "设原始采区边界为 ，边界煤柱约束、区段煤柱约束及保护对象约束对应的缓冲集合分别为 、 和 ，则有效布置域可写为",
        "设原始采区边界为Ω_0，边界煤柱约束、区段煤柱约束及保护对象约束对应的缓冲集合分别为B_b、B_s和D_p，则有效布置域可写为",
    )
    replace_text(
        doc,
        "对给定样点集合 ，煤层厚度等参数场可表示为",
        "对给定样点集合，煤层厚度等参数场可表示为",
    )
    replace_text(
        doc,
        "设候选方案集合为 ，其中每个方案 包括工作面布置、巷道结构、推进方向、尺寸参数和煤柱配置等空间对象。",
        "设候选方案集合为Π，其中每个方案π包括工作面布置、巷道结构、推进方向、尺寸参数和煤柱配置等空间对象。",
    )
    replace_text(
        doc,
        "分别定义方案评价函数 、 和 。其中， 主要表征有效布置域覆盖率、工作面连续性、巷道组织便利性及工程实施稳定性； 主要表征厚度场加权下的可采资源水平与回收潜力； 主要表征方案在 ODI 约束下的扰动控制能力，可由 ODI 均值、高分位值和超限暴露比例综合计算得到。",
        "分别定义方案评价函数S_e(π)、S_r(π)和S_m(π)。其中，S_e(π)主要表征有效布置域覆盖率、工作面连续性、巷道组织便利性及工程实施稳定性；S_r(π)主要表征厚度场加权下的可采资源水平与回收潜力；S_m(π)主要表征方案在ODI约束下的扰动控制能力，可由ODI均值、高分位值和超限暴露比例综合计算得到。",
    )
    replace_text(
        doc,
        "对通过必要安全与几何校核的候选方案，再依据式（8）和式（9）开展模式化排序与优选。",
        "对通过必要安全与几何校核的候选方案，再依据综合目标函数及其权重约束开展模式化排序与优选。",
    )
    replace_text(
        doc,
        "设接续方案为 ，其指标关系可表示为",
        "设接续方案为s，其指标关系可表示为",
    )
    replace_text(
        doc,
        "式中，P_s为产量组织指标，R_s为风险可控性指标，C_s为工期与组织可控性指标；α、β和γ为对应权重，且满足α+β+γ=1。",
        "式中，P_s为产量组织指标，R_s为风险可控性指标，C_s为工期与组织可控性指标；α、β和γ为对应权重，且满足α+β+γ=1。",
    )
    normalize_equations(doc)
    add_score_rows(doc)
    doc.save(str(OUT))

    FIG_DIR.mkdir(parents=True, exist_ok=True)
    fig9 = FIG_DIR / "figure9_economic_input_clean.png"
    draw_economic_input(fig9)
    replace_media(OUT, {"word/media/image10.png": fig9})
    write_report()
    if PDF.exists():
        PDF.unlink()
    print(OUT)


if __name__ == "__main__":
    main()
