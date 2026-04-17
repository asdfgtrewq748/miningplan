import math
import os
import re
import shutil
import statistics
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


DOCX_IN = Path(os.environ["DOCX_IN"])
DOCX_OUT = Path(os.environ["DOCX_OUT"])
DATA_DIR = Path(os.environ["DATA_DIR"])


def load_json_with_keys(*keys):
    import json

    for path in DATA_DIR.glob("*.json"):
        data = json.loads(path.read_text(encoding="utf-8"))
        if isinstance(data, dict) and all(k in data for k in keys):
            return path.name, data
    raise FileNotFoundError(f"No JSON with keys: {keys}")


def percentile(values, pct):
    vals = sorted(values)
    k = (len(vals) - 1) * pct / 100
    f = math.floor(k)
    c = math.ceil(k)
    return vals[f] * (c - k) + vals[c] * (k - f)


def polygon_area(points):
    area = 0.0
    for i, p in enumerate(points):
        q = points[(i + 1) % len(points)]
        area += p["x"] * q["y"] - q["x"] * p["y"]
    return abs(area) / 2


def set_cell(cell, text):
    cell.text = str(text)
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def replace_text(paragraph, replacements):
    text = "".join(run.text for run in paragraph.runs)
    new = text
    for old, repl in replacements:
        new = new.replace(old, repl)
    if new != text:
        paragraph.clear()
        paragraph.add_run(new)


def replace_regex(paragraph, pattern, repl):
    text = "".join(run.text for run in paragraph.runs)
    new = re.sub(pattern, repl, text)
    if new != text:
        paragraph.clear()
        paragraph.add_run(new)


def main():
    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(DOCX_IN, DOCX_OUT)

    _, design = load_json_with_keys("panels", "roadways", "designParams", "miningRules")
    _, odi = load_json_with_keys("field", "weights", "stats")
    _, boundary = load_json_with_keys("boundary")

    boreholes = design["boreholes"]
    thickness = [float(b["coalThickness"]) for b in boreholes]
    design_params = design["designParams"]
    mining_rules = design["miningRules"]
    stats = design["stats"]
    panels = design["panels"]
    roadways = design["roadways"]
    roadway_total = sum(float(r.get("length", 0)) for r in roadways)
    boundary_area = polygon_area(boundary["boundary"])
    coverage = float(stats["totalArea"]) / boundary_area * 100

    field = [float(x) for row in odi["field"] for x in row]
    odi_mean = statistics.mean(field)
    odi_p90 = percentile(field, 90)
    exceed_070 = sum(v > 0.70 for v in field) / len(field) * 100
    exceed_080 = sum(v > 0.80 for v in field) / len(field) * 100
    weights = odi["weights"]
    score_weights = mining_rules["scoreWeights"]

    doc = Document(DOCX_OUT)

    # Priority 1 text tightening: avoid unsupported "optimization/superiority" claims.
    replacements = [
        ("方案比选与优选方法", "方案比选与筛选方法"),
        ("scheme comparison and optimization method", "scheme comparison and selection method"),
        ("scheme optimization in mining-district planning", "scheme comparison and selection in mining-district planning"),
        ("方案优选提供统一的计算框架和技术路径", "方案比选提供统一的计算框架和技术路径"),
        ("形成多模式候选方案生成与比选机制", "形成多模式候选方案生成、筛选与比选机制"),
        ("图2展示了“数据输入—参数体系构建—采区规划输出”的完整过程", "表1汇总了“数据输入—参数体系构建—采区规划输出”的主要参数与约束条件"),
    ]
    for p in doc.paragraphs:
        replace_text(p, replacements)
        replace_text(p, [("当前分析表明，在续采工作面调控过程中，采高降低能够显著压低 ODI 均值及高分位统计量，", "当前样例仅完成调控响应接口展示，采高、工作面宽度和区段煤柱宽度等参数仍需在实矿数据下开展定量敏感性分析；")])
        replace_text(p, [("局部优化", "局部调整")])

    # Normalize figure captions if a user-edit version reuses numbers.
    fig_no = 1
    for p in doc.paragraphs:
        text = "".join(run.text for run in p.runs).strip()
        if re.match(r"^图\s*\d+\s+", text):
            new = re.sub(r"^图\s*\d+", f"图{fig_no}", text)
            p.clear()
            p.add_run(new)
            fig_no += 1

    def replace_table(old_table, rows):
        new_table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        try:
            new_table.style = old_table.style
        except Exception:
            pass
        new_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for r, row in enumerate(rows):
            for c, value in enumerate(row):
                set_cell(new_table.rows[r].cells[c], value)
        old_table._element.addnext(new_table._element)
        old_table._element.getparent().remove(old_table._element)

    # Table 1: fill actual project parameters and remove placeholders.
    table1_rows = [
        ["类别", "参数名称", "符号", "数值/范围", "单位", "说明", "数据来源"],
        ["研究区基础", "钻孔数量", "n", len(boreholes), "个", "钻孔样点总数", "钻孔/设计接口结果"],
        ["研究区基础", "煤层厚度最小值", "h_min", f"{min(thickness):.1f}", "m", "样点统计值", "钻孔数据"],
        ["研究区基础", "煤层厚度最大值", "h_max", f"{max(thickness):.1f}", "m", "样点统计值", "钻孔数据"],
        ["研究区基础", "煤层厚度平均值", "h_avg", f"{statistics.mean(thickness):.4f}", "m", "样点统计值", "钻孔数据"],
        ["几何约束", "边界煤柱宽度", "B_b", f"{design_params['boundaryMargin']:.1f}", "m", "当前样例采用值；规则范围20-50 m", "designParams/miningRules"],
        ["几何约束", "区段煤柱宽度", "B_s", f"{design_params['pillarWidth']:.1f}", "m", "当前样例采用值；规则范围15-30 m", "designParams/miningRules"],
        ["几何约束", "原始边界面积", "A_0", f"{boundary_area:.2f}", "m²", "由6个边界点计算", "边界数据"],
        ["布置参数", "工作面宽度", "W_f", f"{design_params['workfaceWidth']:.1f}", "m", "当前样例采用值", "designParams"],
        ["布置参数", "工作面长度规则", "L_f", f"{mining_rules['faceLength']['min']:.0f}-{mining_rules['faceLength']['max']:.0f}，推荐{mining_rules['faceLength']['preferred']:.0f}", "m", "规则参数", "miningRules"],
        ["布置参数", "最小推进长度", "L_a,min", f"{mining_rules['advanceLength']['min']:.1f}", "m", "当前样例校核阈值", "miningRules"],
        ["布置参数", "推进方向", "D", stats["layoutDirection"], "—", stats["miningMethod"], "设计接口结果"],
        ["风险参数", "ODI权重", "w_d/w_o/w_f", f"{weights['wd']:.2f}/{weights['wo']:.2f}/{weights['wf']:.2f}", "—", "地表沉陷/含水层扰动/上行开采", "ODI场结果"],
        ["风险参数", "ODI统计阈值", "T_ODI", "0.70/0.80", "—", "用于本文超限暴露统计，不作为行业阈值", "本文统计口径"],
        ["评价权重", "工程评分权重", "w", f"厚度{score_weights['coal_thickness']:.2f}，顶板{score_weights['roof_stability']:.2f}，瓦斯{score_weights['gas_content']:.2f}，水{score_weights['water_inflow']:.2f}，构造{score_weights['geological_structure']:.2f}", "—", "设计评分采用的因素权重", "miningRules"],
        ["样例输出", "有效覆盖率", "C", f"{coverage:.2f}", "%", "总布置面积/原始边界面积", "计算值"],
    ]
    replace_table(doc.tables[0], table1_rows)

    # Table 3: fill what is actually exported, and explicitly mark non-exported per-mode metrics.
    table3_rows = [
        ["模式", "工作面数量/个", "巷道数量/条", "巷道总长度/m", "有效覆盖率/%", "可采资源指标", "ODI均值", "ODI P90", "超限暴露比例/%", "接续评分", "经济指标（NPV/万元）", "结果特征"],
        ["工程效率优先", "当前接口未单独导出", "当前接口未单独导出", "当前接口未单独导出", "以连续性、巷道组织为评价口径", "未单独导出", "未单独导出", "未单独导出", "未单独导出", "未单独导出", "未单独导出", "可作为后续模式对比口径"],
        ["资源回收优先", "当前接口未单独导出", "当前接口未单独导出", "当前接口未单独导出", "以厚度场与可采面积为评价口径", f"厚度均值{statistics.mean(thickness):.4f} m", "未单独导出", "未单独导出", "未单独导出", "未单独导出", "未单独导出", "当前仅有参数场支撑"],
        ["覆岩扰动优先", "当前接口未单独导出", "当前接口未单独导出", "当前接口未单独导出", "以ODI均值、P90和超限比例为评价口径", "未单独导出", f"{odi_mean:.4f}", f"{odi_p90:.4f}", f"{exceed_070:.2f}(ODI>0.70)", "未单独导出", "未单独导出", "当前仅有ODI场支撑"],
        ["综合权衡/当前样例输出", str(stats["count"]), str(len(roadways)), f"{roadway_total:.2f}", f"{coverage:.2f}", f"布置面积{float(stats['totalArea']):.2f} m²", f"{odi_mean:.4f}", f"{odi_p90:.4f}", f"{exceed_070:.2f}(ODI>0.70)", "接口已接通，未导出数值", "接口已接通，未导出数值", f"平均评分{float(stats['avgScore']):.1f}，对象链已贯通"],
    ]
    replace_table(doc.tables[2], table3_rows)

    # Table 4: replace empty sensitivity table with actual interface support status.
    table4_rows = [
        ["调控参数", "当前样例取值/范围", "ODI均值", "ODI P90", "超限暴露比例/%", "资源指标变化/%", "评价"],
        ["采高", "接口支持，当前DOCX未单独导出取值序列", f"{odi_mean:.4f}", f"{odi_p90:.4f}", f"{exceed_070:.2f}(ODI>0.70)", "未单独导出", "保留为后续敏感性分析参数，不写强结论"],
        ["工作面宽度", f"{design_params['workfaceWidth']:.1f} m", f"{odi_mean:.4f}", f"{odi_p90:.4f}", f"{exceed_070:.2f}(ODI>0.70)", "未单独导出", "当前样例用于布局生成，未形成宽度序列对比"],
        ["区段煤柱宽度", f"{design_params['pillarWidth']:.1f} m", f"{odi_mean:.4f}", f"{odi_p90:.4f}", f"{exceed_070:.2f}(ODI>0.70)", "未单独导出", "当前样例用于边界/区段隔离，未形成序列对比"],
        ["边界煤柱宽度", f"{design_params['boundaryMargin']:.1f} m", f"{odi_mean:.4f}", f"{odi_p90:.4f}", f"{exceed_070:.2f}(ODI>0.70)", "未单独导出", "当前样例用于有效布置域内缩"],
        ["综合ODI场", f"80×56网格，4480个栅格", f"{odi_mean:.4f}", f"{odi_p90:.4f}", f"{exceed_070:.2f}(ODI>0.70)；{exceed_080:.2f}(ODI>0.80)", "不适用", "支撑风险场描述，不替代实矿参数标定"],
    ]
    replace_table(doc.tables[3], table4_rows)

    # Remove all remaining explicit placeholders.
    placeholder_pairs = [
        ("按实际填写", "当前样例未单独记录"),
        ("如走向/倾向", stats["layoutDirection"]),
        ("绝对最优", "样例条件下最优"),
        ("显著压低 ODI 均值及高分位统计量", "展示调控参数可进入ODI响应分析"),
    ]
    for p in doc.paragraphs:
        replace_text(p, placeholder_pairs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_text(p, placeholder_pairs)

    # Add a compact note after Table 3 if not already present.
    for i, p in enumerate(doc.paragraphs):
        if "".join(r.text for r in p.runs).strip() == "图7 四模式规划指标对比图":
            note = doc.paragraphs[i + 1].insert_paragraph_before(
                "需要说明的是，当前接口结果已导出综合样例的结构化布局、巷道长度、覆盖率与ODI场统计；工程效率优先、资源回收优先和覆岩扰动优先3类模式在本版DOCX中尚未分别导出完整数值，因此表3对其采用评价口径说明，避免将未导出的模式结果误写为实证数值。"
            )
            break

    doc.save(DOCX_OUT)


if __name__ == "__main__":
    main()
