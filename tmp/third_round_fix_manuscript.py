from __future__ import annotations

import json
import math
import shutil
import zipfile
from pathlib import Path

import matplotlib
import numpy as np
from docx import Document

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, Polygon, Rectangle


ROOT = Path.cwd()
PKG = ROOT / "论文" / "重构工作区" / "06_投稿包"
SRC = PKG / "最新版论文4.16_插图版_第二轮精修终版.docx"
OUT = PKG / "最新版论文4.16_插图版_第三轮修订版.docx"
PDF = PKG / "最新版论文4.16_插图版_第三轮修订版.pdf"
DATA_DIR = ROOT / "论文" / "重构工作区" / "05_支撑材料" / "接口结果"
IMG_DIR = ROOT / "tmp" / "third_round_replacement_figures"


def setup_font() -> None:
    plt.rcParams["font.sans-serif"] = [
        "Microsoft YaHei",
        "SimHei",
        "SimSun",
        "Arial Unicode MS",
        "DejaVu Sans",
    ]
    plt.rcParams["axes.unicode_minus"] = False


def load_data() -> tuple[dict, dict]:
    with open(DATA_DIR / "采区设计结果.json", "r", encoding="utf-8") as f:
        design = json.load(f)
    with open(DATA_DIR / "000_mindong_layout_odi_field.json", "r", encoding="utf-8") as f:
        odi = json.load(f)
    return design, odi


def poly_points(items: list[dict]) -> list[tuple[float, float]]:
    return [(p["x"], p["y"]) for p in items]


def bounds_from_design(design: dict) -> tuple[float, float, float, float]:
    xs, ys = [], []
    for p in design["boundary"]:
        xs.append(p["x"])
        ys.append(p["y"])
    for panel in design["panels"]:
        for p in panel["points"]:
            xs.append(p["x"])
            ys.append(p["y"])
    for road in design["roadways"]:
        for p in road["path"]:
            xs.append(p["x"])
            ys.append(p["y"])
    pad = 45
    return min(xs) - pad, max(xs) + pad, min(ys) - pad, max(ys) + pad


def draw_layout(
    out: Path,
    size: tuple[int, int],
    design: dict,
    odi: dict | None = None,
    note: str | None = None,
) -> None:
    w, h = size
    fig, ax = plt.subplots(figsize=(w / 220, h / 220), dpi=220)
    xmin, xmax, ymin, ymax = bounds_from_design(design)

    if odi is not None:
        field = np.array(odi["field"], dtype=float)
        im = ax.imshow(
            field,
            extent=[xmin, xmax, ymin, ymax],
            origin="lower",
            cmap="RdYlBu_r",
            alpha=0.42,
            aspect="auto",
        )
        cb = fig.colorbar(im, ax=ax, shrink=0.78, pad=0.012)
        cb.set_label("ODI", fontsize=8)
        cb.ax.tick_params(labelsize=7)

    boundary = poly_points(design["boundary"])
    ax.add_patch(
        Polygon(
            boundary,
            closed=True,
            fill=False,
            edgecolor="#2f3a45",
            linewidth=2.0,
            label="采区边界",
        )
    )

    colors = ["#b7d7ee", "#bce0c6", "#f6d5a8"]
    for idx, panel in enumerate(design["panels"]):
        pts = poly_points(panel["points"])
        ax.add_patch(
            Polygon(
                pts,
                closed=True,
                facecolor=colors[idx % len(colors)],
                edgecolor="#2b6f9f",
                alpha=0.78,
                linewidth=1.8,
            )
        )
        cx = sum(x for x, _ in pts) / len(pts)
        cy = sum(y for _, y in pts) / len(pts)
        ax.text(
            cx,
            cy,
            f"{panel['id']}\n推进{panel['advanceLength']:.1f} m",
            ha="center",
            va="center",
            fontsize=9,
            color="#1d3343",
            weight="bold",
        )

    road_styles = {
        "main": ("#d34d4d", 2.0, "主运输/主回风"),
        "ventilation": ("#d34d4d", 2.0, None),
        "transport": ("#2f8fa3", 1.4, "运输顺槽"),
        "return": ("#5d8b3d", 1.4, "回风顺槽"),
        "cut": ("#d8a33b", 1.4, "开切眼"),
    }
    used = set()
    for road in design["roadways"]:
        pts = poly_points(road["path"])
        color, lw, label = road_styles.get(road.get("type"), ("#666666", 1.2, None))
        show_label = label if label and label not in used else None
        if label:
            used.add(label)
        ax.plot([p[0] for p in pts], [p[1] for p in pts], color=color, lw=lw, label=show_label)

    for bh in design["boreholes"]:
        ax.scatter(bh["x"], bh["y"], s=18, c="#222222", alpha=0.55, zorder=5)
    ax.scatter([], [], s=18, c="#222222", alpha=0.55, label="钻孔样点")

    ax.set_xlim(xmin, xmax)
    ax.set_ylim(ymin, ymax)
    ax.set_xlabel("X / m", fontsize=9)
    ax.set_ylabel("Y / m", fontsize=9)
    ax.grid(True, color="#d9e2e8", linewidth=0.6, alpha=0.9)
    ax.tick_params(labelsize=8)
    ax.legend(loc="upper right", fontsize=8, framealpha=0.92, ncol=2 if odi else 1)
    if note:
        ax.text(
            0.012,
            0.018,
            note,
            transform=ax.transAxes,
            fontsize=8,
            color="#4b5b66",
            bbox=dict(boxstyle="round,pad=0.35", fc="white", ec="#cbd5dc", alpha=0.92),
        )
    fig.tight_layout(pad=0.4)
    fig.savefig(out, dpi=220)
    plt.close(fig)


def box(ax, xy, width, height, text, fc="#eef5f8", ec="#315a70", fontsize=10):
    x, y = xy
    patch = Rectangle((x, y), width, height, facecolor=fc, edgecolor=ec, linewidth=1.6)
    ax.add_patch(patch)
    ax.text(x + width / 2, y + height / 2, text, ha="center", va="center", fontsize=fontsize, color="#18313f")
    return patch


def arrow(ax, start, end):
    ax.add_patch(
        FancyArrowPatch(
            start,
            end,
            arrowstyle="-|>",
            mutation_scale=14,
            linewidth=1.6,
            color="#51636f",
        )
    )


def draw_transfer(out: Path, size: tuple[int, int], design: dict) -> None:
    w, h = size
    fig, ax = plt.subplots(figsize=(w / 220, h / 220), dpi=220)
    ax.set_axis_off()
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 7)
    box(ax, (0.5, 4.8), 2.3, 1.1, "规划输出对象\n3个工作面 / 11条巷道", "#e9f3f8")
    box(ax, (4.7, 4.8), 2.6, 1.1, "接续任务对象\n边界、顺序、推进长度", "#edf7ed")
    box(ax, (8.9, 4.8), 2.5, 1.1, "后续评价输入\n产量、风险、工期", "#fff4df")
    arrow(ax, (2.85, 5.35), (4.55, 5.35))
    arrow(ax, (7.35, 5.35), (8.75, 5.35))

    y0 = 3.2
    for i, p in enumerate(design["panels"]):
        y = y0 - i * 0.85
        box(ax, (0.7, y), 2.3, 0.55, f"{p['id']}  推进 {p['advanceLength']:.1f} m", "#f7fbfd", fontsize=8.5)
        box(ax, (4.9, y), 2.25, 0.55, f"任务{i + 1}: 几何边界+顺槽", "#f8fcf8", fontsize=8.5)
        box(ax, (9.0, y), 2.25, 0.55, "作为月度测算输入", "#fffaf0", fontsize=8.5)
        arrow(ax, (3.03, y + 0.28), (4.78, y + 0.28))
        arrow(ax, (7.18, y + 0.28), (8.88, y + 0.28))

    ax.text(0.5, 0.55, "注：本图表达对象传递关系，不给出未经独立导出的接续评分。", fontsize=8, color="#5b6870")
    fig.tight_layout(pad=0.25)
    fig.savefig(out, dpi=220)
    plt.close(fig)


def draw_parameter_link(out: Path, size: tuple[int, int]) -> None:
    w, h = size
    fig, ax = plt.subplots(figsize=(w / 220, h / 220), dpi=220)
    ax.set_axis_off()
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 6)
    left = [
        ("工作面宽度", "120.0 m"),
        ("区段煤柱", "20.0 m"),
        ("边界煤柱", "30.0 m"),
        ("推进长度校核", "推荐阈值 800.0 m"),
    ]
    right = [
        ("ODI均值", "0.4669"),
        ("ODI P90", "0.7474"),
        ("ODI>0.70", "15.89%"),
        ("ODI>0.80", "3.55%"),
    ]
    box(ax, (4.2, 4.4), 3.6, 0.8, "工作面级调控变量与ODI统计联动", "#eaf3f7", fontsize=11)
    for i, (k, v) in enumerate(left):
        y = 3.45 - i * 0.75
        box(ax, (0.8, y), 2.7, 0.5, f"{k}\n{v}", "#f7fbfd", fontsize=8.2)
        arrow(ax, (3.55, y + 0.25), (4.95, 2.45))
    box(ax, (4.95, 2.05), 2.1, 0.8, "候选方案\n参数校核与风险统计", "#fff6e8", fontsize=9.2)
    for i, (k, v) in enumerate(right):
        y = 3.45 - i * 0.75
        box(ax, (8.5, y), 2.7, 0.5, f"{k}\n{v}", "#f9fbf4", fontsize=8.2)
        arrow(ax, (7.1, 2.45), (8.45, y + 0.25))
    ax.text(0.8, 0.38, "注：800.0 m为推荐/校核阈值；当前样例低于该阈值，作为后续修正提示，不作为工程定案。", fontsize=8, color="#5b6870")
    fig.tight_layout(pad=0.25)
    fig.savefig(out, dpi=220)
    plt.close(fig)


def draw_economic_input(out: Path, size: tuple[int, int]) -> None:
    w, h = size
    fig, ax = plt.subplots(figsize=(w / 220, h / 220), dpi=220)
    ax.set_axis_off()
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 6)
    nodes = [
        ((0.6, 3.7), "规划对象\n工作面/巷道/面积"),
        ((3.2, 3.7), "接续对象\n顺序/工期/产量口径"),
        ((5.9, 3.7), "月度输入\n收入/成本/风险联动成本"),
        ((8.8, 3.7), "评价指标\n月度净现金流口径"),
    ]
    for xy, text in nodes:
        box(ax, xy, 2.2, 0.95, text, "#eef6f8" if xy[0] < 5 else "#fff5e4", fontsize=9)
    for s, e in [((2.82, 4.18), (3.15, 4.18)), ((5.43, 4.18), (5.82, 4.18)), ((8.15, 4.18), (8.72, 4.18))]:
        arrow(ax, s, e)
    box(ax, (1.2, 1.7), 4.0, 0.8, "NCF_t = Rev_t - Cost_t - RiskCost_t", "#f7fbfd", fontsize=10)
    box(ax, (6.6, 1.7), 4.0, 0.8, "用于后续经济评价，不输出本轮未导出数值", "#f7fbfd", fontsize=9.5)
    arrow(ax, (5.25, 2.1), (6.55, 2.1))
    ax.text(0.6, 0.45, "注：本图仅说明工程经济评价口径；当前样例不报告未经独立导出的现金流曲线或经济评价数值。", fontsize=8, color="#5b6870")
    fig.tight_layout(pad=0.25)
    fig.savefig(out, dpi=220)
    plt.close(fig)


def replace_paragraph(doc: Document, old: str, new: str) -> int:
    count = 0
    for p in doc.paragraphs:
        if p.text.strip() == old:
            p.text = new
            count += 1
    return count


def replace_contains(doc: Document, needle: str, new: str) -> int:
    count = 0
    for p in doc.paragraphs:
        if needle in p.text:
            p.text = new
            count += 1
    return count


def remove_paragraphs_by_xml(doc: Document, needle: str) -> int:
    count = 0
    for p in list(doc.paragraphs):
        if needle in p._p.xml:
            p._element.getparent().remove(p._element)
            count += 1
    return count


def edit_docx() -> None:
    doc = Document(str(SRC))

    replace_contains(
        doc,
        "具体而言，工程效率优先模式以有效布置域覆盖率",
        "具体而言，工程效率、资源回收、覆岩扰动控制和综合权衡可作为同一候选池上的不同评价口径。工程效率口径关注有效布置域覆盖率、工作面连续性、巷道组织便利性和实施稳定性；资源回收口径关注厚度场和可采资源价值；覆岩扰动控制口径关注ODI均值、P90及超限暴露比例；综合权衡口径则在3类目标之间进行权重调节和排序。本文将上述模式作为方法框架和候选方案组织方式说明，当前样例不输出四类模式的独立量化指标。",
    )
    replace_contains(
        doc,
        "从实现流程看，候选方案生成并不是先完全脱离风险场形成几何结果",
        "从实现流程看，候选方案生成并不是先完全脱离风险场形成几何结果，再进行后验校核，而是在有效布置域、参数场和ODI约束共同作用下完成候选池构建。对不满足最小宽度、煤柱隔离和高风险暴露等必要安全底线的候选对象，在生成阶段即予以剔除；对推进长度等受样例边界尺度影响较大的工程推荐阈值，系统记录校核状态并作为方案降级或后续修正依据；对通过必要安全与几何校核的候选方案，再依据式（8）和式（9）开展模式化排序与优选。这样，风险控制不再停留在布局完成后的解释性结论，而是成为候选方案形成阶段的前置筛选条件。",
    )
    replace_contains(
        doc,
        "这意味着规划阶段并不直接产出不可更改的最终工程方案",
        "在调控层面，规划结果可进一步转化为工作面级和生产级控制变量。例如，在含水层扰动场景下，可将工作面采高、工作面宽度、区段煤柱宽度以及推进长度作为主要调控参数，通过比较不同参数条件下ODI均值、高分位值和超限比例的变化，识别对扰动强度最敏感的控制量，并据此开展局部调整。这意味着规划阶段并不直接产出不可更改的最终工程方案，而是先形成满足必要安全与几何底线、并保留推荐阈值校核状态的候选空间对象，再通过参数调控提升其在扰动约束下的适应性。",
    )
    replace_contains(
        doc,
        "在接续评价层面，规划结果还可进一步映射为生产组织对象",
        "在接续评价层面，规划结果还可进一步映射为生产组织对象，并围绕产量、风险和工期形成候选接续方案。设接续方案为 ，其指标关系可表示为",
    )
    replace_contains(
        doc,
        "式中，P_s为产量组织指标",
        "式中，P_s为产量组织指标，R_s为风险可控性指标，C_s为工期与组织可控性指标；α、β和γ为对应权重，且满足α+β+γ=1。该式用于表达规划结果向接续评价传递的指标关系，本文不据此给出未经实证导出的量化评价值。",
    )
    replace_contains(
        doc,
        "进一步地，经济评价通过月度净现金流和净现值指标实现",
        "进一步地，工程经济评价可通过月度净现金流口径承接规划与接续结果。第t月净现金流可表示为",
    )
    replace_contains(
        doc,
        "式中，I_0为初始投入",
        "本文保留工程经济评价的理论传递关系，但不在样例结果中给出未经独立导出的经济评价结论。",
    )
    remove_paragraphs_by_xml(doc, "NPV")
    replace_contains(
        doc,
        "边界煤柱、区段煤柱、工作面宽度、推进长度和ODI统计阈值则共同构成后续规划求解的约束条件。",
        "研究区输入对象包括采区边界、钻孔样点和规划控制参数3类。钻孔样点用于构建煤层厚度等连续参数场；边界对象用于界定有效布置域；边界煤柱、区段煤柱、工作面宽度、推进长度校核阈值和ODI统计阈值则共同构成后续规划求解的约束与校核条件。",
    )
    replace_contains(
        doc,
        "从样例结果看，当前规划共形成3个工作面和11条巷道",
        "从样例结果看，当前规划共形成3个工作面和11条巷道，巷道总长度为4817.50 m，工作面布置面积为176565.72 m²，相对于原始边界的有效覆盖率为69.17%。3个工作面的推进长度分别为WF-01推进515.1 m、WF-02推进536.6 m、WF-03推进551.1 m，平均推进长度为534.3 m。需要说明的是，800.0 m在本样例中作为推荐/校核阈值使用，当前3个工作面的推进长度均低于该阈值，因此该结果用于证明对象生成、风险统计与后续传递链路，不作为真实生产方案的最终工程定案。",
    )
    replace_contains(
        doc,
        "在采掘接续层面，规划阶段形成的工作面边界和推进关系可映射为接续任务对象。",
        "在采掘接续层面，规划阶段形成的工作面边界和推进关系可映射为接续任务对象。当前样例以示意方式给出由采区规划空间结果向采掘接续空间结果的传递路径，其证据重点是“对象可传递、链路可延伸”，而不是未经独立导出的量化优选结论。",
    )
    replace_contains(
        doc,
        "在工程经济评价层面，规划与接续结果可以继续进入收入、成本、风险联动成本和现金流分析过程。",
        "在工程经济评价层面，规划与接续结果可以继续进入收入、成本、风险联动成本和现金流分析过程。当前样例已完成规划对象向后续评价口径的传递，但尚不具备独立导出的真实矿井经济对照数据，因此本文仅说明工程经济评价的输入关系与计算口径，不给出未经数据支撑的现金流曲线或经济优选结论。",
    )

    replace_paragraph(doc, "图7 采掘接续时序安排图", "图7 规划结果向采掘接续传递示意图")
    replace_paragraph(doc, "图8 月产量变化曲线图", "图8 规划对象与评价参数联动示意图")
    replace_paragraph(doc, "图9 采掘接续方案现金流分析图", "图9 工程经济评价输入口径示意图")

    # Table 1: make the 800 m value a check/recommendation threshold, not a hard result constraint.
    tbl = doc.tables[0]
    row = tbl.rows[10].cells
    row[1].text = "推进长度校核阈值"
    row[2].text = "L_a,check"
    row[3].text = "800.0"
    row[5].text = "推荐/校核阈值；当前样例用于提示方案修正，不作为硬性达标结果"

    # Table 3: clarify face width versus advance length.
    tbl = doc.tables[2]
    tbl.rows[4].cells[1].text = "平均工作面宽度"
    tbl.rows[4].cells[4].text = "当前样例工作面宽度统计值"
    tbl.rows[5].cells[4].text = "当前样例统计值；低于800.0 m推荐校核阈值，需后续实矿修正"

    doc.save(str(OUT))


def replace_media(media_map: dict[str, Path]) -> None:
    tmp = OUT.with_suffix(".tmp.docx")
    with zipfile.ZipFile(OUT, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename in media_map:
                data = media_map[item.filename].read_bytes()
            zout.writestr(item, data)
    tmp.replace(OUT)


def write_log() -> None:
    log = ROOT / "论文" / "重构工作区" / "00_过程文档" / "第三轮修订修改记录.md"
    log.write_text(
        "\n".join(
            [
                "# 第三轮修订修改记录",
                "",
                "## 修订对象",
                f"- 源文件：`{SRC}`",
                f"- 修订版：`{OUT}`",
                "",
                "## 已处理问题",
                "- 将表1中的“最小推进长度”调整为“推进长度校核阈值”，明确800.0 m为推荐/校核阈值，不作为当前样例的硬性达标结果。",
                "- 将正文中关于推进长度的硬约束表述调整为“必要安全底线剔除 + 推荐阈值校核提示”的口径。",
                "- 在规划结果段补充说明：3个工作面推进长度低于800.0 m推荐/校核阈值，当前结果用于链路验证，不作为真实生产方案定案。",
                "- 将表3“平均工作面长度”改为“平均工作面宽度”，并补充平均推进长度低于推荐阈值的说明。",
                "- 替换图5和图6为3个工作面版本，消除No.1-No.5与正文3个工作面的冲突。",
                "- 将图7至图9改为传递关系、参数联动和经济评价口径示意，删除对未独立导出的接续评分、月产量曲线和现金流结果的暗示。",
                "- 弱化四模式独立结果表述，明确当前样例不输出四类模式的独立量化指标。",
                "",
            ]
        ),
        encoding="utf-8",
    )


def main() -> None:
    setup_font()
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    design, odi = load_data()

    fig5 = IMG_DIR / "image5_replacement.png"
    fig6 = IMG_DIR / "image6_replacement.png"
    fig7 = IMG_DIR / "image8_replacement.png"
    fig8 = IMG_DIR / "image9_replacement.png"
    fig9 = IMG_DIR / "image10_replacement.png"
    draw_layout(fig5, (2255, 1405), design, odi, "当前样例输出：3个工作面、11条巷道；ODI用于风险叠置与方案校核。")
    draw_layout(fig6, (2135, 1262), design, None, "当前样例输出：WF-01至WF-03；推进长度为推荐/校核阈值项。")
    draw_transfer(fig7, (2379, 1338), design)
    draw_parameter_link(fig8, (2135, 1113))
    draw_economic_input(fig9, (2256, 1172))

    if OUT.exists():
        OUT.unlink()
    if PDF.exists():
        PDF.unlink()
    edit_docx()
    replace_media(
        {
            "word/media/image5.png": fig5,
            "word/media/image6.png": fig6,
            "word/media/image8.png": fig7,
            "word/media/image9.png": fig8,
            "word/media/image10.png": fig9,
        }
    )
    write_log()
    print(OUT)


if __name__ == "__main__":
    main()
