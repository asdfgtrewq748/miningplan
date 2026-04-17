import json
import math
import re
import statistics
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


BASE = Path("D:/xiangmu/miningplan")
DOCX_PATH = BASE / "论文" / "重构工作区" / "06_投稿包" / "最新版论文4.16_插图版_第一优先级修订.docx"
DATA_DIR = BASE / "论文" / "重构工作区" / "05_支撑材料" / "接口结果"


def load_json_with_keys(*keys):
    for path in DATA_DIR.glob("*.json"):
        data = json.loads(path.read_text(encoding="utf-8"))
        if isinstance(data, dict) and all(k in data for k in keys):
            return data
    raise FileNotFoundError(f"No JSON with keys: {keys}")


def polygon_area(points):
    area = 0.0
    for i, p in enumerate(points):
        q = points[(i + 1) % len(points)]
        area += p["x"] * q["y"] - q["x"] * p["y"]
    return abs(area) / 2


def percentile(values, pct):
    vals = sorted(values)
    k = (len(vals) - 1) * pct / 100
    f = math.floor(k)
    c = math.ceil(k)
    return vals[f] * (c - k) + vals[c] * (k - f)


def set_paragraph_text(paragraph, text):
    paragraph.clear()
    paragraph.add_run(text)


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def has_drawing(paragraph):
    return bool(paragraph._element.xpath(".//*[local-name()='drawing']"))


def style_cell(cell, text, font_size=8.5):
    cell.text = str(text)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
            run.font.name = "Times New Roman"


def replace_table(doc, old_table, rows, font_size=8.5):
    new_table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        new_table.style = old_table.style
    except Exception:
        pass
    new_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    new_table.autofit = True
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            style_cell(new_table.rows[r].cells[c], value, font_size=font_size)
    old_table._element.addnext(new_table._element)
    old_table._element.getparent().remove(old_table._element)


def replace_all(doc, replacements):
    for paragraph in doc.paragraphs:
        text = paragraph.text
        new = text
        for old, repl in replacements:
            new = new.replace(old, repl)
        if new != text:
            set_paragraph_text(paragraph, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    new = text
                    for old, repl in replacements:
                        new = new.replace(old, repl)
                    if new != text:
                        set_paragraph_text(paragraph, new)


def delete_four_mode_result_block(doc):
    # Remove the unsupported four-mode result subsection, its chart, and the caveat paragraph.
    targets = []
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text.startswith("3.2 四模式候选规划"):
            targets.extend([i, i + 1, i + 2])
        if text.startswith("图7 四模式规划指标对比图"):
            if i > 0 and has_drawing(doc.paragraphs[i - 1]):
                targets.append(i - 1)
            targets.extend([i, i + 1])
        if text.startswith("同时也应看到，当前样例级结果"):
            targets.append(i)
    for i in sorted(set(t for t in targets if 0 <= t < len(doc.paragraphs)), reverse=True):
        remove_paragraph(doc.paragraphs[i])


def renumber_figure_captions(doc):
    number = 1
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if re.match(r"^图\d+\s+", text):
            set_paragraph_text(paragraph, re.sub(r"^图\d+", f"图{number}", text))
            number += 1


def main():
    design = load_json_with_keys("panels", "roadways", "designParams", "miningRules")
    odi = load_json_with_keys("field", "weights", "stats")
    boundary = load_json_with_keys("boundary")

    stats = design["stats"]
    roadways = design["roadways"]
    design_params = design["designParams"]
    weights = odi["weights"]
    boundary_area = polygon_area(boundary["boundary"])
    coverage = float(stats["totalArea"]) / boundary_area * 100
    roadway_total = sum(float(r.get("length", 0)) for r in roadways)
    field = [float(x) for row in odi["field"] for x in row]
    odi_mean = statistics.mean(field)
    odi_p90 = percentile(field, 90)
    exceed_070 = sum(v > 0.70 for v in field) / len(field) * 100
    exceed_080 = sum(v > 0.80 for v in field) / len(field) * 100

    doc = Document(DOCX_PATH)

    replacements = [
        ("四模式规划结果以及规划结果向采掘接续和工程经济评价的传递过程", "结构化规划结果与ODI风险统计结果"),
        ("四模式候选规划与后续评价链路", "结构化规划结果与ODI风险统计"),
        ("后续四模式规划结果的形成提供了统一的风险底图", "规划结果的风险约束提供了统一底图"),
        ("3 四模式规划结果与方案比选", "3 结构化规划结果与风险统计"),
        ("3.3 规划结果向采掘接续与工程经济评价的传递", "3.2 规划结果向采掘接续与工程经济评价的传递"),
        ("本文进一步开展候选方案生成与多模式规划比选。与传统只给出单一布局结果的做法不同，本文将候选方案组织为工程效率优先、资源回收优先、覆岩扰动优先和综合权衡优化4种规划模式。四种模式并非4套彼此独立的算法，而是面向同一候选方案池的不同评价视角，其目的在于为采区规划提供可比、可选和可解释的方案集合。", "本文进一步开展采区规划结果生成与风险约束检验。与传统只给出静态边界划分的做法不同，本文将钻孔样点、连续参数场、保护煤柱约束和ODI风险场共同纳入同一对象体系，使规划结果能够以工作面、巷道和覆盖范围等结构化对象输出。"),
    ]
    replace_all(doc, replacements)
    delete_four_mode_result_block(doc)

    table3_rows = [
        ["统计类别", "指标", "数值", "单位/口径", "说明"],
        ["工程布置", "工作面数量", stats["count"], "个", "当前样例输出"],
        ["工程布置", "巷道数量", len(roadways), "条", "当前样例输出"],
        ["工程布置", "巷道总长度", f"{roadway_total:.2f}", "m", "按巷道对象长度汇总"],
        ["工程布置", "平均工作面长度", f"{float(stats['avgFaceLength']):.1f}", "m", "当前样例统计值"],
        ["工程布置", "平均推进长度", f"{float(stats['avgAdvanceLength']):.1f}", "m", "当前样例统计值"],
        ["资源覆盖", "布置面积", f"{float(stats['totalArea']):.2f}", "m²", "工作面布置面积汇总"],
        ["资源覆盖", "有效覆盖率", f"{coverage:.2f}", "%", "布置面积/原始边界面积"],
        ["ODI风险", "ODI均值", f"{odi_mean:.4f}", "—", "80×56网格统计"],
        ["ODI风险", "ODI P90", f"{odi_p90:.4f}", "—", "90%分位值"],
        ["ODI风险", "ODI>0.70比例", f"{exceed_070:.2f}", "%", "本文统计口径"],
        ["ODI风险", "ODI>0.80比例", f"{exceed_080:.2f}", "%", "本文统计口径"],
        ["综合评价", "平均评分", f"{float(stats['avgScore']):.1f}", "分", "按设计评分权重汇总"],
    ]
    replace_table(doc, doc.tables[2], table3_rows, font_size=8.5)

    table4_rows = [
        ["对象", "当前样例取值", "关联指标", "统计值", "用途"],
        ["工作面宽度", f"{float(design_params['workfaceWidth']):.1f} m", "几何布置", "用于工作面生成", "约束工作面空间尺度"],
        ["区段煤柱宽度", f"{float(design_params['pillarWidth']):.1f} m", "安全隔离", "规则范围15-30 m", "约束区段间距"],
        ["边界煤柱宽度", f"{float(design_params['boundaryMargin']):.1f} m", "边界保护", "规则范围20-50 m", "形成有效布置域"],
        ["ODI权重", f"{weights['wd']:.2f}/{weights['wo']:.2f}/{weights['wf']:.2f}", "沉陷/含水层/上行开采", "综合风险加权", "形成统一风险场"],
        ["ODI栅格", "80×56，4480个栅格", "风险统计", f"均值{odi_mean:.4f}；P90={odi_p90:.4f}", "描述研究区风险分布"],
        ["高风险暴露", "ODI阈值0.70/0.80", "超限比例", f"{exceed_070:.2f}%/{exceed_080:.2f}%", "识别高扰动敏感区"],
    ]
    replace_table(doc, doc.tables[3], table4_rows, font_size=8.5)

    renumber_figure_captions(doc)
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
