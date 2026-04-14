"""
Generate Chinese SCI Paper First Draft with Figures (Word .docx)
Target: 煤炭科学技术 / Comprehensive (method + cases)
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT = Path(r"D:\xiangmu\miningplan\data\output\submission_package\paper_draft.docx")
PNG = Path(r"D:\xiangmu\miningplan\data\output\all_png")
SUPP = Path(r"D:\xiangmu\miningplan\data\output\supplementary_figures")


def p(doc, text, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
      indent=True, after=6):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = 1.5
    if indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    run = para.add_run(text)
    run.font.name = "宋体"
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return para


def heading(doc, text, level=1):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.name = "黑体"
    run.font.size = Pt({1: 15, 2: 13, 3: 12}.get(level, 12))
    run.bold = True
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    return para


def fig(doc, img_path, caption, width=14.5):
    """Insert a figure with caption."""
    if not Path(img_path).exists():
        p(doc, f"[图片缺失: {img_path}]", size=9, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    # Image
    para_img = doc.add_paragraph()
    para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_img.paragraph_format.space_before = Pt(6)
    para_img.paragraph_format.space_after = Pt(2)
    run = para_img.add_run()
    run.add_picture(str(img_path), width=Cm(width))
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    run_c = cap.add_run(caption)
    run_c.font.name = "宋体"
    run_c.font.size = Pt(9)
    run_c._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def fig_dual(doc, img1, img2, caption, w=7.0):
    """Insert two figures side by side."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run()
    if Path(img1).exists():
        run.add_picture(str(img1), width=Cm(w))
    run2 = para.add_run()
    if Path(img2).exists():
        run2.add_picture(str(img2), width=Cm(w))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    rc = cap.add_run(caption)
    rc.font.name = "宋体"
    rc.font.size = Pt(9)
    rc._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # ═══════════════════════════════════════════════════════════
    #  TITLE
    # ═══════════════════════════════════════════════════════════
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(4)
    r = t.add_run("基于多源地质参数的覆岩扰动指数评价方法与工程应用")
    r.font.name = "黑体"; r.font.size = Pt(18); r.bold = True
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.paragraph_format.space_after = Pt(12)
    r = st.add_run("——七类典型采矿场景的验证与分析")
    r.font.name = "楷体"; r.font.size = Pt(13)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")

    ap = doc.add_paragraph()
    ap.alignment = WD_ALIGN_PARAGRAPH.CENTER; ap.paragraph_format.space_after = Pt(2)
    r = ap.add_run("作者1  作者2  作者3"); r.font.name = "宋体"; r.font.size = Pt(12)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    af = doc.add_paragraph()
    af.alignment = WD_ALIGN_PARAGRAPH.CENTER; af.paragraph_format.space_after = Pt(16)
    r = af.add_run("(1. 单位名称，省份 城市 邮编；2. 单位名称，省份 城市 邮编)")
    r.font.name = "宋体"; r.font.size = Pt(9)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # ═══════════════════════════════════════════════════════════
    #  ABSTRACT
    # ═══════════════════════════════════════════════════════════
    heading(doc, "摘  要", level=1)

    p(doc, (
        "煤矿地下开采引起的覆岩扰动和含水层破坏是制约煤炭绿色开采的关键工程地质问题。"
        "现有评价方法多依赖单一指标或经验公式，难以综合反映地质条件、开采参数和覆岩结构"
        "等多因素耦合作用下扰动程度的空间分布特征。针对上述问题，本文提出了一种基于多源"
        "地质参数的覆岩扰动指数（Overburden Disturbance Index, ODI）综合评价方法。"
        "该方法从钻孔柱状图和开采设计参数中提取岩层厚度Ti、距含水层距离Hi、开采深度Di、"
        "采高Mi等8个评价因子，通过反距离加权（IDW）和自然邻点插值构建连续地质参数场；"
        "建立8×9权重映射矩阵将输入参数转化为覆岩破坏（Smax、DSmax）、应力变形（Ksi、"
        "Dsi、Asi）和导水裂隙（Hf、Kw、Bf、Af）三个维度的9个中间指标，经三级加权聚合"
        "得到ODI值；最后通过归一化和五级分区（I~V级）实现覆岩扰动的空间量化评价。基于"
        "该方法开发了采矿扰动评估可视化系统，并应用于地表下沉、含水层扰动评价、采区规划、"
        "协同调控、采掘接续和全覆岩扰动等7个典型工程案例（评价点32~1976个）。结果表明："
        "（1）ODI空间分布与实测导水裂隙带高度和地表下沉等值线具有良好一致性，Case 0三条"
        "测线的平均误差比为0.13~0.21；（2）含水层扰动场景中导水裂隙维度权重wf=0.60，"
        "距含水层距离Hi是控制导水裂隙发育的关键因子（Pearson r≈-0.6~-0.8）；（3）七类"
        "案例的ODI均值变化范围为0.325~0.567，符合工程实际规律。本研究为煤矿覆岩扰动定量"
        "评价提供了系统化的技术方法，可服务于保水采煤和绿色开采决策。"
    ), size=10)

    kw = doc.add_paragraph()
    kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    kw.paragraph_format.first_line_indent = Cm(0.74)
    r = kw.add_run("关键词："); r.font.name = "黑体"; r.font.size = Pt(10)
    r.bold = True; r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r2 = kw.add_run("覆岩扰动指数；多源地质参数；空间插值；含水层保护；保水采煤；采矿扰动评价")
    r2.font.name = "宋体"; r2.font.size = Pt(10)
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # ═══════════════════════════════════════════════════════════
    #  1. 引言
    # ═══════════════════════════════════════════════════════════
    heading(doc, "1  引言", level=1)

    p(doc, (
        "煤炭地下开采引起上覆岩层移动、变形和破断，进而诱发地表沉陷、含水层破坏和生态"
        "环境损害等工程地质问题[1-3]。我国西部矿区煤炭资源丰富但生态环境脆弱，煤层开采对"
        "地下水系统的扰动尤为显著，保水采煤已成为煤炭绿色开发的核心需求[4-5]。如何科学、"
        "定量地评价覆岩扰动的空间分布特征与强度等级，对于优化开采设计、保护含水层和生态"
        "环境具有重要的理论意义和工程应用价值。"
    ))

    p(doc, (
        "目前覆岩扰动评价方法主要包括经验公式法、数值模拟法和综合指标法三类。经验公式法"
        "以\"三带\"理论为基础，依据《建筑物、水体、铁路及主要井巷煤柱留设与压煤开采规范》"
        "推荐公式计算导水裂隙带发育高度[6]，计算简便但仅适用于标准地质条件，难以反映局部"
        "地质异常。数值模拟法采用有限元或离散元方法模拟覆岩应力场和破坏场演化[7-8]，精度"
        "较高但对地质模型和力学参数要求严格，不便于多方案快速比选。综合指标法将地质条件和"
        "开采参数纳入统一评价框架[9-10]，近年来受到广泛关注，但现有方法在空间插值精度、"
        "多因素权重确定和结果可视化方面仍有不足。"
    ))

    p(doc, (
        "针对上述问题，本文提出了基于多源地质参数的覆岩扰动指数（ODI）综合评价方法。"
        "核心创新包括：（1）建立了8参数输入->9指标中间层->3维度加权聚合的权重映射矩阵，"
        "实现多因素耦合的扰动量化；（2）采用IDW和自然邻点双重空间插值策略构建连续扰动场；"
        "（3）提出了基于分位数截断的统一标尺归一化方法（ODI*），实现跨场景横向可比；"
        "（4）开发了交互式可视化评估系统，并通过7个典型工程案例进行了验证。"
    ))

    # ═══════════════════════════════════════════════════════════
    #  2. 研究方法
    # ═══════════════════════════════════════════════════════════
    heading(doc, "2  研究方法", level=1)

    heading(doc, "2.1  ODI评价模型框架", level=2)

    p(doc, (
        "ODI评价模型采用\"参数提取-空间插值-指标映射-加权聚合-归一化分区\"五步技术路线。"
        "首先从钻孔柱状图和开采设计方案中提取地质与工程参数；通过空间插值将离散钻孔数据"
        "转化为连续参数场；在评价点处采样参数并经权重矩阵映射得到三个维度的中间指标；"
        "最终通过三级加权聚合得到ODI值，并进行归一化和等级分区。"
    ))

    fig(doc, PNG / "00_surface_fig03_spatial_map.png",
        "图1  评价点、钻孔、工作面和边界空间分布（以Case 0地表下沉场景为例）", width=12)

    heading(doc, "2.2  评价参数体系", level=2)

    p(doc, (
        "ODI模型的输入参数包括4类地质参数和4类工程参数，共8个评价因子。"
        "地质参数包括：①岩层厚度Ti——目标层位岩层累计厚度（m），取多目标层最大值（偏安全）；"
        "②距含水层距离Hi——煤层与含水层间隔水层累计厚度（m），取最小值（最不利条件）；"
        "③开采深度Di——目标层位顶板埋深（m）；④弹性模量Ei——由岩性推断的弹性模量（GPa），"
        "取最小值。上述参数通过钻孔柱状图分层信息计算。"
    ))

    p(doc, (
        "工程参数包括：⑤有效采高Mi——取煤层厚度与设计采高的较小值（m），非工作面区域Mi=0；"
        "⑥顶板垮落角δ——工作面顶板断裂角（°）；⑦工作面长度lpi——对中心线点取工作面短边"
        "长度，其他点取2倍到边缘最小距离（m）；⑧区段煤柱宽度lci——相邻工作面中心线间煤柱"
        "宽度（m），仅中心线点赋值。8个参数在钻孔点计算后通过IDW或自然邻点插值生成连续场。"
    ))

    fig(doc, PNG / "00_surface_fig02_geology_cloud.png",
        "图2  地质参数空间插值云图（Case 0：Ti、Hi、Di、Mi四参数viridis配色）", width=14)

    heading(doc, "2.3  权重映射矩阵与ODI计算", level=2)

    p(doc, (
        "ODI计算的核心是将8个输入参数映射为9个中间指标的权重映射矩阵W（8×9）。矩阵基于"
        "层次分析法（AHP）和专家判断确定，每行对应一个输入参数，每列对应一个中间指标。"
        "9个中间指标按物理意义分为三个维度：（1）覆岩破坏维度Dd，含最大下沉量Smax和差异"
        "下沉量DSmax；（2）应力变形维度Do，含应力集中系数Ksi、变形指数Dsi和加速度指数Asi；"
        "（3）导水裂隙维度Df，含导水裂隙带高度Hf、渗透系数变化Kw、裂隙宽度Bf和裂隙面积"
        "密度Af。"
    ))

    p(doc, (
        "映射过程为矩阵乘法：[I1×9] = [P1×8] × [W8×9]。使用前矩阵W每列归一化（列和为1），"
        "并剔除退化行后重新归一化。ODI最终计算公式为："
    ))
    p(doc, "    ODI(p) = wd × Dd(p) + wo × Do(p) + wf × Df(p)",
      indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)

    p(doc, (
        "式中wd、wo、wf分别为三个维度的权重系数，满足wd+wo+wf=1。地表下沉场景取"
        "wd=0.45、wo=0.30、wf=0.25，侧重覆岩破坏评价；含水层扰动场景取wd=0.15、wo=0.25、"
        "wf=0.60，突出导水裂隙维度。"
    ))

    heading(doc, "2.4  归一化与等级分区", level=2)

    p(doc, (
        "ODI原始值经min-max归一化映射到[0,1]区间：ODInorm = (ODIraw - minODI)/(maxODI - minODI)。"
        "跨场景对比时采用基于联合分位数截断的统一标尺（ODI*），将不同评价阶段ODI样本合并，"
        "取5%和95%分位数进行Winsorize处理后再线性映射到[0,1]，消除场景间量纲差异。"
    ))

    p(doc, (
        "归一化后ODI值按阈值分为五级：I级（稳定，ODI<0.40）、II级（轻微，0.40≤ODI<0.65）、"
        "III级（中等，0.65≤ODI<0.85）、IV级（较强，0.85≤ODI<0.90）、V级（强扰动，ODI≥0.90）。"
        "含水层场景采用固定阈值；地表下沉场景采用动态规划（DP）优化阈值，以实测分级为基准"
        "最小化分类误差。"
    ))

    heading(doc, "2.5  空间插值方法", level=2)

    p(doc, (
        "本研究采用两种空间插值方法。反距离加权法（IDW）取幂参数p=2、最近邻k=24，适用于"
        "地表下沉等规则边界场景；自然邻点插值法（Sibson）基于Delaunay三角剖分和Voronoi图"
        "面积比确定权重，取k=32，适用于含水层扰动等不规则边界场景，样本不足时自动回退IDW。"
        "两种方法均支持0~6次平滑处理（默认2次），网格分辨率80×56或60×42。"
    ))

    # ═══════════════════════════════════════════════════════════
    #  3. 系统开发
    # ═══════════════════════════════════════════════════════════
    heading(doc, "3  评价系统开发", level=1)

    p(doc, (
        "基于上述ODI方法开发了Web端采矿扰动评估可视化系统，采用前后端分离架构。前端基于"
        "React框架，提供钻孔数据导入、参数提取、ODI热力图生成、误差分析和等级分区等交互"
        "功能；后端基于Python FastAPI框架，提供文件上传、地质解析、智能规划和数据导出等"
        "API服务。系统核心功能包括：（1）CSV格式钻孔柱状图导入与自动岩层解析；（2）基于"
        "IDW插值的高密度评价点自动生成和8参数提取；（3）ODI热力图实时渲染（cubic插值，"
        "blueRed 5步离散配色）；（4）测线级误差分析（MAE、RMSE、误差比）；（5）I~V级"
        "扰动分区与统计汇总。系统内置7个典型工程案例，可直接加载验证。"
    ))

    # ═══════════════════════════════════════════════════════════
    #  4. 工程案例验证
    # ═══════════════════════════════════════════════════════════
    heading(doc, "4  工程案例验证与分析", level=1)

    heading(doc, "4.1  案例概况", level=2)

    p(doc, (
        "本文选取7个典型采矿工程案例（表2），覆盖煤矿开采扰动评价的主要应用场景。数据来源于"
        "实际工程钻孔资料和开采设计方案，包含钻孔8~40个、工作面1~7个、评价点32~1976个。"
        "Case 0（地表下沉）布设8个钻孔和3条实测线共88个测点，用于预测-实测对比验证；"
        "Case 1/2（含水层预评价/评价）分别采用32和376个评价点；Case 4（协同调控）采用577个"
        "评价点进行多工作面联合评价；Case 5（采掘接续）包含1976个评价点；Case 6（全覆岩扰动）"
        "采用16个评价点评估覆岩全厚综合扰动状态。"
    ))

    heading(doc, "4.2  ODI空间分布特征", level=2)

    p(doc, (
        "图3给出了代表性案例的ODI空间分布热力图。ODI高值区（红色）集中分布于工作面采空区"
        "上方及邻近区域，低值区（蓝色）位于远离采动影响的区域，空间分布规律符合覆岩移动和"
        "破坏的一般理论认识。钻孔附近的ODI等值线呈现明显的局部梯度变化，反映了钻孔地质"
        "参数的空间不均匀性对扰动评价结果的影响。"
    ))

    fig_dual(doc, PNG / "00_surface_fig01_odi_heatmap.png",
             PNG / "04_aquifer_fig01_odi_heatmap.png",
             "图3  ODI空间分布热力图  (a) Case 0 地表下沉（185点）  (b) Case 4 协同调控（577点）", w=7.0)

    p(doc, (
        "表3的统计结果表明，7个案例ODI均值变化范围为0.325~0.567，标准差0.255~0.363。"
        "Case 1（含水层预评价）均值最高（0.567），表明32个评价点整体扰动水平较高；"
        "Case 0均值最低（0.325）但偏度最大，呈右偏分布，说明大部分区域扰动较小但局部存在"
        "强扰动集中区。Case 5（1976点）ODI均值0.441、P90达0.880，反映了多工作面连续开采"
        "的累积扰动效应。"
    ))

    fig(doc, SUPP / "figS3_cross_case_comparison.png",
        "图4  七类工程案例ODI统计量横向对比（均值±标准差、P90、最大值及样本量）", width=14)

    heading(doc, "4.3  参数敏感性分析", level=2)

    p(doc, (
        "图5给出了各案例中4个主要参数与ODI的Pearson相关系数（详细相关矩阵见图S1）。"
        "在含水层扰动场景（Case 1~5）中，距含水层距离Hi与ODI呈强负相关（r≈-0.6~-0.8），"
        "表明隔水层厚度是控制导水裂隙发育的关键因素——Hi越小（隔水层越薄），扰动越强，"
        "这与导水裂隙带发育高度受基岩厚度控制的认识一致[6]。采高Mi与ODI呈正相关，反映"
        "开采强度对扰动的放大效应：Mi越大，采空区空间越大，顶板垮落和裂隙发育越充分。"
        "岩层厚度Ti呈弱正相关，深度Di与ODI的相关性方向因场景而异，反映覆岩扰动受深度效应"
        "和岩性效应的耦合影响。"
    ))

    fig(doc, SUPP / "figS4_sensitivity_analysis.png",
        "图5  参数敏感性分析——各案例中Ti、Hi、Di、Mi与ODI的Pearson相关系数", width=14)

    heading(doc, "4.4  误差分析", level=2)

    p(doc, (
        "以Case 0为验证对象，利用3条实测线88个测点数据对ODI预测结果进行误差分析（表4）。"
        "图6显示测线1和测线2的平均误差比分别为0.132和0.136，最大误差比约0.34；测线3平均"
        "误差比0.212、最大0.460，误差略大，可能与该测线靠近工作面边界、地质条件变化较大"
        "有关。图S2的实测-预测散点图显示ODI归一化值与实测归一化值之间存在线性趋势但离散"
        "性较大，后续可通过引入断层、褶皱等构造约束提高局部预测精度。"
    ))

    fig(doc, PNG / "00_surface_fig07a_error_trend_line1.png",
        "图6  Case 0测线1误差趋势图——ODI归一化值（蓝色实线）、实测值（红色虚线）与误差比（灰色柱）", width=14)

    heading(doc, "4.5  等级分区结果", level=2)

    p(doc, (
        "图7和表5给出了ODI等级分区统计。Case 0采用DP优化阈值，分区为[0, 0.045)、"
        "[0.045, 0.345)、[0.345, 0.825)、[0.825, 0.847)、[0.847, 1.0]，70.8%评价点处于"
        "II级、25.4%处于V级，呈两极分化。其余案例采用等间距阈值，Case 5中58.4%处于II级、"
        "26.4%处于V级；Case 4中57.2%处于I级，反映不同开采方案下扰动分布的显著差异。"
    ))

    fig_dual(doc, PNG / "00_surface_fig04_odi_histogram.png",
             PNG / "00_surface_fig05_odi_level_pie.png",
             "图7  Case 0 ODI频次直方图（左）与等级分区饼图（右）", w=7.0)

    heading(doc, "4.6  权重系数对比", level=2)

    p(doc, (
        "图8和表2对比了各案例的ODI权重配置。地表下沉场景（Case 0、6）采用wd=0.45、wo=0.30、"
        "wf=0.25，侧重覆岩破坏维度；含水层扰动场景（Case 1~5）采用wd=0.15、wo=0.25、wf=0.60，"
        "将导水裂隙维度权重提升至0.60。这种场景化权重调整机制使ODI方法可灵活适应不同评价"
        "目标，但权重确定目前依赖专家经验，未来可引入数据驱动方法进行优化。"
    ))

    fig(doc, PNG / "00_surface_fig08_weight_radar.png",
        "图8  Case 0 ODI权重雷达图（wd=0.45地质, wo=0.30采矿, wf=0.25综合）", width=8)

    # ═══════════════════════════════════════════════════════════
    #  5. 讨论
    # ═══════════════════════════════════════════════════════════
    heading(doc, "5  讨论", level=1)

    heading(doc, "5.1  方法优势与局限性", level=2)

    p(doc, (
        "相较于传统方法，ODI方法具有四方面优势：（1）多因素耦合——通过8×9权重矩阵实现"
        "地质与工程参数耦合评价，克服了单一指标法的局限；（2）空间连续性——基于IDW和自然"
        "邻点插值的连续参数场和ODI场，可反映扰动强度的空间梯度变化；（3）场景适应性——通过"
        "调整权重系数和分区阈值灵活适应不同评价目标；（4）可视化交互——基于Web的系统降低了"
        "使用门槛。局限性在于：权重矩阵基于AHP专家判断，存在主观性；插值精度受钻孔密度"
        "控制；当前模型未考虑断层、褶皱等构造因素的局部放大效应；局部预测精度仍有提升空间。"
    ))

    heading(doc, "5.2  与现有方法比较", level=2)

    p(doc, (
        "与经验公式法相比，ODI方法通过多参数耦合评价全面反映地质条件空间变化，评价结果"
        "具有空间连续性。与数值模拟法相比，ODI方法计算效率高，秒级完成数千评价点的计算和"
        "热力图渲染，适用于多方案快速比选。与综合指标法[9-10]相比，ODI方法建立了系统的"
        "参数提取-权重映射-等级分区技术框架，并开发了完整的可视化系统，工程实用性更强。"
    ))

    heading(doc, "5.3  工程应用建议", level=2)

    p(doc, (
        "基于7个案例验证结果提出以下建议：（1）含水层保护评价中应重点关注Hi<100m区域，"
        "该区导水裂隙发育风险较高；（2）采高Mi>10m时ODI值显著增大，含水层下开采应严格控制"
        "采高；（3）多工作面连续开采（如Case 5）显著增大累积扰动范围，工作面间应预留足够"
        "煤柱宽度；（4）保水采煤设计中建议将ODI评价结果作为工作面布置和参数优化的参考依据。"
    ))

    # ═══════════════════════════════════════════════════════════
    #  6. 结论
    # ═══════════════════════════════════════════════════════════
    heading(doc, "6  结论", level=1)

    p(doc, (
        "（1）提出了基于多源地质参数的覆岩扰动指数（ODI）评价方法，建立了\"8参数输入-"
        "9指标中间层-3维度加权聚合\"的评价框架。该方法通过权重映射矩阵实现地质与工程参数"
        "耦合，通过IDW和自然邻点插值保证空间连续性，通过场景化权重和动态分区阈值灵活适应"
        "不同评价需求。"
    ))

    p(doc, (
        "（2）开发了基于Web的采矿扰动评估可视化系统，集成钻孔数据管理、参数自动提取、"
        "ODI热力图生成、误差分析和等级分区等功能。系统内置7个典型工程案例，涵盖地表下沉、"
        "含水层扰动、采区规划、协同调控、采掘接续和全覆岩扰动等场景。"
    ))

    p(doc, (
        "（3）7个工程案例验证表明：ODI空间分布与实测覆岩扰动规律吻合良好；距含水层距离"
        "Hi是控制导水裂隙发育的关键因子（r≈-0.6~-0.8）；含水层场景中导水裂隙维度权重"
        "wf=0.60，显著高于地质（wd=0.15）和采矿（wo=0.25）维度；Case 0三条测线平均误差比"
        "0.13~0.21，满足工程评价精度要求。"
    ))

    p(doc, (
        "（4）后续研究将从以下方向深化：①引入断层、褶皱等构造因素提高局部精度；②采用"
        "数据驱动方法优化权重矩阵减少主观性；③建立ODI与覆岩破坏高度、导水裂隙带高度的"
        "定量关系模型增强物理可解释性；④开展长期监测数据验证评估时间维度的预测能力。"
    ))

    # ═══════════════════════════════════════════════════════════
    #  SUPPLEMENTARY FIGURES
    # ═══════════════════════════════════════════════════════════
    heading(doc, "附录：补充图", level=1)

    fig(doc, SUPP / "figS1_correlation_heatmap.png",
        "图S1  Pearson相关系数矩阵——Ti、Hi、Di、Mi与ODI（全案例聚合）", width=11)

    fig(doc, SUPP / "figS2_measured_vs_predicted.png",
        "图S2  实测与预测ODI归一化值1:1散点图（Case 0三条测线，含R²和线性拟合）", width=14)

    # ═══════════════════════════════════════════════════════════
    #  REFERENCES
    # ═══════════════════════════════════════════════════════════
    heading(doc, "参考文献", level=1)

    for ref in [
        "[1] 钱鸣高, 石平五. 矿山压力与岩层控制[M]. 徐州: 中国矿业大学出版社, 2003.",
        "[2] 缪协兴, 钱鸣高. 中国煤炭资源绿色开采研究现状与展望[J]. 采矿与安全工程学报, 2009, 26(1): 1-6.",
        "[3] 王双明, 黄庆享, 范立民, 等. 生态脆弱矿区含(隔)水层特征及保水开采分区方法[J]. 煤炭学报, 2010, 35(1): 7-12.",
        "[4] 范立民. 保水采煤研究进展及亟待解决的关键问题[J]. 煤炭科学技术, 2021, 49(9): 1-10.",
        "[5] 黄庆享. 浅埋煤层覆岩隔水性与保水开采分类[J]. 煤炭学报, 2020, 45(1): 138-146.",
        "[6] 国家煤炭工业局. 建筑物、水体、铁路及主要井巷煤柱留设与压煤开采规范[M]. 北京: 煤炭工业出版社, 2017.",
        "[7] 谢和平, 周宏伟, 刘建锋, 等. 岩石破坏中的分形与能量特征[J]. 力学学报, 2008, 40(4): 545-553.",
        "[8] 王家臣, 杨胜利. 综放开采顶煤运移规律与覆岩破坏特征研究[J]. 煤炭学报, 2015, 40(5): 987-993.",
        "[9] 李文平, 叶贵钧, 张莱, 等. 陕北榆神府矿区保水采煤工程地质条件综合评价[J]. 工程地质学报, 2011, 19(5): 711-718.",
        "[10] 马雄德, 王文科, 范立民, 等. 生态脆弱矿区采煤对含水层的影响与保护技术[J]. 煤炭科学技术, 2020, 48(4): 1-8.",
        "[11] SIBSON R. A brief description of natural neighbor interpolation[C]//Interpreting Multivariate Data. New York: Wiley, 1981: 21-36.",
        "[12] SHEPARD D. A two-dimensional interpolation function for irregularly-spaced data[C]//Proceedings of the 1968 23rd ACM National Conference. New York: ACM, 1968: 517-524.",
        "[13] SAATY T L. The analytic hierarchy process[M]. New York: McGraw-Hill, 1980.",
    ]:
        p(doc, ref, size=9, indent=False)

    doc.save(str(OUT))
    print(f"[OK] {OUT}")


if __name__ == "__main__":
    main()
