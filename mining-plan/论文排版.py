"""
论文一键排版脚本

用法示例（需先安装 python-docx）：
  pip install python-docx
  python 论文排版.py --template 样例模版.docx --input 待排版.docx --output 已排版.docx

思路：
1) 读取提供的 Word 模版，提取页面设置与常用样式（正文、标题、图表题注等）。
2) 按样式名将模版的段落/字体/行距/缩进等属性直接复制到目标文档的段落与表格。
3) 复制模版首页的页面尺寸与页边距到目标文档每个节，尽量还原排版。

注意：
- 本脚本不拷贝模版中的实际内容，只同步格式；保持目标文档原有文字顺序。
- 样式名以模版为准；常见中文名（如“标题 1”）也会尝试匹配。
- 若模版缺少某个样式，则跳过该样式的复制。
"""

from __future__ import annotations

import argparse
from pathlib import Path
from typing import Dict, Iterable, Optional

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.shared import OxmlElement


# 需要重点关注的样式名（英文与常见中文别名）
STYLE_CANDIDATES = {
	"Normal": ["Normal", "正文", "正文文本"],
	"Heading 1": ["Heading 1", "标题 1", "标题1"],
	"Heading 2": ["Heading 2", "标题 2", "标题2"],
	"Heading 3": ["Heading 3", "标题 3", "标题3"],
	"Caption": ["Caption", "题注"],
	"Table Grid": ["Table Grid", "表格网格"],
}


def _find_style(doc: Document, aliases: Iterable[str]):
	for name in aliases:
		try:
			style = doc.styles[name]
			if style is not None:
				return style
		except KeyError:
			continue
	return None


def copy_paragraph_style(template_style, paragraph: Paragraph):
	"""将模版样式的段落与字体属性复制到目标段落。"""
	if template_style is None:
		return

	# 段落级格式
	tpl_pf = template_style.paragraph_format
	pf = paragraph.paragraph_format
	for attr in [
		"left_indent",
		"right_indent",
		"first_line_indent",
		"space_before",
		"space_after",
		"line_spacing",
		"line_spacing_rule",
		"keep_together",
		"keep_with_next",
		"widow_control",
		"tab_stops",
		"alignment",
	]:
		try:
			value = getattr(tpl_pf, attr)
			if value is not None:
				setattr(pf, attr, value)
		except Exception:
			continue

	# 文字级格式（对每个 run 覆盖）
	tpl_font = template_style.font
	for run in paragraph.runs:
		rf = run.font
		for attr in [
			"name",
			"size",
			"bold",
			"italic",
			"underline",
			"color",
			"all_caps",
			"small_caps",
		]:
			try:
				value = getattr(tpl_font, attr)
				if value is not None:
					setattr(rf, attr, value)
			except Exception:
				continue

	# 对齐方式如果未设置，保持原状；若模版定义了，优先模版
	if template_style.paragraph_format.alignment is not None:
		paragraph.alignment = template_style.paragraph_format.alignment


def copy_table_style(table: Table, table_style_name: Optional[str]):
	if table_style_name:
		try:
			table.style = table_style_name
		except Exception:
			pass

	# 统一表格单元格的垂直对齐与内边距（尽量还原模版效果）
	for row in table.rows:
		for cell in row.cells:
			tc_pr = cell._tc.get_or_add_tcPr()
			v_align = OxmlElement("w:vAlign")
			v_align.set("w:val", "center")
			tc_pr.append(v_align)
			for paragraph in cell.paragraphs:
				paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def copy_page_setup(src: Document, dst: Document):
	if not src.sections:
		return
	tpl_sec = src.sections[0]
	for sec in dst.sections:
		sec.page_height = tpl_sec.page_height
		sec.page_width = tpl_sec.page_width
		sec.left_margin = tpl_sec.left_margin
		sec.right_margin = tpl_sec.right_margin
		sec.top_margin = tpl_sec.top_margin
		sec.bottom_margin = tpl_sec.bottom_margin
		sec.header_distance = tpl_sec.header_distance
		sec.footer_distance = tpl_sec.footer_distance


def apply_template(template_path: Path, input_path: Path, output_path: Path):
	tpl_doc = Document(str(template_path))
	tgt_doc = Document(str(input_path))

	# 收集模版可用样式
	style_map: Dict[str, object] = {}
	for key, aliases in STYLE_CANDIDATES.items():
		style_map[key] = _find_style(tpl_doc, aliases)

	# 复制页面设置
	copy_page_setup(tpl_doc, tgt_doc)

	# 处理段落
	for para in tgt_doc.paragraphs:
		name = para.style.name if para.style is not None else ""
		match = None
		for key, aliases in STYLE_CANDIDATES.items():
			if name in aliases or name == key:
				match = key
				break
		if match is None:
			# 默认用正文样式定义
			match = "Normal"
		copy_paragraph_style(style_map.get(match), para)

	# 处理表格
	table_style = None
	for key in ("Table Grid",):
		if style_map.get(key):
			table_style = style_map[key].name
			break
	for table in tgt_doc.tables:
		copy_table_style(table, table_style)
		for row in table.rows:
			for cell in row.cells:
				for para in cell.paragraphs:
					copy_paragraph_style(style_map.get("Normal"), para)

	tgt_doc.save(str(output_path))


def parse_args():
	parser = argparse.ArgumentParser(description="论文一键排版：按模版同步格式")
	parser.add_argument("--template", required=True, help="参考模版 Word 路径 (.docx)")
	parser.add_argument("--input", required=True, help="待排版的 Word 路径 (.docx)")
	parser.add_argument("--output", required=True, help="输出文件路径 (.docx)")
	return parser.parse_args()


def main():
	args = parse_args()
	template_path = Path(args.template)
	input_path = Path(args.input)
	output_path = Path(args.output)

	if not template_path.exists():
		raise FileNotFoundError(f"模版不存在: {template_path}")
	if not input_path.exists():
		raise FileNotFoundError(f"输入文件不存在: {input_path}")

	output_path.parent.mkdir(parents=True, exist_ok=True)
	apply_template(template_path, input_path, output_path)
	print(f"✅ 已完成排版: {output_path}")


if __name__ == "__main__":
	main()
