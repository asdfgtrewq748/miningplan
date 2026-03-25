from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from latex2mathml.converter import convert as latex转mathml
from lxml import etree


正文引用模式 = re.compile(r"\[(\d+(?:\s*[-,，]\s*\d+)*)\]")
参考文献行模式 = re.compile(r"^\[(\d+)\]\s*(.+)$|^(\d+)\.\s*(.+)$")
书签计数器 = 0


def 去除行内标记(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def 设置中文字体(run, font_name: str, size: float, bold: bool = False, superscript: bool = False) -> None:
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.superscript = superscript


def 添加文本到段落(
    paragraph,
    text: str,
    font_name: str,
    size: float,
    bold: bool = False,
    superscript: bool = False,
) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    设置中文字体(run, font_name, size, bold, superscript=superscript)


def 设置段落基础格式(paragraph, first_line_indent_cm: float = 0.74, line_spacing_pt: float = 15) -> None:
    pf = paragraph.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_spacing_pt)
    pf.first_line_indent = Cm(first_line_indent_cm)


def 设置分栏(section, num: int = 1, space_twips: int = 720) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        cols_el = cols[0]
    else:
        cols_el = OxmlElement("w:cols")
        sect_pr.append(cols_el)
    cols_el.set(qn("w:num"), str(num))
    cols_el.set(qn("w:space"), str(space_twips))


def 提取标题(md_path: Path) -> str:
    for line in md_path.read_text(encoding="utf-8").splitlines():
        if line.strip().startswith("# "):
            return 去除行内标记(line.strip()[2:].strip())
    raise ValueError(f"未在 {md_path} 中找到一级标题")


def 获取mml2omml样式表路径() -> Path:
    candidates = [
        Path(r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"),
        Path(r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL"),
    ]
    for path in candidates:
        if path.exists():
            return path
    raise FileNotFoundError("未找到 MML2OMML.XSL，无法生成 Word 可编辑公式。")


MML2OMML_TRANSFORM = etree.XSLT(etree.parse(str(获取mml2omml样式表路径())))


def latex转omml元素(latex: str):
    mathml = latex转mathml(latex)
    omml = MML2OMML_TRANSFORM(etree.fromstring(mathml.encode("utf-8")))
    return parse_xml(etree.tostring(omml.getroot(), encoding="unicode"))


def 生成书签编号() -> str:
    global 书签计数器
    书签计数器 += 1
    return str(书签计数器)


def 创建运行元素(
    text: str,
    font_name: str,
    size: float,
    bold: bool = False,
    superscript: bool = False,
):
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), font_name)
    rpr.append(rfonts)

    if bold:
        rpr.append(OxmlElement("w:b"))

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rpr.append(sz)

    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), str(int(size * 2)))
    rpr.append(sz_cs)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rpr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    rpr.append(underline)

    if superscript:
        valign = OxmlElement("w:vertAlign")
        valign.set(qn("w:val"), "superscript")
        rpr.append(valign)

    run.append(rpr)
    text_el = OxmlElement("w:t")
    if text != text.strip():
        text_el.set(qn("xml:space"), "preserve")
    text_el.text = text
    run.append(text_el)
    return run


def 添加内部超链接(paragraph, text: str, anchor: str, font_name: str, size: float, superscript: bool = False) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    hyperlink.append(创建运行元素(text, font_name, size, superscript=superscript))
    paragraph._p.append(hyperlink)


def 添加上角标文本(paragraph, text: str, font_name: str = "Times New Roman", size: float = 7.5) -> None:
    添加文本到段落(paragraph, text, font_name, size, superscript=True)


def 添加上角标引用(paragraph, citation: str, size: float = 7.5) -> None:
    inner = citation[1:-1]
    添加上角标文本(paragraph, "[", size=size)
    for part in re.findall(r"\d+|[^\d]+", inner):
        if part.isdigit():
            添加内部超链接(paragraph, part, f"ref-{part}", "Times New Roman", size, superscript=True)
        else:
            添加上角标文本(paragraph, part, size=size)
    添加上角标文本(paragraph, "]", size=size)


def 添加文本与引用(paragraph, text: str, font_name: str, size: float, bold: bool = False) -> None:
    parts = re.split(r"(\[(?:\d+(?:\s*[-,，]\s*\d+)*)\])", text)
    for part in parts:
        if not part:
            continue
        if 正文引用模式.fullmatch(part):
            添加上角标引用(paragraph, part)
        else:
            添加文本到段落(paragraph, part, font_name, size, bold)


def 添加含行内公式段落内容(paragraph, text: str, font_name: str, size: float, bold: bool = False) -> None:
    parts = re.split(r"(\$[^$]+\$)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("$") and part.endswith("$") and len(part) >= 2:
            latex = part[1:-1].strip()
            try:
                paragraph._p.append(latex转omml元素(latex))
            except Exception:
                添加文本到段落(paragraph, latex, font_name, size, bold)
        else:
            添加文本与引用(paragraph, part, font_name, size, bold)


def 初始化文档(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)
    sec.start_type = WD_SECTION.NEW_PAGE
    设置分栏(sec, 1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(9)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(15)
    pf.first_line_indent = Cm(0.74)
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)


def 添加普通段落(
    doc: Document,
    text: str,
    center: bool = False,
    bold_prefix: str | None = None,
    font_name: str = "宋体",
    size: float = 9,
    first_line_indent_cm: float = 0.74,
    line_spacing_pt: float = 15,
) -> None:
    p = doc.add_paragraph()
    设置段落基础格式(p, first_line_indent_cm=first_line_indent_cm, line_spacing_pt=line_spacing_pt)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
    if bold_prefix and text.startswith(bold_prefix):
        添加文本到段落(p, bold_prefix, "黑体", size, True)
        添加含行内公式段落内容(p, text[len(bold_prefix):], font_name, size, False)
        return
    添加含行内公式段落内容(p, text, font_name, size, False)


def 添加标题(doc: Document, text: str, level: int, is_first_title: bool) -> None:
    p = doc.add_paragraph()
    设置段落基础格式(p, first_line_indent_cm=0, line_spacing_pt=15)
    p.paragraph_format.first_line_indent = Cm(0)
    if is_first_title:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        设置中文字体(run, "黑体", 18, True)
        return
    size = {2: 12, 3: 10.5}.get(level, 9)
    run = p.add_run(text)
    设置中文字体(run, "黑体", size, True)


def 添加列表(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    设置段落基础格式(p, first_line_indent_cm=0, line_spacing_pt=15)
    添加含行内公式段落内容(p, text, "宋体", 9, False)


def 添加公式(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    设置段落基础格式(p, first_line_indent_cm=0, line_spacing_pt=15)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(7.2), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.SPACES)
    formula_text = " ".join(line.strip() for line in lines if line.strip())
    formula_text = formula_text.replace("\\qquad", " ")
    match = re.search(r"（\d+）\s*$", formula_text)
    eq_no = match.group(0) if match else ""
    latex = formula_text[:match.start()].strip() if match else formula_text.strip()
    p._p.append(latex转omml元素(latex))
    if eq_no:
        run = p.add_run(f"\t{eq_no}")
        设置中文字体(run, "Times New Roman", 9, False)


def 添加图片(doc: Document, md_dir: Path, alt: str, rel_path: str, body_two_columns: bool) -> None:
    img_path = (md_dir / rel_path).resolve()
    if not img_path.exists():
        添加普通段落(doc, f"{alt}（图片未找到：{img_path.name}）", center=True, first_line_indent_cm=0, size=8)
        return
    if img_path.suffix.lower() not in {".png", ".jpg", ".jpeg", ".bmp"}:
        添加普通段落(doc, f"{alt}（矢量原稿见配套文件：{img_path.name}）", center=True, first_line_indent_cm=0, size=8)
        return
    p = doc.add_paragraph()
    设置段落基础格式(p, first_line_indent_cm=0, line_spacing_pt=12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(img_path), width=Cm(7.2 if body_two_columns else 15.0))

    cp = doc.add_paragraph()
    设置段落基础格式(cp, first_line_indent_cm=0, line_spacing_pt=12)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.first_line_indent = Cm(0)
    cr = cp.add_run(alt)
    设置中文字体(cr, "宋体", 8, False)


def 解析参考文献行(text: str) -> tuple[str, str] | None:
    match = 参考文献行模式.match(text)
    if not match:
        return None
    if match.group(1):
        return match.group(1), match.group(2)
    return match.group(3), match.group(4)


def 添加参考文献(doc: Document, 编号: str, text: str) -> None:
    p = doc.add_paragraph()
    设置段落基础格式(p, first_line_indent_cm=0, line_spacing_pt=12)
    p.paragraph_format.first_line_indent = Cm(-0.74)
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(0.74), alignment=WD_TAB_ALIGNMENT.LEFT, leader=WD_TAB_LEADER.SPACES)

    书签编号 = 生成书签编号()
    书签开始 = OxmlElement("w:bookmarkStart")
    书签开始.set(qn("w:id"), 书签编号)
    书签开始.set(qn("w:name"), f"ref-{编号}")
    p._p.append(书签开始)

    添加文本到段落(p, f"[{编号}]\t", "Times New Roman", 7.5, False)
    添加含行内公式段落内容(p, text, "宋体", 7.5, False)

    书签结束 = OxmlElement("w:bookmarkEnd")
    书签结束.set(qn("w:id"), 书签编号)
    p._p.append(书签结束)


def markdown转docx(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    初始化文档(doc)

    math_mode = False
    math_lines: list[str] = []
    first_title_done = False
    body_two_columns = False
    in_references = False

    for line in lines:
        raw = line.rstrip()
        stripped = raw.strip()

        if math_mode:
            if stripped == "$$":
                添加公式(doc, math_lines)
                math_lines = []
                math_mode = False
            else:
                math_lines.append(raw)
            continue

        if not stripped:
            continue

        if stripped == "$$":
            math_mode = True
            math_lines = []
            continue

        image_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", stripped)
        if image_match:
            添加图片(doc, md_path.parent, 去除行内标记(image_match.group(1)) or "图件", image_match.group(2), body_two_columns)
            continue

        if stripped.startswith("# "):
            添加标题(doc, 去除行内标记(stripped[2:].strip()), 1, not first_title_done)
            first_title_done = True
            continue

        if stripped.startswith("## "):
            heading_text = 去除行内标记(stripped[3:].strip())
            in_references = heading_text == "参考文献"
            if not body_two_columns and heading_text == "0 引言":
                section = doc.add_section(WD_SECTION.CONTINUOUS)
                section.page_width = Cm(21.0)
                section.page_height = Cm(29.7)
                section.top_margin = Cm(2.5)
                section.bottom_margin = Cm(2.2)
                section.left_margin = Cm(1.8)
                section.right_margin = Cm(1.8)
                设置分栏(section, 2)
                body_two_columns = True
            添加标题(doc, heading_text, 2, False)
            continue

        if stripped.startswith("### "):
            添加标题(doc, 去除行内标记(stripped[4:].strip()), 3, False)
            continue

        if in_references:
            参考文献信息 = 解析参考文献行(去除行内标记(stripped))
            if 参考文献信息:
                编号, 内容 = 参考文献信息
                添加参考文献(doc, 编号, 内容)
                continue

        if re.match(r"^\d+\.\s", stripped):
            添加列表(doc, 去除行内标记(stripped))
            continue

        if stripped.startswith("- "):
            添加列表(doc, "• " + 去除行内标记(stripped[2:]))
            continue

        text_line = 去除行内标记(stripped.replace("  ", " "))
        if text_line.startswith("摘要："):
            添加普通段落(doc, text_line, bold_prefix="摘要：", size=9, first_line_indent_cm=0)
        elif text_line.startswith("关键词："):
            添加普通段落(doc, text_line, bold_prefix="关键词：", size=9, first_line_indent_cm=0)
        elif text_line.startswith("作者："):
            添加普通段落(doc, text_line, center=True, font_name="仿宋", size=10.5, first_line_indent_cm=0)
        elif text_line.startswith("单位："):
            添加普通段落(doc, text_line, center=True, font_name="宋体", size=8, first_line_indent_cm=0, line_spacing_pt=12)
        elif text_line.startswith("基金项目：") or text_line.startswith("中图分类号：") or text_line.startswith("文献标志码："):
            添加普通段落(doc, text_line, center=False, font_name="宋体", size=8, first_line_indent_cm=0, line_spacing_pt=12)
        else:
            添加普通段落(doc, text_line)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main() -> int:
    if len(sys.argv) == 3:
        md_path = Path(sys.argv[1]).resolve()
        docx_path = Path(sys.argv[2]).resolve()
    else:
        base = Path(__file__).resolve().parents[1]
        md_path = base / "04_论文稿件" / "采区智能规划设计一体化方法与系统.md"
        title = 提取标题(md_path)
        docx_path = base / "04_论文稿件" / f"{title}.docx"

    markdown转docx(md_path, docx_path)
    print(f"已生成：{docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
