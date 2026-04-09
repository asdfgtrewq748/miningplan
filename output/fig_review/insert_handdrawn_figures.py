from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph


DOCX_PATH = Path(r"E:\xiangmu\miningplan\论文\重构工作区\06_投稿包\采区智能规划设计一体化方法与系统_终稿收口版.docx")
FLOW_PATH = Path(r"E:\xiangmu\miningplan\论文\流程图.png")
TECH_PATH = Path(r"E:\xiangmu\miningplan\论文\技术路线图.png")
BACKUP_PATH = DOCX_PATH.with_name(DOCX_PATH.stem + "_插图前备份.docx")
FALLBACK_PATH = DOCX_PATH.with_name(DOCX_PATH.stem + "_插图版.docx")


def clear_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def insert_paragraph_after(paragraph: Paragraph, text: str | None = None, style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def set_text(paragraph: Paragraph, text: str) -> None:
    style = paragraph.style
    alignment = paragraph.alignment
    clear_paragraph(paragraph)
    paragraph.style = style
    paragraph.alignment = alignment
    paragraph.add_run(text)


def replace_exact_or_prefix(paragraphs: list[Paragraph], exact: str, new_text: str, prefix: bool = False) -> None:
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        matched = text.startswith(exact) if prefix else text == exact
        if matched:
            set_text(paragraph, new_text)
            return
    raise ValueError(f"Paragraph not found: {exact}")


def main() -> None:
    shutil.copy2(DOCX_PATH, BACKUP_PATH)

    doc = Document(str(DOCX_PATH))
    paragraphs = doc.paragraphs

    caption_style = paragraphs[34].style
    body_style = paragraphs[70].style

    # Replace the existing Figure 1 image with the hand-drawn workflow.
    fig1_image_para = paragraphs[33]
    clear_paragraph(fig1_image_para)
    fig1_image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig1_image_para.add_run().add_picture(str(FLOW_PATH), width=Inches(6.25))

    # Insert the ODI technical-route figure before Table 2.
    anchor = paragraphs[65]
    image_para = insert_paragraph_after(anchor, style=fig1_image_para.style)
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.add_run().add_picture(str(TECH_PATH), width=Inches(6.05))

    caption_para = insert_paragraph_after(image_para, style=caption_style)
    caption_para.alignment = paragraphs[34].alignment
    caption_para.add_run("图5 ODI 风险组织与扰动分级技术路线图")

    body_para = insert_paragraph_after(caption_para, style=body_style)
    body_para.alignment = paragraphs[70].alignment
    body_para.add_run(
        "图5给出了多场景 ODI 风险组织的技术路线。图中将地质因素与开采因素统一映射到输入层，"
        "再经位移响应、力学响应和水力响应构建覆岩扰动综合评价指标，并最终形成可用于规划阶段"
        "分级筛选的 ODI 表征结果。该图进一步说明，本文中的 ODI 并非单一经验指标，而是连接输"
        "入参数、响应特征和风险分级的统一组织框架。"
    )

    spacer = insert_paragraph_after(body_para, style=paragraphs[66].style)
    spacer.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Renumber downstream figures and references.
    paragraphs = doc.paragraphs
    replace_exact_or_prefix(paragraphs, "图5 四模式智能规划协同关系图", "图6 四模式智能规划协同关系图")
    replace_exact_or_prefix(
        paragraphs,
        "图5展示了候选池生成、多目标评分和四模式筛选之间的协同关系。",
        "图6展示了候选池生成、多目标评分和四模式筛选之间的协同关系。与传统“只输出一个结果方案”的方式相比，这种组织方式更符合采区规划中“多方案并行比较、人工与算法协同决策”的工程实际[11-15]，也为后续接续评价和经济比较保留了更充分的方案空间。",
        prefix=True,
    )
    replace_exact_or_prefix(paragraphs, "图6 规划-接续-经济闭环评价图", "图7 规划-接续-经济闭环评价图")
    replace_exact_or_prefix(
        paragraphs,
        "图6表明，本文方法并未把几何布局视为终点，",
        "图7表明，本文方法并未把几何布局视为终点，而是把规划结果继续组织为接续分析和经济评价可直接调用的上游对象。由此，传统上后置的风险校核与经济比较被提前纳入同一条规划闭环之中，评价结果也能够反向作用于方案修订，而不再只是末端解释性结论。",
        prefix=True,
    )
    replace_exact_or_prefix(paragraphs, "图7 采区综合空间信息与规划结果可视化界面", "图8 采区综合空间信息与规划结果可视化界面")
    replace_exact_or_prefix(
        paragraphs,
        "图7给出了更贴近论文论证口径的组合式结果图。",
        "图8给出了更贴近论文论证口径的组合式结果图。其中，子图 (a) 将采区边界、钻孔样点、16-3煤 厚度场与布局阶段 ODI 风险等值线置于同一空间底图中，说明原始采样信息、连续参数场与风险表达已经能够在统一坐标语义下叠合；子图 (b) 则以工作面组和关键巷道为主，展示候选规划结果如何以结构化几何对象形式输出。与单独的布局截图相比，这张组合图更能支撑本文的核心论点，即规划链条中的空间场、风险场和几何对象并非彼此分离的中间产物，而是能够在同一表达界面中保持一致口径，并继续传递到人工复核、接续分析与经济评价环节。需要同时说明的是，接口校核仍给出了“推进长度小于最小值 800.0 m”的提示，因此该结果更适合被解释为“候选设计对象可生成、可表达、可传递”的样例证据，而非真实矿井条件下的最终工程终判结果。",
        prefix=True,
    )

    saved_to = DOCX_PATH
    try:
        doc.save(str(DOCX_PATH))
    except PermissionError:
        doc.save(str(FALLBACK_PATH))
        saved_to = FALLBACK_PATH

    print(f"backup={BACKUP_PATH}")
    print(f"saved={saved_to}")


if __name__ == "__main__":
    main()
