from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
import collections
import json
import re


BASE = Path(__file__).resolve().parents[1]
OUT = BASE / "format_audit_20260524"
FILES = {
    "template": OUT / "template_from_download.docx",
    "paper": BASE / "煤科论文5.23格式改稿.docx",
}


def twips_to_cm(value):
    return round((value or 0) / 567, 2) if value is not None else None


def font_name(run):
    rpr = run._element.rPr
    if rpr is not None and rpr.rFonts is not None:
        return (
            rpr.rFonts.get(qn("w:eastAsia"))
            or rpr.rFonts.get(qn("w:ascii"))
            or rpr.rFonts.get(qn("w:hAnsi"))
        )
    return None


def para_outline(paragraph):
    runs = []
    for run in paragraph.runs:
        if not run.text:
            continue
        runs.append(
            {
                "text": run.text[:80],
                "font": font_name(run),
                "size": round(run.font.size.pt, 1) if run.font.size else None,
                "bold": run.bold,
                "italic": run.italic,
                "underline": bool(run.underline)
                if run.underline is not None
                else None,
            }
        )
    fmt = paragraph.paragraph_format
    return {
        "text": paragraph.text.strip(),
        "style": paragraph.style.name if paragraph.style else None,
        "alignment": str(paragraph.alignment),
        "left_indent_cm": twips_to_cm(fmt.left_indent.twips if fmt.left_indent else None),
        "first_indent_cm": twips_to_cm(
            fmt.first_line_indent.twips if fmt.first_line_indent else None
        ),
        "space_before_pt": round(fmt.space_before.pt, 1) if fmt.space_before else None,
        "space_after_pt": round(fmt.space_after.pt, 1) if fmt.space_after else None,
        "line_spacing": fmt.line_spacing,
        "runs": runs[:6],
    }


def collect_references(paragraphs):
    ref_start = None
    for idx, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if re.fullmatch(r"(参考文献|References)\s*[:：]?", text, re.I) or text.startswith(
            "参考文献"
        ):
            ref_start = idx
            break

    refs = []
    if ref_start is None:
        return refs

    current = None
    for paragraph in paragraphs[ref_start + 1 :]:
        text = paragraph.text.strip()
        if not text:
            continue
        match = re.match(r"^\s*\[?(\d{1,3})\]?\s*[\.、]?\s*(.*)", text)
        starts_numbered = text.startswith("[") or bool(re.match(r"^\d{1,3}[\.、]\s", text))
        if match and starts_numbered:
            if current:
                refs.append(current)
            current = {
                "num": int(match.group(1)),
                "text": text,
                "style": paragraph.style.name if paragraph.style else None,
                "outline": para_outline(paragraph),
            }
        elif current:
            current["text"] += " " + text
        else:
            refs.append(
                {
                    "num": None,
                    "text": text,
                    "style": paragraph.style.name if paragraph.style else None,
                    "outline": para_outline(paragraph),
                }
            )
    if current:
        refs.append(current)
    return refs


def find_reference_start(paragraphs):
    for idx, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if re.fullmatch(r"(参考文献|References)\s*[:：]?", text, re.I) or text.startswith(
            "参考文献"
        ):
            return idx
    return None


def scan_intext_citations(paragraphs):
    full_text = "\n".join(paragraph.text for paragraph in paragraphs)
    citations = []
    for match in re.finditer(r"(?<!图)(?<!表)(?<!式)\[(\d+(?:\s*[-–,，]\s*\d+)*)\]", full_text):
        token = match.group(1)
        nums = []
        for part in re.split(r"[,，]", token):
            part = part.strip()
            range_match = re.match(r"^(\d+)\s*[-–]\s*(\d+)$", part)
            if range_match:
                start, end = map(int, range_match.groups())
                nums.extend(range(start, end + 1))
            elif part.isdigit():
                nums.append(int(part))
        context = full_text[max(0, match.start() - 60) : match.end() + 60].replace("\n", " ")
        citations.append({"token": "[" + token + "]", "nums": nums, "context": context})
    return citations


def doc_report(path):
    doc = Document(path)
    sections = []
    for section in doc.sections:
        sections.append(
            {
                "page_w_cm": twips_to_cm(section.page_width.twips),
                "page_h_cm": twips_to_cm(section.page_height.twips),
                "top_cm": twips_to_cm(section.top_margin.twips),
                "bottom_cm": twips_to_cm(section.bottom_margin.twips),
                "left_cm": twips_to_cm(section.left_margin.twips),
                "right_cm": twips_to_cm(section.right_margin.twips),
                "header_cm": twips_to_cm(section.header_distance.twips),
                "footer_cm": twips_to_cm(section.footer_distance.twips),
                "different_first_page": section.different_first_page_header_footer,
            }
        )

    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    styles = collections.Counter(p.style.name if p.style else "" for p in paragraphs)
    markers = []
    marker_terms = [
        "摘要",
        "关键词",
        "Abstract",
        "Keywords",
        "引言",
        "结论",
        "参考文献",
        "References",
        "基金项目",
        "作者简介",
        "中图分类号",
        "文献标志码",
    ]
    for idx, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if any(term in text for term in marker_terms) or re.match(
            r"^\d+(\.\d+)*\s*[^\d\s]", text
        ):
            markers.append(
                {
                    "idx": idx + 1,
                    "text": text[:180],
                    "style": paragraph.style.name if paragraph.style else None,
                    "outline": para_outline(paragraph),
                }
            )

    tables = []
    for table_idx, table in enumerate(doc.tables, start=1):
        row_count = len(table.rows)
        col_count = max((len(row.cells) for row in table.rows), default=0)
        first = ""
        if row_count:
            first = " | ".join(
                cell.text.strip().replace("\n", " ")[:50] for cell in table.rows[0].cells
            )
        tables.append(
            {
                "idx": table_idx,
                "rows": row_count,
                "cols": col_count,
                "first_row": first,
            }
        )

    ref_start = find_reference_start(paragraphs)
    body_paragraphs = paragraphs[:ref_start] if ref_start is not None else paragraphs

    return {
        "path": str(path),
        "sections": sections,
        "paragraph_count": len(paragraphs),
        "styles": styles.most_common(),
        "sample": [
            {"idx": idx + 1, **para_outline(paragraph)}
            for idx, paragraph in enumerate(paragraphs[:100])
        ],
        "markers": markers[:160],
        "references": collect_references(paragraphs),
        "intext_citations": scan_intext_citations(body_paragraphs),
        "tables": tables,
    }


def main():
    reports = {name: doc_report(path) for name, path in FILES.items()}
    (OUT / "docx_audit_raw.json").write_text(
        json.dumps(reports, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    for name, report in reports.items():
        lines = [
            f"== {name} ==",
            f"paragraphs: {report['paragraph_count']}",
            f"sections: {report['sections']}",
            "styles: " + str(report["styles"][:25]),
            "",
            "-- markers --",
        ]
        for marker in report["markers"]:
            lines.append(f"{marker['idx']}: {marker['style']} | {marker['text']}")
        lines.append("")
        lines.append("-- references --")
        for ref in report["references"]:
            lines.append(f"{ref.get('num')}: {ref['style']} | {ref['text'][:500]}")
        lines.append("")
        lines.append("-- intext citation tokens --")
        for citation in report["intext_citations"][:240]:
            lines.append(f"{citation['token']} | {citation['context']}")
        (OUT / f"{name}_audit_summary.txt").write_text(
            "\n".join(lines), encoding="utf-8"
        )

    print(
        json.dumps(
            {
                name: {
                    "paragraphs": report["paragraph_count"],
                    "refs": len(report["references"]),
                    "citations": len(report["intext_citations"]),
                    "sections": report["sections"][:1],
                    "top_styles": report["styles"][:8],
                    "tables": len(report["tables"]),
                }
                for name, report in reports.items()
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
