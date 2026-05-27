from copy import deepcopy
from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED
import re
import shutil
import tempfile

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph
from lxml import etree


BASE = Path(__file__).resolve().parents[1]
SOURCE = BASE / "煤科论文5.23格式改稿.docx"
OUT_DIR = BASE / "format_audit_20260524"
OUTPUT = BASE / "煤科论文5.23格式改稿_格式问题修正版.docx"
TMP = OUT_DIR / "_format_fix_stage.docx"


TEMPLATE_MARGINS = {
    "top_margin": Cm(2.3),
    "bottom_margin": Cm(2.4),
    "left_margin": Cm(1.7),
    "right_margin": Cm(1.7),
    "header_distance": Cm(1.7),
    "footer_distance": Cm(1.75),
}


REFERENCE_REPLACEMENTS = {
    2: "[2] WANG Guofa, REN Huaiwei, ZHAO Guorui, et al. Research and practice of intelligent coal mine technology systems in China[J]. International Journal of Coal Science & Technology, 2022, 9(1): 24. DOI: 10.1007/s40789-022-00491-3.",
    4: "[4] 许家林. 煤矿绿色开采20年研究及进展[J]. 煤炭科学技术, 2020, 48(9): 1-15. XU Jialin. Twenty years of research and progress in green mining of coal mines[J]. Coal Science and Technology, 2020, 48(9): 1-15. DOI: 10.13199/j.cnki.cst.2020.09.001.",
    5: "[5] 王国法, 杜毅博. 煤矿智能化标准体系框架与建设思路[J]. 煤炭科学技术, 2020, 48(1): 1-9. WANG Guofa, DU Yibo. Framework and construction ideas of intelligent coal mine standard system[J]. Coal Science and Technology, 2020, 48(1): 1-9. DOI: 10.13199/j.cnki.cst.2020.01.001.",
    6: "[6] 王国法, 任怀伟, 庞义辉, 等. 煤矿智能化(初级阶段)技术体系研究与工程进展[J]. 煤炭科学技术, 2020, 48(7): 1-27. WANG Guofa, REN Huaiwei, PANG Yihui, et al. Research and engineering progress of the technical system for intelligent coal mines at the primary stage[J]. Coal Science and Technology, 2020, 48(7): 1-27. DOI: 10.13199/j.cnki.cst.2020.07.001.",
    7: "[7] 李首滨. 智能化开采研究进展与发展趋势[J]. 煤炭科学技术, 2019, 47(10): 102-110. LI Shoubin. Research progress and development trend of intelligent mining[J]. Coal Science and Technology, 2019, 47(10): 102-110. DOI: 10.13199/j.cnki.cst.2019.10.012.",
    8: "[8] 王存飞, 荣耀. 透明工作面的概念、架构与关键技术[J]. 煤炭科学技术, 2019, 47(7): 156-163. WANG Cunfei, RONG Yao. Concept, architecture and key technologies of transparent working face[J]. Coal Science and Technology, 2019, 47(7): 156-163. DOI: 10.13199/j.cnki.cst.2019.07.019.",
    9: "[9] 程建远, 朱梦博, 王云宏, 等. 煤炭智能精准开采工作面地质模型梯级构建及其关键技术[J]. 煤炭学报, 2019, 44(8): 2285-2295. CHENG Jianyuan, ZHU Mengbo, WANG Yunhong, et al. Hierarchical construction of geological models for intelligent precise coal mining working faces and its key technologies[J]. Journal of China Coal Society, 2019, 44(8): 2285-2295. DOI: 10.13225/j.cnki.jccs.KJ19.0510.",
    10: "[10] 刘万里, 张学亮, 王世博. 采煤工作面煤层三维模型构建及动态修正技术[J]. 煤炭学报, 2020, 45(6): 1973-1983. LIU Wanli, ZHANG Xueliang, WANG Shibo. Construction and dynamic correction technology of a three-dimensional coal seam model for coal mining face[J]. Journal of China Coal Society, 2020, 45(6): 1973-1983. DOI: 10.13225/j.cnki.jccs.ZN20.0364.",
    11: "[11] YANG Yi, LI Yingchun, WANG Lujun, et al. On strata damage and stress disturbance induced by coal mining based on physical similarity simulation experiments[J]. Scientific Reports, 2023, 13(1): 15458. DOI: 10.1038/s41598-023-42148-4.",
    14: "[14] ZHANG Jie, WANG Li, YANG Tao, et al. Study on overburden failure characteristics and ground pressure behavior in shallow coal seam mining underneath the gully[J]. Frontiers in Earth Science, 2024, 12: 1375979. DOI: 10.3389/feart.2024.1375979.",
    15: "[15] CAO Jian, HUANG Qingxiang, GUO Lingfei. Subsidence prediction of overburden strata and ground surface in shallow coal seam mining[J]. Scientific Reports, 2021, 11(1): 18972. DOI: 10.1038/s41598-021-98520-9.",
    19: "[19] XUE Sen, WANG Qiqing, SONG Zhen. Analysis of water-permeable fractured zone in weakly cemented overburden considering rock strain-softening[J]. Scientific Reports, 2026, 16(1): 10776. DOI: 10.1038/s41598-026-45413-4.",
    20: "[20] ZHU Xiaojun, ZHA Feng, GUO Guangli, et al. Mechanical prediction method of strata movement and surface subsidence in backfill-strip mining[J]. Scientific Reports, 2024, 14(1): 31331. DOI: 10.1038/s41598-024-82761-5.",
    21: "[21] KRATZSCH Helmut. Strata movement at the mining horizon[M]//Mining Subsidence Engineering. Berlin, Heidelberg: Springer, 1983: 7-40. DOI: 10.1007/978-3-642-81923-0_2.",
    22: "[22] 徐智敏, 孙亚军, 高尚, 等. 干旱矿区采动顶板导水裂隙的演化规律及保水采煤意义[J]. 煤炭学报, 2019, 44(3): 767-776. XU Zhimin, SUN Yajun, GAO Shang, et al. Evolution law of mining-induced roof water-conducting fractures in arid mining areas and its significance for water-preserved coal mining[J]. Journal of China Coal Society, 2019, 44(3): 767-776. DOI: 10.13225/j.cnki.jccs.2018.6041.",
    23: "[23] 曹志国, 鞠金峰, 许家林. 采动覆岩导水裂隙主通道分布模型及其水流动特性[J]. 煤炭学报, 2019, 44(12): 3719-3728. CAO Zhiguo, JU Jinfeng, XU Jialin. Distribution model of main water-conducting fracture channels in mining-induced overburden and its water-flow characteristics[J]. Journal of China Coal Society, 2019, 44(12): 3719-3728. DOI: 10.13225/j.cnki.jccs.SH19.0446.",
    24: "[24] 郭文兵, 娄高中. 覆岩破坏充分采动程度定义及判别方法[J]. 煤炭学报, 2019, 44(3): 755-766. GUO Wenbing, LOU Gaozhong. Definition and discrimination method of full mining degree for overburden failure[J]. Journal of China Coal Society, 2019, 44(3): 755-766. DOI: 10.13225/j.cnki.jccs.2018.6038.",
    25: "[25] 鞠金峰, 马祥, 赵富强, 等. 东胜煤田导水裂缝发育及其分区特征研究[J]. 煤炭科学技术, 2022, 50(2): 202-212. JU Jinfeng, MA Xiang, ZHAO Fuqiang, et al. Development and zoning characteristics of water-conducting fissures in Dongsheng coalfield[J]. Coal Science and Technology, 2022, 50(2): 202-212. DOI: 10.13199/j.cnki.cst.2021-0169.",
    26: "[26] 郭小铭, 王皓, 周麟晟. 煤层顶板巨厚基岩含水层空间富水性评价[J]. 煤炭科学技术, 2021, 49(9): 167-175. GUO Xiaoming, WANG Hao, ZHOU Linsheng. Evaluation of spatial water abundance of extra-thick bedrock aquifer in coal seam roof[J]. Coal Science and Technology, 2021, 49(9): 167-175. DOI: 10.13199/j.cnki.cst.2021.09.024.",
    27: "[27] 王玉涛, 刘震. 深部煤层非充分采动下覆岩裂隙场可视化探测研究[J]. 煤炭科学技术, 2020, 48(3): 197-204. WANG Yutao, LIU Zhen. Visual detection of overburden fracture field under insufficient mining in deep coal seams[J]. Coal Science and Technology, 2020, 48(3): 197-204. DOI: 10.13199/j.cnki.cst.2020.03.024.",
    28: "[28] 董江鑫, 王飞. 地质雷达和高密度电法联合探测底板含水性的应用[J]. 煤炭科学技术, 2022, 50(5): 222-231. DONG Jiangxin, WANG Fei. Application of ground-penetrating radar and high-density electrical method in detecting floor water-bearing property[J]. Coal Science and Technology, 2022, 50(5): 222-231. DOI: 10.13199/j.cnki.cst.2020-0418.",
    29: "[29] 张玉军, 宋业杰, 樊振丽, 等. 鄂尔多斯盆地侏罗系煤田保水开采技术与应用[J]. 煤炭科学技术, 2021, 49(4): 159-168. ZHANG Yujun, SONG Yejie, FAN Zhenli, et al. Technology and application of water-preserved mining in Jurassic coalfield of Ordos Basin[J]. Coal Science and Technology, 2021, 49(4): 159-168. DOI: 10.13199/j.cnki.cst.2021.04.019.",
    30: "[30] 董书宁, 姬亚东, 王皓, 等. 鄂尔多斯盆地侏罗纪煤田典型顶板水害防控技术与应用[J]. 煤炭学报, 2020, 45(7): 2367-2375. DONG Shuning, JI Yadong, WANG Hao, et al. Typical roof water disaster prevention and control technology and application in Jurassic coalfield of Ordos Basin[J]. Journal of China Coal Society, 2020, 45(7): 2367-2375. DOI: 10.13225/j.cnki.jccs.DZ20.0697.",
    31: "[31] 范立民, 马雄德, 蒋泽泉, 等. 保水采煤研究30年回顾与展望[J]. 煤炭科学技术, 2019, 47(7): 1-30. FAN Limin, MA Xiongde, JIANG Zequan, et al. Review and prospect of water-preserved coal mining research over 30 years[J]. Coal Science and Technology, 2019, 47(7): 1-30. DOI: 10.13199/j.cnki.cst.2019.07.001.",
    32: "[32] 许家林, 朱卫兵, 王晓振. 基于关键层位置的导水裂隙带高度预计方法[J]. 煤炭学报, 2012, 37(5): 762-769. XU Jialin, ZHU Weibing, WANG Xiaozhen. Method for predicting the height of water-conducting fractured zone based on key strata location[J]. Journal of China Coal Society, 2012, 37(5): 762-769. DOI: 10.13225/j.cnki.jccs.2012.05.002.",
    33: "[33] 郭小铭, 董书宁. 深埋煤层开采顶板基岩含水层渗流规律及保水技术[J]. 煤炭学报, 2019, 44(3): 804-811. GUO Xiaoming, DONG Shuning. Seepage law of roof bedrock aquifer in deep buried coal seam mining and water-preserved technology[J]. Journal of China Coal Society, 2019, 44(3): 804-811. DOI: 10.13225/j.cnki.jccs.2018.6025.",
    39: "[39] 王苏健, 金声尧. 类孤岛工作面复采方案优选分析研究[J]. 煤炭科学技术, 2021, 49(9): 9-16. WANG Sujian, JIN Shengyao. Analysis and study on optimization of remining scheme for island-like working face[J]. Coal Science and Technology, 2021, 49(9): 9-16. DOI: 10.13199/j.cnki.cst.2021.09.002.",
    46: "[46] 刘晓丽, 曹志国, 陈苏社, 等. 煤矿分布式地下水库渗流场分析及优化调度[J]. 煤炭学报, 2019, 44(12): 3693-3699. LIU Xiaoli, CAO Zhiguo, CHEN Sushe, et al. Analysis and optimal regulation of seepage field in distributed underground reservoirs in coal mines[J]. Journal of China Coal Society, 2019, 44(12): 3693-3699. DOI: 10.13225/j.cnki.jccs.SH19.1167.",
    47: "[47] 庞义辉, 李全生, 曹光明, 等. 煤矿地下水库储水空间构成分析及计算方法[J]. 煤炭学报, 2019, 44(2): 557-566. PANG Yihui, LI Quansheng, CAO Guangming, et al. Analysis and calculation method of storage space composition of underground reservoirs in coal mines[J]. Journal of China Coal Society, 2019, 44(2): 557-566. DOI: 10.13225/j.cnki.jccs.2018.0417.",
    48: "[48] 李全生, 鞠金峰, 曹志国, 等. 基于导水裂隙带高度的地下水库适应性评价[J]. 煤炭学报, 2017, 42(8): 2116-2124. LI Quansheng, JU Jinfeng, CAO Zhiguo, et al. Applicability evaluation of underground reservoirs based on the height of water-conducting fractured zone[J]. Journal of China Coal Society, 2017, 42(8): 2116-2124. DOI: 10.13225/j.cnki.jccs.2016.1871.",
    49: "[49] 鞠金峰, 许家林, 朱卫兵. 西部缺水矿区地下水库保水的库容研究[J]. 煤炭学报, 2017, 42(2): 381-387. JU Jinfeng, XU Jialin, ZHU Weibing. Study on storage capacity of underground reservoirs for water conservation in western water-shortage mining areas[J]. Journal of China Coal Society, 2017, 42(2): 381-387. DOI: 10.13225/j.cnki.jccs.2016.6016.",
    50: "[50] 徐智敏, 高尚, 孙亚军, 等. 西部典型侏罗系富煤区含水介质条件与水动力学特征[J]. 煤炭学报, 2017, 42(2): 444-451. XU Zhimin, GAO Shang, SUN Yajun, et al. Water-bearing medium conditions and hydrodynamic characteristics in typical Jurassic coal-rich areas of western China[J]. Journal of China Coal Society, 2017, 42(2): 444-451. DOI: 10.13225/j.cnki.jccs.2016.6024.",
}


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def set_text(paragraph, text, size_pt=None):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    if size_pt:
        run.font.size = Pt(size_pt)
    return run


def insert_paragraph_after(paragraph, text, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    new_para.add_run(text)
    return new_para


def apply_docx_level_fixes():
    shutil.copy2(SOURCE, OUTPUT)
    doc = Document(OUTPUT)

    for idx, section in enumerate(doc.sections):
        for attr, value in TEMPLATE_MARGINS.items():
            setattr(section, attr, value)
        section.different_first_page_header_footer = idx == 0

    paragraphs = [p for p in doc.paragraphs if p.text.strip()]

    for paragraph in paragraphs:
        if "文献标志码：A" in paragraph.text:
            for run in paragraph.runs:
                if "文献标志码：A" in run.text:
                    run.text = run.text.replace("文献标志码：A", "文献标识码：A")
            if "文献标志码：A" in paragraph.text:
                set_text(paragraph, paragraph.text.replace("文献标志码：A", "文献标识码：A"))

    # Split mixed Chinese/English figure captions into the template's two-line paragraph form.
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if "\nFig." in text and (paragraph.style and paragraph.style.name == "14图题"):
            cn, en = text.split("\n", 1)
            set_text(paragraph, cn)
            insert_paragraph_after(paragraph, en, paragraph.style)

    # Restore missing English table numbers from the preceding Chinese table caption.
    last_table_number = None
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        match = re.match(r"^表\s*(\d+)", text)
        if match:
            last_table_number = match.group(1)
            continue
        if paragraph.style and paragraph.style.name == "15表题":
            if text.startswith("Table  ") and last_table_number:
                set_text(paragraph, text.replace("Table  ", f"Table{last_table_number}  ", 1))

    in_refs = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "参考文献":
            in_refs = True
            continue
        if not in_refs:
            continue
        match = re.match(r"^\[(\d{1,3})\]\s*(.*)", text)
        if not match:
            continue
        ref_num = int(match.group(1))
        new_text = REFERENCE_REPLACEMENTS.get(ref_num, text.replace("\t", " "))
        set_text(paragraph, new_text, size_pt=9)
        fmt = paragraph.paragraph_format
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = Pt(14)
        fmt.first_line_indent = None
        fmt.left_indent = None
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)

    doc.save(TMP)


def remove_redundant_section_breaks(src, dst):
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        with ZipFile(src) as zin:
            zin.extractall(tmp_path)

        document_xml = tmp_path / "word" / "document.xml"
        tree = etree.parse(str(document_xml))
        root = tree.getroot()

        for sect in root.xpath(".//w:body/w:p/w:pPr/w:sectPr", namespaces=ns):
            parent = sect.getparent()
            parent.remove(sect)

        body_sect = root.xpath(".//w:body/w:sectPr", namespaces=ns)[0]
        pg_mar = body_sect.find(qn("w:pgMar"))
        if pg_mar is None:
            pg_mar = OxmlElement("w:pgMar")
            body_sect.append(pg_mar)
        pg_mar.set(qn("w:top"), "1304")
        pg_mar.set(qn("w:bottom"), "1361")
        pg_mar.set(qn("w:left"), "964")
        pg_mar.set(qn("w:right"), "964")
        pg_mar.set(qn("w:header"), "964")
        pg_mar.set(qn("w:footer"), "992")

        if body_sect.find(qn("w:titlePg")) is None:
            body_sect.append(OxmlElement("w:titlePg"))

        tree.write(str(document_xml), encoding="UTF-8", xml_declaration=True, standalone="yes")

        with ZipFile(dst, "w", ZIP_DEFLATED) as zout:
            for file in tmp_path.rglob("*"):
                if file.is_file():
                    zout.write(file, file.relative_to(tmp_path).as_posix())


def main():
    apply_docx_level_fixes()
    remove_redundant_section_breaks(TMP, OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
