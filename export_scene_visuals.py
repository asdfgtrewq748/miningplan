from __future__ import annotations

import argparse
import csv
import json
import math
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Iterable, Sequence

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from matplotlib.colors import BoundaryNorm, LinearSegmentedColormap
from matplotlib.patches import Patch
from scipy.interpolate import griddata


SCENE_SLUGS = {
    "0-地表下沉.miningplan": "00_surface_subsidence",
    "1-含水层扰动预评价.miningplan": "01_aquifer_pre_eval",
    "2-含水层扰动评价.miningplan": "02_aquifer_eval",
    "3-采区规划案例.miningplan": "03_mining_planning",
    "4-协同调控-突水点.miningplan": "04_cocontrol_water_inrush",
    "5-采掘接续.miningplan": "05_mining_succession",
    "6-全覆岩扰动.miningplan": "06_full_overburden",
}

DEFAULT_INPUT_CANDIDATES = (
    Path("软件案例附件") / "工程文件案例",
    Path("mining-plan") / "frontend" / "public" / "demo",
)

PARAM_KEYS = ("Ti", "Hi", "Di", "Mi")
PARAM_LABELS = {
    "Ti": "目标层厚度 Ti (m)",
    "Hi": "煤层与目标层间距 Hi (m)",
    "Di": "目标层埋深 Di (m)",
    "Mi": "目标层煤层厚度 Mi (m)",
}

BLUE_RED_STOPS = ["#1d4ed8", "#60a5fa", "#f8fafc", "#fca5a5", "#dc2626"]
VIRIDIS_STOPS = ["#440154", "#3b528b", "#21918c", "#5ec962", "#fde725"]
LEVEL_COLORS = ["#10b981", "#84cc16", "#eab308", "#f97316", "#ef4444"]
LEVEL_LABELS = ["I级", "II级", "III级", "IV级", "V级"]
ERR_ODI_COLOR = "#3b82f6"
ERR_MEASURED_COLOR = "#ef4444"
ERR_RATIO_COLOR = "#cbd5e1"
CAT_COLORS = {
    "geo": "#111827",
    "gray": "#94a3b8",
    "blue": "#2563eb",
    "pink": "#ec4899",
    "green": "#16a34a",
    "red": "#ef4444",
}

plt.rcParams.update(
    {
        "font.family": "sans-serif",
        "font.sans-serif": [
            "Microsoft YaHei",
            "SimHei",
            "Noto Sans CJK SC",
            "Source Han Sans SC",
            "DejaVu Sans",
        ],
        "axes.unicode_minus": False,
        "font.size": 9,
        "axes.titlesize": 11,
        "axes.labelsize": 9,
        "legend.fontsize": 8,
        "xtick.labelsize": 8,
        "ytick.labelsize": 8,
        "savefig.dpi": 300,
        "figure.dpi": 150,
        "savefig.facecolor": "white",
        "axes.facecolor": "white",
    }
)

FIGURE_NAME_MAP = {
    "fig01_odi_distribution": "01-ODI分布图",
    "fig02_geology_clouds": "02-地质参数云图",
    "fig03_spatial_map": "03-空间分布图",
    "fig04_odi_histogram": "04-ODI频率分布",
    "fig05_odi_levels": "05-ODI分级占比",
    "fig06_weight_radar": "06-权重雷达图",
    "fig01_workface_plan_layout": "01-采区规划布局",
    "fig02_planning_mode_scores": "02-四模式规划指标对比",
    "fig03_weighted_top_candidates": "03-加权优选候选方案",
    "fig01_monthly_production": "01-月产曲线",
    "fig02_schedule_gantt": "02-采掘接续甘特图",
    "fig03_stage3_candidate_scores": "03-接续方案对比",
    "fig01_cashflow": "01-现金流分析",
    "fig02_revenue_cost": "02-收入成本结构",
    "fig03_cost_structure": "03-成本构成",
}


@dataclass
class ExportSummary:
    scene_name: str
    scene_slug: str
    source_file: Path
    figure_count: int = 0
    files: list[Path] = field(default_factory=list)


def build_default_output_dir(base_dir: Path | None = None) -> Path:
    root = Path(base_dir or Path.cwd())
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return root / "output" / "scene_visual_exports" / stamp


def choose_default_input_dir(base_dir: Path | None = None) -> Path:
    root = Path(base_dir or Path.cwd())
    for candidate in DEFAULT_INPUT_CANDIDATES:
        path = root / candidate
        if path.is_dir():
            return path
    raise FileNotFoundError(
        "未找到默认场景目录，请显式传入 --input，或确认软件案例附件/工程文件案例 是否存在。"
    )


def discover_scene_files(inputs: Sequence[str | Path] | None = None, base_dir: Path | None = None) -> list[Path]:
    root = Path(base_dir or Path.cwd())
    candidates = list(inputs or [choose_default_input_dir(root)])
    discovered: dict[Path, Path] = {}
    for item in candidates:
        path = Path(item)
        if not path.is_absolute():
            path = root / path
        if path.is_dir():
            for file_path in sorted(path.glob("*.miningplan.json")):
                discovered[file_path.resolve()] = file_path.resolve()
        elif path.is_file() and path.name.endswith(".miningplan.json"):
            discovered[path.resolve()] = path.resolve()
    return [discovered[key] for key in sorted(discovered)]


def strip_double_suffix(path: Path) -> str:
    name = path.name
    if name.endswith(".miningplan.json"):
        return name[: -len(".json")]
    return path.stem


def slugify_ascii(text: str) -> str:
    cleaned = re.sub(r"[^0-9A-Za-z]+", "_", text).strip("_").lower()
    return cleaned or "scene"


def scene_slug(path: Path) -> str:
    key = strip_double_suffix(path)
    if key in SCENE_SLUGS:
        return SCENE_SLUGS[key]
    stem = path.stem
    match = re.match(r"(\d+)-(.+)$", stem)
    if match:
        prefix = int(match.group(1))
        return f"{prefix:02d}_{slugify_ascii(match.group(2))}"
    return slugify_ascii(stem)


def load_scene(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def iter_tabs_with_data(scene: dict) -> list[str]:
    tabs: list[str] = []
    scenario_params = scene.get("scenarioParamsById") or {}
    for tab_id, tab_data in scenario_params.items():
        if not isinstance(tab_data, dict):
            continue
        if any(
            [
                get_odi_points(scene, tab_id),
                get_param_points(scene, tab_id),
                tab_data.get("drillholeData"),
                tab_data.get("boundaryData"),
                tab_data.get("workingFaceData"),
                tab_data.get("measuredConstraintData"),
                tab_data.get("errorAnalysisByLineId"),
            ]
        ):
            tabs.append(tab_id)
    return tabs


def get_tab(scene: dict, tab_id: str) -> dict:
    return (scene.get("scenarioParamsById") or {}).get(tab_id) or {}


def get_odi_points(scene: dict, tab_id: str) -> list[dict]:
    tab = get_tab(scene, tab_id)
    odi_result = tab.get("odiResult") or {}
    points = odi_result.get("points") or []
    if points:
        return points
    union = (((scene.get("cocontrol") or {}).get("results") or {}).get("odiUnionResult") or {})
    return union.get("points") or []


def get_param_points(scene: dict, tab_id: str) -> list[dict]:
    tab = get_tab(scene, tab_id)
    result = tab.get("paramExtractionResult") or {}
    return result.get("points") or []


def get_level_ranges(tab: dict) -> list[dict]:
    measured = tab.get("measuredZoningResult") or {}
    bins = measured.get("bins") or []
    if len(bins) == 5:
        return [
            {"lo": float(bin_["odiLo"]), "hi": float(bin_["odiHi"]), "include_hi": index == 4}
            for index, bin_ in enumerate(bins)
        ]
    return [
        {"lo": 0.0, "hi": 0.2, "include_hi": False},
        {"lo": 0.2, "hi": 0.4, "include_hi": False},
        {"lo": 0.4, "hi": 0.6, "include_hi": False},
        {"lo": 0.6, "hi": 0.8, "include_hi": False},
        {"lo": 0.8, "hi": 1.0, "include_hi": True},
    ]


def get_workface_plan_loops(scene: dict) -> list[dict]:
    workface_plan = scene.get("workfacePlan") or {}
    return workface_plan.get("plannedWorkfaceLoopsWorld") or []


def blue_red_cmap() -> LinearSegmentedColormap:
    return LinearSegmentedColormap.from_list("blue_red", BLUE_RED_STOPS, N=256)


def viridis_cmap() -> LinearSegmentedColormap:
    return LinearSegmentedColormap.from_list("viridis_like", VIRIDIS_STOPS, N=256)


def normalize_workface_groups(workfaces: list[dict]) -> dict[str, list[tuple[float, float]]]:
    grouped: dict[str, list[tuple[float, float]]] = {}
    for point in workfaces or []:
        x = float(point.get("x", math.nan))
        y = float(point.get("y", math.nan))
        if not (math.isfinite(x) and math.isfinite(y)):
            continue
        key = str(point.get("id") or point.get("faceIndex") or "工作面")
        grouped.setdefault(key, []).append((x, y))
    return grouped


def classify_odi_level(value: float, level_ranges: list[dict]) -> int:
    for index, spec in enumerate(level_ranges):
        lo = float(spec["lo"])
        hi = float(spec["hi"])
        include_hi = bool(spec["include_hi"])
        if value >= lo and (value <= hi if include_hi else value < hi):
            return index
    return len(level_ranges) - 1


def interpolate_grid(points: list[dict], value_key: str, resolution: int = 120) -> tuple[np.ndarray, np.ndarray, np.ndarray] | None:
    if len(points) < 4:
        return None
    xs = np.array([float(point["x"]) for point in points], dtype=float)
    ys = np.array([float(point["y"]) for point in points], dtype=float)
    vals = np.array([float(point.get(value_key, np.nan)) for point in points], dtype=float)
    mask = np.isfinite(xs) & np.isfinite(ys) & np.isfinite(vals)
    xs = xs[mask]
    ys = ys[mask]
    vals = vals[mask]
    if len(xs) < 4:
        return None
    xi = np.linspace(xs.min(), xs.max(), resolution)
    yi = np.linspace(ys.min(), ys.max(), resolution)
    xi_grid, yi_grid = np.meshgrid(xi, yi)
    try:
        zi = griddata((xs, ys), vals, (xi_grid, yi_grid), method="cubic")
    except Exception:
        try:
            zi = griddata((xs, ys), vals, (xi_grid, yi_grid), method="linear")
        except Exception:
            zi = griddata((xs, ys), vals, (xi_grid, yi_grid), method="nearest")
    if zi is None:
        return None
    return xi_grid, yi_grid, np.ma.masked_invalid(zi)


def add_common_map_overlays(ax: plt.Axes, drillholes: list[dict], boundary: list[dict], workfaces: list[dict]) -> None:
    if boundary:
        bx = [float(point["x"]) for point in boundary] + [float(boundary[0]["x"])]
        by = [float(point["y"]) for point in boundary] + [float(boundary[0]["y"])]
        ax.plot(bx, by, color="#1e3a8a", linewidth=1.2, alpha=0.9)
    for label, points in normalize_workface_groups(workfaces).items():
        if len(points) < 2:
            continue
        closed = points + [points[0]]
        ax.plot(
            [point[0] for point in closed],
            [point[1] for point in closed],
            color="#ec4899",
            linewidth=1.1,
            linestyle="--",
            alpha=0.9,
        )
        cx = sum(point[0] for point in points) / len(points)
        cy = sum(point[1] for point in points) / len(points)
        ax.text(cx, cy, label, fontsize=7, color="#be185d", ha="center", va="center")
    if drillholes:
        ax.scatter(
            [float(point["x"]) for point in drillholes],
            [float(point["y"]) for point in drillholes],
            s=24,
            c="#111827",
            marker="o",
            zorder=4,
            label="钻孔",
        )


def finalize_map_axes(ax: plt.Axes, title: str) -> None:
    ax.set_title(title, fontweight="bold")
    ax.set_xlabel("X / m")
    ax.set_ylabel("Y / m")
    ax.set_aspect("equal", adjustable="box")
    ax.grid(color="#e2e8f0", linewidth=0.5, linestyle=":", alpha=0.8)


def plot_odi_heatmap(
    odi_points: list[dict],
    drillholes: list[dict],
    boundary: list[dict],
    workfaces: list[dict],
    title: str,
) -> plt.Figure | None:
    grid = interpolate_grid(odi_points, "odiNorm", resolution=140)
    if grid is None:
        return None
    xi_grid, yi_grid, zi = grid
    levels = np.linspace(0.0, 1.0, 11)
    fig, ax = plt.subplots(figsize=(7.2, 5.6))
    contour = ax.contourf(
        xi_grid,
        yi_grid,
        zi,
        levels=levels,
        cmap=blue_red_cmap(),
        antialiased=True,
    )
    for point in odi_points:
        x = float(point["x"])
        y = float(point["y"])
        ax.scatter(x, y, s=14, facecolors="white", edgecolors="#ef4444", linewidths=0.9, zorder=5)
    add_common_map_overlays(ax, drillholes, boundary, workfaces)
    cbar = fig.colorbar(contour, ax=ax, pad=0.02)
    cbar.set_label("ODI (归一化)")
    finalize_map_axes(ax, title)
    fig.tight_layout()
    return fig


def plot_geology_clouds(param_points: list[dict], drillholes: list[dict], title: str) -> plt.Figure | None:
    if len(param_points) < 4:
        return None
    fig, axes = plt.subplots(2, 2, figsize=(8.0, 6.8))
    cmap = viridis_cmap()
    created = False
    for axis, key in zip(axes.flat, PARAM_KEYS):
        grid = interpolate_grid(param_points, key, resolution=110)
        if grid is None:
            axis.text(0.5, 0.5, "无足够数据", ha="center", va="center", transform=axis.transAxes)
            axis.set_axis_off()
            continue
        xi_grid, yi_grid, zi = grid
        levels = 12
        contour = axis.contourf(xi_grid, yi_grid, zi, levels=levels, cmap=cmap, antialiased=True)
        if drillholes:
            axis.scatter(
                [float(point["x"]) for point in drillholes],
                [float(point["y"]) for point in drillholes],
                s=16,
                c="#334155",
                zorder=4,
            )
        axis.set_title(PARAM_LABELS[key], fontsize=9, fontweight="bold")
        axis.set_aspect("equal", adjustable="box")
        axis.grid(color="#e2e8f0", linewidth=0.4, linestyle=":", alpha=0.7)
        fig.colorbar(contour, ax=axis, pad=0.01, shrink=0.82)
        created = True
    if not created:
        plt.close(fig)
        return None
    fig.suptitle(title, fontsize=11, fontweight="bold")
    fig.tight_layout(rect=(0, 0, 1, 0.97))
    return fig


def plot_spatial_map(
    odi_points: list[dict],
    drillholes: list[dict],
    boundary: list[dict],
    workfaces: list[dict],
    title: str,
) -> plt.Figure | None:
    if not any([odi_points, drillholes, boundary, workfaces]):
        return None
    fig, ax = plt.subplots(figsize=(7.0, 5.4))
    add_common_map_overlays(ax, drillholes, boundary, workfaces)
    for category, color in CAT_COLORS.items():
        points = [point for point in odi_points if str(point.get("cat")) == category]
        if not points:
            continue
        ax.scatter(
            [float(point["x"]) for point in points],
            [float(point["y"]) for point in points],
            s=16,
            c=color,
            alpha=0.75,
            label=category,
        )
    if odi_points:
        ax.legend(loc="lower left", ncol=3, frameon=False)
    finalize_map_axes(ax, title)
    fig.tight_layout()
    return fig


def plot_odi_histogram(odi_points: list[dict], title: str) -> plt.Figure | None:
    if not odi_points:
        return None
    values = [float(point.get("odiNorm", point.get("odi", 0.0))) for point in odi_points]
    fig, ax = plt.subplots(figsize=(4.8, 3.4))
    ax.hist(values, bins=min(16, max(6, len(values))), color="#2563eb", alpha=0.85, edgecolor="white")
    ax.axvline(float(np.mean(values)), color="#dc2626", linewidth=1.1, linestyle="--", label="均值")
    ax.set_title(title, fontweight="bold")
    ax.set_xlabel("ODI (归一化)")
    ax.set_ylabel("频数")
    ax.legend(frameon=False)
    ax.grid(axis="y", color="#e2e8f0", linewidth=0.5, linestyle=":")
    fig.tight_layout()
    return fig


def plot_odi_level_pie(odi_points: list[dict], level_ranges: list[dict], title: str) -> plt.Figure | None:
    if not odi_points:
        return None
    counts = [0, 0, 0, 0, 0]
    for point in odi_points:
        level = classify_odi_level(float(point.get("odiNorm", point.get("odi", 0.0))), level_ranges)
        counts[level] += 1
    if not any(counts):
        return None
    fig, ax = plt.subplots(figsize=(4.6, 4.0))
    ax.pie(
        counts,
        labels=LEVEL_LABELS,
        colors=LEVEL_COLORS,
        autopct=lambda pct: f"{pct:.1f}%" if pct > 0 else "",
        startangle=90,
        counterclock=False,
        wedgeprops={"linewidth": 0.8, "edgecolor": "white"},
        textprops={"fontsize": 8},
    )
    ax.set_title(title, fontweight="bold")
    fig.tight_layout()
    return fig


def plot_error_trend(error_points: list[dict], title: str) -> plt.Figure | None:
    if not error_points:
        return None
    x = np.arange(len(error_points))
    labels = [str(point.get("id", index + 1)) for index, point in enumerate(error_points)]
    odi_values = [float(point.get("odiRenorm", 0.0)) for point in error_points]
    measured_values = [float(point.get("measured", 0.0)) for point in error_points]
    error_values = [abs(float(point.get("errorRatioChart", point.get("errorRatio", 0.0)))) for point in error_points]

    fig, ax_left = plt.subplots(figsize=(7.4, 3.6))
    ax_left.bar(x, error_values, color=ERR_RATIO_COLOR, alpha=0.85, label="误差比")
    ax_left.plot(x, odi_values, color=ERR_ODI_COLOR, marker="o", linewidth=1.5, label="ODI")
    ax_left.set_ylabel("ODI / 误差")
    ax_left.set_xticks(x)
    ax_left.set_xticklabels(labels, rotation=45, ha="right")
    ax_left.grid(axis="y", color="#e2e8f0", linewidth=0.5, linestyle=":")

    ax_right = ax_left.twinx()
    ax_right.plot(
        x,
        measured_values,
        color=ERR_MEASURED_COLOR,
        marker="s",
        linewidth=1.4,
        linestyle="--",
        label="实测值",
    )
    ax_right.set_ylabel("实测值 / m", color=ERR_MEASURED_COLOR)
    ax_right.tick_params(axis="y", colors=ERR_MEASURED_COLOR)

    handles_left, labels_left = ax_left.get_legend_handles_labels()
    handles_right, labels_right = ax_right.get_legend_handles_labels()
    ax_left.legend(handles_left + handles_right, labels_left + labels_right, loc="upper left", frameon=False)
    ax_left.set_title(title, fontweight="bold")
    fig.tight_layout()
    return fig


def plot_weight_radar(weights: dict, title: str) -> plt.Figure | None:
    wd = weights.get("wd")
    wo = weights.get("wo")
    wf = weights.get("wf")
    if not all(isinstance(value, (int, float)) for value in (wd, wo, wf)):
        return None
    labels = ["地质因子", "开采因子", "综合因子"]
    values = [float(wd), float(wo), float(wf)]
    angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
    angles += angles[:1]
    values += values[:1]
    fig, ax = plt.subplots(figsize=(4.4, 4.0), subplot_kw={"polar": True})
    ax.plot(angles, values, color="#2563eb", linewidth=1.6, marker="o")
    ax.fill(angles, values, color="#60a5fa", alpha=0.25)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels)
    ax.set_ylim(0, max(values) * 1.2 if max(values) > 0 else 1.0)
    ax.set_title(title, fontweight="bold", pad=18)
    fig.tight_layout()
    return fig


def plot_workface_layout(plan_loops: list[dict], title: str) -> plt.Figure | None:
    if not plan_loops:
        return None
    fig, ax = plt.subplots(figsize=(7.0, 5.4))
    cmap = plt.colormaps.get_cmap("Blues").resampled(max(3, len(plan_loops) + 2))
    for index, face in enumerate(plan_loops, start=1):
        loop = face.get("loop") or []
        if len(loop) < 3:
            continue
        xs = [float(point["x"]) for point in loop] + [float(loop[0]["x"])]
        ys = [float(point["y"]) for point in loop] + [float(loop[0]["y"])]
        ax.fill(xs, ys, color=cmap(index), alpha=0.35)
        ax.plot(xs, ys, color=cmap(index), linewidth=1.4)
        cx = sum(float(point["x"]) for point in loop) / len(loop)
        cy = sum(float(point["y"]) for point in loop) / len(loop)
        label = f"No.{face.get('faceIndex', index)}"
        ax.text(cx, cy, label, ha="center", va="center", fontsize=8, color="#0f172a", fontweight="bold")
    finalize_map_axes(ax, title)
    fig.tight_layout()
    return fig


def safe_float(value: object, default: float | None = None) -> float | None:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return default
    return number if math.isfinite(number) else default


def clamp01(value: object, default: float = 0.0) -> float:
    number = safe_float(value, default)
    if number is None:
        return default
    return max(0.0, min(1.0, number))


def parse_bool(value: object, default: bool = False) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        text = value.strip().lower()
        if text in {"1", "true", "yes", "y", "on"}:
            return True
        if text in {"0", "false", "no", "n", "off"}:
            return False
    return default


def pick_numeric(row: dict, keys: Sequence[str]) -> float | None:
    for key in keys:
        value = safe_float(row.get(key))
        if value is not None:
            return value
    return None


PLANNING_MODE_META = {
    "efficiency": {
        "label": "工程效率",
        "score_keys": ("efficiencyScore", "effScore", "score"),
        "higher_is_better": True,
    },
    "recovery": {
        "label": "资源回收",
        "score_keys": ("recoveryScore", "recScore", "score"),
        "higher_is_better": True,
    },
    "disturbance": {
        "label": "最小扰动",
        "score_keys": ("distScore", "_distScore", "score"),
        "higher_is_better": False,
    },
    "weighted": {
        "label": "加权优选",
        "score_keys": ("totalScore", "combinedScore", "score"),
        "higher_is_better": True,
    },
}


def get_planning_rows(scene: dict, mode: str) -> list[dict]:
    result = (((scene.get("planningResults") or {}).get(mode) or {}).get("result") or {})
    if mode == "disturbance":
        disturbance_pack = result.get("disturbance") or {}
        rows = ((disturbance_pack.get("table") or {}).get("rows") or [])
        if rows:
            return [row for row in rows if isinstance(row, dict)]
    rows = (((result.get("table") or {}).get("rows")) or [])
    return [row for row in rows if isinstance(row, dict)]


def get_best_planning_row(rows: list[dict], mode: str) -> dict | None:
    if not rows:
        return None
    meta = PLANNING_MODE_META[mode]
    score_keys = meta["score_keys"]
    higher_is_better = bool(meta["higher_is_better"])
    ranked: list[tuple[float, dict]] = []
    for row in rows:
        score = pick_numeric(row, score_keys)
        if score is None:
            continue
        ranked.append((score, row))
    if not ranked:
        return rows[0]
    ranked.sort(key=lambda item: item[0], reverse=higher_is_better)
    return ranked[0][1]


def plot_planning_mode_scores(scene: dict, title: str) -> plt.Figure | None:
    labels: list[str] = []
    values: list[float] = []
    tonnages: list[float] = []
    for mode in ("efficiency", "recovery", "disturbance", "weighted"):
        rows = get_planning_rows(scene, mode)
        if not rows:
            continue
        best = get_best_planning_row(rows, mode)
        if not best:
            continue
        score = pick_numeric(best, PLANNING_MODE_META[mode]["score_keys"])
        if score is None:
            continue
        labels.append(PLANNING_MODE_META[mode]["label"])
        values.append(score)
        tonnages.append((pick_numeric(best, ("tonnageTotal",)) or 0.0) / 10000.0)
    if not values:
        return None

    fig, ax = plt.subplots(figsize=(6.8, 4.2))
    colors = ["#2563eb", "#16a34a", "#f59e0b", "#7c3aed"][: len(values)]
    x = np.arange(len(values))
    bars = ax.bar(x, values, color=colors, width=0.58)
    ax.set_xticks(x)
    ax.set_xticklabels(labels)
    ax.set_ylabel("目标函数值")
    ax.set_title(title, fontweight="bold")
    ax.grid(axis="y", color="#e2e8f0", linewidth=0.5, linestyle=":")
    for bar, value in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height(), f"{value:.3g}", ha="center", va="bottom", fontsize=8)

    if any(tonnages):
        ax2 = ax.twinx()
        ax2.plot(x, tonnages, color="#111827", marker="o", linewidth=1.4, linestyle="--")
        ax2.set_ylabel("储量 / 万t")
        ax2.tick_params(axis="y", labelsize=8)

    fig.tight_layout()
    return fig


def plot_weighted_top_candidates(scene: dict, title: str) -> plt.Figure | None:
    rows = get_planning_rows(scene, "weighted")
    if not rows:
        return None
    data = []
    for row in rows:
        score = pick_numeric(row, ("totalScore", "combinedScore", "score"))
        if score is None:
            continue
        label = str(row.get("signature") or row.get("key") or f"cand-{len(data) + 1}")
        data.append((label, score))
    if not data:
        return None
    data = sorted(data, key=lambda item: item[1], reverse=True)[:8]

    fig, ax = plt.subplots(figsize=(7.2, 4.4))
    labels = [item[0] for item in data][::-1]
    values = [item[1] for item in data][::-1]
    bars = ax.barh(labels, values, color="#4f46e5", alpha=0.9)
    ax.set_xlabel("综合得分")
    ax.set_title(title, fontweight="bold")
    ax.grid(axis="x", color="#e2e8f0", linewidth=0.5, linestyle=":")
    for bar, value in zip(bars, values):
        ax.text(bar.get_width(), bar.get_y() + bar.get_height() / 2, f" {value:.2f}", va="center", ha="left", fontsize=8)
    fig.tight_layout()
    return fig


def compute_workface_dims_from_loop(loop: list[dict]) -> dict:
    points = []
    for point in loop or []:
        x = safe_float(point.get("x"))
        y = safe_float(point.get("y"))
        if x is None or y is None:
            continue
        points.append({"x": x, "y": y})
    if len(points) < 3:
        return {"center": {"x": 0.0, "y": 0.0}, "widthM": 0.0, "advanceLengthM": 0.0}

    center = {
        "x": sum(point["x"] for point in points) / len(points),
        "y": sum(point["y"] for point in points) / len(points),
    }
    cov_xx = cov_xy = cov_yy = 0.0
    for point in points:
        dx = point["x"] - center["x"]
        dy = point["y"] - center["y"]
        cov_xx += dx * dx
        cov_xy += dx * dy
        cov_yy += dy * dy

    angle = 0.5 * math.atan2(2 * cov_xy, cov_xx - cov_yy)
    u = {"x": math.cos(angle), "y": math.sin(angle)}
    u_len = math.hypot(u["x"], u["y"]) or 1.0
    u = {"x": u["x"] / u_len, "y": u["y"] / u_len}
    v = {"x": -u["y"], "y": u["x"]}

    def project(axis: dict) -> tuple[float, float]:
        values = []
        for point in points:
            dx = point["x"] - center["x"]
            dy = point["y"] - center["y"]
            values.append(dx * axis["x"] + dy * axis["y"])
        return min(values), max(values)

    u_min, u_max = project(u)
    v_min, v_max = project(v)
    if (u_max - u_min) >= (v_max - v_min):
        advance_axis = u
        width_axis = v
    else:
        advance_axis = v
        width_axis = {"x": -v["y"], "y": v["x"]}

    a_min, a_max = project(advance_axis)
    w_min, w_max = project(width_axis)
    return {
        "center": center,
        "widthM": max(0.0, w_max - w_min),
        "advanceLengthM": max(0.0, a_max - a_min),
    }


def get_scene_mining_height(scene: dict) -> float:
    planning_params = scene.get("planningParams") or {}
    seam = safe_float(planning_params.get("seamThickness"))
    if seam is not None and seam > 0:
        return seam
    planned_by_face = ((scene.get("cocontrol") or {}).get("plannedParamsByFaceIndex") or {})
    heights = []
    for pack in planned_by_face.values():
        height = safe_float(((pack or {}).get("production") or {}).get("miningHeightM"))
        if height is not None and height > 0:
            heights.append(height)
    return float(np.mean(heights)) if heights else 4.5


def build_raw_succession_panels(scene: dict) -> list[dict]:
    panels: list[dict] = []
    for face in get_workface_plan_loops(scene):
        loop = face.get("loop") or []
        dims = compute_workface_dims_from_loop(loop)
        face_index = int(face.get("faceIndex", len(panels) + 1))
        panels.append(
            {
                "id": f"No.{face_index}",
                "faceIndex": face_index,
                "widthM": dims["widthM"],
                "advanceLengthM": dims["advanceLengthM"],
                "center_x": dims["center"]["x"],
                "center_y": dims["center"]["y"],
                "loop": loop,
            }
        )
    return [panel for panel in panels if panel["advanceLengthM"] > 0]


def order_succession_panels(
    panels: list[dict],
    mode: str,
    yard_dir: str = "NE",
    yard_offset_m: float = 120.0,
) -> list[dict]:
    ordered = [dict(panel) for panel in panels]
    if not ordered:
        return ordered
    if mode == "yardConfirmed":
        dir_user = str(yard_dir or "NE")
        dir_world = {
            "N": "S",
            "S": "N",
            "E": "W",
            "W": "E",
            "NE": "SW",
            "NW": "SE",
            "SE": "NW",
            "SW": "NE",
        }.get(dir_user, "SW")
        xs = [safe_float(panel.get("center_x"), 0.0) or 0.0 for panel in ordered]
        ys = [safe_float(panel.get("center_y"), 0.0) or 0.0 for panel in ordered]
        min_x, max_x = min(xs), max(xs)
        min_y, max_y = min(ys), max(ys)
        mid_x = (min_x + max_x) / 2.0
        mid_y = (min_y + max_y) / 2.0
        pad = safe_float(yard_offset_m, 120.0) or 120.0
        anchor = {
            "N": {"x": mid_x, "y": max_y + pad},
            "S": {"x": mid_x, "y": min_y - pad},
            "E": {"x": max_x + pad, "y": mid_y},
            "W": {"x": min_x - pad, "y": mid_y},
            "NE": {"x": max_x + pad, "y": max_y + pad},
            "NW": {"x": min_x - pad, "y": max_y + pad},
            "SE": {"x": max_x + pad, "y": min_y - pad},
            "SW": {"x": min_x - pad, "y": min_y - pad},
        }.get(dir_world, {"x": min_x - pad, "y": min_y - pad})
        ordered.sort(
            key=lambda panel: (
                (safe_float(panel.get("center_x"), 0.0) - anchor["x"]) ** 2
                + (safe_float(panel.get("center_y"), 0.0) - anchor["y"]) ** 2,
                safe_float(panel.get("faceIndex"), 0.0),
            )
        )
        return ordered
    ordered.sort(key=lambda panel: safe_float(panel.get("faceIndex"), 0.0))
    return ordered


def build_succession_stage1_plan(panels: list[dict], params: dict) -> dict:
    days_per_month = max(1, int(round(safe_float(params.get("daysPerMonth"), 25.0) or 25.0)))
    utilization = clamp01(params.get("utilization"), 0.85)
    shear_advance_rate = max(0.0, safe_float(params.get("shearAdvanceRate"), 6.0) or 6.0)
    drive_rate = max(0.0, safe_float(params.get("driveRate"), 15.0) or 15.0)
    install_days = max(0.0, safe_float(params.get("installDays"), 15.0) or 15.0)
    relocation_days = max(0.0, safe_float(params.get("relocationDays"), 10.0) or 10.0)
    single_face_mining = parse_bool(params.get("singleFaceMining"), True)
    drive_parallel = parse_bool(params.get("driveParallelWithMining"), True)
    drive_crews = max(1, int(round(safe_float(params.get("driveCrews"), 1.0) or 1.0)))

    tasks: list[dict] = []
    last_mining_end = 0.0
    last_relocation_end = 0.0
    drive_crew_end = [0.0 for _ in range(drive_crews)]

    def take_earliest_crew() -> tuple[int, float]:
        best_idx = 0
        best_day = drive_crew_end[0]
        for index, day in enumerate(drive_crew_end[1:], start=1):
            if day < best_day:
                best_idx = index
                best_day = day
        return best_idx, best_day

    for index, panel in enumerate(panels):
        workface = str(panel.get("id") or f"WF-{index + 1}")
        length_m = max(0.0, safe_float(panel.get("advanceLengthM"), 0.0) or 0.0)
        drive_days = (length_m / drive_rate) / max(1e-6, utilization) if drive_rate > 1e-9 else 0.0
        crew_index, crew_day = take_earliest_crew()
        drive_start = max(crew_day, last_mining_end if (not drive_parallel and index > 0) else 0.0)
        drive_end = drive_start + drive_days
        drive_crew_end[crew_index] = drive_end
        tasks.append({"type": "drive", "workface": workface, "startDay": drive_start, "endDay": drive_end})

        install_start = max(drive_end, last_relocation_end)
        install_end = install_start + install_days
        tasks.append({"type": "install", "workface": workface, "startDay": install_start, "endDay": install_end})

        mining_start = max(install_end, last_mining_end if (single_face_mining and index > 0) else 0.0)
        mining_days = (length_m / shear_advance_rate) / max(1e-6, utilization) if shear_advance_rate > 1e-9 else 0.0
        mining_end = mining_start + mining_days
        tasks.append({"type": "mining", "workface": workface, "startDay": mining_start, "endDay": mining_end, "lengthM": length_m})

        if index < len(panels) - 1:
            relocation_start = mining_end
            relocation_end = relocation_start + relocation_days
            tasks.append({"type": "relocation", "workface": workface, "startDay": relocation_start, "endDay": relocation_end})
            last_relocation_end = relocation_end
        last_mining_end = mining_end

    coal_density = max(0.0, safe_float(params.get("coalDensity"), 1.35) or 1.35)
    mining_height_m = max(0.0, safe_float(params.get("miningHeightM"), 4.5) or 4.5)
    rr_min = clamp01(params.get("recoveryRateMin"), 0.85)
    rr_max = clamp01(params.get("recoveryRateMax"), 0.95)
    recovery_rate = clamp01((rr_min + rr_max) / 2.0, 0.90)
    panels_by_id = {str(panel.get("id")): panel for panel in panels}

    mining_tasks = [task for task in tasks if task.get("type") == "mining"]
    monthly: list[dict] = []
    if mining_tasks:
        total_end_day = max(task["endDay"] for task in mining_tasks)
        total_months = max(1, int(math.ceil(total_end_day / days_per_month)))
        for month in range(1, total_months + 1):
            month_start = (month - 1) * days_per_month
            month_end = month * days_per_month
            tonnage = 0.0
            mined_len = 0.0
            overlap_by_workface: dict[str, float] = {}
            for task in mining_tasks:
                overlap = max(0.0, min(task["endDay"], month_end) - max(task["startDay"], month_start))
                if overlap <= 1e-9:
                    continue
                wf_id = str(task.get("workface") or "")
                overlap_by_workface[wf_id] = overlap_by_workface.get(wf_id, 0.0) + overlap
                panel = panels_by_id.get(wf_id) or {}
                width_m = max(0.0, safe_float(panel.get("widthM"), 0.0) or 0.0)
                length = overlap * shear_advance_rate * utilization
                mined_len += length
                volume = length * width_m * mining_height_m
                tonnage += volume * coal_density * recovery_rate
            workface = ""
            if overlap_by_workface:
                workface = ",".join(
                    key for key, _ in sorted(overlap_by_workface.items(), key=lambda item: (-item[1], item[0]))
                )
            monthly.append({"month": month, "tonnage": tonnage, "minedLen": mined_len, "workface": workface})

    total_end_day = max((task["endDay"] for task in tasks), default=0.0)
    return {
        "ok": True,
        "computedAt": datetime.now().isoformat(),
        "tasks": tasks,
        "daysPerMonth": days_per_month,
        "totalMonths": int(math.ceil(total_end_day / days_per_month)) if total_end_day > 0 else 0,
        "monthly": monthly,
    }


def compute_target_tons_per_month(mine_capacity_wan_per_year: object) -> float | None:
    capacity = safe_float(mine_capacity_wan_per_year)
    if capacity is None or capacity <= 0:
        return None
    return (capacity * 10000.0) / 12.0


def compute_production_kpis(monthly_rows: list[dict], target_tons_per_month: float | None) -> dict:
    if not monthly_rows:
        return {"months": 0, "minTonnage": None, "meanTonnage": None, "maxDeficit": None, "hitRate": None}
    tonnages = [max(0.0, safe_float(row.get("tonnage"), 0.0) or 0.0) for row in monthly_rows]
    deficits = []
    hits = 0
    for tonnage in tonnages:
        if target_tons_per_month is None:
            continue
        deficit = max(0.0, target_tons_per_month - tonnage)
        deficits.append(deficit)
        if tonnage >= target_tons_per_month:
            hits += 1
    return {
        "months": len(monthly_rows),
        "minTonnage": min(tonnages),
        "meanTonnage": float(np.mean(tonnages)),
        "maxDeficit": max(deficits) if deficits else None,
        "hitRate": (hits / len(monthly_rows)) if (target_tons_per_month is not None and monthly_rows) else None,
    }


def build_stage3_candidates(base_params: dict, current_order_mode: str) -> list[dict]:
    crews = max(1, int(round(safe_float(base_params.get("driveCrews"), 1.0) or 1.0)))
    install_days = max(0.0, safe_float(base_params.get("installDays"), 15.0) or 15.0)
    relocation_days = max(0.0, safe_float(base_params.get("relocationDays"), 10.0) or 10.0)
    shear = max(0.0, safe_float(base_params.get("shearAdvanceRate"), 6.0) or 6.0)
    order_mode = str(current_order_mode or "faceIndex")
    return [
        {"key": "base", "label": "当前参数", "patch": {}, "orderMode": order_mode},
        {"key": "crew+1", "label": "掘进队+1", "patch": {"driveCrews": min(4, crews + 1)}, "orderMode": order_mode},
        {"key": "install-20%", "label": "安装工期-20%", "patch": {"installDays": max(0, round(install_days * 0.8))}, "orderMode": order_mode},
        {"key": "reloc-20%", "label": "搬家工期-20%", "patch": {"relocationDays": max(0, round(relocation_days * 0.8))}, "orderMode": order_mode},
        {
            "key": "reloc-pack",
            "label": "安装/搬家协同提效",
            "patch": {
                "installDays": max(0, round(install_days * 0.8)),
                "relocationDays": max(0, round(relocation_days * 0.8)),
            },
            "orderMode": order_mode,
        },
        {"key": "shear+10%", "label": "推进速度+10%", "patch": {"shearAdvanceRate": round(shear * 1.10, 2)}, "orderMode": order_mode},
        {
            "key": "crew+1-install-20%",
            "label": "掘进队+1 + 安装-20%",
            "patch": {"driveCrews": min(4, crews + 1), "installDays": max(0, round(install_days * 0.8))},
            "orderMode": order_mode,
        },
    ]


def score_stage3_scenario(prod_kpis: dict, risk_rows: list[dict] | None, weights: dict) -> dict:
    w_prod = safe_float(weights.get("wProd"), 1.0) or 1.0
    w_risk = safe_float(weights.get("wRisk"), 1.0) or 1.0
    w_months = safe_float(weights.get("wMonths"), 0.15) or 0.15
    hit_rate = safe_float(prod_kpis.get("hitRate"), 0.0) or 0.0
    max_deficit = safe_float(prod_kpis.get("maxDeficit"), 0.0) or 0.0
    months = safe_float(prod_kpis.get("months"), 0.0) or 0.0
    risk_values = [safe_float(row.get("value")) for row in (risk_rows or [])]
    risk_values = [value for value in risk_values if value is not None]
    risk_max = max(risk_values) if risk_values else None
    s_prod = hit_rate * 100.0 - (max_deficit / 1000.0)
    s_risk = ((1.0 - clamp01(risk_max, 0.0)) * 50.0) if risk_max is not None else 0.0
    s_months = -months
    return {"score": w_prod * s_prod + w_risk * s_risk + w_months * s_months, "riskMax": risk_max}


def compute_succession_artifacts(scene: dict) -> dict | None:
    raw_panels = build_raw_succession_panels(scene)
    if not raw_panels:
        return None
    succession = scene.get("succession") or {}
    stage1_params = dict((succession.get("stage1Params") or {}))
    planning_params = scene.get("planningParams") or {}
    yard_dir = str(succession.get("yardDir") or (succession.get("yardConfirmed") or {}).get("dir") or "NE")
    yard_offset = safe_float((succession.get("yardConfirmed") or {}).get("offsetM"), safe_float(succession.get("yardOffsetM"), 120.0) or 120.0) or 120.0
    ordered_panels = order_succession_panels(raw_panels, str(succession.get("panelOrderMode") or "faceIndex"), yard_dir, yard_offset)
    plan_params = {
        **stage1_params,
        "coalDensity": planning_params.get("coalDensity", 1.35),
        "recoveryRateMin": planning_params.get("recoveryRateMin", 0.85),
        "recoveryRateMax": planning_params.get("recoveryRateMax", 0.95),
        "miningHeightM": get_scene_mining_height(scene),
    }
    plan = build_succession_stage1_plan(ordered_panels, plan_params)
    target = compute_target_tons_per_month(planning_params.get("mineCapacity"))
    stage3_params = dict((succession.get("stage3Params") or {}))
    stage3_results = []
    for candidate in build_stage3_candidates(stage1_params, str(succession.get("panelOrderMode") or "faceIndex")):
        candidate_panels = order_succession_panels(raw_panels, str(candidate.get("orderMode") or "faceIndex"), yard_dir, yard_offset)
        candidate_plan = build_succession_stage1_plan(candidate_panels, {**plan_params, **(candidate.get("patch") or {})})
        prod_kpis = compute_production_kpis(candidate_plan.get("monthly") or [], target)
        stage3_score = score_stage3_scenario(prod_kpis, None, stage3_params)
        stage3_results.append(
            {
                "key": candidate.get("key"),
                "label": candidate.get("label"),
                "orderMode": candidate.get("orderMode"),
                "patch": candidate.get("patch") or {},
                "planSummary": {
                    "months": prod_kpis.get("months"),
                    "hitRate": prod_kpis.get("hitRate"),
                    "maxDeficit": prod_kpis.get("maxDeficit"),
                    "minTonnage": prod_kpis.get("minTonnage"),
                    "riskMax": stage3_score.get("riskMax"),
                },
                "score": stage3_score.get("score"),
                "riskSource": None,
            }
        )
    stage3_results.sort(key=lambda row: safe_float(row.get("score"), -1e18) or -1e18, reverse=True)
    return {
        "panels": ordered_panels,
        "planParams": plan_params,
        "plan": plan,
        "targetTonsPerMonth": target,
        "stage3Results": stage3_results[:8],
    }


def plot_monthly_production(plan: dict, target_tons_per_month: float | None, title: str) -> plt.Figure | None:
    monthly = plan.get("monthly") or []
    if not monthly:
        return None
    months = [int(row.get("month", index + 1)) for index, row in enumerate(monthly)]
    tonnage = [(safe_float(row.get("tonnage"), 0.0) or 0.0) / 10000.0 for row in monthly]
    target = (target_tons_per_month / 10000.0) if target_tons_per_month is not None else None
    fig, ax = plt.subplots(figsize=(7.2, 3.8))
    ax.plot(months, tonnage, color="#015699", linewidth=2.0, marker="o", markersize=3, label="月产量")
    if target is not None:
        ax.plot(months, [target] * len(months), color="#111827", linewidth=1.4, linestyle="--", label="目标产量")
    ax.set_xlabel("月份")
    ax.set_ylabel("产量 / 万t")
    ax.set_title(title, fontweight="bold")
    ax.grid(color="#e2e8f0", linewidth=0.5, linestyle=":")
    ax.legend(frameon=False)
    fig.tight_layout()
    return fig


def plot_schedule_gantt(plan: dict, title: str) -> plt.Figure | None:
    tasks = [task for task in (plan.get("tasks") or []) if isinstance(task, dict)]
    if not tasks:
        return None
    workfaces = []
    for task in tasks:
        workface = str(task.get("workface") or "")
        if workface and workface not in workfaces:
            workfaces.append(workface)
    if not workfaces:
        return None
    y_map = {workface: index for index, workface in enumerate(workfaces)}
    colors = {"drive": "#4f596d", "install": "#fac00f", "mining": "#015699", "relocation": "#f3764a"}

    fig_height = max(3.8, 0.55 * len(workfaces) + 1.8)
    fig, ax = plt.subplots(figsize=(8.0, fig_height))
    for task in tasks:
        workface = str(task.get("workface") or "")
        if workface not in y_map:
            continue
        start = safe_float(task.get("startDay"), 0.0) or 0.0
        end = safe_float(task.get("endDay"), start) or start
        width = max(0.0, end - start)
        task_type = str(task.get("type") or "other")
        ax.barh(y_map[workface], width, left=start, height=0.56, color=colors.get(task_type, "#94a3b8"), edgecolor="white", alpha=0.92)

    days_per_month = int(plan.get("daysPerMonth") or 25)
    max_day = max((safe_float(task.get("endDay"), 0.0) or 0.0) for task in tasks)
    for day in range(days_per_month, int(math.ceil(max_day)) + 1, days_per_month):
        ax.axvline(day, color="#cbd5e1", linewidth=0.8, linestyle=":")
    ax.set_yticks(list(y_map.values()))
    ax.set_yticklabels(workfaces)
    ax.set_xlabel("工期 / 天")
    ax.set_title(title, fontweight="bold")
    ax.grid(axis="x", color="#e2e8f0", linewidth=0.5, linestyle=":")
    ax.legend(
        handles=[Patch(facecolor=color, edgecolor="white", label=label) for label, color in {"掘进": colors["drive"], "安装": colors["install"], "回采": colors["mining"], "搬家": colors["relocation"]}.items()],
        frameon=False,
        ncol=4,
        loc="upper right",
    )
    ax.invert_yaxis()
    fig.tight_layout()
    return fig


def plot_stage3_candidate_scores(results: list[dict], title: str) -> plt.Figure | None:
    if not results:
        return None
    labels = [str(item.get("label") or item.get("key")) for item in results[:8]][::-1]
    values = [(safe_float(item.get("score"), 0.0) or 0.0) for item in results[:8]][::-1]
    fig, ax = plt.subplots(figsize=(7.2, 4.2))
    bars = ax.barh(labels, values, color="#6366f1", alpha=0.92)
    ax.set_xlabel("综合评分")
    ax.set_title(title, fontweight="bold")
    ax.grid(axis="x", color="#e2e8f0", linewidth=0.5, linestyle=":")
    for bar, value in zip(bars, values):
        ax.text(bar.get_width(), bar.get_y() + bar.get_height() / 2, f" {value:.2f}", va="center", ha="left", fontsize=8)
    fig.tight_layout()
    return fig


def compute_economics_from_plan(plan: dict, risk: dict | None, params: dict) -> dict:
    monthly = plan.get("monthly") or []
    if not monthly:
        return {"ok": False, "reason": "missing monthly production"}

    price = max(0.0, safe_float(params.get("coalPriceYuanPerTon"), 800.0) or 800.0)
    sales_ratio = clamp01(params.get("salesRatio"), 1.0)
    var_cost = max(0.0, safe_float(params.get("opexVarYuanPerTon"), 320.0) or 320.0)
    fixed_cost_wan = max(0.0, safe_float(params.get("opexFixedWanPerMonth"), 300.0) or 300.0)
    initial_capex_wan = max(0.0, safe_float(params.get("capexInitialWan"), 30000.0) or 30000.0)
    sustain_capex_wan_per_year = max(0.0, safe_float(params.get("capexSustainWanPerYear"), 0.0) or 0.0)
    discount_rate = clamp01(params.get("discountRate"), 0.10)
    monthly_discount_rate = (1.0 + discount_rate) ** (1.0 / 12.0) - 1.0
    risk_link_enabled = parse_bool(params.get("riskLinkEnabled"), True)
    risk_metric_key = str(params.get("riskMetricKey") or ((risk or {}).get("metric") or "p90"))
    risk_threshold = clamp01(params.get("riskImpactThreshold"), 0.85)
    risk_downtime_ratio = clamp01(params.get("riskDowntimeRatio"), 0.10)
    risk_extra_wan = max(0.0, safe_float(params.get("riskExtraCostWanPerHighRiskMonth"), 0.0) or 0.0)

    risk_row_by_month = {}
    for row in ((risk or {}).get("rows") or []):
        month = int(safe_float(row.get("month"), 0.0) or 0.0)
        if month >= 1:
            risk_row_by_month[month] = row

    rows = []
    cumulative_cash_wan = 0.0
    for month_row in monthly:
        month = int(safe_float(month_row.get("month"), 0.0) or 0.0)
        if month < 1:
            continue
        tonnage = max(0.0, safe_float(month_row.get("tonnage"), 0.0) or 0.0)
        workface = str(month_row.get("workface") or "")
        risk_row = risk_row_by_month.get(month) or {}
        risk_value = safe_float(risk_row.get(risk_metric_key))
        is_high_risk = bool(risk_link_enabled and risk_value is not None and risk_value >= risk_threshold)
        tonnage_adj = tonnage * (1.0 - risk_downtime_ratio) if is_high_risk else tonnage
        revenue_wan = (tonnage_adj * price * sales_ratio) / 10000.0
        var_cost_wan = (tonnage_adj * var_cost) / 10000.0
        sustain_wan = sustain_capex_wan_per_year / 12.0
        risk_extra_cost_wan = risk_extra_wan if is_high_risk else 0.0
        initial_wan = initial_capex_wan if month == 1 else 0.0
        net_cash_wan = revenue_wan - var_cost_wan - fixed_cost_wan - sustain_wan - risk_extra_cost_wan - initial_wan
        cumulative_cash_wan += net_cash_wan
        rows.append(
            {
                "month": month,
                "workface": workface,
                "tonnage": tonnage,
                "tonnageAdj": tonnage_adj,
                "risk": risk_value,
                "isHighRisk": is_high_risk,
                "revenueWan": revenue_wan,
                "varCostWan": var_cost_wan,
                "fixedCostWan": fixed_cost_wan,
                "sustainCapexWan": sustain_wan,
                "riskExtraCostWan": risk_extra_cost_wan,
                "capexInitialWan": initial_wan,
                "netCashWan": net_cash_wan,
                "cumCashWan": cumulative_cash_wan,
            }
        )

    npv_wan = 0.0
    for index, row in enumerate(rows, start=1):
        discount = (1.0 + monthly_discount_rate) ** index
        npv_wan += row["netCashWan"] / discount if discount > 0 else row["netCashWan"]
    payback_month = next((row["month"] for row in rows if row["cumCashWan"] >= 0), None)
    total_revenue_wan = sum(row["revenueWan"] for row in rows)
    total_cost_wan = sum(row["varCostWan"] + row["fixedCostWan"] + row["sustainCapexWan"] + row["riskExtraCostWan"] + row["capexInitialWan"] for row in rows)
    total_net_cash_wan = sum(row["netCashWan"] for row in rows)
    total_tonnage = sum(row["tonnageAdj"] for row in rows)
    unit_cost = ((total_cost_wan * 10000.0) / total_tonnage) if total_tonnage > 1e-9 else None
    return {
        "ok": True,
        "computedAt": datetime.now().isoformat(),
        "rows": rows,
        "summary": {
            "months": len(rows),
            "npvWan": npv_wan,
            "discountRateYear": discount_rate,
            "discountRateMonth": monthly_discount_rate,
            "paybackMonth": payback_month,
            "totalRevenueWan": total_revenue_wan,
            "totalCostWan": total_cost_wan,
            "totalNetCashWan": total_net_cash_wan,
            "unitCostYuanPerTon": unit_cost,
            "unitMarginYuanPerTon": (price - unit_cost) if unit_cost is not None else None,
            "highRiskMonths": sum(1 for row in rows if row["isHighRisk"]),
            "lastCumCashWan": rows[-1]["cumCashWan"] if rows else 0.0,
        },
    }


def build_economics_chart_rows(econ_result: dict) -> list[dict]:
    chart_rows = []
    for index, row in enumerate(econ_result.get("rows") or []):
        year = index // 12 + 1
        month_in_year = index % 12 + 1
        total_cost_wan = row["varCostWan"] + row["fixedCostWan"] + row["sustainCapexWan"] + row["riskExtraCostWan"] + row["capexInitialWan"]
        margin_wan = row["revenueWan"] - total_cost_wan
        tonnage_adj = safe_float(row.get("tonnageAdj"), 0.0) or 0.0
        chart_rows.append(
            {
                **row,
                "year": year,
                "monthInYear": month_in_year,
                "tonnageWanT": (safe_float(row.get("tonnage"), 0.0) or 0.0) / 10000.0,
                "tonnageAdjWanT": tonnage_adj / 10000.0,
                "unitTotalCost": ((total_cost_wan * 10000.0) / tonnage_adj) if tonnage_adj > 1e-9 else None,
                "unitMargin": ((margin_wan * 10000.0) / tonnage_adj) if tonnage_adj > 1e-9 else None,
                "unitNetCash": ((row["netCashWan"] * 10000.0) / tonnage_adj) if tonnage_adj > 1e-9 else None,
            }
        )
    return chart_rows


def build_cashflow_chart_rows(econ_result: dict) -> list[dict]:
    rows = build_economics_chart_rows(econ_result)
    monthly_discount_rate = safe_float((econ_result.get("summary") or {}).get("discountRateMonth"), 0.0) or 0.0
    discounted_running = 0.0
    out = []
    for row in rows:
        month = int(row["month"])
        discount = (1.0 + monthly_discount_rate) ** month
        pv = row["netCashWan"] / discount if discount > 0 else row["netCashWan"]
        discounted_running += pv
        out.append({**row, "discountedCumWan": discounted_running})
    return out


def plot_cashflow(econ_result: dict, title: str) -> plt.Figure | None:
    rows = build_cashflow_chart_rows(econ_result)
    if not rows:
        return None
    x = [row["month"] for row in rows]
    fig, ax = plt.subplots(figsize=(7.6, 4.0))
    ax.bar(x, [row["netCashWan"] for row in rows], color="#4d85bd", alpha=0.92, label="净现金流")
    ax.plot(x, [row["cumCashWan"] for row in rows], color="#f7903d", linewidth=2.0, marker="o", markersize=2.8, label="累计现金流")
    ax.plot(x, [row["discountedCumWan"] for row in rows], color="#59a95a", linewidth=1.8, linestyle="--", label="折现累计")
    ax.set_xlabel("月份")
    ax.set_ylabel("万元")
    ax.set_title(title, fontweight="bold")
    ax.grid(axis="y", color="#e2e8f0", linewidth=0.5, linestyle=":")
    ax.legend(frameon=False)
    fig.tight_layout()
    return fig


def plot_revenue_cost(econ_result: dict, title: str) -> plt.Figure | None:
    rows = build_economics_chart_rows(econ_result)
    if not rows:
        return None
    x = np.arange(len(rows))
    labels = [row["month"] for row in rows]
    var_cost = np.array([row["varCostWan"] for row in rows], dtype=float)
    sustain = np.array([row["sustainCapexWan"] for row in rows], dtype=float)
    risk_extra = np.array([row["riskExtraCostWan"] for row in rows], dtype=float)
    initial = np.array([row["capexInitialWan"] for row in rows], dtype=float)
    revenue = np.array([row["revenueWan"] for row in rows], dtype=float)

    fig, ax = plt.subplots(figsize=(8.0, 4.2))
    ax.bar(x, var_cost, color="#d22027", label="变动成本")
    ax.bar(x, sustain, bottom=var_cost, color="#4d85bd", label="维持投资")
    ax.bar(x, risk_extra, bottom=var_cost + sustain, color="#fe817d", label="风险附加")
    ax.bar(x, initial, bottom=var_cost + sustain + risk_extra, color="#385989", label="初始投资")
    ax.plot(x, revenue, color="#111827", linewidth=2.0, marker="o", markersize=2.8, label="收入")
    ax.set_xticks(x)
    ax.set_xticklabels(labels)
    ax.set_xlabel("月份")
    ax.set_ylabel("万元")
    ax.set_title(title, fontweight="bold")
    ax.grid(axis="y", color="#e2e8f0", linewidth=0.5, linestyle=":")
    ax.legend(frameon=False, ncol=3)
    fig.tight_layout()
    return fig


def plot_cost_structure(econ_result: dict, title: str) -> plt.Figure | None:
    rows = econ_result.get("rows") or []
    if not rows:
        return None
    parts = [
        ("变动成本", sum(row["varCostWan"] for row in rows), "#4d85bd"),
        ("固定成本", sum(row["fixedCostWan"] for row in rows), "#81b8df"),
        ("维持投资", sum(row["sustainCapexWan"] for row in rows), "#59a95a"),
        ("风险附加", sum(row["riskExtraCostWan"] for row in rows), "#fe817d"),
        ("初始投资", sum(row["capexInitialWan"] for row in rows), "#385989"),
    ]
    parts = [part for part in parts if abs(part[1]) > 1e-9]
    if not parts:
        return None
    fig, ax = plt.subplots(figsize=(5.2, 4.2))
    ax.pie([part[1] for part in parts], labels=[part[0] for part in parts], colors=[part[2] for part in parts], autopct=lambda pct: f"{pct:.1f}%" if pct > 0 else "", startangle=90, counterclock=False, wedgeprops={"linewidth": 0.8, "edgecolor": "white"}, textprops={"fontsize": 8})
    ax.set_title(title, fontweight="bold")
    fig.tight_layout()
    return fig


def build_scene_figure_lookup(output_root: Path, summaries: Sequence[ExportSummary]) -> dict[str, list[str]]:
    lookup: dict[str, list[str]] = {}
    for summary in summaries:
        scene_dir = output_root / summary.scene_slug
        figures = sorted(str(path.relative_to(output_root)).replace("\\", "/") for path in scene_dir.rglob("*.svg"))
        lookup[summary.scene_slug] = figures
    return lookup


def _pick_matching_paths(figures: dict[str, list[str]], scene_slug: str, tokens: Sequence[str]) -> list[str]:
    matches = []
    for path in figures.get(scene_slug, []):
        if all(token in path for token in tokens):
            matches.append(path)
    return matches


def write_paper_figure_guide(
    output_root: Path,
    summaries: Sequence[ExportSummary],
    heading_lines: Sequence[str] | None = None,
    paper_docx_path: Path | None = None,
) -> Path:
    headings = list(heading_lines or [])
    if not headings and paper_docx_path:
        headings = extract_docx_heading_lines(paper_docx_path)
    if not headings:
        headings = ["2.2 连续参数场构建结果", "2.3 多场景 ODI 风险表征结果", "3.1 结构化规划结果", "3.2 四模式候选规划与方案比选", "3.3 规划结果向采掘接续与工程经济评价的传递"]

    figure_lookup = build_scene_figure_lookup(output_root, summaries)
    scene_list = ", ".join(summary.scene_slug for summary in summaries)
    lines = ["# 论文插图建议", "", f"- 导出目录：`{output_root}`", f"- 场景范围：{scene_list}"]
    if paper_docx_path:
        lines.append(f"- 论文文件：`{paper_docx_path}`")
    lines.extend(["", "以下建议按论文章节组织，优先选择已经导出的矢量图（SVG/PDF）。", ""])

    for heading in headings:
        text = str(heading).strip()
        if not text:
            continue
        lines.append(f"## {text}")
        recs: list[str] = []
        if "2.2" in text:
            for slug in ("00_surface_subsidence", "01_aquifer_pre_eval", "02_aquifer_eval", "06_full_overburden"):
                recs.extend(_pick_matching_paths(figure_lookup, slug, (FIGURE_NAME_MAP["fig02_geology_clouds"],)))
        elif "2.3" in text:
            for slug in ("00_surface_subsidence", "01_aquifer_pre_eval", "02_aquifer_eval", "04_cocontrol_water_inrush", "06_full_overburden"):
                recs.extend(_pick_matching_paths(figure_lookup, slug, (FIGURE_NAME_MAP["fig01_odi_distribution"],)))
                recs.extend(_pick_matching_paths(figure_lookup, slug, (FIGURE_NAME_MAP["fig05_odi_levels"],)))
        elif "3.1" in text:
            for slug in ("03_mining_planning", "05_mining_succession"):
                recs.extend(_pick_matching_paths(figure_lookup, slug, ("overview/" + FIGURE_NAME_MAP["fig01_workface_plan_layout"],)))
                recs.extend(_pick_matching_paths(figure_lookup, slug, ("overview/" + FIGURE_NAME_MAP["fig02_planning_mode_scores"],)))
        elif "3.2" in text:
            for slug in ("03_mining_planning", "05_mining_succession"):
                recs.extend(_pick_matching_paths(figure_lookup, slug, ("overview/" + FIGURE_NAME_MAP["fig03_weighted_top_candidates"],)))
        elif "3.3" in text:
            slug = "05_mining_succession"
            recs.extend(_pick_matching_paths(figure_lookup, slug, ("succession/" + FIGURE_NAME_MAP["fig01_monthly_production"],)))
            recs.extend(_pick_matching_paths(figure_lookup, slug, ("succession/" + FIGURE_NAME_MAP["fig02_schedule_gantt"],)))
            recs.extend(_pick_matching_paths(figure_lookup, slug, ("succession/" + FIGURE_NAME_MAP["fig03_stage3_candidate_scores"],)))
            recs.extend(_pick_matching_paths(figure_lookup, slug, ("economics/" + FIGURE_NAME_MAP["fig01_cashflow"],)))
            recs.extend(_pick_matching_paths(figure_lookup, slug, ("economics/" + FIGURE_NAME_MAP["fig02_revenue_cost"],)))
            recs.extend(_pick_matching_paths(figure_lookup, slug, ("economics/" + FIGURE_NAME_MAP["fig03_cost_structure"],)))

        if recs:
            for rec in recs:
                lines.append(f"- 建议插入：`{rec}`")
        else:
            lines.append("- 暂未匹配到对应图件，可结合本节文字说明从同场景相邻图件中补选。")
        lines.append("")

    lines.extend(["## 插图编排建议", "- 参数场与 ODI 主图建议放在第 2 章，优先展示方法输入与多场景风险差异。", "- 四模式规划与加权优选图建议放在第 3 章前半部分，作为方案比选核心证据。", "- 采掘接续与工程经济图建议放在第 3.3 节，形成“规划 -> 排程 -> 经济”链条。", "- 若版面有限，优先保留主图、对比图和一张经济性主图，其他图作为补充材料。", ""])
    guide_path = output_root / "paper_figure_guide.md"
    guide_path.write_text("\n".join(lines), encoding="utf-8")
    return guide_path


def extract_docx_heading_lines(docx_path: Path) -> list[str]:
    try:
        from docx import Document  # type: ignore
    except Exception:
        return []
    document = Document(str(docx_path))
    headings: list[str] = []
    for paragraph in document.paragraphs:
        text = str(paragraph.text or "").strip()
        if not text:
            continue
        style_name = str(getattr(paragraph.style, "name", "") or "")
        if style_name.lower().startswith("heading"):
            headings.append(text)
            continue
        if re.match(r"^\d+(\.\d+)*\s+\S+", text):
            headings.append(text)
    return headings


def write_paper_figure_guide_docx(markdown_path: Path, paper_docx_path: Path) -> Path | None:
    try:
        from docx import Document  # type: ignore
    except Exception:
        return None
    content = markdown_path.read_text(encoding="utf-8").splitlines()
    document = Document()
    title = document.add_heading("论文插图建议", level=0)
    title.alignment = 1
    document.add_paragraph(f"对应论文：{paper_docx_path}")
    document.add_paragraph(f"图件目录：{markdown_path.parent}")
    for line in content:
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        if line.strip():
            document.add_paragraph(line.strip())
    output_path = paper_docx_path.with_name(f"{paper_docx_path.stem}_插图建议.docx")
    document.save(str(output_path))
    return output_path


def _csv_content(rows: Iterable[Sequence[object]]) -> str:
    buffer: list[str] = []
    for row in rows:
        cells = []
        for cell in row:
            value = str(cell if cell is not None else "")
            value = value.replace('"', '""')
            if any(token in value for token in [",", '"', "\n", "\r"]):
                value = f'"{value}"'
            cells.append(value)
        buffer.append(",".join(cells))
    return "\ufeff" + "\n".join(buffer) + "\n"


def write_csv(path: Path, rows: Iterable[Sequence[object]]) -> Path:
    path.write_text(_csv_content(rows), encoding="utf-8")
    return path


def write_json(path: Path, payload: dict) -> Path:
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def save_figure(fig: plt.Figure | None, stem: Path, formats: Sequence[str]) -> list[Path]:
    if fig is None:
        return []
    stem.parent.mkdir(parents=True, exist_ok=True)
    written: list[Path] = []
    for fmt in formats:
        out_path = stem.with_suffix(f".{fmt}")
        fig.savefig(out_path, format=fmt, bbox_inches="tight", facecolor="white")
        written.append(out_path)
    plt.close(fig)
    return written


def figure_stem_path(parent: Path, figure_key: str) -> Path:
    return parent / FIGURE_NAME_MAP.get(figure_key, figure_key)


def error_trend_stem(line_pack: dict, index: int, line_key: str) -> str:
    label = str(line_pack.get("label") or line_key or f"测线{index}")
    label = re.sub(r'[\\\\/:*?"<>|]+', "_", label).strip()
    if not label:
        label = f"测线{index}"
    return f"{index + 6:02d}-误差趋势-{label}"


def export_scene_file(scene_file: Path, output_root: Path, formats: Sequence[str] = ("svg", "pdf", "png")) -> ExportSummary:
    scene = load_scene(scene_file)
    summary = ExportSummary(
        scene_name=strip_double_suffix(scene_file),
        scene_slug=scene_slug(scene_file),
        source_file=scene_file,
    )
    scene_dir = output_root / summary.scene_slug
    scene_dir.mkdir(parents=True, exist_ok=True)
    metadata_path = scene_dir / "scene_metadata.json"
    summary.files.append(
        write_json(
            metadata_path,
            {
                "sceneName": summary.scene_name,
                "sceneSlug": summary.scene_slug,
                "sourceFile": str(scene_file),
                "exportedAt": datetime.now().isoformat(),
            },
        )
    )

    tabs = iter_tabs_with_data(scene)
    for tab_id in tabs:
        tab = get_tab(scene, tab_id)
        tab_dir = scene_dir / tab_id
        tab_dir.mkdir(parents=True, exist_ok=True)

        odi_points = get_odi_points(scene, tab_id)
        param_points = get_param_points(scene, tab_id)
        drillholes = list(tab.get("drillholeData") or [])
        boundary = list(tab.get("boundaryData") or [])
        workfaces = list(tab.get("workingFaceData") or [])
        measured = list(tab.get("measuredConstraintData") or [])
        level_ranges = get_level_ranges(tab)
        error_by_line = tab.get("errorAnalysisByLineId") or {}
        weights = ((tab.get("odiResult") or {}).get("weights") or {})

        figure_builders = [
            ("fig01_odi_distribution", plot_odi_heatmap(odi_points, drillholes, boundary, workfaces, f"{tab_id} - ODI 分布")),
            ("fig02_geology_clouds", plot_geology_clouds(param_points, drillholes, f"{tab_id} - 地质参数分布")),
            ("fig03_spatial_map", plot_spatial_map(odi_points, drillholes, boundary, workfaces, f"{tab_id} - 空间分布")),
            ("fig04_odi_histogram", plot_odi_histogram(odi_points, f"{tab_id} - ODI 频数分布")),
            ("fig05_odi_levels", plot_odi_level_pie(odi_points, level_ranges, f"{tab_id} - ODI 分级")),
            ("fig06_weight_radar", plot_weight_radar(weights, f"{tab_id} - 权重雷达图")),
        ]
        for stem_name, figure in figure_builders:
            written = save_figure(figure, figure_stem_path(tab_dir, stem_name), formats)
            if written:
                summary.files.extend(written)
                summary.figure_count += 1

        for index, (line_key, line_pack) in enumerate(error_by_line.items(), start=1):
            error_fig = plot_error_trend(line_pack.get("data") or [], f"{tab_id} - 误差趋势 - {line_pack.get('label', line_key)}")
            written = save_figure(error_fig, tab_dir / error_trend_stem(line_pack, index, line_key), formats)
            if written:
                summary.files.extend(written)
                summary.figure_count += 1

        if odi_points:
            rows = [
                ["ID", "Cat", "X", "Y", "ODI", "ODI_norm", "Ti", "Hi", "Di", "Mi"],
                *[
                    [
                        point.get("id", ""),
                        point.get("cat", ""),
                        point.get("x", ""),
                        point.get("y", ""),
                        point.get("odi", ""),
                        point.get("odiNorm", ""),
                        point.get("Ti", ""),
                        point.get("Hi", ""),
                        point.get("Di", ""),
                        point.get("Mi", ""),
                    ]
                    for point in odi_points
                ],
            ]
            summary.files.append(write_csv(tab_dir / "data_odi_points.csv", rows))
            summary.files.append(
                write_json(
                    tab_dir / "odi_summary.json",
                    {
                        "pointCount": len(odi_points),
                        "minOdi": (tab.get("odiResult") or {}).get("minOdi"),
                        "maxOdi": (tab.get("odiResult") or {}).get("maxOdi"),
                        "weights": weights,
                    },
                )
            )

        if param_points:
            rows = [
                ["ID", "X", "Y", "Ti", "Hi", "Di", "Mi"],
                *[
                    [point.get("id", ""), point.get("x", ""), point.get("y", ""), point.get("Ti", ""), point.get("Hi", ""), point.get("Di", ""), point.get("Mi", "")]
                    for point in param_points
                ],
            ]
            summary.files.append(write_csv(tab_dir / "data_parameters.csv", rows))

        if measured:
            rows = [
                ["ID", "X", "Y", "Measured"],
                *[[point.get("id", ""), point.get("x", ""), point.get("y", ""), point.get("measured", "")] for point in measured],
            ]
            summary.files.append(write_csv(tab_dir / "data_measured.csv", rows))

    overview_dir = scene_dir / "overview"
    overview_dir.mkdir(parents=True, exist_ok=True)
    plan_loops = get_workface_plan_loops(scene)
    written = save_figure(
        plot_workface_layout(plan_loops, f"{summary.scene_name} - 采区规划布局"),
        figure_stem_path(overview_dir, "fig01_workface_plan_layout"),
        formats,
    )
    if written:
        summary.files.extend(written)
        summary.figure_count += 1

    for mode in ("efficiency", "recovery", "disturbance", "weighted"):
        rows = get_planning_rows(scene, mode)
        if rows:
            header = sorted({key for row in rows for key in row.keys()})
            csv_rows = [header] + [[row.get(key, "") for key in header] for row in rows]
            summary.files.append(write_csv(overview_dir / f"planning_{mode}_table.csv", csv_rows))

    for stem_name, figure in [
        ("fig02_planning_mode_scores", plot_planning_mode_scores(scene, f"{summary.scene_name} - Planning Mode Scores")),
        ("fig03_weighted_top_candidates", plot_weighted_top_candidates(scene, f"{summary.scene_name} - Weighted Top Candidates")),
    ]:
        written = save_figure(figure, figure_stem_path(overview_dir, stem_name), formats)
        if written:
            summary.files.extend(written)
            summary.figure_count += 1

    succession_artifacts = compute_succession_artifacts(scene)
    if succession_artifacts:
        succession_dir = scene_dir / "succession"
        succession_dir.mkdir(parents=True, exist_ok=True)
        plan = succession_artifacts["plan"]
        target = succession_artifacts["targetTonsPerMonth"]

        for stem_name, figure in [
            ("fig01_monthly_production", plot_monthly_production(plan, target, f"{summary.scene_name} - Monthly Production")),
            ("fig02_schedule_gantt", plot_schedule_gantt(plan, f"{summary.scene_name} - Succession Gantt")),
            ("fig03_stage3_candidate_scores", plot_stage3_candidate_scores(succession_artifacts["stage3Results"], f"{summary.scene_name} - Stage3 Candidate Scores")),
        ]:
            written = save_figure(figure, figure_stem_path(succession_dir, stem_name), formats)
            if written:
                summary.files.extend(written)
                summary.figure_count += 1

        monthly_rows = [["month", "tonnage", "minedLen", "workface"]]
        monthly_rows.extend([[row.get("month", ""), row.get("tonnage", ""), row.get("minedLen", ""), row.get("workface", "")] for row in (plan.get("monthly") or [])])
        summary.files.append(write_csv(succession_dir / "stage1_monthly.csv", monthly_rows))

        task_rows = [["type", "workface", "startDay", "endDay", "lengthM"]]
        task_rows.extend([[row.get("type", ""), row.get("workface", ""), row.get("startDay", ""), row.get("endDay", ""), row.get("lengthM", "")] for row in (plan.get("tasks") or [])])
        summary.files.append(write_csv(succession_dir / "stage1_tasks.csv", task_rows))

        stage3_rows = [["key", "label", "orderMode", "score", "months", "hitRate", "maxDeficit", "minTonnage", "riskMax"]]
        stage3_rows.extend(
            [
                [
                    row.get("key", ""),
                    row.get("label", ""),
                    row.get("orderMode", ""),
                    row.get("score", ""),
                    (row.get("planSummary") or {}).get("months", ""),
                    (row.get("planSummary") or {}).get("hitRate", ""),
                    (row.get("planSummary") or {}).get("maxDeficit", ""),
                    (row.get("planSummary") or {}).get("minTonnage", ""),
                    (row.get("planSummary") or {}).get("riskMax", ""),
                ]
                for row in succession_artifacts["stage3Results"]
            ]
        )
        summary.files.append(write_csv(succession_dir / "stage3_candidates.csv", stage3_rows))

        economics_result = compute_economics_from_plan(plan, None, scene.get("economicsParams") or {})
        if economics_result.get("ok"):
            economics_dir = scene_dir / "economics"
            economics_dir.mkdir(parents=True, exist_ok=True)
            for stem_name, figure in [
                ("fig01_cashflow", plot_cashflow(economics_result, f"{summary.scene_name} - Cashflow")),
                ("fig02_revenue_cost", plot_revenue_cost(economics_result, f"{summary.scene_name} - Revenue Cost Structure")),
                ("fig03_cost_structure", plot_cost_structure(economics_result, f"{summary.scene_name} - Cost Structure")),
            ]:
                written = save_figure(figure, figure_stem_path(economics_dir, stem_name), formats)
                if written:
                    summary.files.extend(written)
                    summary.figure_count += 1

            econ_rows = [["month", "workface", "tonnage", "tonnageAdj", "risk", "revenueWan", "varCostWan", "fixedCostWan", "sustainCapexWan", "riskExtraCostWan", "capexInitialWan", "netCashWan", "cumCashWan"]]
            econ_rows.extend(
                [
                    [
                        row.get("month", ""),
                        row.get("workface", ""),
                        row.get("tonnage", ""),
                        row.get("tonnageAdj", ""),
                        row.get("risk", ""),
                        row.get("revenueWan", ""),
                        row.get("varCostWan", ""),
                        row.get("fixedCostWan", ""),
                        row.get("sustainCapexWan", ""),
                        row.get("riskExtraCostWan", ""),
                        row.get("capexInitialWan", ""),
                        row.get("netCashWan", ""),
                        row.get("cumCashWan", ""),
                    ]
                    for row in (economics_result.get("rows") or [])
                ]
            )
            summary.files.append(write_csv(economics_dir / "economics_monthly.csv", econ_rows))
            summary.files.append(write_json(economics_dir / "economics_summary.json", economics_result.get("summary") or {}))

    return summary


def write_index(output_root: Path, summaries: Sequence[ExportSummary]) -> Path:
    lines = [
        "# Scene Visual Export Index",
        "",
        f"Generated: {datetime.now().isoformat()}",
        "",
        "| Scene | Figures | Source |",
        "| --- | ---: | --- |",
    ]
    for summary in summaries:
        lines.append(f"| {summary.scene_slug} | {summary.figure_count} | `{summary.source_file}` |")
    path = output_root / "INDEX.md"
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Batch-export miningplan visuals with vector-first outputs.")
    parser.add_argument("--input", nargs="*", help="Input scene JSON files or directories.")
    parser.add_argument("--output-dir", help="Output directory. Defaults to output/scene_visual_exports/<timestamp>.")
    parser.add_argument("--formats", default="svg,pdf,png", help="Comma-separated formats, default: svg,pdf,png.")
    parser.add_argument("--paper-docx", help="Optional paper DOCX path used to generate a figure insertion guide.")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    output_root = Path(args.output_dir) if args.output_dir else build_default_output_dir()
    formats = tuple(part.strip().lower() for part in args.formats.split(",") if part.strip())
    scene_files = discover_scene_files(args.input)
    output_root.mkdir(parents=True, exist_ok=True)

    summaries = [export_scene_file(scene_file, output_root, formats=formats) for scene_file in scene_files]
    index_path = write_index(output_root, summaries)
    guide_path = write_paper_figure_guide(
        output_root,
        summaries,
        paper_docx_path=Path(args.paper_docx) if args.paper_docx else None,
    )
    guide_docx_path = None
    if args.paper_docx:
        guide_docx_path = write_paper_figure_guide_docx(guide_path, Path(args.paper_docx))

    print(f"Output directory: {output_root}")
    print(f"Scene count: {len(summaries)}")
    for summary in summaries:
        print(f"  - {summary.scene_slug}: {summary.figure_count} figures")
    print(f"Index file: {index_path}")
    print(f"Paper guide: {guide_path}")
    if guide_docx_path:
        print(f"Paper guide DOCX: {guide_docx_path}")


if __name__ == "__main__":
    main()
